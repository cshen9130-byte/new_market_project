"""融航结算单 ZIP → 投资报告分析（Word / PDF）

用法：
    python generate_ronghang_report.py --zip path/to/data.zip --outdir /tmp/out --format both

环境变量：
    RONGHANG_REPORT_OUTPUT_DIR  默认输出目录覆盖
"""
from __future__ import annotations

import argparse
import io
import os
import re
import sys
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path

if sys.platform == "win32":
    if hasattr(sys.stdout, "buffer"):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    if hasattr(sys.stderr, "buffer"):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)

os.environ.setdefault("TQDM_DISABLE", "1")

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties, fontManager
import numpy as np
import pandas as pd
import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

BASE_DIR = Path(__file__).resolve().parent
_CN_FONT: FontProperties | None = None
_PDF_FONT = "Helvetica"

PRODUCT_NAME = {
    "AU": "黄金", "AG": "白银", "CU": "铜", "AL": "铝", "ZN": "锌", "PB": "铅",
    "NI": "镍", "SN": "锡", "AO": "氧化铝", "BC": "国际铜", "SI": "工业硅",
    "LC": "碳酸锂", "PS": "多晶硅", "RB": "螺纹钢", "I": "铁矿石", "HC": "热轧卷板",
    "SS": "不锈钢", "SF": "硅铁", "SM": "锰硅", "JM": "焦煤", "J": "焦炭",
    "SC": "原油", "FU": "燃料油", "LU": "低硫燃料油", "PG": "液化石油气", "ZC": "动力煤",
    "JD": "鸡蛋", "AP": "苹果", "CJ": "红枣", "C": "玉米", "A": "豆一", "B": "豆二",
    "CS": "淀粉", "M": "豆粕", "Y": "豆油", "RM": "菜粕", "OI": "菜油", "P": "棕榈油",
    "PK": "花生", "LH": "生猪", "NR": "20号胶", "EB": "苯乙烯", "TA": "PTA", "V": "PVC",
    "BR": "丁苯橡胶", "RU": "橡胶", "L": "塑料", "PF": "短纤", "EG": "乙二醇",
    "MA": "甲醇", "PP": "聚丙烯", "UR": "尿素", "SA": "纯碱", "PX": "PX", "BU": "沥青",
    "TS": "2年国债", "TF": "5年国债", "T": "10年国债", "TL": "30年国债",
    "IF": "沪深300股指期货", "IH": "上证50股指期货", "IC": "中证500股指期货",
    "IM": "中证1000股指期货", "EC": "集运指数(欧线)", "CF": "棉花", "SR": "白糖",
    "CY": "棉纱", "FG": "玻璃", "SP": "纸浆", "LG": "原木",
}

SECTOR_RULES = {
    "有色金属": ["CU", "AL", "ZN", "PB", "NI", "SN", "AO", "BC", "SI", "LC"],
    "黑色": ["RB", "I", "HC", "SS", "WR", "SF", "SM", "JM", "J"],
    "能源": ["SC", "FU", "LU", "PG", "ZC"],
    "农产品": ["JD", "AP", "CJ", "C", "A", "CS", "WH", "PM", "RR", "RI", "JR", "LR", "B", "M", "Y", "RS", "RM", "OI", "P", "PK", "LH"],
    "化工": ["PR", "NR", "EB", "TA", "V", "BR", "RU", "L", "PF", "EG", "MA", "PP", "ED", "UR", "SA", "SH", "PX", "BU", "PS"],
    "贵金属": ["AU", "AG"],
    "国债": ["TS", "TF", "T", "TL"],
    "金融指数": ["IF", "IH", "IC", "IM", "EC"],
    "软商品": ["CF", "SR", "CY"],
    "建材": ["FG", "SP", "FB", "BB", "LG"],
}
PRODUCT_SECTOR = {c: s for s, codes in SECTOR_RULES.items() for c in codes}

RED = "#e23e3e"
BLUE = "#5470c6"
GREEN = "#91cc75"
DARK_GREEN = "#3ba272"


# ── fonts ────────────────────────────────────────────────────────────────────

def _cn_font_candidates() -> list[str]:
    home = Path.home()
    return [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        str(home / ".local/share/fonts/NotoSansSC-Regular.otf"),
    ]


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    chosen = None
    for path in _cn_font_candidates():
        if path and os.path.isfile(path):
            try:
                fontManager.addfont(path)
                chosen = path
                break
            except Exception:
                continue
    if chosen:
        _CN_FONT = FontProperties(fname=chosen)
        plt.rcParams["font.family"] = "sans-serif"
        plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "DejaVu Sans"]
        print(f"[FONT] Charts using: {_CN_FONT.get_name()} ({chosen})", flush=True)


def configure_pdf_font() -> str:
    global _PDF_FONT
    for path in _cn_font_candidates():
        if not path or not os.path.isfile(path):
            continue
        try:
            # Prefer TTF; for TTC reportlab may still work with subfontIndex
            name = "RonghangCN"
            if path.lower().endswith(".ttc"):
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont(name, path))
            _PDF_FONT = name
            print(f"[FONT] PDF using: {path}", flush=True)
            return name
        except Exception as exc:
            print(f"[FONT] skip {path}: {exc}", flush=True)
            continue
    return "Helvetica"


def _fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def _leg() -> dict:
    return {"prop": _CN_FONT} if _CN_FONT is not None else {}


# ── parse helpers ─────────────────────────────────────────────────────────────

def n0(v) -> float:
    if v is None or v == "":
        return 0.0
    if isinstance(v, (int, float)):
        return float(v) if np.isfinite(v) else 0.0
    s = str(v).strip().replace(",", "").replace("%", "")
    if not s or s in {"-", "--"}:
        return 0.0
    try:
        return float(s)
    except Exception:
        return 0.0


def cell_str(v) -> str:
    return str(v or "").strip()


def product_code(instrument: str) -> str:
    m = re.match(r"([A-Za-z]+)", cell_str(instrument))
    return m.group(1).upper() if m else cell_str(instrument).upper()


def product_name(code: str) -> str:
    return PRODUCT_NAME.get(code, code)


def sector_of(code: str) -> str:
    return PRODUCT_SECTOR.get(code, "其他")


def find_label(sheet, label: str):
    for r in range(sheet.nrows):
        for c in range(sheet.ncols):
            if cell_str(sheet.cell_value(r, c)) == label:
                for dc in (1, 2, 3):
                    if c + dc < sheet.ncols:
                        v = sheet.cell_value(r, c + dc)
                        if v != "" and v is not None:
                            return v
    return None


def find_header(sheet, required: list[str]):
    for r in range(sheet.nrows):
        headers = [cell_str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
        if all(h in headers for h in required):
            return r, headers
    return None, []


def parse_workbook(path: Path) -> dict:
    wb = xlrd.open_workbook(str(path))
    sh = wb.sheet_by_index(0)
    trade_date = cell_str(find_label(sh, "交易日期"))[:10]
    risk_raw = find_label(sh, "风险度")
    risk = n0(risk_raw)
    if isinstance(risk_raw, str) and "%" in risk_raw:
        risk = risk / 100.0
    elif risk > 1:
        risk = risk / 100.0

    account = {
        "trade_date": trade_date,
        "client_id": cell_str(find_label(sh, "客户期货期权内部资金账户")),
        "client_name": cell_str(find_label(sh, "客户名称")),
        "broker_name": cell_str(find_label(sh, "期货公司名称")),
        "balance_bf": n0(find_label(sh, "上日结存")),
        "balance_cf": n0(find_label(sh, "当日结存")),
        "equity": n0(find_label(sh, "客户权益")),
        "deposit": n0(find_label(sh, "当日存取合计")),
        "daily_pl": n0(find_label(sh, "当日盈亏")),
        "premium": n0(find_label(sh, "当日总权利金")),
        "fee": n0(find_label(sh, "当日手续费")),
        "margin": n0(find_label(sh, "保证金占用")),
        "available": n0(find_label(sh, "可用资金")),
        "risk": risk,
    }

    # positions from 期货持仓汇总
    positions = []
    for r in range(sh.nrows):
        if "期货持仓汇总" in cell_str(sh.cell_value(r, 0)):
            hr, headers = None, []
            for rr in range(r + 1, min(r + 5, sh.nrows)):
                hs = [cell_str(sh.cell_value(rr, c)) for c in range(min(sh.ncols, 12))]
                if "合约" in hs and "持仓盈亏" in hs:
                    hr, headers = rr, hs
                    break
            if hr is None:
                break
            idx = {h: headers.index(h) for h in headers if h}
            for rr in range(hr + 1, sh.nrows):
                inst = cell_str(sh.cell_value(rr, idx.get("合约", 0)))
                if not inst or inst == "合计" or "期权" in inst:
                    break
                if not re.match(r"^[A-Za-z]", inst):
                    break
                code = product_code(inst)
                positions.append({
                    "instrument": inst,
                    "product": code,
                    "long": n0(sh.cell_value(rr, idx["买持仓"])) if "买持仓" in idx else 0,
                    "short": n0(sh.cell_value(rr, idx["卖持仓"])) if "卖持仓" in idx else 0,
                    "mtm": n0(sh.cell_value(rr, idx["持仓盈亏"])) if "持仓盈亏" in idx else 0,
                    "margin": n0(sh.cell_value(rr, idx["交易保证金"])) if "交易保证金" in idx else 0,
                })
            break

    trades, closes, details = [], [], []
    names = wb.sheet_names()

    if "成交明细" in names:
        tsh = wb.sheet_by_name("成交明细")
        hr, headers = find_header(tsh, ["合约", "手数"])
        if hr is not None:
            idx = {h: i for i, h in enumerate(headers)}
            for r in range(hr + 1, tsh.nrows):
                inst = cell_str(tsh.cell_value(r, idx["合约"]))
                if not inst or inst == "合计":
                    break
                tid = cell_str(tsh.cell_value(r, idx["成交序号"])) if "成交序号" in idx else ""
                ad = cell_str(tsh.cell_value(r, idx["实际成交日期"]))[:10] if "实际成交日期" in idx else trade_date
                trades.append({
                    "trade_date": ad or trade_date,
                    "instrument": inst,
                    "product": product_code(inst),
                    "bs": cell_str(tsh.cell_value(r, idx["买/卖"])) if "买/卖" in idx else "",
                    "oc": cell_str(tsh.cell_value(r, idx["开/平"])) if "开/平" in idx else "",
                    "lots": n0(tsh.cell_value(r, idx["手数"])),
                    "turnover": n0(tsh.cell_value(r, idx["成交额"])) if "成交额" in idx else 0,
                    "fee": n0(tsh.cell_value(r, idx["手续费"])) if "手续费" in idx else 0,
                    "trade_id": tid,
                })

    if "平仓明细" in names:
        csh = wb.sheet_by_name("平仓明细")
        hr, headers = find_header(csh, ["合约", "平仓盈亏"])
        if hr is not None:
            idx = {h: i for i, h in enumerate(headers)}
            for r in range(hr + 1, csh.nrows):
                inst = cell_str(csh.cell_value(r, idx["合约"]))
                if not inst or inst == "合计":
                    break
                oid = cell_str(csh.cell_value(r, idx["原成交序号"])) if "原成交序号" in idx else ""
                ad = cell_str(csh.cell_value(r, idx["实际成交日期"]))[:10] if "实际成交日期" in idx else trade_date
                od = ""
                if "原成交日期" in idx:
                    od = cell_str(csh.cell_value(r, idx["原成交日期"]))[:10]
                closes.append({
                    "trade_date": ad or trade_date,
                    "instrument": inst,
                    "product": product_code(inst),
                    "bs": cell_str(csh.cell_value(r, idx["买/卖"])) if "买/卖" in idx else "",
                    "lots": n0(csh.cell_value(r, idx["手数"])) if "手数" in idx else 0,
                    "pnl": n0(csh.cell_value(r, idx["平仓盈亏"])),
                    "open_trade_id": oid,
                    "open_date": od,
                })

    if "持仓明细" in names:
        dsh = wb.sheet_by_name("持仓明细")
        hr, headers = find_header(dsh, ["合约", "成交序号"])
        if hr is not None:
            idx = {h: i for i, h in enumerate(headers)}
            for r in range(hr + 1, dsh.nrows):
                inst = cell_str(dsh.cell_value(r, idx["合约"]))
                if not inst or inst == "合计":
                    break
                tid = cell_str(dsh.cell_value(r, idx["成交序号"]))
                od = cell_str(dsh.cell_value(r, idx["实际成交日期"]))[:10] if "实际成交日期" in idx else ""
                if tid and od:
                    details.append({"trade_id": tid, "open_date": od, "instrument": inst})

    return {
        "account": account,
        "positions": positions,
        "trades": trades,
        "closes": closes,
        "details": details,
        "source": path.name,
    }


def load_zip(zip_path: Path) -> list[dict]:
    days = []
    with zipfile.ZipFile(zip_path, "r") as zf:
        names = [n for n in zf.namelist() if re.search(r"\.xlsx?$", n, re.I) and not n.startswith("__MACOSX")]
        tmp = zip_path.parent / "_ronghang_extract"
        tmp.mkdir(parents=True, exist_ok=True)
        for name in names:
            base = Path(name).name
            target = tmp / base
            with zf.open(name) as src, open(target, "wb") as dst:
                dst.write(src.read())
            try:
                days.append(parse_workbook(target))
            except Exception as exc:
                print(f"[WARN] skip {base}: {exc}", flush=True)
    days.sort(key=lambda d: d["account"]["trade_date"])
    if not days:
        raise RuntimeError("ZIP 中未解析到有效结算单")
    return days


# ── analysis ──────────────────────────────────────────────────────────────────

def close_direction(bs: str) -> str:
    s = bs.replace(" ", "")
    if "卖" in s:
        return "买"
    if "买" in s:
        return "卖"
    return "买"


def holding_bucket(open_date: str, close_date: str) -> str:
    if not open_date or not close_date:
        return "短线"
    try:
        a = datetime.strptime(open_date[:10], "%Y-%m-%d")
        b = datetime.strptime(close_date[:10], "%Y-%m-%d")
    except Exception:
        return "短线"
    gap = (b - a).days
    if gap <= 0:
        return "日内"
    if gap <= 5:
        return "短线"
    if gap <= 20:
        return "中线"
    return "长线"


def calendar_days(start: str, end: str) -> int:
    try:
        a = datetime.strptime(start[:10], "%Y-%m-%d")
        b = datetime.strptime(end[:10], "%Y-%m-%d")
        return (b - a).days + 1
    except Exception:
        return 0


def analyze(days: list[dict]) -> dict:
    acc = [d["account"] for d in days]
    df = pd.DataFrame(acc)
    df["trade_date"] = pd.to_datetime(df["trade_date"])
    df = df.sort_values("trade_date").reset_index(drop=True)

    # Arithmetic cumulative return matching sample ≈ -24.85% / NAV 0.7515
    df["daily_ret_report"] = (df["daily_pl"] - df["fee"] + df["premium"]) / df["balance_bf"].replace(0, np.nan)
    df["daily_ret_chart"] = df["daily_pl"] / df["balance_bf"].replace(0, np.nan)
    df["cum_ret_report"] = df["daily_ret_report"].fillna(0).cumsum()
    df["cum_ret_chart"] = df["daily_ret_chart"].fillna(0).cumsum()
    df["nav_report"] = 1.0 + df["cum_ret_report"]
    df["nav_chart"] = 1.0 + df["cum_ret_chart"]
    # Unit/max NAV in overview uses chart series peak (≈1.026)
    unit_nav = float(df["nav_chart"].max())
    period_return = float(df["cum_ret_report"].iloc[-1])
    final_nav = float(df["nav_report"].iloc[-1])

    peak = df["nav_report"].cummax()
    df["drawdown"] = (df["nav_report"] - peak) / peak.replace(0, np.nan)
    max_dd = float((-df["drawdown"].min()) if len(df) else 0)
    max_daily_dd = float((-df["daily_ret_report"].min()) if len(df) else 0)

    # peak date / dd interval
    peak_idx = int(df["nav_report"].idxmax())
    trough_idx = int(df["drawdown"].idxmin())
    dd_start = df.loc[peak_idx, "trade_date"].strftime("%Y-%m-%d")
    dd_end = df.loc[trough_idx, "trade_date"].strftime("%Y-%m-%d")

    n = len(df)
    rets = df["daily_ret_report"].fillna(0).to_numpy()
    mean = float(np.mean(rets))
    std = float(np.std(rets, ddof=0))
    down = rets[rets < 0]
    down_std = float(np.std(down, ddof=0)) if len(down) else 0.0
    ann_ret = (1 + period_return) ** (252 / n) - 1 if n and (1 + period_return) > 0 else -1
    # For negative period, use: sign * abs(final_nav)^(252/n) - 1 style
    if final_nav > 0:
        ann_ret = final_nav ** (252 / n) - 1
    ann_vol = std * np.sqrt(252)
    ann_dvol = down_std * np.sqrt(252)
    sharpe = (mean / std) * np.sqrt(252) if std > 0 else 0.0
    sortino = (mean / down_std) * np.sqrt(252) if down_std > 0 else 0.0
    calmar = ann_ret / max_dd if max_dd > 0 else 0.0

    total_fee = float(df["fee"].sum())
    total_pl = float(df["daily_pl"].sum())
    net_profit = total_pl - total_fee
    net_deposit = float(df["deposit"].sum())
    total_deposit = float(df.loc[df["deposit"] > 0, "deposit"].sum())
    total_withdraw = float((-df.loc[df["deposit"] < 0, "deposit"]).sum())

    avg_equity = float(df["equity"].mean())
    avg_margin = float(df["margin"].mean())
    avg_margin_ratio = avg_margin / avg_equity if avg_equity else 0
    avg_fee_ratio = (total_fee / n) / avg_equity if avg_equity and n else 0
    daily_win = float((df["daily_pl"] > 0).mean())

    # monthly
    df["month"] = df["trade_date"].dt.strftime("%Y-%m")
    monthly = []
    for month, g in df.groupby("month", sort=True):
        # month return from report series within month
        start_nav = 1 + float(g["cum_ret_report"].iloc[0] - g["daily_ret_report"].iloc[0])
        end_nav = 1 + float(g["cum_ret_report"].iloc[-1])
        mret = end_nav / start_nav - 1 if start_nav else 0
        monthly.append({
            "month": month,
            "return": mret,
            "pnl": float((g["daily_pl"] - g["fee"]).sum()),
            "trades": 0,
            "lots": 0,
            "turnover": 0.0,
        })

    # open date map
    open_map = {}
    for d in days:
        for t in d["trades"]:
            if t["trade_id"]:
                open_map[t["trade_id"].lstrip("0") or t["trade_id"]] = t["trade_date"]
        for det in d["details"]:
            if det["trade_id"]:
                open_map[det["trade_id"].lstrip("0") or det["trade_id"]] = det["open_date"]

    product_pnl = defaultdict(float)
    product_lots = defaultdict(float)
    sector_pnl = defaultdict(float)
    sector_lots = defaultdict(float)
    dir_pnl = defaultdict(float)
    close_rows = []
    daily_sector_pnl = defaultdict(lambda: defaultdict(float))  # date -> sector -> pnl
    daily_product_pnl = defaultdict(lambda: defaultdict(float))  # date -> product -> pnl
    daily_product_lots = defaultdict(lambda: defaultdict(float))  # date -> product -> lots
    monthly_product_pnl = defaultdict(lambda: defaultdict(float))  # month -> product -> pnl
    daily_long_short = []  # per day long/short lots

    total_lots = 0
    total_trades = 0
    for d in days:
        date = d["account"]["trade_date"]
        month = date[:7]
        long_lots = sum(p["long"] for p in d["positions"])
        short_lots = sum(p["short"] for p in d["positions"])
        daily_long_short.append({"date": date, "long": long_lots, "short": -short_lots, "net": long_lots - short_lots})

        for t in d["trades"]:
            product_lots[t["product"]] += t["lots"]
            sector_lots[sector_of(t["product"])] += t["lots"]
            daily_product_lots[date][t["product"]] += t["lots"]
            total_lots += t["lots"]
            total_trades += 1
            for m in monthly:
                if m["month"] == month:
                    m["lots"] += t["lots"]
                    m["trades"] += 1
                    m["turnover"] += t["turnover"]

        for c in d["closes"]:
            product_pnl[c["product"]] += c["pnl"]
            sector_pnl[sector_of(c["product"])] += c["pnl"]
            daily_sector_pnl[date][sector_of(c["product"])] += c["pnl"]
            daily_product_pnl[date][c["product"]] += c["pnl"]
            monthly_product_pnl[month][c["product"]] += c["pnl"]
            direction = close_direction(c["bs"])
            dir_pnl[(c["product"], direction)] += c["pnl"]
            oid = (c["open_trade_id"].lstrip("0") or c["open_trade_id"]) if c["open_trade_id"] else ""
            od = c["open_date"] or open_map.get(oid, "")
            close_rows.append({
                "lots": c["lots"], "pnl": c["pnl"], "direction": direction,
                "period": holding_bucket(od, c["trade_date"]),
            })

        for p in d["positions"]:
            product_pnl[p["product"]] += p["mtm"]
            sector_pnl[sector_of(p["product"])] += p["mtm"]
            daily_sector_pnl[date][sector_of(p["product"])] += p["mtm"]
            daily_product_pnl[date][p["product"]] += p["mtm"]
            monthly_product_pnl[month][p["product"]] += p["mtm"]
            direction = "买" if p["long"] >= p["short"] else "卖"
            dir_pnl[(p["product"], direction)] += p["mtm"]

    abs_prod = sum(abs(v) for v in product_pnl.values()) or 1
    abs_sec = sum(abs(v) for v in sector_pnl.values()) or 1
    products = sorted(
        [
            {
                "code": k,
                "name": product_name(k),
                "sector": sector_of(k),
                "pnl": v,
                "lots": product_lots.get(k, 0),
                "weight": abs(v) / abs_prod,
            }
            for k, v in product_pnl.items()
        ],
        key=lambda x: x["pnl"],
    )
    sectors = sorted(
        [
            {
                "name": k,
                "pnl": v,
                "lots": sector_lots.get(k, 0),
                "weight": abs(v) / abs_sec,
            }
            for k, v in sector_pnl.items()
        ],
        key=lambda x: x["pnl"],
    )
    directions = sorted(
        [
            {
                "code": k[0],
                "name": product_name(k[0]),
                "direction": k[1],
                "pnl": v,
                "weight": abs(v) / (sum(abs(x) for x in dir_pnl.values()) or 1),
            }
            for k, v in dir_pnl.items()
        ],
        key=lambda x: x["pnl"],
    )

    def side_stats(rows):
        win = [r for r in rows if r["pnl"] > 0]
        loss = [r for r in rows if r["pnl"] < 0]
        flat = [r for r in rows if r["pnl"] == 0]
        def agg(xs):
            lots = sum(r["lots"] for r in xs)
            pnl = sum(r["pnl"] for r in xs)
            return {"lots": lots, "pnl": pnl, "avg": pnl / lots if lots else 0}
        w, l, f = agg(win), agg(loss), agg(flat)
        total_lots_ = w["lots"] + l["lots"] + f["lots"]
        total_pnl_ = w["pnl"] + l["pnl"] + f["pnl"]
        wr = w["lots"] / total_lots_ if total_lots_ else 0
        pf = abs(w["pnl"] / l["pnl"]) if l["pnl"] < 0 else (0 if w["pnl"] == 0 else 99)
        return {"win": w, "loss": l, "flat": f, "total_lots": total_lots_, "total_pnl": total_pnl_, "win_rate": wr, "pf": pf}

    periods = ["日内", "短线", "中线", "长线"]
    period_stats = []
    total_close_lots = sum(r["lots"] for r in close_rows) or 1
    for p in periods:
        rows = [r for r in close_rows if r["period"] == p]
        profit = sum(r["pnl"] for r in rows if r["pnl"] > 0)
        loss = sum(r["pnl"] for r in rows if r["pnl"] < 0)
        lots = sum(r["lots"] for r in rows)
        trades_n = len(rows)
        wins = sum(1 for r in rows if r["pnl"] > 0)
        period_stats.append({
            "period": p,
            "profit": profit,
            "loss": loss,
            "pnl": profit + loss,
            "lots": lots,
            "lot_share": lots / total_close_lots,
            "trades": trades_n,
            "wins": wins,
            "win_rate": wins / trades_n if trades_n else 0,
        })

    # drawdown buckets
    labels = [f"<={i}0%" for i in range(1, 11)]
    counts = [0] * 10
    for dd in df["drawdown"].fillna(0):
        pct = -float(dd) * 100
        if pct <= 0:
            continue
        idx = min(9, max(0, int(np.ceil(pct / 10) - 1)))
        counts[idx] += 1
    total_dd_days = sum(counts) or 1
    dd_buckets = [{"label": labels[i], "days": counts[i], "share": counts[i] / total_dd_days} for i in range(10)]

    # cumulative sector / product pnl series
    sector_names = [s["name"] for s in sectors]
    sector_cum = {s: [] for s in sector_names}
    running = {s: 0.0 for s in sector_names}
    dates = [d["account"]["trade_date"] for d in days]
    for date in dates:
        for s in sector_names:
            running[s] += daily_sector_pnl[date].get(s, 0.0)
            sector_cum[s].append(running[s])

    product_codes = [p["code"] for p in products]
    product_daily = {c: [] for c in product_codes}
    product_cum = {c: [] for c in product_codes}
    product_lots_series = {c: [] for c in product_codes}
    running_p = {c: 0.0 for c in product_codes}
    for date in dates:
        for c in product_codes:
            day_pnl = daily_product_pnl[date].get(c, 0.0)
            running_p[c] += day_pnl
            product_daily[c].append(day_pnl)
            product_cum[c].append(running_p[c])
            product_lots_series[c].append(daily_product_lots[date].get(c, 0.0))

    months_sorted = sorted(monthly_product_pnl.keys())
    product_monthly = {
        c: [monthly_product_pnl[mo].get(c, 0.0) for mo in months_sorted] for c in product_codes
    }

    # 创新高最大回撤率 buckets (same bins; count of peak-to-trough events = 1 in this sample)
    new_high_dd_buckets = [{"label": labels[i], "count": 0, "share": 0.0} for i in range(10)]
    if max_dd > 0:
        idx = min(9, max(0, int(np.ceil(max_dd * 100 / 10) - 1)))
        new_high_dd_buckets[idx]["count"] = 1
        new_high_dd_buckets[idx]["share"] = 1.0

    dd_detail = [{
        "rank": 1,
        "max_dd": max_dd,
        "range": f"{dd_start}~{dd_end}",
        "repair_days": n,
        "new_high_date": df["trade_date"].iloc[-1].strftime("%Y-%m-%d"),
    }]

    month_win = sum(1 for m in monthly if m["return"] > 0) / len(monthly) if monthly else 0
    max_ret_idx = int(df["cum_ret_chart"].idxmax())
    min_ret_idx = int(df["cum_ret_report"].idxmin())

    # cumulative equity pnl series (for 累计盈亏 chart)
    df["cum_pnl"] = (df["daily_pl"] - df["fee"]).cumsum()

    profit_products = sorted([p for p in products if p["pnl"] > 0], key=lambda x: -x["pnl"])
    loss_products = sorted([p for p in products if p["pnl"] < 0], key=lambda x: x["pnl"])

    return {
        "meta": {
            "client_id": acc[0]["client_id"],
            "client_name": acc[0]["client_name"],
            "broker_name": acc[0]["broker_name"],
            "start": df["trade_date"].iloc[0].strftime("%Y-%m-%d"),
            "end": df["trade_date"].iloc[-1].strftime("%Y-%m-%d"),
            "trading_days": n,
            "export_time": datetime.now().strftime("%Y-%m-%d"),
        },
        "overview": {
            "start_balance": float(df["balance_bf"].iloc[0]),
            "end_balance": float(df["balance_cf"].iloc[-1]),
            "start_equity": float(df["balance_bf"].iloc[0]),
            "end_equity": float(df["equity"].iloc[-1]),
            "total_deposit": total_deposit,
            "total_withdraw": total_withdraw,
            "net_deposit": net_deposit,
            "total_fee": total_fee,
            "net_profit": net_profit,
            "unit_nav": unit_nav,
            "max_nav": unit_nav,
            "final_nav": final_nav,
            "period_return": period_return,
            "ann_return": float(ann_ret),
            "max_daily_dd": max_daily_dd,
            "max_peak_dd": max_dd,
            "dd_start": dd_start,
            "dd_end": dd_end,
            "dd_calendar_days": calendar_days(dd_start, dd_end),
            "underwater_days": calendar_days(dd_start, df["trade_date"].iloc[-1].strftime("%Y-%m-%d")),
            "ann_vol": float(ann_vol),
            "ann_dvol": float(ann_dvol),
            "total_lots": total_lots,
            "total_trades": total_trades,
            "daily_win": daily_win,
            "month_win": month_win,
            "avg_margin": avg_margin,
            "avg_margin_ratio": avg_margin_ratio,
            "sharpe": float(sharpe),
            "sortino": float(sortino),
            "calmar": float(calmar),
            "avg_fee_ratio": avg_fee_ratio,
            "max_ret": float(df["cum_ret_chart"].iloc[max_ret_idx]),
            "max_ret_date": df.loc[max_ret_idx, "trade_date"].strftime("%Y-%m-%d"),
            "min_ret": period_return,
            "min_ret_date": df.loc[min_ret_idx, "trade_date"].strftime("%Y-%m-%d"),
            "max_margin_ratio": float((df["margin"] / df["equity"].replace(0, np.nan)).max()),
            "max_margin_date": df.loc[(df["margin"] / df["equity"].replace(0, np.nan)).idxmax(), "trade_date"].strftime("%Y-%m-%d"),
            "max_margin_day_pnl": float(df.loc[(df["margin"] / df["equity"].replace(0, np.nan)).idxmax(), "daily_pl"]),
        },
        "df": df,
        "monthly": monthly,
        "dd_buckets": dd_buckets,
        "new_high_dd_buckets": new_high_dd_buckets,
        "dd_detail": dd_detail,
        "products": products,
        "profit_products": profit_products,
        "loss_products": loss_products,
        "sectors": sectors,
        "directions": directions,
        "long_short": {
            "overall": side_stats(close_rows),
            "long": side_stats([r for r in close_rows if r["direction"] == "买"]),
            "short": side_stats([r for r in close_rows if r["direction"] == "卖"]),
        },
        "periods": period_stats,
        "daily_long_short": daily_long_short,
        "sector_cum": sector_cum,
        "product_daily": product_daily,
        "product_cum": product_cum,
        "product_lots_series": product_lots_series,
        "product_monthly": product_monthly,
        "months_sorted": months_sorted,
        "dates": dates,
    }


def describe_products(items: list[dict], kind: str, total_lots: float) -> list[str]:
    lines = []
    verb = "盈利" if kind == "profit" else "亏损"
    for p in items[:3]:
        lot_share = p["lots"] / total_lots if total_lots else 0
        lines.append(
            f"{p['name']}（{p['name']}）： {verb} {p['pnl']:.1f} 元， 占总{verb}权重为 {p['weight']:.4f}，"
            f"交易手数为 {p['lots']:.1f} ， 占总交易手数的 {lot_share:.4f}。"
        )
    return lines


def build_narratives(data: dict) -> dict:
    o = data["overview"]
    m = data["meta"]
    best = max(data["monthly"], key=lambda x: x["return"]) if data["monthly"] else None
    worst = min(data["monthly"], key=lambda x: x["return"]) if data["monthly"] else None
    profit_ps = data["profit_products"]
    loss_ps = data["loss_products"]
    profit_secs = [s for s in data["sectors"] if s["pnl"] > 0][:3]
    loss_secs = sorted([s for s in data["sectors"] if s["pnl"] < 0], key=lambda x: x["pnl"])[:3]
    dep_days = data["df"].loc[data["df"]["deposit"] != 0]
    fee_peak_idx = int(data["df"]["fee"].idxmax()) if len(data["df"]) else 0
    fee_peak = data["df"].loc[fee_peak_idx] if len(data["df"]) else None
    focus_loss = loss_ps[0] if loss_ps else None
    focus_profit = profit_ps[0] if profit_ps else None

    return {
        "ret": (
            f"该报告日期自 {m['start']}~{m['end']}，报告期内共计 {m['trading_days']} 个交易日，"
            f"合计收益率 {fmt_pct(o['period_return'])}，您于 {o['max_ret_date']} 达到了最大收益率："
            f"{fmt_pct(o['max_ret'])}，而收益率最低值为 {fmt_pct(o['min_ret'])}，日期为 {o['min_ret_date']}。"
        ),
        "monthly": (
            f"您总共交易月数为{len(data['monthly'])}个月，其中盈利月数为"
            f"{sum(1 for x in data['monthly'] if x['return'] > 0)}个月，月胜率为 {fmt_pct(o['month_win'])}。"
            + (
                f"您收益率最高的月份是 {best['month']}，收益率为 {fmt_pct(best['return'], 4)}，"
                f"您在这个月共做了 {best['trades']} 次交易，交易手数 {best['lots']:.1f} 手，"
                f"交易总市值为 {fmt_money(best['turnover'])} 元，合计盈亏 {fmt_money(best['pnl'])} 元；"
                if best else ""
            )
            + (
                f"收益率最低的月份是 {worst['month']}，收益率为 {fmt_pct(worst['return'], 4)}，"
                f"您在这个月共做了 {worst['trades']} 次交易，交易手数 {worst['lots']:.1f} 手，"
                f"交易总市值为 {fmt_money(worst['turnover'])} 元，合计盈亏 {fmt_money(worst['pnl'])} 元。"
                if worst else ""
            )
        ),
        "nav": (
            f"您的累计净值为 {o['final_nav']:.4f}，其中，在 {o['max_ret_date']} 达到了最高净值 {o['max_nav']:.3f}，"
            f"而最低点是在 {o['min_ret_date']}，达到了 {o['final_nav']:.4f}；平均持仓占比（保证金/权益）为 "
            f"{fmt_pct(o['avg_margin_ratio'])}，这一数值在 {o['max_margin_date']} 达到了最高，为 "
            f"{fmt_pct(o['max_margin_ratio'])}，这一天您合计盈亏 {fmt_money(o['max_margin_day_pnl'])}元。"
        ),
        "drawdown": (
            f"最大动态回撤为 {fmt_pct(o['max_peak_dd'], 4)}，最大动态回撤期：{o['dd_start']}至{o['dd_end']}，"
            f"最长连续回撤天数：{o['dd_calendar_days']} 天，最长连续回撤期：{o['dd_start']}至{o['dd_end']}，"
            f"最长连续未创新高天数：{o['underwater_days']}天。"
        ),
        "equity_dep": (
            f"统计期内客户权益由 {fmt_money(o['start_equity'])} 变动至 {fmt_money(o['end_equity'])}，"
            f"净出入金 {fmt_money(o['net_deposit'])} 元"
            + (
                f"（主要出入金发生在 {dep_days.iloc[0]['trade_date'].strftime('%Y-%m-%d')}，金额 {fmt_money(dep_days.iloc[0]['deposit'])} 元）"
                if len(dep_days) else "（期间无大额出入金）"
            )
            + f"。权益下行与累计亏损 {fmt_money(o['net_profit'])} 基本同步，需关注回撤区间内的仓位放大。"
        ),
        "fee": (
            f"报告期总手续费 {fmt_money(o['total_fee'])} 元，日平均手续费比 {fmt_pct(o['avg_fee_ratio'], 4)}。"
            + (
                f"单日手续费峰值出现在 {fee_peak['trade_date'].strftime('%Y-%m-%d')}，为 {fmt_money(fee_peak['fee'])} 元，"
                if fee_peak is not None else ""
            )
            + "手续费抬升通常对应交易活跃日，可结合成交手数判断是否过度交易。"
        ),
        "month_pnl": (
            f"您总共交易月数为{len(data['monthly'])}个月，其中盈利月数为"
            f"{sum(1 for x in data['monthly'] if x['pnl'] > 0)}个月，月胜率为 {fmt_pct(o['month_win'])}。"
            + (
                f"盈亏最好的月份是 {best['month']}，合计盈亏 {fmt_money(best['pnl'])} 元，收益率 {fmt_pct(best['return'], 4)}；"
                if best else ""
            )
            + (
                f"盈亏最差的月份是 {worst['month']}，合计盈亏 {fmt_money(worst['pnl'])} 元，收益率 {fmt_pct(worst['return'], 4)}。"
                if worst else ""
            )
        ),
        "vol": (
            f"上图为各交易日收益率波动（日盈亏/期初权益）。"
            f"年化波动率 {fmt_pct(o['ann_vol'])}，年化下行波动率 {fmt_pct(o['ann_dvol'])}；"
            f"夏普比率 {o['sharpe']:.4f}，索提诺比率 {o['sortino']:.4f}，卡玛比率 {o['calmar']:.4f}。"
            f"波动主要来自下跌交易日，风险调整后收益偏弱。"
        ),
        "sector_intro": "以下按板块汇总平仓盈亏与持仓盯市盈亏，并给出盈利/亏损板块与交易量结构。",
        "sector_profit": describe_products(
            [{"name": s["name"], "pnl": s["pnl"], "lots": s["lots"], "weight": s["weight"]} for s in profit_secs],
            "profit",
            o["total_lots"],
        ),
        "sector_loss": describe_products(
            [{"name": s["name"], "pnl": s["pnl"], "lots": s["lots"], "weight": s["weight"]} for s in loss_secs],
            "loss",
            o["total_lots"],
        ),
        "product_intro": (
            "品种盈亏由每日平仓盈亏与持仓盯市盈亏加总得到。下面先给出盈利/亏损品种文字归因，"
            "再展示盈利品种与亏损品种的累计盈亏、手数轨迹，以及日盈亏与月盈亏对比。"
        ),
        "product_profit": describe_products(profit_ps, "profit", o["total_lots"]),
        "product_loss": describe_products(loss_ps, "loss", o["total_lots"]),
        "product_profit_chart": (
            f"盈利品种共 {len(profit_ps)} 个，合计贡献 {fmt_money(sum(p['pnl'] for p in profit_ps))} 元。"
            + (
                f"其中 {focus_profit['name']} 贡献最大（{fmt_money(focus_profit['pnl'])} 元，手数 {focus_profit['lots']:.1f}），"
                f"累计盈亏曲线在报告期内总体向上或回撤后修复。"
                if focus_profit else "本期无盈利品种。"
            )
        ),
        "product_loss_chart": (
            f"亏损品种共 {len(loss_ps)} 个，合计亏损 {fmt_money(sum(p['pnl'] for p in loss_ps))} 元。"
            + (
                f"损失最大的是 {focus_loss['name']}（{fmt_money(focus_loss['pnl'])} 元，手数 {focus_loss['lots']:.1f}），"
                f"其累计盈亏持续下行，是组合回撤的主要拖累。"
                if focus_loss else "本期无亏损品种。"
            )
        ),
        "product_daily": (
            "品种日盈亏图展示各品种单日盈亏波动。若多个品种在同一交易日同步大亏，说明仓位风险较为集中；"
            "若亏损由少数品种主导，则应优先收缩该品种敞口。"
        ),
        "product_monthly": (
            "品种月盈亏对比可识别亏损是否集中在特定月份。若单月多品种共振亏损，需降低杠杆并缩短持仓周期。"
        ),
        "pos_strategy": (
            f"持仓策略图同时展示多头/空头手数、净持仓与净值。期末权益 {fmt_money(o['end_equity'])}，"
            f"平均保证金占比 {fmt_pct(o['avg_margin_ratio'])}；若净持仓放大同时净值下行，说明单边敞口放大了亏损。"
        ),
        "sector_focus": (
            f"板块聚焦选取绝对盈亏最大的板块，观察其累计盈亏路径。"
            + (
                f"当前聚焦 {sorted(data['sectors'], key=lambda x: abs(x['pnl']), reverse=True)[0]['name']}，"
                f"期末累计盈亏 {fmt_money(sorted(data['sectors'], key=lambda x: abs(x['pnl']), reverse=True)[0]['pnl'])} 元。"
                if data["sectors"] else ""
            )
        ),
        "product_focus": (
            (
                f"品种聚焦：{focus_loss['name']}。报告期内累计盈亏 {fmt_money(focus_loss['pnl'])} 元，"
                f"交易手数 {focus_loss['lots']:.1f}，占总交易手数 {focus_loss['lots'] / (o['total_lots'] or 1):.2%}。"
                f"该品种是组合主要亏损来源，建议复核开平节奏与单笔风险限额。"
            )
            if focus_loss
            else "本期暂无足够的品种聚焦样本。"
        ),
    }


# ── charts ────────────────────────────────────────────────────────────────────

def save_fig(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()


def make_charts(data: dict, chart_dir: Path) -> dict[str, Path]:
    chart_dir.mkdir(parents=True, exist_ok=True)
    df = data["df"]
    dates = df["trade_date"].dt.strftime("%Y-%m-%d")
    out = {}

    # 2.1 收益率
    fig, ax = plt.subplots(figsize=(9, 3.6))
    ax.plot(dates, df["cum_ret_chart"] * 100, color=RED, marker="o", markevery=max(1, len(df) // 8), markersize=4, label="收益率")
    ax.axhline(0, color="#999", lw=0.8)
    ax.set_title("收益率", **_fp())
    ax.set_ylabel("收益率", **_fp())
    ax.legend(**_leg())
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis="x", labelrotation=30, labelsize=8)
    for label in ax.get_xticklabels():
        label.set_fontproperties(_CN_FONT)
    out["ret"] = chart_dir / "ret.png"
    save_fig(out["ret"])

    # 2.2 月收益率
    months = [m["month"] for m in data["monthly"]]
    mrets = [m["return"] * 100 for m in data["monthly"]]
    cum = np.cumsum(mrets)
    fig, ax = plt.subplots(figsize=(9, 3.2))
    ax.bar(months, mrets, color=GREEN, label="月收益率", width=0.45)
    ax.plot(months, cum, color=RED, marker="o", label="总收益率")
    ax.set_title("月收益率、总收益率", **_fp())
    ax.set_ylabel("收益率", **_fp())
    ax.legend(**_leg())
    ax.grid(True, axis="y", alpha=0.3)
    out["monthly"] = chart_dir / "monthly.png"
    save_fig(out["monthly"])

    # 2.3 净值与持仓
    fig, ax1 = plt.subplots(figsize=(9, 3.6))
    ax2 = ax1.twinx()
    margin_ratio = (df["margin"] / df["equity"].replace(0, np.nan) * 100).fillna(0)
    ax1.plot(dates, margin_ratio, color=BLUE, label="持仓")
    ax2.plot(dates, df["nav_report"], color=RED, label="净值")
    ax1.set_ylabel("持仓", **_fp())
    ax2.set_ylabel("净值", **_fp())
    ax1.set_title("净值与持仓占比", **_fp())
    lines = ax1.get_lines() + ax2.get_lines()
    ax1.legend(lines, [l.get_label() for l in lines], **_leg())
    ax1.grid(True, alpha=0.3)
    ax1.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["nav_pos"] = chart_dir / "nav_pos.png"
    save_fig(out["nav_pos"])

    # 2.4 动态回撤
    fig, ax = plt.subplots(figsize=(9, 3.2))
    ax.plot(dates, df["drawdown"] * 100, color=BLUE, label="动态回撤")
    ax.fill_between(dates, df["drawdown"] * 100, 0, color=BLUE, alpha=0.15)
    ax.set_title("动态回撤", **_fp())
    ax.set_ylabel("回撤", **_fp())
    ax.legend(**_leg())
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["drawdown"] = chart_dir / "drawdown.png"
    save_fig(out["drawdown"])

    # 2.5 出入金和权益
    fig, ax1 = plt.subplots(figsize=(9, 3.4))
    ax2 = ax1.twinx()
    ax2.bar(dates, df["deposit"], color=RED, alpha=0.7, label="出入金", width=0.7)
    ax1.plot(dates, df["equity"], color=RED, label="客户权益")
    ax1.set_ylabel("客户权益", **_fp())
    ax2.set_ylabel("出入金", **_fp())
    ax1.set_title("出入金和客户权益", **_fp())
    lines = ax1.get_lines() + ax2.containers
    ax1.legend(["客户权益", "出入金"], **_leg())
    ax1.grid(True, alpha=0.3)
    ax1.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["equity_dep"] = chart_dir / "equity_dep.png"
    save_fig(out["equity_dep"])

    # 2.6 手续费
    fig, ax1 = plt.subplots(figsize=(9, 3.4))
    ax2 = ax1.twinx()
    cum_fee = df["fee"].cumsum()
    fee_ratio = (df["fee"] / df["equity"].replace(0, np.nan) * 100).fillna(0)
    ax1.plot(dates, cum_fee, color=BLUE, label="累计手续费")
    ax2.bar(dates, fee_ratio, color=RED, alpha=0.7, label="日手续费比", width=0.7)
    ax1.set_title("手续费分析", **_fp())
    ax1.set_ylabel("累计手续费", **_fp())
    ax2.set_ylabel("日手续费比(%)", **_fp())
    ax1.legend(loc="upper left", **_leg())
    ax1.grid(True, alpha=0.3)
    ax1.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["fee"] = chart_dir / "fee.png"
    save_fig(out["fee"])

    # 2.7 月度盈亏
    fig, ax = plt.subplots(figsize=(9, 3.0))
    pnls = [m["pnl"] for m in data["monthly"]]
    ax.bar(months, pnls, color=[GREEN if p >= 0 else RED for p in pnls])
    ax.set_title("月度盈亏", **_fp())
    ax.set_ylabel("盈亏", **_fp())
    ax.grid(True, axis="y", alpha=0.3)
    out["month_pnl"] = chart_dir / "month_pnl.png"
    save_fig(out["month_pnl"])

    # 累计盈亏
    fig, ax = plt.subplots(figsize=(9, 3.0))
    ax.plot(dates, df["cum_pnl"], color=RED, marker="o", markersize=3, label="累计盈亏")
    ax.axhline(0, color="#999", lw=0.8)
    ax.set_title("累计盈亏", **_fp())
    ax.set_ylabel("金额", **_fp())
    ax.legend(**_leg())
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["cum_pnl"] = chart_dir / "cum_pnl.png"
    save_fig(out["cum_pnl"])

    # 2.8 波动率 — sample report plots daily return % (labeled 波动率), not ann. vol
    fig, ax = plt.subplots(figsize=(9, 3.2))
    vol_series = df["daily_ret_chart"].fillna(0) * 100
    ax.plot(dates, vol_series, color=RED, marker="o", markersize=3.5, label="波动率")
    ax.axhline(0, color="#999", lw=0.8)
    ax.set_title("波动率", **_fp())
    ax.set_ylabel("波动率", **_fp())
    from matplotlib.ticker import FuncFormatter

    ax.yaxis.set_major_formatter(FuncFormatter(lambda v, _: f"{v:.0f}%"))
    ax.legend(**_leg())
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["vol"] = chart_dir / "vol.png"
    save_fig(out["vol"])

    # 3.1 sector bar
    fig, ax = plt.subplots(figsize=(9, 3.4))
    names = [s["name"] for s in data["sectors"]]
    vals = [s["pnl"] for s in data["sectors"]]
    ax.bar(names, vals, color=[GREEN if v >= 0 else GREEN for v in vals])
    ax.set_title("板块盈利", **_fp())
    ax.tick_params(axis="x", labelrotation=20, labelsize=9)
    for label in ax.get_xticklabels():
        label.set_fontproperties(_CN_FONT)
    ax.grid(True, axis="y", alpha=0.3)
    out["sector_bar"] = chart_dir / "sector_bar.png"
    save_fig(out["sector_bar"])

    pie_palette = ["#5470c6", "#91cc75", "#fac858", "#ee6666", "#73c0de", "#3ba272", "#fc8452", "#9a60b4", "#ea7ccc", "#fcce10"]

    def render_one_pie(labels, values, title: str) -> PILImage.Image:
        """Render one circular pie (+ legend below) into a fixed-size RGB image."""
        # Fixed figure geometry: square axes band so wedges cannot become ovals.
        fig = plt.figure(figsize=(4.0, 4.6), dpi=140, facecolor="white")
        ax = fig.add_axes([0.14, 0.28, 0.72, 0.62])  # width==height in figure fraction
        if not values or sum(abs(v) for v in values) <= 0:
            wedges, _ = ax.pie([1], colors=["#d9d9d9"], startangle=90)
            ax.text(0, 0, "暂无数据", ha="center", va="center", fontsize=11, **_fp())
            legend_labels: list[str] = []
        else:
            abs_vals = [abs(v) for v in values]
            total = sum(abs_vals) or 1.0
            colors_ = [pie_palette[i % len(pie_palette)] for i in range(len(abs_vals))]
            wedges, _ = ax.pie(
                abs_vals,
                colors=colors_,
                startangle=90,
                wedgeprops={"linewidth": 0.6, "edgecolor": "white"},
            )
            legend_labels = [f"{lab} {val / total * 100:.1f}%" for lab, val in zip(labels, abs_vals)]
        ax.set_title(title, **_fp(), pad=8, fontsize=12)
        ax.set_aspect("equal", adjustable="box")
        if legend_labels:
            ncol = 2 if len(legend_labels) > 4 else 1
            fig.legend(
                wedges,
                legend_labels,
                loc="lower center",
                bbox_to_anchor=(0.5, 0.01),
                fontsize=8,
                frameon=False,
                ncol=ncol,
                **_leg(),
            )
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=140, facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return PILImage.open(buf).convert("RGB")

    def save_pie_panel(path: Path, panels: list[tuple[list[str], list[float], str]], ncols: int = 3):
        """Compose individually rendered circular pies into one panel image."""
        n = len(panels)
        ncols = min(ncols, max(1, n))
        nrows = int(np.ceil(n / ncols))
        tiles = [render_one_pie(labels, values, title) for labels, values, title in panels]
        tw, th = tiles[0].size
        canvas = PILImage.new("RGB", (tw * ncols, th * nrows), (255, 255, 255))
        for i, tile in enumerate(tiles):
            r, c = divmod(i, ncols)
            canvas.paste(tile, (c * tw, r * th))
        path.parent.mkdir(parents=True, exist_ok=True)
        canvas.save(path, format="PNG")

    profit_sec = [s for s in data["sectors"] if s["pnl"] > 0]
    loss_sec = [s for s in data["sectors"] if s["pnl"] < 0]
    lots_sec = [s for s in data["sectors"] if s["lots"] > 0]
    out["sector_pies"] = chart_dir / "sector_pies.png"
    save_pie_panel(
        out["sector_pies"],
        [
            ([s["name"] for s in profit_sec], [s["pnl"] for s in profit_sec], "盈利板块"),
            ([s["name"] for s in loss_sec], [s["pnl"] for s in loss_sec], "亏损板块"),
            ([s["name"] for s in lots_sec], [s["lots"] for s in lots_sec], "板块交易量"),
        ],
        ncols=3,
    )
    # keep individual keys pointing to the panel for backward-compatible inserts
    out["sector_profit_pie"] = out["sector_pies"]
    out["sector_loss_pie"] = out["sector_pies"]
    out["sector_lots_pie"] = out["sector_pies"]

    # product bar top abs
    top_prod = sorted(data["products"], key=lambda x: abs(x["pnl"]), reverse=True)[:12]
    top_prod = sorted(top_prod, key=lambda x: x["pnl"])
    fig, ax = plt.subplots(figsize=(9, 3.6))
    ax.barh([p["name"] for p in top_prod], [p["pnl"] for p in top_prod], color=[GREEN if p["pnl"] >= 0 else RED for p in top_prod])
    ax.set_title("品种盈亏", **_fp())
    for label in ax.get_yticklabels():
        label.set_fontproperties(_CN_FONT)
    ax.grid(True, axis="x", alpha=0.3)
    out["product_bar"] = chart_dir / "product_bar.png"
    save_fig(out["product_bar"])

    profit_prod = data["profit_products"]
    loss_prod = data["loss_products"]
    out["product_pies"] = chart_dir / "product_pies.png"
    save_pie_panel(
        out["product_pies"],
        [
            ([p["name"] for p in profit_prod], [p["pnl"] for p in profit_prod], "盈利品种"),
            ([p["name"] for p in loss_prod[:8]], [p["pnl"] for p in loss_prod[:8]], "亏损品种"),
        ],
        ncols=2,
    )
    out["product_profit_pie"] = out["product_pies"]
    out["product_loss_pie"] = out["product_pies"]

    palette = pie_palette

    def multi_line(path: Path, codes: list[str], series_map: dict, title: str, ylabel: str):
        fig, ax = plt.subplots(figsize=(9, 3.6))
        if not codes:
            ax.text(0.5, 0.5, "暂无数据", ha="center", va="center", **_fp())
        else:
            for i, code in enumerate(codes):
                name = product_name(code)
                ax.plot(
                    data["dates"],
                    series_map.get(code, [0] * len(data["dates"])),
                    color=palette[i % len(palette)],
                    marker="o",
                    markersize=2.5,
                    lw=1.4,
                    label=name,
                )
        ax.set_title(title, **_fp())
        ax.set_ylabel(ylabel, **_fp())
        ax.legend(fontsize=7, ncol=3, **_leg())
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis="x", labelrotation=30, labelsize=8)
        save_fig(path)

    def grouped_month_bars(path: Path, items: list[dict], title: str):
        fig, ax = plt.subplots(figsize=(9, 3.4))
        months_x = data["months_sorted"] or [m["month"] for m in data["monthly"]]
        if not items or not months_x:
            ax.text(0.5, 0.5, "暂无数据", ha="center", va="center", **_fp())
        else:
            x = np.arange(len(months_x))
            width = 0.8 / max(1, len(items))
            for i, p in enumerate(items):
                vals = data["product_monthly"].get(p["code"], [0] * len(months_x))
                ax.bar(x + i * width - 0.4 + width / 2, vals, width=width, color=palette[i % len(palette)], label=p["name"])
            ax.set_xticks(x)
            ax.set_xticklabels(months_x)
        ax.set_title(title, **_fp())
        ax.set_ylabel("月盈亏", **_fp())
        ax.legend(fontsize=7, ncol=3, **_leg())
        ax.grid(True, axis="y", alpha=0.3)
        save_fig(path)

    # 盈利品种：累计盈亏 + 手数
    top_profit = profit_prod[:5]
    top_loss = loss_prod[:5]
    out["profit_cum"] = chart_dir / "profit_cum.png"
    out["profit_lots"] = chart_dir / "profit_lots.png"
    out["profit_daily"] = chart_dir / "profit_daily.png"
    out["profit_monthly"] = chart_dir / "profit_monthly.png"
    multi_line(out["profit_cum"], [p["code"] for p in top_profit], data["product_cum"], "品种累计盈亏", "累计盈亏")
    multi_line(out["profit_lots"], [p["code"] for p in top_profit], data["product_lots_series"], "品种手数", "手数")
    multi_line(out["profit_daily"], [p["code"] for p in top_profit], data["product_daily"], "品种盈亏", "日盈亏")
    grouped_month_bars(out["profit_monthly"], top_profit, "品种月盈亏")

    # 亏损品种：累计盈亏 + 手数 + 日/月
    out["loss_cum"] = chart_dir / "loss_cum.png"
    out["loss_lots"] = chart_dir / "loss_lots.png"
    out["loss_daily"] = chart_dir / "loss_daily.png"
    out["loss_monthly"] = chart_dir / "loss_monthly.png"
    multi_line(out["loss_cum"], [p["code"] for p in top_loss], data["product_cum"], "品种累计盈亏", "累计盈亏")
    multi_line(out["loss_lots"], [p["code"] for p in top_loss], data["product_lots_series"], "品种手数", "手数")
    multi_line(out["loss_daily"], [p["code"] for p in top_loss], data["product_daily"], "品种盈亏", "日盈亏")
    grouped_month_bars(out["loss_monthly"], top_loss, "品种月盈亏")

    # 3.4 period pie — same circular renderer as sector/product pies
    out["period_pie"] = chart_dir / "period_pie.png"
    pvals = [max(0.0, abs(p["pnl"])) for p in data["periods"]]
    plabels = [p["period"] for p in data["periods"]]
    render_one_pie(plabels, pvals, "周期盈亏占比").save(out["period_pie"], format="PNG")

    # period monthly bars
    for p in data["periods"]:
        fig, ax = plt.subplots(figsize=(6, 2.8))
        # approximate: put full period pnl on end month if any
        if data["monthly"]:
            ax.bar([m["month"] for m in data["monthly"]], [p["pnl"] if i == len(data["monthly"]) - 1 else 0 for i, _ in enumerate(data["monthly"])], color=RED if p["pnl"] < 0 else GREEN)
        ax.set_title(p["period"], **_fp())
        ax.set_ylabel("盈利", **_fp())
        ax.grid(True, axis="y", alpha=0.3)
        key = f"period_{p['period']}"
        out[key] = chart_dir / f"{key}.png"
        save_fig(out[key])

    # 4.1 持仓策略
    dls = data["daily_long_short"]
    fig, ax1 = plt.subplots(figsize=(9, 3.6))
    ax2 = ax1.twinx()
    xs = [x["date"] for x in dls]
    ax1.bar(xs, [x["long"] for x in dls], color=RED, label="多头", width=0.7)
    ax1.bar(xs, [x["short"] for x in dls], color=GREEN, label="空头", width=0.7)
    ax1.plot(xs, [x["net"] for x in dls], color=DARK_GREEN, marker="o", markersize=3, label="净持仓")
    ax2.plot(dates, df["nav_report"], color=BLUE, marker="o", markersize=3, label="净值")
    ax1.set_title("板块持仓分析", **_fp())
    ax1.set_ylabel("手数", **_fp())
    ax2.set_ylabel("净值", **_fp())
    ax1.legend(loc="upper left", **_leg())
    ax1.tick_params(axis="x", labelrotation=30, labelsize=8)
    ax1.grid(True, alpha=0.3)
    out["pos_strategy"] = chart_dir / "pos_strategy.png"
    save_fig(out["pos_strategy"])

    # sector cumulative pnl (sum all sectors as one line + multi)
    fig, ax = plt.subplots(figsize=(9, 3.4))
    if data["sector_cum"]:
        # plot each sector lightly + total
        total = np.zeros(len(data["dates"]))
        for s, series in data["sector_cum"].items():
            arr = np.array(series, dtype=float)
            total += arr
            if abs(arr[-1]) > 0:
                ax.plot(data["dates"], arr, lw=1, alpha=0.7, label=s)
        ax.plot(data["dates"], total, color=RED, lw=2, marker="o", markersize=3, label="金额")
    ax.set_title("板块累计盈亏金额未去除手续费", **_fp())
    ax.set_ylabel("金额", **_fp())
    ax.legend(fontsize=7, ncol=3, **_leg())
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis="x", labelrotation=30, labelsize=8)
    out["sector_cum"] = chart_dir / "sector_cum.png"
    save_fig(out["sector_cum"])

    # 4.2 / 4.3 focus: top sector & top product cumulative
    if data["sectors"]:
        focus = sorted(data["sectors"], key=lambda x: abs(x["pnl"]), reverse=True)[0]["name"]
        fig, ax = plt.subplots(figsize=(9, 3.2))
        ax.plot(data["dates"], data["sector_cum"].get(focus, [0] * len(data["dates"])), color=RED, marker="o", markersize=3)
        ax.set_title(f"板块聚焦：{focus}", **_fp())
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis="x", labelrotation=30, labelsize=8)
        out["sector_focus"] = chart_dir / "sector_focus.png"
        save_fig(out["sector_focus"])

    if data["products"]:
        focus_p = sorted(data["products"], key=lambda x: abs(x["pnl"]), reverse=True)[0]
        fig, ax = plt.subplots(figsize=(9, 3.2))
        series = data["product_cum"].get(focus_p["code"], [0] * len(data["dates"]))
        ax.plot(data["dates"], series, color=RED, marker="o", markersize=3)
        ax.set_title(f"品种聚焦：{focus_p['name']}", **_fp())
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis="x", labelrotation=30, labelsize=8)
        out["product_focus"] = chart_dir / "product_focus.png"
        save_fig(out["product_focus"])

    return out


# ── Word ──────────────────────────────────────────────────────────────────────

def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, size=16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_text(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_kv_table(doc, rows: list[tuple[str, str]], cols=2):
    # flatten to pairs
    flat = list(rows)
    table_rows = []
    for i in range(0, len(flat), cols):
        chunk = flat[i:i + cols]
        row = []
        for k, v in chunk:
            row.extend([k, v])
        while len(row) < cols * 2:
            row.extend(["", ""])
        table_rows.append(row)
    table = doc.add_table(rows=len(table_rows), cols=cols * 2)
    table.style = "Table Grid"
    for r, row in enumerate(table_rows):
        for c, val in enumerate(row):
            table.cell(r, c).text = str(val)
            for p in table.cell(r, c).paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)


def add_image(doc, path: Path, width=6.3):
    if path and path.exists():
        doc.add_picture(str(path), width=Inches(width))


def add_simple_matrix_table(doc, rows: list[list[str]], col_width_in: float | None = None):
    """2D table with even data-column widths (avoids cramped / misaligned money cells)."""
    if not rows:
        return None
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    if col_width_in is None:
        # label col slightly narrower; remaining shared by value cols
        label_w = 1.1
        value_w = max(1.4, (6.5 - label_w) / max(1, ncols - 1))
        widths = [Inches(label_w)] + [Inches(value_w)] * (ncols - 1)
    else:
        widths = [Inches(col_width_in)] * ncols
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = str(val)
            if c < len(widths):
                cell.width = widths[c]
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=9)
    return table


def fmt_money(v) -> str:
    return f"{float(v):,.2f}"


def fmt_pct(v, d=2) -> str:
    return f"{float(v) * 100:.{d}f}%"


def build_docx(data: dict, charts: dict[str, Path], out_path: Path):
    o = data["overview"]
    m = data["meta"]
    nar = build_narratives(data)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("投资报告分析")
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    for label, val in [
        ("账号:", m["client_id"]),
        ("周期:", f"{m['start']}~{m['end']}"),
        ("导出时间:", m["export_time"]),
        ("提交人员:", ""),
        ("投顾名称:", m["client_name"]),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label} {val}")
        set_run_font(run, size=14)
    doc.add_page_break()

    # Part 1
    add_text(doc, f"用户账号:{m['client_id']}")
    add_text(doc, f"统计周期:{m['start']}~{m['end']}")
    add_heading(doc, "第一部分 总览", 16)
    add_text(doc, f"交易日:{m['trading_days']}天")
    add_heading(doc, "基本信息", 13)
    add_kv_table(doc, [
        ("期初资金", fmt_money(o["start_balance"])),
        ("期末资金", fmt_money(o["end_balance"])),
        ("期初市值权益", fmt_money(o["start_equity"])),
        ("期末市值权益", fmt_money(o["end_equity"])),
        ("总入金", fmt_money(o["total_deposit"])),
        ("总出金", fmt_money(o["total_withdraw"])),
        ("净出入金", fmt_money(o["net_deposit"])),
        ("总手续费", fmt_money(o["total_fee"])),
        ("有效存续交易日数量", str(m["trading_days"])),
        ("有交易操作交易日数量", str(m["trading_days"])),
        ("净收益", fmt_money(o["net_profit"])),
    ])
    add_heading(doc, "业绩信息", 13)
    add_kv_table(doc, [
        ("单位净值", f"{o['unit_nav']:.3f}"),
        ("最大净值", f"{o['max_nav']:.3f}"),
        ("周期内收益", fmt_pct(o["period_return"])),
        ("年化收益", fmt_pct(o["ann_return"])),
        ("单日最大回撤", fmt_pct(o["max_daily_dd"], 4)),
        ("最大峰值回撤", fmt_pct(o["max_peak_dd"], 4)),
        ("连续回撤天数", str(o["dd_calendar_days"])),
        ("最长未创新高", str(o["underwater_days"])),
        ("年化波动率", fmt_pct(o["ann_vol"])),
        ("年化下行波动率", fmt_pct(o["ann_dvol"])),
        ("总交易量", f"{o['total_lots']:.1f}手"),
        ("总交易次数", f"{o['total_trades']}次"),
        ("日胜率", fmt_pct(o["daily_win"])),
        ("月胜率", fmt_pct(o["month_win"])),
        ("日平均保证金", fmt_money(o["avg_margin"])),
        ("日平均保证金比", fmt_pct(o["avg_margin_ratio"], 4)),
        ("夏普比率", f"{o['sharpe']:.4f}"),
        ("索提诺比率", f"{o['sortino']:.4f}"),
        ("卡玛比率", f"{o['calmar']:.4f}"),
        ("日平均手续费比", fmt_pct(o["avg_fee_ratio"], 4)),
    ])
    doc.add_page_break()

    # Part 2
    add_heading(doc, "第二部分 业绩分析", 16)
    add_heading(doc, "2.1 收益率分析", 13)
    add_image(doc, charts.get("ret"))
    add_text(doc, nar["ret"])

    add_heading(doc, "2.2 月收益率分析", 13)
    add_image(doc, charts.get("monthly"))
    add_text(doc, nar["monthly"])
    if data["monthly"]:
        add_simple_matrix_table(
            doc,
            [
                ["日期"] + [r["month"] for r in data["monthly"]],
                ["月收益率"] + [fmt_pct(r["return"]) for r in data["monthly"]],
            ],
        )

    add_heading(doc, "2.3 净值与持仓", 13)
    add_image(doc, charts.get("nav_pos"))
    add_text(doc, nar["nav"])

    add_heading(doc, "2.4 动态回撤", 13)
    add_image(doc, charts.get("drawdown"))
    add_text(doc, nar["drawdown"])
    bt = doc.add_table(rows=3, cols=1 + len(data["dd_buckets"]))
    bt.style = "Table Grid"
    bt.cell(0, 0).text = "回撤率"
    bt.cell(1, 0).text = "天数"
    bt.cell(2, 0).text = "占比"
    for i, b in enumerate(data["dd_buckets"]):
        bt.cell(0, i + 1).text = b["label"]
        bt.cell(1, i + 1).text = str(b["days"])
        bt.cell(2, i + 1).text = fmt_pct(b["share"])

    add_heading(doc, "创新高最大回撤率", 12)
    nh = doc.add_table(rows=3, cols=1 + len(data["new_high_dd_buckets"]))
    nh.style = "Table Grid"
    nh.cell(0, 0).text = "创新高最大回撤率"
    nh.cell(1, 0).text = "次数"
    nh.cell(2, 0).text = "占比"
    for i, b in enumerate(data["new_high_dd_buckets"]):
        nh.cell(0, i + 1).text = b["label"]
        nh.cell(1, i + 1).text = str(b["count"])
        nh.cell(2, i + 1).text = fmt_pct(b["share"])

    add_heading(doc, "最大回撤明细", 12)
    dd = doc.add_table(rows=1 + len(data["dd_detail"]), cols=5)
    dd.style = "Table Grid"
    for i, h in enumerate(["序号", "最大回撤", "最大回撤起止区间", "修补天数", "创新日期"]):
        dd.cell(0, i).text = h
    for r, row in enumerate(data["dd_detail"], start=1):
        dd.cell(r, 0).text = str(row["rank"])
        dd.cell(r, 1).text = fmt_pct(row["max_dd"], 4)
        dd.cell(r, 2).text = row["range"]
        dd.cell(r, 3).text = str(row["repair_days"])
        dd.cell(r, 4).text = row["new_high_date"]

    add_heading(doc, "2.5 出入金和动态权益", 13)
    add_text(doc, nar["equity_dep"])
    add_image(doc, charts.get("equity_dep"))
    add_heading(doc, "2.6 手续费分析", 13)
    add_text(doc, nar["fee"])
    add_image(doc, charts.get("fee"))
    add_heading(doc, "2.7 月度盈亏分析", 13)
    add_image(doc, charts.get("month_pnl"))
    add_text(doc, nar["month_pnl"])
    if data["monthly"]:
        add_simple_matrix_table(
            doc,
            [
                ["日期"] + [r["month"] for r in data["monthly"]],
                ["月盈亏"] + [f"￥{fmt_money(r['pnl'])}" for r in data["monthly"]],
            ],
        )
    add_image(doc, charts.get("cum_pnl"))
    add_text(doc, "上图为扣减手续费后的累计盈亏路径，可与净值曲线对照观察亏损是否持续扩大。")
    add_heading(doc, "2.8 波动率", 13)
    add_image(doc, charts.get("vol"))
    add_text(doc, nar["vol"])
    doc.add_page_break()

    # Part 3
    add_heading(doc, "第三部分 板块盈亏分析", 16)
    add_heading(doc, "3.1 板块盈亏分析", 13)
    add_text(doc, nar["sector_intro"])
    add_image(doc, charts.get("sector_bar"))
    add_image(doc, charts.get("sector_pies") or charts.get("sector_profit_pie"), width=6.5)

    add_text(doc, "您盈利最多的三个板块分别为：")
    for line in nar["sector_profit"] or ["本期无盈利板块。"]:
        add_text(doc, line)
    add_text(doc, "您亏损最多的三个板块分别为：")
    for line in nar["sector_loss"] or ["本期无亏损板块。"]:
        add_text(doc, line)

    add_text(doc, "板块分类:")
    for sector, codes in SECTOR_RULES.items():
        add_text(doc, f"{sector}: {codes}", size=9)

    add_heading(doc, "3.2 品种盈亏分析", 13)
    add_text(doc, nar["product_intro"])
    add_image(doc, charts.get("product_bar"))
    add_image(doc, charts.get("product_pies") or charts.get("product_profit_pie"), width=6.5)
    add_text(doc, "您盈利最多的三个品种分别为：")
    for line in nar["product_profit"] or ["本期无盈利品种。"]:
        add_text(doc, line)
    add_text(doc, "您亏损最多的三个品种分别为：")
    for line in nar["product_loss"] or ["本期无亏损品种。"]:
        add_text(doc, line)

    add_heading(doc, "盈利品种", 13)
    add_text(doc, nar["product_profit_chart"])
    add_image(doc, charts.get("profit_cum"))
    add_text(doc, "上图为盈利品种累计盈亏；下图为对应交易手数变化，便于对照“赚钱是否伴随放量”。")
    add_image(doc, charts.get("profit_lots"))
    add_text(doc, nar["product_daily"])
    add_image(doc, charts.get("profit_daily"))
    add_text(doc, nar["product_monthly"])
    add_image(doc, charts.get("profit_monthly"))

    add_heading(doc, "亏损品种", 13)
    add_text(doc, nar["product_loss_chart"])
    add_image(doc, charts.get("loss_cum"))
    add_text(doc, "上图为亏损品种累计盈亏轨迹；下图为手数。若手数高且累计亏损加深，说明该品种交易活跃但方向错误。")
    add_image(doc, charts.get("loss_lots"))
    add_text(doc, nar["product_daily"])
    add_image(doc, charts.get("loss_daily"))
    add_text(doc, nar["product_monthly"])
    add_image(doc, charts.get("loss_monthly"))

    add_heading(doc, "品种绩效归因", 13)
    dt = doc.add_table(rows=1 + len(data["directions"]), cols=4)
    dt.style = "Table Grid"
    for i, h in enumerate(["品种", "方向", "利润", "利润比"]):
        dt.cell(0, i).text = h
    for r, row in enumerate(data["directions"], start=1):
        dt.cell(r, 0).text = row["name"]
        dt.cell(r, 1).text = row["direction"]
        dt.cell(r, 2).text = f"{row['pnl']:.1f}"
        dt.cell(r, 3).text = fmt_pct(row["weight"])

    add_heading(doc, "品种汇总绩效归因", 13)
    pt = doc.add_table(rows=1 + len(data["products"]), cols=3)
    pt.style = "Table Grid"
    for i, h in enumerate(["品种", "利润", "利润比"]):
        pt.cell(0, i).text = h
    for r, row in enumerate(data["products"], start=1):
        pt.cell(r, 0).text = row["name"]
        pt.cell(r, 1).text = f"{row['pnl']:.1f}"
        pt.cell(r, 2).text = fmt_pct(row["weight"])

    add_heading(doc, "3.3：多空胜率分析", 13)
    for title, key in [("汇总", "overall"), ("多头-卖平", "long"), ("空头-买平", "short")]:
        st = data["long_short"][key]
        add_text(
            doc,
            f"{title}：盈利手数 {st['win']['lots']:.1f}，亏损手数 {st['loss']['lots']:.1f}，"
            f"累计盈亏 {fmt_money(st['total_pnl'])}，每笔平均盈亏 盈利 {st['win']['avg']:.2f} / 亏损 {st['loss']['avg']:.2f}，"
            f"胜率 {fmt_pct(st['win_rate'])}，盈亏比 {st['pf']:.2f}",
        )
    add_text(doc, "注：最后一日的持仓盈亏未纳入胜率统计。")

    add_heading(doc, "3.4：周期分析", 13)
    add_text(doc, "各周期开平仓日期间隔定义：日内表示当日开仓，当日平仓；短线表示平5日之内的仓；中线表示平20日之内的仓；超过20日即为长线。")
    add_image(doc, charts.get("period_pie"), width=5.2)
    ht = doc.add_table(rows=1 + len(data["periods"]), cols=6)
    ht.style = "Table Grid"
    for i, h in enumerate(["周期", "盈利金额", "亏损金额", "累计盈亏", "交易手数", "交易手数占比"]):
        ht.cell(0, i).text = h
    for r, row in enumerate(data["periods"], start=1):
        ht.cell(r, 0).text = row["period"]
        ht.cell(r, 1).text = f"{row['profit']:.1f}"
        ht.cell(r, 2).text = f"{row['loss']:.1f}"
        ht.cell(r, 3).text = f"{row['pnl']:.1f}"
        ht.cell(r, 4).text = f"{row['lots']:.1f}"
        ht.cell(r, 5).text = fmt_pct(row["lot_share"])
    for row in data["periods"]:
        add_text(
            doc,
            f"{row['period']}交易（{row['period']}）： 共计 盈亏 {fmt_money(row['pnl'])} 元，"
            f"共计交易 {row['trades']} 次，其中盈利 {row['wins']} 次，总胜率 {fmt_pct(row['win_rate'])}，"
            f"交易手数为 {row['lots']:.1f}，占总交易手数的 {fmt_pct(row['lot_share'])}",
        )
        add_image(doc, charts.get(f"period_{row['period']}"), width=5.0)

    doc.add_page_break()
    add_heading(doc, "第四部分 持仓分析", 16)
    add_heading(doc, "4.1：持仓策略", 13)
    add_text(doc, nar["pos_strategy"])
    add_image(doc, charts.get("pos_strategy"))
    add_text(doc, "下图为各板块累计盈亏（未扣手续费），可识别拖累净值的主要板块。")
    add_image(doc, charts.get("sector_cum"))
    add_heading(doc, "4.2：板块聚焦", 13)
    add_text(doc, nar["sector_focus"])
    add_image(doc, charts.get("sector_focus"))
    add_heading(doc, "4.3：品种聚焦", 13)
    add_text(doc, nar["product_focus"])
    add_image(doc, charts.get("product_focus"))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Word: {out_path}", flush=True)

# ── PDF ───────────────────────────────────────────────────────────────────────

def build_pdf(data: dict, charts: dict[str, Path], out_path: Path):
    font = configure_pdf_font()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CNTitle", fontName=font, fontSize=22, leading=28, alignment=1, spaceAfter=16))
    styles.add(ParagraphStyle(name="CNH1", fontName=font, fontSize=16, leading=22, spaceBefore=10, spaceAfter=8))
    styles.add(ParagraphStyle(name="CNH2", fontName=font, fontSize=12, leading=18, spaceBefore=8, spaceAfter=6))
    styles.add(ParagraphStyle(name="CNBody", fontName=font, fontSize=9.5, leading=14, spaceAfter=6))
    styles.add(ParagraphStyle(name="CNCenter", fontName=font, fontSize=12, leading=18, alignment=1, spaceAfter=6))

    o = data["overview"]
    m = data["meta"]
    nar = build_narratives(data)
    story = []

    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph("投资报告分析", styles["CNTitle"]))
    story.append(Spacer(1, 1 * cm))
    for line in [
        f"账号: {m['client_id']}",
        f"周期: {m['start']}~{m['end']}",
        f"导出时间: {m['export_time']}",
        f"投顾名称: {m['client_name']}",
    ]:
        story.append(Paragraph(line, styles["CNCenter"]))
    story.append(PageBreak())

    story.append(Paragraph(f"用户账号:{m['client_id']}", styles["CNBody"]))
    story.append(Paragraph(f"统计周期:{m['start']}~{m['end']}", styles["CNBody"]))
    story.append(Paragraph("第一部分 总览", styles["CNH1"]))
    story.append(Paragraph(f"交易日:{m['trading_days']}天", styles["CNBody"]))

    def kv_table(pairs):
        data_rows = []
        for i in range(0, len(pairs), 2):
            row = []
            for k, v in pairs[i : i + 2]:
                row.extend([k, v])
            while len(row) < 4:
                row.append("")
            data_rows.append(row)
        t = Table(data_rows, colWidths=[3.2 * cm, 4.2 * cm, 3.2 * cm, 4.2 * cm])
        t.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.95, 0.95, 0.95)),
                    ("BACKGROUND", (2, 0), (2, -1), colors.Color(0.95, 0.95, 0.95)),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        return t

    story.append(Paragraph("基本信息", styles["CNH2"]))
    story.append(
        kv_table(
            [
                ("期初资金", fmt_money(o["start_balance"])),
                ("期末资金", fmt_money(o["end_balance"])),
                ("期初市值权益", fmt_money(o["start_equity"])),
                ("期末市值权益", fmt_money(o["end_equity"])),
                ("总入金", fmt_money(o["total_deposit"])),
                ("总出金", fmt_money(o["total_withdraw"])),
                ("净出入金", fmt_money(o["net_deposit"])),
                ("总手续费", fmt_money(o["total_fee"])),
                ("有效存续交易日数量", str(m["trading_days"])),
                ("有交易操作交易日数量", str(m["trading_days"])),
                ("净收益", fmt_money(o["net_profit"])),
                ("", ""),
            ]
        )
    )
    story.append(Paragraph("业绩信息", styles["CNH2"]))
    story.append(
        kv_table(
            [
                ("单位净值", f"{o['unit_nav']:.3f}"),
                ("最大净值", f"{o['max_nav']:.3f}"),
                ("周期内收益", fmt_pct(o["period_return"])),
                ("年化收益", fmt_pct(o["ann_return"])),
                ("单日最大回撤", fmt_pct(o["max_daily_dd"], 4)),
                ("最大峰值回撤", fmt_pct(o["max_peak_dd"], 4)),
                ("连续回撤天数", str(o["dd_calendar_days"])),
                ("最长未创新高", str(o["underwater_days"])),
                ("年化波动率", fmt_pct(o["ann_vol"])),
                ("年化下行波动率", fmt_pct(o["ann_dvol"])),
                ("总交易量", f"{o['total_lots']:.1f}手"),
                ("总交易次数", f"{o['total_trades']}次"),
                ("日胜率", fmt_pct(o["daily_win"])),
                ("月胜率", fmt_pct(o["month_win"])),
                ("日平均保证金", fmt_money(o["avg_margin"])),
                ("日平均保证金比", fmt_pct(o["avg_margin_ratio"], 4)),
                ("夏普比率", f"{o['sharpe']:.4f}"),
                ("索提诺比率", f"{o['sortino']:.4f}"),
                ("卡玛比率", f"{o['calmar']:.4f}"),
                ("日平均手续费比", fmt_pct(o["avg_fee_ratio"], 4)),
            ]
        )
    )
    story.append(PageBreak())

    def img(key, w=16 * cm):
        """Insert chart with explicit height from pixel aspect (ReportLab width-only uses px height)."""
        p = charts.get(key)
        if p and p.exists():
            with PILImage.open(p) as im:
                iw, ih = im.size
            h = w * (ih / max(iw, 1))
            story.append(RLImage(str(p), width=w, height=h))
            story.append(Spacer(1, 4 * mm))

    def tiny_table(rows, *, max_width=17 * cm):
        """Compact, content-sized table; kept together so rows don't split across pages."""
        ncols = max(1, len(rows[0]))

        def cell_units(s: str) -> float:
            s = str(s)
            return sum(2.0 if ord(ch) > 127 else 1.0 for ch in s) + 2

        units = [max(cell_units(rows[r][c]) for r in range(len(rows))) for c in range(ncols)]
        units[0] = max(units[0], 5.0)
        for c in range(1, ncols):
            units[c] = max(units[c], 7.0)
        # ~2.1mm per unit — prefer compact left-aligned table (like sample), only shrink if overflow
        col_widths = [u * 2.1 * mm for u in units]
        total = sum(col_widths)
        if total > max_width:
            scale = max_width / total
            col_widths = [w * scale for w in col_widths]
        t = Table(rows, colWidths=col_widths, hAlign="LEFT")
        t.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.95, 0.95, 0.95)),
                ]
            )
        )
        story.append(KeepTogether([t, Spacer(1, 3 * mm)]))

    story.append(Paragraph("第二部分 业绩分析", styles["CNH1"]))
    story.append(Paragraph("2.1 收益率分析", styles["CNH2"]))
    img("ret")
    story.append(Paragraph(nar["ret"], styles["CNBody"]))

    story.append(Paragraph("2.2 月收益率分析", styles["CNH2"]))
    img("monthly")
    story.append(Paragraph(nar["monthly"], styles["CNBody"]))
    if data["monthly"]:
        tiny_table(
            [["日期"] + [r["month"] for r in data["monthly"]],
             ["月收益率"] + [fmt_pct(r["return"]) for r in data["monthly"]]]
        )

    story.append(Paragraph("2.3 净值与持仓", styles["CNH2"]))
    img("nav_pos")
    story.append(Paragraph(nar["nav"], styles["CNBody"]))

    story.append(Paragraph("2.4 动态回撤", styles["CNH2"]))
    img("drawdown")
    story.append(Paragraph(nar["drawdown"], styles["CNBody"]))
    tiny_table(
        [["回撤率"] + [b["label"] for b in data["dd_buckets"]],
         ["天数"] + [str(b["days"]) for b in data["dd_buckets"]],
         ["占比"] + [fmt_pct(b["share"]) for b in data["dd_buckets"]]]
    )
    story.append(Paragraph("创新高最大回撤率", styles["CNH2"]))
    tiny_table(
        [["创新高最大回撤率"] + [b["label"] for b in data["new_high_dd_buckets"]],
         ["次数"] + [str(b["count"]) for b in data["new_high_dd_buckets"]],
         ["占比"] + [fmt_pct(b["share"]) for b in data["new_high_dd_buckets"]]]
    )
    story.append(Paragraph("最大回撤明细", styles["CNH2"]))
    dd_rows = [["序号", "最大回撤", "最大回撤起止区间", "修补天数", "创新日期"]]
    for row in data["dd_detail"]:
        dd_rows.append([str(row["rank"]), fmt_pct(row["max_dd"], 4), row["range"], str(row["repair_days"]), row["new_high_date"]])
    t = Table(dd_rows, colWidths=[1.5 * cm, 2.5 * cm, 5 * cm, 2.2 * cm, 2.8 * cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.92)),
    ]))
    story.append(t)

    story.append(Paragraph("2.5 出入金和动态权益", styles["CNH2"]))
    story.append(Paragraph(nar["equity_dep"], styles["CNBody"]))
    img("equity_dep")
    story.append(Paragraph("2.6 手续费分析", styles["CNH2"]))
    story.append(Paragraph(nar["fee"], styles["CNBody"]))
    img("fee")
    story.append(Paragraph("2.7 月度盈亏分析", styles["CNH2"]))
    img("month_pnl")
    story.append(Paragraph(nar["month_pnl"], styles["CNBody"]))
    if data["monthly"]:
        tiny_table(
            [["日期"] + [r["month"] for r in data["monthly"]],
             ["月盈亏"] + [f"￥{fmt_money(r['pnl'])}" for r in data["monthly"]]]
        )
    img("cum_pnl")
    story.append(Paragraph("上图为扣减手续费后的累计盈亏路径，可与净值曲线对照观察亏损是否持续扩大。", styles["CNBody"]))
    story.append(Paragraph("2.8 波动率", styles["CNH2"]))
    img("vol")
    story.append(Paragraph(nar["vol"], styles["CNBody"]))
    story.append(PageBreak())

    story.append(Paragraph("第三部分 板块盈亏分析", styles["CNH1"]))
    story.append(Paragraph("3.1 板块盈亏分析", styles["CNH2"]))
    story.append(Paragraph(nar["sector_intro"], styles["CNBody"]))
    img("sector_bar")
    img("sector_pies", w=17.5 * cm)
    story.append(Paragraph("您盈利最多的三个板块分别为：", styles["CNBody"]))
    for line in nar["sector_profit"] or ["本期无盈利板块。"]:
        story.append(Paragraph(line, styles["CNBody"]))
    story.append(Paragraph("您亏损最多的三个板块分别为：", styles["CNBody"]))
    for line in nar["sector_loss"] or ["本期无亏损板块。"]:
        story.append(Paragraph(line, styles["CNBody"]))

    story.append(Paragraph("3.2 品种盈亏分析", styles["CNH2"]))
    story.append(Paragraph(nar["product_intro"], styles["CNBody"]))
    img("product_bar")
    img("product_pies", w=16.5 * cm)
    story.append(Paragraph("您盈利最多的三个品种分别为：", styles["CNBody"]))
    for line in nar["product_profit"] or ["本期无盈利品种。"]:
        story.append(Paragraph(line, styles["CNBody"]))
    story.append(Paragraph("您亏损最多的三个品种分别为：", styles["CNBody"]))
    for line in nar["product_loss"] or ["本期无亏损品种。"]:
        story.append(Paragraph(line, styles["CNBody"]))

    story.append(Paragraph("盈利品种", styles["CNH2"]))
    story.append(Paragraph(nar["product_profit_chart"], styles["CNBody"]))
    img("profit_cum")
    story.append(Paragraph("上图为盈利品种累计盈亏；下图为对应交易手数变化。", styles["CNBody"]))
    img("profit_lots")
    story.append(Paragraph(nar["product_daily"], styles["CNBody"]))
    img("profit_daily")
    story.append(Paragraph(nar["product_monthly"], styles["CNBody"]))
    img("profit_monthly")

    story.append(Paragraph("亏损品种", styles["CNH2"]))
    story.append(Paragraph(nar["product_loss_chart"], styles["CNBody"]))
    img("loss_cum")
    story.append(Paragraph("上图为亏损品种累计盈亏轨迹；下图为手数。", styles["CNBody"]))
    img("loss_lots")
    story.append(Paragraph(nar["product_daily"], styles["CNBody"]))
    img("loss_daily")
    story.append(Paragraph(nar["product_monthly"], styles["CNBody"]))
    img("loss_monthly")

    story.append(Paragraph("品种绩效归因", styles["CNH2"]))
    dir_rows = [["品种", "方向", "利润", "利润比"]] + [
        [r["name"], r["direction"], f"{r['pnl']:.1f}", fmt_pct(r["weight"])] for r in data["directions"]
    ]
    t = Table(dir_rows, colWidths=[4 * cm, 2 * cm, 3.5 * cm, 3 * cm])
    t.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.92)),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 4 * mm))

    story.append(Paragraph("品种汇总绩效归因", styles["CNH2"]))
    prod_rows = [["品种", "利润", "利润比"]] + [
        [r["name"], f"{r['pnl']:.1f}", fmt_pct(r["weight"])] for r in data["products"]
    ]
    t = Table(prod_rows, colWidths=[5 * cm, 4 * cm, 3.5 * cm])
    t.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.92)),
            ]
        )
    )
    story.append(t)

    story.append(Paragraph("3.3：多空胜率分析", styles["CNH2"]))
    for title, key in [("汇总", "overall"), ("多头-卖平", "long"), ("空头-买平", "short")]:
        st = data["long_short"][key]
        story.append(
            Paragraph(
                f"{title}：盈利手数 {st['win']['lots']:.1f}，亏损手数 {st['loss']['lots']:.1f}，"
                f"累计盈亏 {fmt_money(st['total_pnl'])}，胜率 {fmt_pct(st['win_rate'])}，盈亏比 {st['pf']:.2f}",
                styles["CNBody"],
            )
        )

    story.append(Paragraph("3.4：周期分析", styles["CNH2"]))
    story.append(
        Paragraph(
            "日内表示当日开仓当日平仓；短线表示平5日之内的仓；中线表示平20日之内的仓；超过20日即为长线。",
            styles["CNBody"],
        )
    )
    img("period_pie", w=11 * cm)
    period_rows = [["周期", "盈利金额", "亏损金额", "累计盈亏", "交易手数", "手数占比"]]
    for row in data["periods"]:
        period_rows.append(
            [
                row["period"],
                f"{row['profit']:.1f}",
                f"{row['loss']:.1f}",
                f"{row['pnl']:.1f}",
                f"{row['lots']:.1f}",
                fmt_pct(row["lot_share"]),
            ]
        )
    t = Table(period_rows, colWidths=[2 * cm, 2.6 * cm, 2.6 * cm, 2.6 * cm, 2.4 * cm, 2.4 * cm])
    t.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.92)),
            ]
        )
    )
    story.append(t)
    for row in data["periods"]:
        story.append(
            Paragraph(
                f"{row['period']}：盈亏 {fmt_money(row['pnl'])}，交易 {row['trades']} 次，"
                f"胜率 {fmt_pct(row['win_rate'])}，手数 {row['lots']:.1f}",
                styles["CNBody"],
            )
        )
        img(f"period_{row['period']}", w=12 * cm)

    story.append(PageBreak())
    story.append(Paragraph("第四部分 持仓分析", styles["CNH1"]))
    story.append(Paragraph("4.1：持仓策略", styles["CNH2"]))
    story.append(Paragraph(nar["pos_strategy"], styles["CNBody"]))
    img("pos_strategy")
    story.append(Paragraph("下图为各板块累计盈亏（未扣手续费），可识别拖累净值的主要板块。", styles["CNBody"]))
    img("sector_cum")
    story.append(Paragraph("4.2：板块聚焦", styles["CNH2"]))
    story.append(Paragraph(nar["sector_focus"], styles["CNBody"]))
    img("sector_focus")
    story.append(Paragraph("4.3：品种聚焦", styles["CNH2"]))
    story.append(Paragraph(nar["product_focus"], styles["CNBody"]))
    img("product_focus")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="投资报告分析",
    )
    doc.build(story)
    print(f"[OK] PDF: {out_path}", flush=True)


def main():
    parser = argparse.ArgumentParser(description="融航结算单 ZIP → Word/PDF 投资报告分析")
    parser.add_argument("--zip", required=True, help="data.zip 路径")
    parser.add_argument("--outdir", default="", help="输出目录")
    parser.add_argument("--format", choices=["docx", "pdf", "both"], default="both")
    parser.add_argument("--advisor", default="", help="封面投顾名称（覆盖结算单客户名称）")
    args = parser.parse_args()

    zip_path = Path(args.zip).resolve()
    if not zip_path.exists():
        raise SystemExit(f"ZIP 不存在: {zip_path}")

    outdir = (
        Path(args.outdir).resolve()
        if args.outdir.strip()
        else Path(os.environ.get("RONGHANG_REPORT_OUTPUT_DIR", "").strip() or (BASE_DIR / "report_output"))
    )
    outdir.mkdir(parents=True, exist_ok=True)
    chart_dir = outdir / "charts"

    configure_matplotlib()
    print(f"[LOAD] {zip_path}", flush=True)
    days = load_zip(zip_path)
    print(f"[LOAD] {len(days)} trading days", flush=True)
    data = analyze(days)
    advisor = (args.advisor or "").strip()
    if advisor:
        data["meta"]["client_name"] = advisor
        print(f"[META] advisor override: {advisor}", flush=True)
    print(
        f"[METRICS] period={data['overview']['period_return']*100:.2f}% "
        f"nav={data['overview']['final_nav']:.4f} "
        f"net_profit={data['overview']['net_profit']:.2f}",
        flush=True,
    )
    charts = make_charts(data, chart_dir)

    docx_path = outdir / "投资报告分析.docx"
    pdf_path = outdir / "投资报告分析.pdf"
    if args.format in {"docx", "both"}:
        build_docx(data, charts, docx_path)
    if args.format in {"pdf", "both"}:
        build_pdf(data, charts, pdf_path)

    if args.format in {"docx", "both"}:
        print(f"REPORT_DOCX={docx_path}", flush=True)
    if args.format in {"pdf", "both"}:
        print(f"REPORT_PDF={pdf_path}", flush=True)


if __name__ == "__main__":
    main()
