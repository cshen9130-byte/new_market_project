"""银河期货账户策略报告 — 从 PostgreSQL yinhe_* 表生成 Word

用法：
    python generate_yinhe_word_report_db.py

环境变量：
    DATABASE_URL  postgresql://user:pass@host:port/dbname
"""
from __future__ import annotations

import io
import os
import sys
from pathlib import Path

if sys.platform == "win32":
    if hasattr(sys.stdout, "buffer"):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    if hasattr(sys.stderr, "buffer"):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)

import os as _os
_os.environ.setdefault("TQDM_DISABLE", "1")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties, fontManager
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_CN_FONT: FontProperties | None = None

try:
    import psycopg2
except ImportError:
    sys.exit("请先安装 psycopg2: pip install psycopg2-binary")

BASE_DIR = Path(__file__).resolve().parent
_output_override = os.environ.get("YINHE_REPORT_OUTPUT_DIR", "").strip()
OUTPUT_DIR = Path(_output_override) if _output_override else (BASE_DIR / "report_output")
CHART_DIR = OUTPUT_DIR / "charts"
REPORT_PATH = OUTPUT_DIR / "银河期货交易策略分析报告.docx"


def parse_database_url(url: str) -> dict:
    from urllib.parse import urlparse, unquote
    u = urlparse(url)
    return {
        "host": u.hostname or "127.0.0.1",
        "port": u.port or 5432,
        "dbname": (u.path or "/").lstrip("/") or "postgres",
        "user": unquote(u.username or ""),
        "password": unquote(u.password or ""),
    }


def load_frames(conn) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    account = pd.read_sql(
        """
        SELECT trade_date, client_id, client_name, balance_bf, deposit_withdrawal,
               realized_pl, mtm_pl, commission, balance_cf, client_equity,
               fund_avail, risk_degree, margin_occupied
        FROM yinhe_account_summary
        ORDER BY trade_date
        """,
        conn,
        parse_dates=["trade_date"],
    )
    trades = pd.read_sql(
        """
        SELECT trade_date, settlement_date, product, instrument, bs, oc,
               lots, price, turnover, fee, realized_pl
        FROM yinhe_transaction_records
        ORDER BY trade_date
        """,
        conn,
        parse_dates=["trade_date", "settlement_date"],
    )
    positions = pd.read_sql(
        """
        SELECT settlement_date, product, instrument, long_pos, short_pos,
               mtm_pl, margin_occupied
        FROM yinhe_position_summary
        ORDER BY settlement_date
        """,
        conn,
        parse_dates=["settlement_date"],
    )
    closed = pd.read_sql(
        """
        SELECT settlement_date, product, instrument, bs, lots, realized_pl
        FROM yinhe_position_closed
        ORDER BY settlement_date
        """,
        conn,
        parse_dates=["settlement_date"],
    )
    return account, trades, positions, closed


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str, *, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def fmt_money(v) -> str:
    try:
        return f"{float(v):,.2f}"
    except Exception:
        return str(v)


def fmt_pct(v) -> str:
    try:
        return f"{float(v) * 100:.2f}%"
    except Exception:
        return str(v)


def _cn_font_candidates() -> list[str]:
    """Prefer single-face TTF/OTF over TTC — matplotlib often tofu-renders CJK from TTC."""
    home = Path.home()
    return [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/wqy-microhei/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        str(home / ".local/share/fonts/NotoSansSC-Regular.otf"),
        str(home / ".fonts/NotoSansSC-Regular.otf"),
    ]


def configure_matplotlib() -> None:
    """Register a real CJK font file so chart Chinese is not rendered as □□□."""
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False

    chosen_path: str | None = None
    for path in _cn_font_candidates():
        if not path or not os.path.isfile(path):
            continue
        try:
            fontManager.addfont(path)
            chosen_path = path
            break
        except Exception:
            continue

    if chosen_path is None:
        for name in ("Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Arial Unicode MS"):
            try:
                path = fontManager.findfont(FontProperties(family=name), fallback_to_default=False)
                if path and "dejavu" not in path.lower():
                    chosen_path = path
                    break
            except Exception:
                continue

    if chosen_path is None:
        print("[WARN] No CJK font found; chart Chinese labels may render as boxes.", flush=True)
        _CN_FONT = None
        return

    try:
        _CN_FONT = FontProperties(fname=chosen_path)
        family = _CN_FONT.get_name()
    except Exception:
        _CN_FONT = None
        family = "SimHei"

    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = [family, "DejaVu Sans"]
    print(f"[FONT] Charts using: {family} ({chosen_path})", flush=True)


def _cn_fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def _risk_degree_pct(series: pd.Series) -> pd.Series:
    """Yinhe stores risk_degree as a 0–1 ratio; plot as percent when needed."""
    s = pd.to_numeric(series, errors="coerce")
    if s.dropna().empty:
        return s
    return s * 100.0 if float(s.max()) <= 1.5 else s


def _norm_bs(value) -> str | None:
    s = str(value or "").strip()
    if not s:
        return None
    if s in {"买", "B", "b"} or "买" in s:
        return "买"
    if s in {"卖", "S", "s"} or "卖" in s:
        return "卖"
    return None


def _norm_oc(value) -> str | None:
    s = str(value or "").strip()
    if not s:
        return None
    if s.startswith("开") or s in {"开", "O", "o"}:
        return "开"
    if s.startswith("平") or s in {"平", "C", "c"}:
        return "平"
    return None


_ORDER_STYLES = {
    ("买", "开"): {"marker": "^", "c": "#15803d", "label": "买开", "zorder": 6},
    ("卖", "开"): {"marker": "v", "c": "#b91c1c", "label": "卖开", "zorder": 6},
    ("买", "平"): {
        "marker": "o",
        "facecolors": "#86efac",
        "edgecolors": "#15803d",
        "label": "买平",
        "zorder": 5,
    },
    ("卖", "平"): {
        "marker": "o",
        "facecolors": "#fecaca",
        "edgecolors": "#b91c1c",
        "label": "卖平",
        "zorder": 5,
    },
}


def prepare_order_points(trades: pd.DataFrame) -> pd.DataFrame:
    if trades is None or trades.empty:
        return pd.DataFrame(columns=["trade_date", "bs", "oc", "lots", "signed_lots"])
    frame = trades.copy()
    frame["trade_date"] = pd.to_datetime(frame["trade_date"]).dt.normalize()
    frame["bs"] = frame["bs"].map(_norm_bs)
    frame["oc"] = frame["oc"].map(_norm_oc)
    frame["lots"] = pd.to_numeric(frame["lots"], errors="coerce").fillna(0.0)
    frame = frame.dropna(subset=["bs", "oc"])
    frame = frame[frame["lots"] > 0]
    if frame.empty:
        return pd.DataFrame(columns=["trade_date", "bs", "oc", "lots", "signed_lots"])
    agg = frame.groupby(["trade_date", "bs", "oc"], as_index=False).agg(lots=("lots", "sum"))
    agg["signed_lots"] = np.where(agg["bs"] == "买", agg["lots"], -agg["lots"])
    return agg


def _order_marker_sizes(lots: pd.Series):
    return np.clip(pd.to_numeric(lots, errors="coerce").fillna(0).to_numpy(dtype=float) * 4.0, 28, 220)


def _scatter_orders(ax, x, y, style: dict, sizes, legend_label: str | None):
    kwargs = {k: v for k, v in style.items() if k != "label"}
    if legend_label is not None:
        kwargs["label"] = legend_label
    ax.scatter(x, y, s=sizes, alpha=0.88, linewidths=0.8, **kwargs)


def overlay_orders_on_equity(ax, account: pd.DataFrame, order_points: pd.DataFrame) -> None:
    if order_points.empty or account.empty:
        return
    equity = (
        account.sort_values("trade_date")
        .assign(trade_date=lambda d: pd.to_datetime(d["trade_date"]).dt.normalize())
        .drop_duplicates("trade_date", keep="last")
        .set_index("trade_date")["client_equity"]
    )
    offsets = {("买", "开"): 0.012, ("卖", "开"): -0.012, ("买", "平"): 0.004, ("卖", "平"): -0.004}
    seen: set[str] = set()
    for key, style in _ORDER_STYLES.items():
        sub = order_points[(order_points["bs"] == key[0]) & (order_points["oc"] == key[1])]
        if sub.empty:
            continue
        y = sub["trade_date"].map(equity)
        valid = y.notna()
        if not valid.any():
            continue
        y = y[valid].astype(float) * (1.0 + offsets[key])
        label = style["label"] if style["label"] not in seen else None
        if label:
            seen.add(label)
        _scatter_orders(ax, sub.loc[valid, "trade_date"], y, style, _order_marker_sizes(sub.loc[valid, "lots"]), label)
    if seen:
        ax.legend(loc="best", fontsize=8, framealpha=0.9, prop=_CN_FONT)


def make_charts(account: pd.DataFrame, turnover: pd.DataFrame, trades: pd.DataFrame | None = None) -> dict[str, Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    configure_matplotlib()
    fp = _cn_fp()
    paths: dict[str, Path] = {}
    order_points = prepare_order_points(trades if trades is not None else pd.DataFrame())

    risk_pct = _risk_degree_pct(account["risk_degree"])
    fig, ax1 = plt.subplots(figsize=(10, 4.2))
    ax1.plot(account["trade_date"], account["client_equity"], color="#1d4ed8", linewidth=1.8, label="客户权益")
    overlay_orders_on_equity(ax1, account, order_points)
    ax1.set_title("账户权益与风险度（含订单点）", **fp)
    ax1.set_ylabel("客户权益", **fp)
    ax1.grid(True, alpha=0.25)
    ax2 = ax1.twinx()
    ax2.plot(account["trade_date"], risk_pct, color="#dc2626", linestyle="--", linewidth=1.2, label="风险度")
    ax2.set_ylabel("风险度(%)", **fp)
    fig.tight_layout()
    p = CHART_DIR / "equity_risk.png"
    fig.savefig(p, dpi=140)
    plt.close(fig)
    paths["equity"] = p

    if not order_points.empty:
        fig, ax = plt.subplots(figsize=(10, 4.2))
        seen: set[str] = set()
        for key, style in _ORDER_STYLES.items():
            sub = order_points[(order_points["bs"] == key[0]) & (order_points["oc"] == key[1])]
            if sub.empty:
                continue
            label = style["label"] if style["label"] not in seen else None
            if label:
                seen.add(label)
            _scatter_orders(
                ax,
                sub["trade_date"],
                sub["signed_lots"],
                style,
                _order_marker_sizes(sub["lots"]),
                label,
            )
        ax.axhline(0, color="#9ca3af", linewidth=0.9)
        ax.set_title("订单买卖开平时序（买为正 / 卖为负）", **fp)
        ax.set_ylabel("成交手数", **fp)
        ax.grid(True, alpha=0.25)
        if seen:
            ax.legend(loc="best", fontsize=8, framealpha=0.9, prop=_CN_FONT)
        fig.tight_layout()
        p = CHART_DIR / "order_timeline.png"
        fig.savefig(p, dpi=140)
        plt.close(fig)
        paths["orders"] = p

    fig, ax = plt.subplots(figsize=(10, 3.6))
    ax.bar(account["trade_date"].astype(str), account["margin_occupied"], color="#64748b")
    ax.set_ylabel("保证金占用", **fp)
    ax.tick_params(axis="x", labelrotation=45, labelsize=7)
    fig.tight_layout()
    p = CHART_DIR / "margin.png"
    fig.savefig(p, dpi=140)
    plt.close(fig)
    paths["margin"] = p

    if not turnover.empty:
        top = turnover.head(12)
        fig, ax = plt.subplots(figsize=(9, 4))
        ax.barh(top["product"][::-1], top["turnover"][::-1], color="#0f766e")
        ax.set_xlabel("成交额", **fp)
        if _CN_FONT is not None:
            for label in ax.get_yticklabels():
                label.set_fontproperties(_CN_FONT)
        fig.tight_layout()
        p = CHART_DIR / "turnover.png"
        fig.savefig(p, dpi=140)
        plt.close(fig)
        paths["turnover"] = p

    return paths


def write_report(account: pd.DataFrame, trades: pd.DataFrame, positions: pd.DataFrame, closed: pd.DataFrame):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    start_eq = float(account.iloc[0]["client_equity"] or 0)
    end_eq = float(account.iloc[-1]["client_equity"] or 0)
    ret = (end_eq - start_eq) / start_eq if start_eq else 0.0
    fee_total = float(account["commission"].fillna(0).sum())
    max_risk = float(account["risk_degree"].fillna(0).max())
    start_d = account.iloc[0]["trade_date"].date()
    end_d = account.iloc[-1]["trade_date"].date()
    client_id = str(account.iloc[-1].get("client_id") or "")
    client_name = str(account.iloc[-1].get("client_name") or "")

    if trades.empty:
        turnover = pd.DataFrame(columns=["product", "turnover", "lots", "share_pct"])
    else:
        turnover = (
            trades.groupby("product", dropna=False)
            .agg(turnover=("turnover", "sum"), lots=("lots", "sum"))
            .reset_index()
            .sort_values("turnover", ascending=False)
        )
        total = float(turnover["turnover"].sum()) or 1.0
        turnover["share_pct"] = turnover["turnover"] / total * 100

    charts = make_charts(account, turnover, trades)

    # Netting sample: latest day
    latest_pos = positions[positions["settlement_date"] == positions["settlement_date"].max()] if not positions.empty else positions
    if not latest_pos.empty:
        netting = (
            latest_pos.groupby("product")
            .agg(long=("long_pos", "sum"), short=("short_pos", "sum"), mtm=("mtm_pl", "sum"))
            .reset_index()
        )
        netting["net"] = netting["long"] - netting["short"]
    else:
        netting = pd.DataFrame(columns=["product", "long", "short", "net", "mtm"])

    closed_pl = float(closed["realized_pl"].fillna(0).sum()) if not closed.empty else float(trades.loc[trades["oc"] == "平", "realized_pl"].fillna(0).sum() if not trades.empty else 0)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("银河期货账户策略分析报告")
    set_run_font(run, size=18, bold=True)

    add_para(doc, f"账户：{client_id} {client_name}".strip())
    add_para(doc, f"样本区间：{start_d} 至 {end_d}（共 {len(account)} 个交易日）")
    add_para(doc, "数据来源：邮箱结算附件（Daily Account Statement / 持仓 / 成交）自动解析入库。")

    doc.add_heading("一、结论摘要", level=1)
    top_products = "、".join(turnover.head(3)["product"].astype(str).tolist()) if not turnover.empty else "暂无"
    net_zero_days = 0
    if not positions.empty:
        daily = positions.groupby(["settlement_date", "product"]).agg(long=("long_pos", "sum"), short=("short_pos", "sum")).reset_index()
        daily["net"] = daily["long"] - daily["short"]
        # share of product-days with near-zero net
        net_zero_days = float((daily["net"].abs() < 1e-6).mean()) if len(daily) else 0.0

    add_bullets(
        doc,
        [
            f"期间收益率 {fmt_pct(ret)}，期初权益 {fmt_money(start_eq)}，期末权益 {fmt_money(end_eq)}。",
            f"累计手续费 {fmt_money(fee_total)}，最高风险度 {fmt_pct(max_risk)}，已实现/平仓盈亏累计约 {fmt_money(closed_pl)}。",
            f"成交主要集中在：{top_products}。",
            f"品种级净敞口接近 0 的占比约 {fmt_pct(net_zero_days)}，可用于判断是否以对冲/价差为主。",
        ],
    )

    doc.add_heading("二、账户收益与风险", level=1)
    add_bullets(
        doc,
        [
            f"交易天数：{len(account)}；成交品种数：{turnover.shape[0]}；成交笔数：{len(trades)}。",
            f"期末保证金占用：{fmt_money(account.iloc[-1]['margin_occupied'])}；期末可用资金：{fmt_money(account.iloc[-1]['fund_avail'])}。",
        ],
    )
    if "equity" in charts:
        doc.add_picture(str(charts["equity"]), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "orders" in charts:
        doc.add_picture(str(charts["orders"]), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_bullets(
            doc,
            [
                "订单点图例：▲买开、▼卖开、○买平、○卖平；点大小按当日汇总手数缩放，便于观察开平仓批次与买卖方向是否成对出现。",
            ],
        )
    if "margin" in charts:
        doc.add_picture(str(charts["margin"]), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("三、品种成交结构", level=1)
    if turnover.empty:
        add_para(doc, "区间内无成交记录。")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["品种", "成交额", "手数", "占比%"]):
            hdr[i].text = h
        for _, row in turnover.head(20).iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["product"])
            cells[1].text = fmt_money(row["turnover"])
            cells[2].text = fmt_money(row["lots"])
            cells[3].text = f"{float(row['share_pct']):.2f}"
        if "turnover" in charts:
            doc.add_paragraph()
            doc.add_picture(str(charts["turnover"]), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("四、最新持仓净敞口", level=1)
    if netting.empty:
        add_para(doc, "暂无持仓汇总。")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for i, h in enumerate(["品种", "多头", "空头", "净手数", "盯市盈亏"]):
            table.rows[0].cells[i].text = h
        for _, row in netting.sort_values("mtm").iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["product"])
            cells[1].text = fmt_money(row["long"])
            cells[2].text = fmt_money(row["short"])
            cells[3].text = fmt_money(row["net"])
            cells[4].text = fmt_money(row["mtm"])

    doc.add_heading("五、说明", level=1)
    add_bullets(
        doc,
        [
            "本报告由系统从银河期货结算邮件附件自动生成，结论基于历史结算数据统计，不构成投资建议。",
            "若某日仅有持仓/成交 XLS 而缺少完整 TXT 资金状况，当日权益类指标可能缺失。",
        ],
    )

    doc.save(REPORT_PATH)
    print(f"Wrote {REPORT_PATH}")


def main():
    db_url = os.environ.get("DATABASE_URL", "").strip()
    if not db_url:
        sys.exit("缺少 DATABASE_URL")
    params = parse_database_url(db_url)
    conn = psycopg2.connect(**params)
    try:
        account, trades, positions, closed = load_frames(conn)
    finally:
        conn.close()
    if account.empty:
        sys.exit("yinhe_account_summary 无数据，请先拉取并解析邮件结算单")
    write_report(account, trades, positions, closed)


if __name__ == "__main__":
    main()
