# -*- coding: utf-8 -*-
"""Generate a Word methodology note for MOM decision signals (量化 vs 主观)."""
from __future__ import annotations

import os
from datetime import date
from pathlib import Path
from xml.sax.saxutils import escape

import matplotlib
import numpy as np

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from matplotlib.font_manager import FontProperties, fontManager
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "scripts" / "ma" / "_mom_signal_method_output"
CHART_DIR = OUT_DIR / "charts"
REPORT_PATH = ROOT / "MOM决策信号计算方法说明.docx"
REPORT_PATH_ASCII = ROOT / "MOM_decision_signal_methodology.docx"

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xB4, 0x53, 0x09)
VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SKY = RGBColor(0x03, 0x69, 0xA1)

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_CN_FONT: FontProperties | None = None


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    for path in (
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ):
        if os.path.isfile(path):
            try:
                fontManager.addfont(path)
                _CN_FONT = FontProperties(fname=path)
                plt.rcParams["font.family"] = "sans-serif"
                plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "Microsoft YaHei", "SimHei"]
                return
            except Exception:
                continue
    _CN_FONT = FontProperties(family="Microsoft YaHei")


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None, italic=False, name="微软雅黑"):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic, name=name)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0, italic=False, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.28
    if first_line and align is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color or TEXT, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    size = {1: 16, 2: 13, 3: 12}.get(level, 11)
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=NAVY)
    return p


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color or TEXT)
    set_cell_border(cell)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=9, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "F7FAFC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            align = "left" if c_i == 0 else "center"
            cell_text(table.rows[r_i + 1].cells[c_i], val, size=8, align=align)
            shade(table.rows[r_i + 1].cells[c_i], fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def caption(doc, text):
    para(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_before=2, space_after=12)


# ---------- Office Math (OMML) ----------

def mt(s: str) -> str:
    return f'<m:r><m:rPr><m:sty m:val="p"/></m:rPr><m:t xml:space="preserve">{escape(str(s))}</m:t></m:r>'


def mi(s: str) -> str:
    return f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">{escape(str(s))}</m:t></m:r>'


def mfrac(num: str, den: str) -> str:
    return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def msup(base: str, exp: str) -> str:
    return f"<m:sSup><m:e>{base}</m:e><m:sup>{exp}</m:sup></m:sSup>"


def msub(base: str, sub: str) -> str:
    return f"<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>"


def msubsup(base: str, sub: str, sup: str) -> str:
    return f"<m:sSubSup><m:e>{base}</m:e><m:sub>{sub}</m:sub><m:sup>{sup}</m:sup></m:sSubSup>"


def mabs(inner: str) -> str:
    return (
        '<m:d><m:dPr><m:begChr m:val="|"/><m:endChr m:val="|"/></m:dPr>'
        f"<m:e>{inner}</m:e></m:d>"
    )


def mparen(inner: str) -> str:
    return (
        '<m:d><m:dPr><m:begChr m:val="("/><m:endChr m:val=")"/></m:dPr>'
        f"<m:e>{inner}</m:e></m:d>"
    )


def msqrt(inner: str) -> str:
    return f'<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr><m:deg/><m:e>{inner}</m:e></m:rad>'


def msum(inner: str, sub: str | None = None, sup: str | None = None) -> str:
    sub_xml = f"<m:sub>{sub}</m:sub>" if sub else "<m:sub/>"
    sup_xml = f"<m:sup>{sup}</m:sup>" if sup else "<m:sup/>"
    return (
        '<m:nary><m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/></m:naryPr>'
        f"{sub_xml}{sup_xml}<m:e>{inner}</m:e></m:nary>"
    )


def add_eq(doc, inner: str, *, number: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    xml = (
        f'<m:oMath xmlns:m="{M_NS}" xmlns:w="{W_NS}">'
        f"{inner}"
        f"</m:oMath>"
    )
    omath = etree.fromstring(xml.encode("utf-8"))
    p._p.append(omath)
    if number:
        add_text(p, f"    ({number})", size=10, color=MUTED)
    return p


def add_eq_note(doc, text: str):
    para(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=10, space_before=0)


# ---------- charts ----------

def draw_clean_tree() -> Path:
    """A cleaner decision flowchart without overlapping arrows."""
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DIR / "decision_tree.png"
    fig, ax = plt.subplots(figsize=(12.2, 7.6), dpi=180)
    ax.set_xlim(0, 12.2)
    ax.set_ylim(0, 7.6)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def box(x, y, w, h, text, fc, ec="#1A365D", fs=9.5, tc="#1A202C", weight="normal"):
        patch = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.03,rounding_size=0.1",
            facecolor=fc, edgecolor=ec, linewidth=1.15, zorder=2,
        )
        ax.add_patch(patch)
        ax.text(
            x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fs, color=tc, zorder=3, fontweight=weight, **fp(),
        )
        return x + w / 2, y, y + h

    def arrow(x1, y1, x2, y2, label="", lp=(0, 0), lc="#4A5568"):
        ax.annotate(
            "", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="-|>", color="#4A5568", lw=1.15),
            zorder=1,
        )
        if label:
            ax.text(x1 + lp[0], y1 + lp[1], label, fontsize=8, color=lc, **fp())

    cx, _, top = box(3.85, 6.7, 4.5, 0.65, "风险%：量化 $R_q$，主观 $R_s$", "#EBF4FF", fs=11)
    _, _, _ = box(3.6, 5.45, 5.0, 0.7, "同向？且 $|R_q|\\geq 3\\%$ 且 $|R_s|\\geq 3\\%$", "#FFFBEB", fs=10)
    arrow(cx, 6.7, cx, 6.15)

    # yes -> crowding
    box(8.55, 4.15, 3.3, 0.7, "$|R_q|+|R_s|\\geq 25\\%$ ?", "#FFFAF0", fs=10)
    arrow(8.6, 5.8, 10.2, 4.85, "是", (0.35, -0.35), "#C53030")

    box(8.55, 2.85, 3.3, 0.55, "控拥挤", "#FEF3C7", "#B45309", 12, "#B45309", "bold")
    arrow(9.4, 4.15, 9.4, 3.4, "是", (-0.55, -0.22), "#B45309")

    box(8.55, 1.55, 3.3, 0.7, "加码\n两边都 ≥ 8% → 重仓共识", "#FEE2E2", "#C53030", 10, "#C53030", "bold")
    arrow(11.0, 4.15, 11.0, 2.25, "否", (0.12, -0.9), "#C53030")

    # no same-dir -> divergence
    box(0.35, 4.15, 3.3, 0.7, "反向？且两边都 ≥ 3%", "#FAF5FF", fs=10)
    arrow(3.6, 5.8, 2.0, 4.85, "否", (-1.4, -0.35), "#6D28D9")

    box(0.35, 2.85, 3.3, 0.55, "观望", "#EDE9FE", "#6D28D9", 12, "#6D28D9", "bold")
    arrow(2.0, 4.15, 2.0, 3.4, "是", (0.12, -0.25), "#6D28D9")

    box(0.35, 1.55, 3.3, 0.7, "一侧 ≥ 8% 且\n另一侧 < 1.5% ?", "#F0F9FF", fs=10)
    arrow(2.0, 4.15, 2.0, 2.25, "否", (0.12, -1.15), "#0369A1")

    box(0.35, 0.35, 3.3, 0.55, "补风格", "#E0F2FE", "#0369A1", 12, "#0369A1", "bold")
    arrow(2.0, 1.55, 2.0, 0.9, "是", (0.12, -0.25), "#0369A1")

    box(4.45, 0.35, 3.3, 0.55, "中性（列表不展示）", "#F1F5F9", "#64748B", 11, "#475569")
    arrow(3.65, 1.9, 5.4, 0.9, "否", (0.55, -0.35), "#64748B")

    ax.set_title("图 1  MOM 决策信号判定树（优先级从上到下、从右到左）", fontsize=13, color="#1A365D", pad=6, **fp())
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def draw_threshold_map() -> Path:
    path = CHART_DIR / "threshold_map.png"
    fig, ax = plt.subplots(figsize=(10.8, 6.4), dpi=180)
    fig.patch.set_facecolor("white")

    # q on x, s on y, first quadrant (absolute values), plus a note for signs
    ax.set_xlim(0, 30)
    ax.set_ylim(0, 30)
    ax.fill_between([0, 1.5], 8, 30, color="#BAE6FD", alpha=0.7, label="补风格（一侧重仓）")
    ax.fill_between([8, 30], 0, 1.5, color="#BAE6FD", alpha=0.7)
    ax.fill_between([3, 30], 3, 30, color="#FECACA", alpha=0.45, label="加码候选（同向） / 观望（反向）")
    xs = np.linspace(3, 30, 400)
    lower = np.maximum(3.0, 25.0 - xs)
    ax.fill_between(xs, lower, 30, color="#FDE68A", alpha=0.72, label="控拥挤区（同向且 q+s≥25）")
    ax.plot([3, 22], [22, 3], color="#B45309", lw=1.4, ls="--")
    ax.axvline(3, color="#C53030", lw=1, ls=":")
    ax.axhline(3, color="#C53030", lw=1, ls=":")
    ax.axvline(8, color="#1A365D", lw=1, ls=":")
    ax.axhline(8, color="#1A365D", lw=1, ls=":")
    ax.axvline(1.5, color="#0369A1", lw=0.8, ls=":")
    ax.axhline(1.5, color="#0369A1", lw=0.8, ls=":")
    ax.text(14, 10, "加码 / 观望", fontsize=12, color="#9B2C2C", **fp())
    ax.text(16, 16, "重仓共识\n(两边≥8%)", fontsize=10, color="#9B2C2C", **fp())
    ax.text(18, 8.5, "弱共识", fontsize=10, color="#C53030", **fp())
    ax.text(4, 22, "补风格\n主观重、量化空", fontsize=9, color="#0369A1", **fp())
    ax.text(20, 0.4, "补风格  量化重、主观空", fontsize=9, color="#0369A1", **fp())
    ax.text(1, 1, "中性\n噪声区", fontsize=9, color="#64748B", **fp())
    ax.text(12, 20, "控拥挤  q+s≥25", fontsize=11, color="#B45309", **fp())
    ax.set_xlabel(r"$q = |R_q|$  量化 |风险%|", fontsize=11, **fp())
    ax.set_ylabel(r"$s = |R_s|$  主观 |风险%|", fontsize=11, **fp())
    ax.set_title("图 2  阈值在 (q, s) 平面上的分区（符号另外决定加码 vs 观望）", fontsize=12, color="#1A365D", **fp())
    ax.set_xticks([0, 1.5, 3, 8, 12.5, 25, 30])
    ax.set_yticks([0, 1.5, 3, 8, 12.5, 25, 30])
    ax.grid(True, alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def insert_picture(doc, path: Path, width_cm=16.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def set_header_footer(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "MOM 每日风控  ·  决策信号计算方法", size=8, color=MUTED)
    footer = section.footer
    fp_ = footer.paragraphs[0]
    fp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(fp_, "内部方法说明  ·  与系统实现 lib/ma/quant-vs-subjective-signals.ts 一致  ·  ", size=8, color=MUTED)
    # page number field
    run = fp_.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=8, color=MUTED)


def build() -> Path:
    configure_matplotlib()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    tree_path = draw_clean_tree()
    map_path = draw_threshold_map()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    set_header_footer(doc)

    # Cover
    para(doc, "MOM 每日风控", size=12, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=4)
    para(doc, "量化 vs 主观  ·  决策信号计算方法说明", size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=8)
    para(
        doc,
        "How MOM Decision Signals Are Calculated  ·  Risk%, Thresholds, and Action Logic",
        size=11, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=16,
    )
    para(doc, f"版本日期：{date.today().isoformat()}　　口径：风险敞口（Risk%）　　实现：quant-vs-subjective-signals.ts", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=18)

    para(
        doc,
        "本说明对应风控页「量化 vs 主观」中的 MOM 决策信号（加码 / 观望 / 补风格 / 控拥挤 / 扩容）。"
        "页面上的问号弹窗只列了阈值；本文把数字从持仓明细一路算到信号，并解释 3%、8%、1.5%、25% 为什么这样设。"
        "信号不是对市场涨跌的预测，而是 MOM 在两种风格之间分配资金、引进投顾时的对照规则。",
        space_after=10,
    )

    heading(doc, "一、结论先行", 1)
    para(
        doc,
        "决策信号只用风险%，不用保证金%。界面上的「保证金」开关会改柱状图和表格数字，但不会改信号列表。"
        "量化柱和主观柱的分母不同，两边的百分比不能相加当成组合权重。",
    )
    add_table(
        doc,
        ["动作", "判定条件（风险%）", "MOM 含义"],
        [
            ["加码", "同向，且两边 |风险%| 都 ≥ 3%；若都 ≥ 8% 标为重仓共识", "方向已被两种风格同时确认，可增加该板块/品种 beta"],
            ["控拥挤", "先满足加码，再叠加 |量化%| + |主观%| ≥ 25%", "方向对，但两边已经打得很重，加码前先设总量上限"],
            ["观望", "反向，且两边 |风险%| 都 ≥ 3%", "风格打架，暂不加该方向，也不要再堆同一边的新投顾"],
            ["补风格", "一侧 ≥ 8%，另一侧 < 1.5%", "只有一种风格重仓，另一侧几乎空白，可考虑补覆盖或只给重仓侧加钱"],
            ["扩容", "量化户数占比与量化保证金占比缺口过大", "新资金先扩哪一侧，与单个板块方向无关"],
            ["中性", "以上都不满足", "噪声或仓位不够明确，信号列表不展示"],
        ],
        col_widths=[2.2, 8.2, 6.0],
    )
    caption(doc, "表 1  五种动作与中性。控拥挤是加码的子集，KPI「共识加码信号」不计入控拥挤。")

    heading(doc, "二、从持仓到风险%：数字怎么来的", 1)
    para(
        doc,
        "风险%不是结算单上的「风险度」，也不是保证金占用比。它衡量的是：在量化（或主观）这一组账户里，"
        "某个板块或品种的净方向风险，占该组全部品种净方向风险绝对值之和的百分之几。正数表示净多，负数表示净空。",
    )

    heading(doc, "2.1 账户分组与数据范围", 2)
    para(
        doc,
        "默认量化账户为 319、324、334、339、346、350、356（rx319 等写法会先抽出数字再匹配）。其余账户记为主观。"
        "页面右侧可以拖动账户，全页按新划分重算。排除国信/GUOXIN/GUOSEN 及账户 665300200077。"
        "期权合约不计入（合约名匹配数字+C/P+数字，或含两段连字符）。同一品种的不同合约先合成一个品种净仓。",
    )
    para(
        doc,
        "持仓来自 mom_position_details：买持仓对应持仓市值计入多头市值，卖持仓计入空头市值，保证金按合约加总。"
        "账户层保证金占用、客户权益来自 mom_daily_reports。品种日涨跌来自 raw_akshare_futures_daily，用于估波动率。",
        space_after=8,
    )

    heading(doc, "2.2 品种净市值", 2)
    para(doc, "对量化组或主观组 g、品种 p，把组内各账户、各合约的多头市值、空头市值分别加总，再取净额：")
    add_eq(
        doc,
        msub(mi("N"), mi("g,p")) + mt(" = ") + msub(mi("L"), mi("g,p")) + mt(" − ") + msub(mi("S"), mi("g,p")),
        number="1",
    )
    add_eq_note(doc, "N 为净市值，L 为多头持仓市值合计，S 为空头持仓市值合计。板块净市值是板块内各品种净市值之和，板块内多空可以对冲。")

    heading(doc, "2.3 品种波动率 σ", 2)
    para(
        doc,
        "每个品种用近 20 个交易日的日收益率样本标准差。收益率取连续主力合约的涨跌幅。"
        "换月跳空会把 σ 抬得很高，所以先做稳健过滤：在过去 40 日绝对收益上取中位数与 MAD，"
        "阈值至少 6%，超过 med + 12 × 1.4826 × MAD 的点记为 0，再在去掉 0 之后的窗口里算标准差。",
    )
    add_eq(
        doc,
        msub(mi("r"), mt("p,t")) + mt("  =  日涨跌幅（换月尖峰已置零）"),
        number="2",
    )
    add_eq(
        doc,
        msub(mi("σ"), mi("p")) + mt(" = ") + msqrt(
            mfrac(
                msum(msup(mparen(msub(mi("r"), mi("i")) + mt(" − ") + mi("r̄")), mt("2")), mi("i ∈ W"), None),
                mt("|W| − 1"),
            )
        ),
        number="3",
    )
    add_eq_note(doc, "W 为近 20 日中非零收益。分母 |W|−1 是样本标准差。若该品种没有有效行情，则改用当日所有品种 σ 的中位数。年化波动展示为 σ×√252。")
    para(
        doc,
        "尖峰过滤阈值本身是：",
        first_line=False,
    )
    add_eq(
        doc,
        mi("θ") + mt(" = ") + mi("max") + mparen(mt("6%,  med(|r|) + 12 × 1.4826 × MAD")),
        number="4",
    )
    para(
        doc,
        "1.4826 把 MAD 换成与正态分布标准差同量纲的尺度。K=12 很严，目的是只砍换月跳空，不砍普通大波动日。",
    )

    heading(doc, "2.4 品种净风险与本组风险预算", 2)
    para(
        doc,
        "风险敞口把「仓位有多大」和「这个品种一天能晃多少」乘在一起。国债名义市值可以很大，但 σ 很小，风险贡献就小；"
        "碳酸锂、股指相反。这正是信号不用保证金%的原因。",
    )
    add_eq(
        doc,
        msub(mi("Risk"), mi("g,p")) + mt(" = ") + msub(mi("σ"), mi("p")) + mt(" × ") + msub(mi("N"), mi("g,p")),
        number="5",
    )
    para(doc, "本组风险预算不是权益，也不是保证金合计，而是组内各品种净风险的绝对值之和：")
    add_eq(
        doc,
        msub(mi("B"), mi("g")) + mt(" = ") + msum(mabs(msub(mi("Risk"), mi("g,p"))), mi("p"), None),
        number="6",
    )
    add_eq_note(doc, "绝对值加在品种层。若改在合约层取绝对值，日历价差会被算成两倍风险，分母会被撑大。")
    para(
        doc,
        "全书风险预算是量化预算加主观预算。板块净风险是板块内各品种净风险相加（先净后加，所以板块内对冲会抵消）：",
    )
    add_eq(
        doc,
        msub(mi("Risk"), mi("g,k")) + mt(" = ") + msum(msub(mi("Risk"), mi("g,p")), mt("p ∈ 板块 k"), None),
        number="7",
    )

    heading(doc, "2.5 风险% 与 保证金%", 2)
    para(doc, "决策用的风险%（可正可负）是：")
    add_eq(
        doc,
        msub(mi("R"), mi("g,x")) + mt(" = 100 × ") + mfrac(msub(mi("Risk"), mi("g,x")), msub(mi("B"), mi("g"))),
        number="8",
    )
    add_eq_note(doc, "x 可以是板块或品种。量化用 B_quant，主观用 B_subj。展示时四舍五入到 1 位小数。")
    para(doc, "另有一个「全书」口径，只进入强弱分数里的 book 项，不进入阈值判定：")
    add_eq(
        doc,
        msubsup(mi("R"), mi("g,x"), mt("book")) + mt(" = 100 × ") + mfrac(
            msub(mi("Risk"), mi("g,x")),
            msub(mi("B"), mi("quant")) + mt(" + ") + msub(mi("B"), mi("subj")),
        ),
        number="9",
    )
    para(doc, "保证金%只用于图和表，信号不用：")
    add_eq(
        doc,
        msub(mi("M"), mi("g,x")) + mt(" = 100 × ") + mfrac(msub(mt("持仓保证金"), mi("g,x")), msub(mt("本组持仓保证金"), mi("g"))),
        number="10",
    )
    para(
        doc,
        "保证金%在组内合计约 100%，且 ≥ 0。风险%一般到不了 ±100%，因为分母是全组 |净风险| 之和，分子只是其中一块。"
        "量化柱和主观柱分母不同，不能把两边的风险%加总解释成「组合有多少仓」。",
    )

    heading(doc, "三、决策逻辑：从两个百分比到一个动作", 1)
    para(
        doc,
        "对每一行板块、每一行品种，记量化风险% 为 R_q、主观风险% 为 R_s。先取绝对值，再看是否同号。",
    )
    add_eq(
        doc,
        mi("q") + mt(" = ") + mabs(msub(mi("R"), mi("q"))) + mt(" ,    ") + mi("s") + mt(" = ") + mabs(msub(mi("R"), mi("s"))),
        number="11",
    )
    add_eq(
        doc,
        mt("同向  ⇔  ") + mparen(msub(mi("R"), mi("q")) + mt(" > 0 且 ") + msub(mi("R"), mi("s")) + mt(" > 0"))
        + mt(" 或 ")
        + mparen(msub(mi("R"), mi("q")) + mt(" < 0 且 ") + msub(mi("R"), mi("s")) + mt(" < 0")),
        number="12",
    )
    para(
        doc,
        "零不当作有方向：一边为 0、另一边为 +5%，既不是同向也不是反向共识。判定按固定优先级，先匹配的规则生效，后面不再看。",
    )
    insert_picture(doc, tree_path, 16.0)
    caption(doc, "图 1  判定顺序。同向且两边都过 3% 的，再检查是否拥挤；过不了 3% 的，才轮到补风格或中性。")

    heading(doc, "3.1 规则的数学形式", 2)
    para(doc, "把代码里的四个常数写进来：CONSENSUS_MIN = 3，HEAVY_MIN = 8，LIGHT_MAX = 1.5，CROWD_SUM = 25。", first_line=False)

    para(doc, "（1）加码 / 重仓共识 / 控拥挤", bold=True, first_line=False, space_after=4)
    add_eq(
        doc,
        mt("同向 且  q ≥ 3 且  s ≥ 3"),
        number="13",
    )
    para(
        doc,
        "先得到「共识做多」或「共识做空」。若再满足 q + s ≥ 25，动作改记为控拥挤，不再叫加码。"
        "若未拥挤且 q ≥ 8 且 s ≥ 8，文案标成重仓共识；否则是弱共识加码。KPI「共识加码信号」只数动作仍为加码的条数。",
    )

    para(doc, "（2）观望", bold=True, first_line=False, space_after=4)
    add_eq(
        doc,
        mt("反向 且  q ≥ 3 且  s ≥ 3"),
        number="14",
    )
    para(doc, "两边都认真地站在对面。MOM 不要在这个方向上再加 beta，也不要再引进同一方向的新投顾。")

    para(doc, "（3）补风格", bold=True, first_line=False, space_after=4)
    add_eq(
        doc,
        mparen(mi("q") + mt(" ≥ 8 且 ") + mi("s") + mt(" < 1.5")) + mt("  或  ") + mparen(mi("s") + mt(" ≥ 8 且 ") + mi("q") + mt(" < 1.5")),
        number="15",
    )
    para(
        doc,
        "注意：这一条排在 3% 共识之后。若一边 12%、一边 4% 且同向，已经在加码里处理完了，不会变成补风格。"
        "补风格要求另一侧几乎空白，不是「比较轻」。",
    )

    para(doc, "（4）中性", bold=True, first_line=False, space_after=4)
    para(
        doc,
        "其余全部为中性，包括：两边都只有 2%；一边 10%、另一边 2%（2% 既不够 3% 的共识，也不够 <1.5% 的空白）；"
        "一边很大、另一边中等但反向却不到 3%。中性行不进入信号列表，避免把噪声当成决策。",
    )

    para(doc, "（5）扩容（账户层，不是板块层）", bold=True, first_line=False, space_after=4)
    para(doc, "设量化户数为 n_q、主观户数为 n_s，量化保证金占全书保证金为 m_q%。户数占比为：")
    add_eq(
        doc,
        msub(mi("c"), mi("q")) + mt(" = 100 × ") + mfrac(msub(mi("n"), mi("q")), msub(mi("n"), mi("q")) + mt(" + ") + msub(mi("n"), mi("s"))),
        number="16",
    )
    add_eq(
        doc,
        mt("若  ") + msub(mi("m"), mi("q")) + mt(" + 8 < ") + msub(mi("c"), mi("q")) + mt("  →  量化侧保证金占比偏低，优先给量化扩容"),
        number="17",
    )
    add_eq(
        doc,
        mt("若  ") + msub(mi("m"), mi("q")) + mt(" > ") + msub(mi("c"), mi("q")) + mt(" + 15  →  量化资金已偏集中，不宜再单边加量化"),
        number="18",
    )
    para(
        doc,
        "两个缺口不对称：户数明显多于资金（差 8 个百分点）就提示扩容；反过来要资金高出户数 15 个百分点才提示过密。"
        "过密比「人多钱少」更危险，所以门槛更高。扩容只出现在列表里，不进入板块对照表的解读列。",
    )

    heading(doc, "3.2 信号强弱分数", 2)
    para(
        doc,
        "列表按强弱降序，最多保留 18 条。分数封顶 100。book 是该行占全书的两边风险%绝对值之和，用来区分「组内比例高但全书很小」和「全书也重」。",
    )
    add_eq(
        doc,
        mi("book") + mt(" = ") + mabs(msubsup(mi("R"), mi("q"), mt("book"))) + mt(" + ") + mabs(msubsup(mi("R"), mi("s"), mt("book"))),
        number="19",
    )
    add_table(
        doc,
        ["情形", "分数（再与 100 取小，保留 1 位小数）", "设计意图"],
        [
            ["控拥挤", "q + s + book", "两边已经很重，分数要直接反映合计拥挤"],
            ["加码且两边 ≥ 8%", "2 × min(q, s) + book", "短板原则：共识强度取决于较轻的一边，再加倍强调重仓"],
            ["加码（弱共识）", "min(q, s) + 0.5 × book", "方向一致但未重仓，全书权重只算一半，避免小品种被放大"],
            ["观望", "1.5 × min(q, s) + book", "分歧比弱共识更需要被看见，所以乘数 1.5"],
            ["补风格", "重仓一侧的 |风险%| + book", "空白侧不贡献共识，只看已表达观点的那一侧"],
            ["中性", "max(q, s)", "不进列表；若使用也只反映较大一侧"],
            ["扩容（人多钱少）", "c_q − m_q + 20", "加 20 是为了让扩容能排进前 18，不被板块信号全部挤掉"],
            ["扩容（钱多人少）", "m_q − c_q", "只按缺口本身排序"],
        ],
        col_widths=[3.4, 6.4, 6.6],
    )
    caption(doc, "表 2  强弱分数。档位：≥20 强，≥8 中，其余弱。")
    para(
        doc,
        "共识分数 consensusScore 是带符号的 min(q, s)：同向为正（做多为正、做空为负），反向为负的 min(q, s)。"
        "它给散点图着色，不决定动作名称。",
    )

    heading(doc, "四、为什么是 3%、8%、1.5%、25%", 1)
    para(
        doc,
        "这四个数不是从回归或回测里估出来的最优解，而是按「这本书大概有多宽、什么叫噪声、什么叫重仓」定的启发式阈值。"
        "它们嵌套在一起，故意留出空白带，避免临界点附近动作来回跳。",
    )
    insert_picture(doc, map_path, 15.6)
    caption(doc, "图 2  第一象限是绝对值平面。同号走加码/控拥挤，异号走观望。补风格贴在坐标轴附近。")

    heading(doc, "4.1 为什么共识门槛是 3%，不是 1% 或 5%", 2)
    para(
        doc,
        "风险%的分母是该组全部品种 |σ×净市值| 之和。一组账户通常覆盖约 10 个板块、二三十个有仓品种。"
        "若风险完全均分到 10 个板块，每块大约 10%；均分到 30 个品种，每个大约 3.3%。"
        "因此 3% 大致等于「一个中等偏上的品种权」、或「一个板块正常配置的三成」。"
        "到了这个量级，可以认为该侧不是换月残仓、不是四舍五入，而是一次有意识的暴露。",
    )
    para(doc, "更具体的噪声来源有四类，3% 把它们挡在门外：", first_line=False)
    para(doc, "第一，展示四舍五入到 0.1%。2.96% 会显示成 3.0%，3% 是屏幕上能稳定看见的第一档整数。", first_line=False)
    para(doc, "第二，零散手数。CTA 调仓后常留下 1–2 手，折成风险%往往在 0.x% 到 2% 之间。", first_line=False)
    para(doc, "第三，σ 估计误差。20 日窗口、换月过滤之后，σ 仍会抖。把门槛放到 1% 会把估计噪声当成「共识」。", first_line=False)
    para(doc, "第四，品种过滤本身已经丢掉 |q|+|s|<0.3% 的行。3% 是在「行已经出现在表里」之后，再要求两边都足够大。", first_line=False)
    para(
        doc,
        "若改成 1%，几乎每个略有重叠的品种都会亮加码或观望，列表失去重点。"
        "若改成 5%，会漏掉 3%–5% 的真实同向（对品种层来说 4% 已经是前排权重）。"
        "3% 卡在「平均品种权」附近：到了平均值，才配谈共识或分歧。",
    )

    heading(doc, "4.2 为什么重仓线是 8%", 2)
    para(
        doc,
        "8% 约为一个典型板块均分权重（10%）的八成，或两到三个头部品种的量级。"
        "一边到 8%，说明该风格把相当一部分风险预算压在这个主题上，不是「顺便带了一点」。"
        "两边同时 ≥ 8%，才叫重仓共识：两种完全不同的决策流程都把真金白银放在同一方向。"
        "弱共识（3%–8%）只支持小幅加暴露或再观察；重仓共识才作为给同向账户加钱、或引进同向投顾的依据。",
    )
    para(
        doc,
        "8% 也是补风格里「重」的定义。只有一侧重到 8%，另一侧才有必要讨论要不要补人。"
        "若把重仓线降到 5%，补风格会变得很多，MOM 会误以为每个中等仓位都缺另一半风格。",
    )

    heading(doc, "4.3 为什么空白线是 1.5%，而且中间故意空着", 2)
    para(
        doc,
        "补风格要求另一侧 < 1.5%，不是 ≤ 3%。1.5% 大约是 3% 的一半，也大约是四舍五入噪声和残仓的上沿。"
        "代码用严格小于：s < 1.5。等于 1.5% 仍不算空白。",
    )
    para(
        doc,
        "1.5% 与 3% 之间是缓冲区。例如量化 10%、主观 2%：主观已经不是空白，但还没到「认真表态」的 3%，"
        "所以既不加码、也不补风格，记中性。若没有这条缝，2% 会被时而当成共识、时而当成空白，动作会随 σ 的小抖动来回切。"
        "这是阈值设计里最容易被忽略、却最有用的一块。",
    )

    heading(doc, "4.4 为什么拥挤线是两边之和 25%", 2)
    para(
        doc,
        "q 和 s 各自相对自己的风险预算。q + s ≥ 25 并不是「组合 25% 的风险在这个板块」，"
        "因为两个分母不同，不能当组合权重加。它是一个启发式拥挤分：两种风格都已经把很大一块自己的预算打进同一主题。",
    )
    para(
        doc,
        "25% 可以读成「约 2.5 个平均板块」。若两边都把四分之一的风险预算堆进同一方向，MOM 再按共识去加码，"
        "会在同一拥挤点上加倍。所以规则是：方向仍然算共识，但动作从加码改成控拥挤——先设总量上限，而不是继续加 beta。",
    )
    para(
        doc,
        "对称的解读：若两边都是 12.5%，已经同时过了重仓线 8%，共识很强，但合计触及 25%，拥挤优先于加码。"
        "这保证「越一致越要加」不会在高拥挤区失控。KPI 共识加码也不计入这些行，避免把该减速的地方数成该加速。",
    )

    heading(doc, "4.5 扩容的 8 与 15 为什么不对称", 2)
    para(
        doc,
        "户数占比高于资金占比 8 个点：人已经在量化侧就位，但钱没跟上，容量偏紧，新资金可以优先进量化。"
        "反过来，资金占比高于户数占比要 15 个点才报警：这是「钱已经堆在少数量化户上」。"
        "过密的伤害是拥挤和单账户穿透，所以门槛更高；人多钱少只是效率问题，8 个点就提示可以扩。",
    )

    heading(doc, "4.6 为什么信号锁定风险%，不跟保证金开关走", 2)
    para(
        doc,
        "保证金占比衡量的是开仓占用了多少资金，不衡量这个仓位明天能亏多少。"
        "国债、股指、工业品的保证金率和波动率差一个数量级。若用保证金%做共识，会把「占了很多保证金的低波动品种」"
        "误判成重仓共识。风险% = σ×净市值 / 组内预算，已经把波动换算进去，两边才有可比的「打得有多重」。",
    )
    para(
        doc,
        "同向、反向看的是风险%的符号。保证金%恒为非负，没有多空，不能定义观望。"
        "所以即使用户把图切到保证金，判定函数仍传入 metric = \"risk\"。",
    )

    heading(doc, "五、数值例子（按代码逐步算）", 1)
    para(doc, "下面七个例子只使用四舍五入后的风险%，与页面一致。book 假设为 2.0，便于对照分数公式。", first_line=False)

    add_table(
        doc,
        ["例子", "R_q", "R_s", "同向?", "q, s", "命中规则", "动作", "分数"],
        [
            ["A 弱共识", "+5.0%", "+4.0%", "是", "5.0, 4.0", "同向且都 ≥3%，和 9<25，未到双 8", "加码", "min=4；4+0.5×2=5.0"],
            ["B 重仓共识", "+12.0%", "+10.0%", "是", "12, 10", "同向、都 ≥8%，和 22<25", "加码（重仓）", "2×10+2=22.0  强"],
            ["C 控拥挤", "+15.0%", "+14.0%", "是", "15, 14", "先是共识，再 29≥25", "控拥挤", "15+14+2=31.0  强"],
            ["D 观望", "+8.0%", "−6.0%", "否", "8, 6", "反向且都 ≥3%", "观望", "1.5×6+2=11.0  中"],
            ["E 补风格", "+12.0%", "+0.4%", "是*", "12, 0.4", "未过双 3%；量化≥8 且主观<1.5", "补风格", "12+2=14.0  中"],
            ["F 中性（都轻）", "+2.0%", "+2.0%", "是", "2, 2", "同向但都 <3%，也不是补风格", "中性", "max=2.0"],
            ["G 中性（缓冲带）", "+10.0%", "+2.0%", "是", "10, 2", "2 不够 3%，也不 <1.5", "中性", "max=10.0"],
        ],
        col_widths=[2.4, 1.6, 1.6, 1.3, 1.8, 4.0, 2.0, 2.5],
    )
    caption(doc, "表 3  例 E 的「同向」对分类无用，因为没过 3% 门槛，会落到补风格。例 G 是 1.5%–3% 缓冲带，刻意不发信号。")

    para(doc, "把例 C 写开。量化 15%、主观 14%，同号且都大于 3%，进入共识分支；15+14=29≥25，动作从加码改成控拥挤。强弱：")
    add_eq(
        doc,
        mi("strength") + mt(" = min") + mparen(mt("100,  15 + 14 + 2")) + mt(" = 31"),
        number="20",
    )
    para(doc, "例 B 未拥挤、双侧重仓：")
    add_eq(
        doc,
        mi("strength") + mt(" = min") + mparen(mt("100,  2 × min(12,10) + 2")) + mt(" = 22"),
        number="21",
    )
    para(
        doc,
        "若把例 B 的主观改成 −10%，则反向且都 ≥3%，动作变成观望，分数 1.5×min(12,10)+2=17。"
        "同一对绝对值，符号一翻，从「可以加码」变成「不要加」。这就是为什么必须用带符号的风险%，不能用保证金%。",
    )

    heading(doc, "六、实现上的细节（和页面数字对得上）", 1)
    add_table(
        doc,
        ["项目", "规则"],
        [
            ["信号口径", "永远 metric = risk。与柱状图开关无关。"],
            ["板块集合", "农产、生鲜、贵金属、有色、新能源、黑色、能源化工、航运、股指、国债、其他。"],
            ["品种入表", "|量化风险%| + |主观风险%| ≥ 0.3，或两边保证金%之和 ≥ 0.5。"],
            ["送入信号的品种", "按 |q|+|s| 排序后只取前 40 个，避免长尾品种刷屏。"],
            ["输出条数", "按强弱降序，最多 18 条（含扩容）。"],
            ["期权", "不计入净市值、保证金和风险预算。"],
            ["合约合成", "同一品种不同月份先净仓，再乘 σ。"],
            ["σ 缺失", "该品种用当日其他品种 σ 的中位数。"],
            ["百分比精度", "乘 100 后四舍五入到 1 位，再和 3 / 8 / 1.5 / 25 比较。"],
            ["方向标签", "|风险%| > 0.15 才标「多/空」，否则标「平」（只影响文案）。"],
            ["KPI 共识加码", "动作 = 加码的板块+品种条数；控拥挤另行计数，不计入该 KPI。"],
            ["点击信号", "筛选下方多空持仓图，不改计算。"],
        ],
        col_widths=[3.6, 12.8],
    )
    caption(doc, "表 4  与 app/ma/api/mom-analysis/quant-vs-subjective 及 lib/ma/quant-vs-subjective-signals.ts 对齐的实现约定。")

    heading(doc, "七、这套规则能做什么、不能做什么", 1)
    para(
        doc,
        "能做的：在同一个结算日截面上，把两种风格的风险预算地图叠在一起，标出「两边都认真且同向」"
        "「两边都认真但反向」「只有一边认真」和「两边已经过密」。供 MOM 决定加钱、引进投顾、或先设上限。",
    )
    para(
        doc,
        "不能做的：它不是预测下一日涨跌的模型，没有样本外检验，阈值没有用历史胜率标定。"
        "q + s 不是组合权重。板块内对冲之后的净风险，会掩盖「组内多空都很大但净额不大」的结构；那种结构要看下方多空持仓图。"
        "扩容只用户数和保证金，不管投顾质量。账户从量化拖到主观会重算全页，信号随之变，这是分组定义问题，不是行情变了。",
    )
    para(
        doc,
        "若以后要改阈值，建议仍保持嵌套：空白 < 共识 < 重仓，并保留 1.5%–3% 缓冲带；拥挤线应明显高于 2×8%=16%，"
        "否则几乎所有重仓共识都会被改写成控拥挤。改之前应用若干历史截面数一下各动作的条数，避免列表从「每天十几条」变成「每天一条」或「每天四十条」。",
    )

    heading(doc, "附录  符号表", 1)
    add_table(
        doc,
        ["符号", "含义"],
        [
            ["g", "风格组：quant 或 subjective"],
            ["p, k, x", "品种、板块、或二者之一"],
            ["L, S, N", "多头市值、空头市值、净市值 L−S"],
            ["σ_p", "品种 p 近 20 日样本标准差（换月尖峰已剔除）"],
            ["Risk_g,x", "σ × 净市值；板块为品种风险之和"],
            ["B_g", "本组风险预算 Σ_p |Risk_g,p|"],
            ["R_q, R_s", "量化 / 主观风险%，可正可负"],
            ["q, s", "|R_q|、|R_s|"],
            ["book", "该行两边全书风险%绝对值之和"],
            ["m_q, c_q", "量化保证金占比、量化户数占比"],
            ["3 / 8 / 1.5 / 25", "共识 / 重仓 / 空白 / 拥挤 四个阈值（百分数）"],
        ],
        col_widths=[3.2, 13.2],
    )

    para(
        doc,
        "本文公式与判定顺序以当前生产代码为准。页面问号弹窗是同一套规则的缩写。",
        size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_before=16,
    )

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))
    doc.save(str(REPORT_PATH_ASCII))
    return REPORT_PATH


if __name__ == "__main__":
    out = build()
    print(out)
