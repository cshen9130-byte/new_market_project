# -*- coding: utf-8 -*-
"""Login habit + deploy-window Word report from public.auth_login_history."""
from __future__ import annotations

import os
import sys
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import psycopg2
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.font_manager import FontProperties, fontManager
from psycopg2.extras import RealDictCursor

DB_URL = os.environ.get("DATABASE_URL") or (
    "postgresql://market_user:2026SmartDashboard!@127.0.0.1:5433/market_data"
)
OUT_DIR = Path(__file__).resolve().parents[2] / "reports"
CHART_DIR = OUT_DIR / "_login_report_charts"
REPORT_PATH = OUT_DIR / "Login_Usage_and_Deploy_Window_Report.docx"

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_NAVY = "#1A365D"
C_TEAL = "#2B6CB0"
C_GOLD = "#C9A227"
C_GREEN = "#2F855A"
C_RED = "#C53030"
C_ORANGE = "#DD6B20"
C_GRAY = "#718096"
C_PURPLE = "#6B46C1"
PALETTE = [C_NAVY, C_TEAL, C_GOLD, C_GREEN, C_ORANGE, C_PURPLE, C_RED, C_GRAY]

_CN_FONT: FontProperties | None = None

SELFTEST = "__login_history_selftest__"
BURST_MINUTES = 15
TZ_LABEL = "Asia/Shanghai (UTC+8)"

ROLE_BY_NAME = {}


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    plt.rcParams["axes.edgecolor"] = "#CBD5E0"
    plt.rcParams["axes.grid"] = False
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.isfile(path):
            try:
                fontManager.addfont(path)
                _CN_FONT = FontProperties(fname=path)
                plt.rcParams["font.family"] = "sans-serif"
                plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "Microsoft YaHei", "SimHei"]
                return
            except Exception:
                continue
    _CN_FONT = FontProperties(family="Microsoft YaHei")


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.2
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=9, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "F7FAFC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            cell_text(
                table.rows[r_i + 1].cells[c_i],
                val,
                size=8,
                color=TEXT,
                align="center" if c_i else "left",
            )
            shade(table.rows[r_i + 1].cells[c_i], fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_picture(doc, path: Path, width_in=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    return p


def caption(doc, text: str):
    para(doc, text, size=8, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def classify_device(ua: str | None) -> str:
    s = (ua or "").lower()
    if "curl" in s:
        return "Script"
    if "iphone" in s or "android" in s or "mobile" in s or "huaweibrowser" in s:
        return "Mobile"
    if "edg/" in s:
        return "Edge"
    if "quark" in s:
        return "Quark"
    if "chrome" in s:
        return "Chrome"
    return "Other"


def classify_network(ip: str | None) -> str:
    ip = (ip or "").strip()
    if ip in ("::1", "127.0.0.1"):
        return "Localhost (dev)"
    if ip.startswith("203.0.113."):
        return "Test / reserved"
    if ip.startswith("116.237.193."):
        return "Office network A"
    if ip.startswith("116.234.199."):
        return "Evening / network B"
    if ip.startswith("116.234.86."):
        return "Network C"
    if ip.startswith("111.187."):
        return "Network D"
    if ip.startswith("39.144."):
        return "Mobile carrier"
    if ip.startswith("124.79."):
        return "Network E"
    if ip.startswith("61.170."):
        return "Network F"
    return "Other"


def collapse_sessions(g: pd.DataFrame, minutes: int = BURST_MINUTES) -> pd.DataFrame:
    g = g.sort_values("ts")
    sessions = []
    current = None
    for _, row in g.iterrows():
        if current is None or (row["ts"] - current["end"]).total_seconds() > minutes * 60:
            if current is not None:
                sessions.append(current)
            current = {
                "who": row["who"],
                "start": row["ts"],
                "end": row["ts"],
                "n": 1,
                "ok": int(bool(row["success"])),
                "fail": int(not bool(row["success"])),
                "devices": {row["device"]},
                "networks": {row["network"]},
            }
        else:
            current["end"] = row["ts"]
            current["n"] += 1
            current["ok"] += int(bool(row["success"]))
            current["fail"] += int(not bool(row["success"]))
            current["devices"].add(row["device"])
            current["networks"].add(row["network"])
    if current is not None:
        sessions.append(current)
    return pd.DataFrame(sessions)


def fetch() -> tuple[pd.DataFrame, pd.DataFrame]:
    conn = psycopg2.connect(DB_URL)
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute(
        """
        SELECT id, name, email, role, created_at
          FROM public.auth_users
         ORDER BY created_at
        """
    )
    users = pd.DataFrame([dict(r) for r in cur.fetchall()])
    cur.execute(
        """
        SELECT
          COALESCE(NULLIF(name, ''), identifier) AS who,
          logged_at AT TIME ZONE 'Asia/Shanghai' AS ts,
          success,
          fail_reason,
          user_agent,
          ip
        FROM public.auth_login_history
        ORDER BY logged_at
        """
    )
    hist = pd.DataFrame([dict(r) for r in cur.fetchall()])
    conn.close()
    hist["ts"] = pd.to_datetime(hist["ts"])
    hist["device"] = hist["user_agent"].map(classify_device)
    hist["network"] = hist["ip"].map(classify_network)
    hist["hour"] = hist["ts"].dt.hour
    hist["date"] = hist["ts"].dt.date
    hist["weekday"] = hist["ts"].dt.day_name()
    hist["is_selftest"] = hist["who"] == SELFTEST
    hist["is_localhost"] = hist["network"].eq("Localhost (dev)")
    hist["is_prod"] = (~hist["is_selftest"]) & (~hist["is_localhost"])
    return users, hist


def style_ax(ax, title: str, xlabel: str, ylabel: str):
    ax.set_title(title, **fp(), fontsize=12, color=C_NAVY, pad=10)
    ax.set_xlabel(xlabel, **fp(), fontsize=9, color="#4A5568")
    ax.set_ylabel(ylabel, **fp(), fontsize=9, color="#4A5568")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(colors="#4A5568", labelsize=8)
    for lab in ax.get_xticklabels() + ax.get_yticklabels():
        lab.set_fontproperties(_CN_FONT)


def save_fig(fig, name: str) -> Path:
    path = CHART_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def make_charts(hist: pd.DataFrame, sessions: pd.DataFrame) -> dict[str, Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    people = hist.loc[~hist["is_selftest"]].copy()
    prod = people.loc[people["is_prod"]].copy()
    charts = {}

    # 1. Daily volume stacked by person
    daily = (
        people.groupby(["date", "who"])
        .size()
        .unstack(fill_value=0)
        .sort_index()
    )
    fig, ax = plt.subplots(figsize=(8.2, 3.8))
    bottom = np.zeros(len(daily))
    x = np.arange(len(daily))
    names = list(daily.columns)
    for i, name in enumerate(names):
        vals = daily[name].to_numpy()
        ax.bar(x, vals, bottom=bottom, color=PALETTE[i % len(PALETTE)], width=0.62, label=name)
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([d.strftime("%a %m-%d") for d in daily.index], **fp())
    ax.legend(prop=_CN_FONT, fontsize=8, frameon=False, ncol=3, loc="upper left")
    style_ax(ax, "Daily login events by person", "Date (CST)", "Login events")
    charts["daily"] = save_fig(fig, "daily_by_person.png")

    # 2. Hour-of-day: all vs production
    hours = np.arange(24)
    all_h = people.groupby("hour").size().reindex(hours, fill_value=0)
    prod_h = prod.groupby("hour").size().reindex(hours, fill_value=0)
    fig, ax = plt.subplots(figsize=(8.2, 3.6))
    ax.bar(hours - 0.18, all_h.to_numpy(), width=0.36, color=C_NAVY, label="All (incl. localhost)")
    ax.bar(hours + 0.18, prod_h.to_numpy(), width=0.36, color=C_TEAL, label="Production only")
    ax.set_xticks(hours)
    ax.legend(prop=_CN_FONT, fontsize=8, frameon=False)
    style_ax(ax, "Login traffic by hour of day", "Hour (CST)", "Login events")
    charts["hourly"] = save_fig(fig, "hourly.png")

    # 3. Unique people by hour (production)
    uniq_h = prod.groupby("hour")["who"].nunique().reindex(hours, fill_value=0)
    fig, ax = plt.subplots(figsize=(8.2, 3.4))
    colors = [C_RED if v >= 3 else (C_GOLD if v == 2 else (C_TEAL if v == 1 else "#E2E8F0")) for v in uniq_h]
    ax.bar(hours, uniq_h.to_numpy(), color=colors, width=0.7)
    ax.set_xticks(hours)
    ax.set_yticks(range(0, int(uniq_h.max()) + 2))
    style_ax(ax, "Unique people online (production logins) by hour", "Hour (CST)", "Unique people")
    charts["unique_hour"] = save_fig(fig, "unique_hour.png")

    # 4. Weekday x hour heatmap (production)
    wd_order = ["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday"]
    mat = (
        prod.groupby(["weekday", "hour"])
        .size()
        .unstack(fill_value=0)
        .reindex(index=wd_order, columns=hours, fill_value=0)
    )
    fig, ax = plt.subplots(figsize=(8.4, 3.2))
    im = ax.imshow(mat.to_numpy(), aspect="auto", cmap="Blues", vmin=0)
    ax.set_xticks(hours)
    ax.set_yticks(range(len(wd_order)))
    ax.set_yticklabels(["Sun 8/23", "Mon 8/24", "Tue 8/25", "Wed 8/26", "Thu 8/27"], **fp())
    ax.set_xlabel("Hour (CST)", **fp(), fontsize=9, color="#4A5568")
    ax.set_title("Production login heatmap (day × hour)", **fp(), fontsize=12, color=C_NAVY, pad=10)
    ax.tick_params(colors="#4A5568", labelsize=8)
    cbar = fig.colorbar(im, ax=ax, fraction=0.03, pad=0.02)
    cbar.ax.tick_params(labelsize=7)
    cbar.set_label("Events", **fp(), fontsize=8)
    fig.tight_layout()
    charts["heatmap"] = save_fig(fig, "heatmap.png")

    # 5. Person volume
    vol = people.groupby("who").size().sort_values(ascending=True)
    fig, ax = plt.subplots(figsize=(8.2, 3.4))
    ax.barh(vol.index, vol.values, color=C_NAVY, height=0.55)
    for i, v in enumerate(vol.values):
        ax.text(v + 0.3, i, str(v), va="center", **fp(), fontsize=8, color="#4A5568")
    style_ax(ax, "Login events per person", "Login events", "")
    charts["volume"] = save_fig(fig, "volume.png")

    # 6. Device mix
    device_counts = people["device"].value_counts()
    fig, ax = plt.subplots(figsize=(5.6, 3.6))
    wedges, texts, autotexts = ax.pie(
        device_counts.values,
        labels=device_counts.index,
        colors=PALETTE[: len(device_counts)],
        autopct=lambda p: f"{p:.0f}%" if p >= 6 else "",
        startangle=90,
        wedgeprops={"width": 0.45, "edgecolor": "white"},
        pctdistance=0.75,
    )
    for t in list(texts) + list(autotexts):
        t.set_fontproperties(_CN_FONT)
        t.set_fontsize(8)
        t.set_color("#2D3748")
    ax.set_title("Device / browser mix", **fp(), fontsize=12, color=C_NAVY)
    charts["device"] = save_fig(fig, "device.png")

    # 7. Sessions vs raw logins
    sess_vol = sessions.groupby("who").size().sort_index()
    raw_vol = people.groupby("who").size().reindex(sess_vol.index, fill_value=0)
    fig, ax = plt.subplots(figsize=(8.2, 3.5))
    idx = np.arange(len(sess_vol))
    ax.bar(idx - 0.18, raw_vol.to_numpy(), width=0.36, color=C_NAVY, label="Raw login events")
    ax.bar(idx + 0.18, sess_vol.to_numpy(), width=0.36, color=C_GOLD, label=f"Visits ({BURST_MINUTES}-min sessions)")
    ax.set_xticks(idx)
    ax.set_xticklabels(list(sess_vol.index), **fp(), rotation=20)
    ax.legend(prop=_CN_FONT, fontsize=8, frameon=False)
    style_ax(ax, "Raw logins vs collapsed visits", "Person", "Count")
    charts["sessions"] = save_fig(fig, "sessions.png")

    # 8. First login of day (production, per person-day)
    firsts = prod.sort_values("ts").groupby(["who", "date"])["ts"].first().reset_index()
    firsts["hour"] = firsts["ts"].dt.hour + firsts["ts"].dt.minute / 60
    fig, ax = plt.subplots(figsize=(8.2, 3.6))
    people_names = sorted(firsts["who"].unique())
    for i, name in enumerate(people_names):
        sub = firsts.loc[firsts["who"] == name, "hour"]
        ax.scatter(
            sub,
            np.full(len(sub), i),
            s=48,
            color=PALETTE[i % len(PALETTE)],
            zorder=3,
            label=name,
        )
        ax.plot(sub, np.full(len(sub), i), color=PALETTE[i % len(PALETTE)], alpha=0.35, lw=1.5)
    ax.set_yticks(range(len(people_names)))
    ax.set_yticklabels(people_names, **fp())
    ax.set_xlim(8, 22)
    ax.set_xticks(range(8, 23))
    style_ax(ax, "First login of each calendar day", "Hour (CST)", "")
    charts["first"] = save_fig(fig, "first_login.png")

    return charts


def hour_risk_table(prod: pd.DataFrame) -> list[list[str]]:
    hours = list(range(24))
    rows = []
    for h in hours:
        sub = prod.loc[prod["hour"] == h]
        n = int(len(sub))
        people = sorted(sub["who"].unique())
        if n == 0:
            risk = "Lowest"
        elif n <= 2 and len(people) == 1:
            risk = "Low"
        elif len(people) >= 3 or n >= 8:
            risk = "Highest"
        elif len(people) == 2 or n >= 4:
            risk = "High"
        else:
            risk = "Medium"
        rows.append(
            [
                f"{h:02d}:00–{h:02d}:59",
                str(n),
                str(len(people)),
                ", ".join(people) if people else "—",
                risk,
            ]
        )
    return rows


def person_profile(name: str, hist: pd.DataFrame, sessions: pd.DataFrame, users: pd.DataFrame) -> dict:
    h = hist.loc[(hist["who"] == name) & (~hist["is_selftest"])].sort_values("ts")
    s = sessions.loc[sessions["who"] == name] if len(sessions) else pd.DataFrame()
    role = ROLE_BY_NAME.get(name, "unknown")
    urow = users.loc[users["name"] == name]
    created = ""
    if len(urow):
        created = pd.to_datetime(urow.iloc[0]["created_at"]).strftime("%Y-%m-%d")
        role = str(urow.iloc[0]["role"])
    firsts = h.groupby("date")["ts"].first()
    lasts = h.groupby("date")["ts"].last()
    typical_first = firsts.dt.strftime("%H:%M").tolist()
    devices = h["device"].value_counts()
    nets = h["network"].value_counts()
    ok = int(h["success"].sum())
    fail = int((~h["success"]).sum())
    return {
        "name": name,
        "role": role,
        "created": created,
        "events": int(len(h)),
        "visits": int(len(s)),
        "ok": ok,
        "fail": fail,
        "days_active": int(h["date"].nunique()),
        "first_times": typical_first,
        "earliest": h["ts"].min().strftime("%Y-%m-%d %H:%M") if len(h) else "—",
        "latest": h["ts"].max().strftime("%Y-%m-%d %H:%M") if len(h) else "—",
        "median_hour": float(h["hour"].median()) if len(h) else None,
        "devices": devices,
        "networks": nets,
        "hours": sorted(h["hour"].unique().tolist()),
        "fail_reasons": h.loc[~h["success"], "fail_reason"].value_counts().to_dict(),
        "localhost_share": float(h["is_localhost"].mean()) if len(h) else 0.0,
        "weekday_only": all(d.weekday() < 5 for d in h["date"]) if len(h) else True,
        "rows": h,
        "sessions": s,
        "firsts": firsts,
        "lasts": lasts,
    }


def habit_paragraph(p: dict) -> str:
    name = p["name"]
    parts = []
    if p["events"] == 0:
        return f"{name} has no recorded login in this window."

    first_txt = ", ".join(p["first_times"]) if p["first_times"] else "—"
    device_txt = ", ".join(f"{k} {int(v)}" for k, v in p["devices"].items())
    net_txt = ", ".join(f"{k} ({int(v)})" for k, v in p["networks"].items())

    if name == "cshen":
        parts.append(
            f"{name} (admin) accounts for {p['events']} of the recorded events "
            f"({p['visits']} distinct visits after collapsing {BURST_MINUTES}-minute bursts). "
            f"Weekday first login is extremely stable at about 09:06–09:09 CST "
            f"(observed first-of-day times: {first_txt}). "
            f"A second cluster usually appears through 09:30–10:22, including localhost "
            f"({p['localhost_share']:.0%} of events), which is local development rather than "
            f"production traffic. Afternoon use is lighter (13:40–14:49). Evening work is "
            f"regular: 18:25–21:55 on Edge from a second network, plus occasional Huawei "
            f"mobile at lunch/commute (18:25 Mon, 12:14 Tue). "
            f"This is a builder/operator pattern: morning open, mid-day checks, night coding. "
            f"Any production deploy between 09:00–11:00 or 18:30–22:00 will almost certainly "
            f"interrupt this user."
        )
    elif name == "sunjie":
        parts.append(
            f"{name} logged in {p['events']} times across {p['days_active']} weekdays, "
            f"all successful, all from Quark on the office network. First-of-day times: {first_txt}. "
            f"Activity sits in a mid-morning to early-afternoon band (10:31–14:23) and never "
            f"appears at night in this sample. Habit: office-hours desktop user, arrives later "
            f"than the admin, typically one or two visits per working day."
        )
    elif name == "G.Wave":
        parts.append(
            f"{name} has {p['events']} events ({p['ok']} success / {p['fail']} fail) on "
            f"{p['days_active']} days. On Tue 8/25 a mobile login at 13:29 was followed two "
            f"minutes later by two invalid-credential failures on Chrome, then a success at 13:32 — "
            f"a password-retry cluster, not a security scan. A later desktop login at 19:43 Thu "
            f"shows evening use. Network family is distinct from the main office block, so this "
            f"person is usually remote. Devices mix mobile + Chrome/Edge. Habit: afternoon and "
            f"evening, not a 09:00 starter."
        )
    elif name == "luoshuang":
        parts.append(
            f"{name} has {p['events']} successful Chrome desktop logins on {p['days_active']} days "
            f"(13:23 Tue; 10:35 and 14:25 Wed) from a single stable network. First-of-day: {first_txt}. "
            f"Habit: weekday office hours only, one to two visits, no evening or mobile in this window."
        )
    elif name == "chenpeifeng":
        parts.append(
            f"{name} is the most off-hours user. Three visits: Sunday 21:38 mobile, Tuesday 13:07 "
            f"mobile from the office network, Tuesday 22:11 Edge from another network. "
            f"First-of-day times: {first_txt}. Habit: mobile-first, irregular, including Sunday night "
            f"and late evening. A late deploy (22:00+) can still hit this person even when the "
            f"office is empty."
        )
    elif name == "liuyamin":
        parts.append(
            f"{name} appears once: Tue 8/25 10:23 on Edge from a unique network. "
            f"Too little data to call a habit; the single observation is a weekday morning visit."
        )
    else:
        parts.append(
            f"{name} ({p['role']}): {p['events']} events, {p['visits']} visits, "
            f"{p['ok']} success / {p['fail']} fail, active {p['days_active']} days. "
            f"Devices: {device_txt}. Networks: {net_txt}. First-of-day: {first_txt}."
        )
    return " ".join(parts)


def build_doc(users: pd.DataFrame, hist: pd.DataFrame, charts: dict[str, Path], sessions: pd.DataFrame) -> None:
    people = hist.loc[~hist["is_selftest"]].copy()
    prod = people.loc[people["is_prod"]].copy()
    tmin = hist["ts"].min()
    tmax = hist["ts"].max()
    n_all = len(hist)
    n_ok = int(hist["success"].sum())
    n_fail = n_all - n_ok
    n_prod = len(prod)
    n_local = int(people["is_localhost"].sum())
    names_active = sorted(people["who"].unique())
    names_all = list(users["name"])
    inactive = [n for n in names_all if n not in set(names_active)]

    profiles = {n: person_profile(n, hist, sessions, users) for n in names_active}

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    para(doc, "Market Dashboard  ·  Internal operations", size=10, color=GOLD, space_after=2)
    p = para(doc, "Login Usage, User Habits, and Deploy-Window Analysis", size=22, bold=True, color=NAVY, space_after=4)
    p.runs[0].font.size = Pt(22)
    para(
        doc,
        f"Source: public.auth_login_history  ·  Timezone: {TZ_LABEL}  ·  "
        f"Window: {tmin.strftime('%Y-%m-%d %H:%M')} – {tmax.strftime('%Y-%m-%d %H:%M')}  ·  "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        size=9,
        color=MUTED,
        space_after=14,
    )

    heading(doc, "1. Executive summary", 1)
    para(
        doc,
        "Login recording started on the evening of Sunday 23 August 2026, so this is a four-day "
        "weekday-heavy sample (Sun evening plus Mon–Thu), not a full month. Treat the deploy "
        "windows as a first cut to be re-checked after two more weeks of data. Within that limit "
        "the pattern is already sharp.",
    )
    para(
        doc,
        f"There are {n_all} recorded attempts ({n_ok} success, {n_fail} failure). "
        f"{n_local} events are localhost (admin development) and 1 is a self-test; "
        f"{n_prod} are production-facing. Six of twelve accounts logged in. "
        f"cshen generated {profiles['cshen']['events']} events "
        f"({profiles['cshen']['events'] / max(len(people), 1):.0%} of person traffic) but only "
        f"{profiles['cshen']['visits']} collapsed visits — many clicks are test bursts, not extra people online.",
    )
    para(doc, "Deploy recommendation (CST), ranked:", size=11, bold=True, color=NAVY, space_after=4)
    para(
        doc,
        "1) Best: 00:00–08:00. Zero logins in the sample. Overnight is the only empty band.",
        space_after=3,
    )
    para(
        doc,
        "2) Same-day low risk: 15:00–17:30. The only afternoon event in this window was a "
        "localhost login at 16:24 Thursday — not production users.",
        space_after=3,
    )
    para(
        doc,
        "3) Acceptable: 11:00–12:50. Lunch lull (sunjie once at 11:24, cshen mobile once at 12:14).",
        space_after=3,
    )
    para(
        doc,
        "4) Do not deploy: 09:00–10:45 (morning open) and 13:00–14:30 (highest unique-user overlap). "
        "Evening 18:30–22:00 is also busy because the admin works late and two others appear after 19:40.",
        space_after=12,
    )

    heading(doc, "2. Data coverage and caveats", 1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Table", "public.auth_login_history"],
            ["Timezone used in this report", TZ_LABEL],
            ["First record", tmin.strftime("%Y-%m-%d %H:%M:%S") + " (self-test)"],
            ["Last record", tmax.strftime("%Y-%m-%d %H:%M:%S")],
            ["Calendar span", "4.0 days (Sun evening + Mon–Thu)"],
            ["Events / success / fail", f"{n_all} / {n_ok} / {n_fail}"],
            ["Production events (excl. localhost & self-test)", str(n_prod)],
            ["Localhost (dev) events", str(n_local)],
            ["Registered accounts", str(len(users))],
            ["Accounts seen in this window", str(len(names_active))],
            ["Accounts with no login", ", ".join(inactive) if inactive else "—"],
            ["Visit definition", f"consecutive events ≤ {BURST_MINUTES} min apart, same person"],
        ],
        col_widths=[6.5, 10.5],
    )
    para(doc, "", space_after=6)
    para(
        doc,
        "A login row is an authentication event, not a session length. The site can force a new "
        "login after a deploy, a cookie expiry, or a local-dev restart — which is why cshen’s raw "
        "count is inflated. Friday, Saturday, and a full Sunday are not in the sample. "
        "Failed logins are rare (two password typos + one incomplete self-test).",
        size=10,
        color=MUTED,
    )

    heading(doc, "3. Traffic overview", 1)
    para(
        doc,
        "Volume is dominated by one operator, then a small office cohort on weekdays, then a "
        "remote/mobile tail in the afternoon and evening. Daily totals rise from Monday into "
        "Tuesday–Wednesday (more unique people) and stay high on Thursday because of admin bursts.",
    )
    add_picture(doc, charts["daily"], 6.3)
    caption(doc, "Figure 1. Daily login events stacked by person. Source: auth_login_history · 23–27 Aug 2026 CST.")

    add_picture(doc, charts["volume"], 6.3)
    caption(doc, "Figure 2. Raw login events per person (not collapsed).")

    add_picture(doc, charts["sessions"], 6.3)
    caption(
        doc,
        f"Figure 3. Raw events vs {BURST_MINUTES}-minute visits. cshen’s gap is the development/test burst effect.",
    )

    heading(doc, "4. When people are actually online", 1)
    para(
        doc,
        "Three peaks show up in clock time: morning open (09–10), post-lunch (13–14), and evening "
        "(18–21). The empty overnight band is absolute in this sample. Unique-person count, not "
        "raw events, is the right metric for “will a restart kick someone out”.",
    )
    add_picture(doc, charts["hourly"], 6.3)
    caption(doc, "Figure 4. Events by hour. Navy = all person events; teal = production only (excludes ::1).")

    add_picture(doc, charts["unique_hour"], 6.3)
    caption(
        doc,
        "Figure 5. Unique people with a production login in that hour, summed across the four days. "
        "Red ≥ 3 people, gold = 2, teal = 1.",
    )
    add_picture(doc, charts["heatmap"], 6.3)
    caption(
        doc,
        "Figure 6. Production heatmap. Sunday 8/23 is evening-only (recording started 19:22). "
        "No Friday or Saturday yet.",
    )

    heading(doc, "5. Recommended deploy windows", 1)
    para(
        doc,
        "Score each clock hour by production event count and unique people. Localhost is ignored "
        "here: a local Next.js restart is not a production outage. If a release also bounces the "
        "admin’s laptop server, treat 09:30–10:30 and 20:00–21:40 as extra-sensitive for that "
        "person only.",
    )

    add_table(
        doc,
        ["Window (CST)", "Events", "Unique people", "Who appeared", "Risk"],
        hour_risk_table(prod),
    )
    caption(doc, "Table. Production logins only, pooled across 23–27 Aug 2026. Risk is relative inside this sample.")

    para(doc, "How to use this on a release day", size=12, bold=True, color=NAVY, space_before=8, space_after=4)
    para(
        doc,
        "Overnight (00:00–08:00). Preferred for schema migrations, Node/PM2 restarts, and Nginx "
        "reloads. Nobody in the sample was authenticated in this band. Aim to be done before 08:50 "
        "so cshen’s 09:06 open hits a healthy site.",
        space_after=6,
    )
    para(
        doc,
        "Weekday 15:00–17:30. Best same-day slot. Office morning and after-lunch waves are over; "
        "evening workers have not started. Verify after 18:30 when cshen is typically back.",
        space_after=6,
    )
    para(
        doc,
        "Weekday 11:00–12:50. Usable if 15:00 is blocked. Keep the change short. sunjie may be on "
        "the site around 11:24; cshen sometimes checks from a phone at lunch.",
        space_after=6,
    )
    para(
        doc,
        "Avoid 09:00–10:45. This is the team’s open. cshen is on every weekday at ~09:07; "
        "sunjie, luoshuang and liuyamin all have morning observations in this band.",
        space_after=6,
    )
    para(
        doc,
        "Avoid 13:00–14:30. Highest unique-user overlap in the sample: five different people "
        "(cshen, sunjie, luoshuang, chenpeifeng, G.Wave) on Tuesday 13:07–13:32 alone.",
        space_after=6,
    )
    para(
        doc,
        "Avoid 18:30–22:00 unless the release is expected to be watched by the admin. cshen works "
        "this band most nights; G.Wave logged in at 19:43 Thursday; chenpeifeng at 21:38 Sunday "
        "and 22:11 Tuesday. A “late quiet deploy” is not actually quiet.",
        space_after=10,
    )

    para(doc, "Suggested release playbook", size=12, bold=True, color=NAVY, space_after=4)
    add_table(
        doc,
        ["Change type", "Preferred window", "Why"],
        [
            ["Hotfix needing same-day verify", "15:00–17:00, then check 18:30+", "Empty production hour; admin returns in evening"],
            ["Routine app / Nginx restart", "01:00–06:00", "Zero observed users; finish before 08:50"],
            ["DB migration with brief lock", "02:00–05:00", "No logins; extra buffer either side"],
            ["UI-only change", "11:15–12:30 or 15:00–17:00", "Low unique-user count; easy rollback"],
            ["Change that forces re-login", "00:00–08:00", "Avoid kicking the 09:07 and 13:00 cohorts"],
        ],
        col_widths=[5.5, 5.0, 7.0],
    )
    para(doc, "", space_after=8)

    heading(doc, "6. Person-by-person habits", 1)
    para(
        doc,
        "Profiles below use both raw events and collapsed visits. Device and network labels are "
        "grouped (office / evening / mobile / localhost) rather than listing every address.",
    )
    add_picture(doc, charts["first"], 6.3)
    caption(doc, "Figure 7. First production login of each calendar day. cshen’s 09:06–09:09 line is the tightest habit.")

    add_picture(doc, charts["device"], 5.2)
    caption(doc, "Figure 8. Device/browser mix across person events. Edge is the admin’s default; Quark is sunjie.")

    summary_rows = []
    for n in sorted(profiles, key=lambda x: -profiles[x]["events"]):
        p = profiles[n]
        firsts = ", ".join(p["first_times"]) if p["first_times"] else "—"
        device = p["devices"].idxmax() if len(p["devices"]) else "—"
        summary_rows.append(
            [
                n,
                p["role"],
                str(p["events"]),
                str(p["visits"]),
                f"{p['ok']}/{p['fail']}",
                str(p["days_active"]),
                firsts,
                device,
            ]
        )
    add_table(
        doc,
        ["Person", "Role", "Events", "Visits", "OK/Fail", "Days", "First login each day", "Main device"],
        summary_rows,
    )
    caption(doc, "Table. Active accounts in the recording window.")

    for n in sorted(profiles, key=lambda x: -profiles[x]["events"]):
        p = profiles[n]
        heading(doc, f"{p['name']}  ·  {p['role']}", 2)
        para(doc, habit_paragraph(p))
        add_table(
            doc,
            ["Metric", "Value"],
            [
                ["Account created", p["created"] or "—"],
                ["Events / visits", f"{p['events']} / {p['visits']}"],
                ["Success / fail", f"{p['ok']} / {p['fail']}"],
                ["Days active in window", str(p["days_active"])],
                ["Earliest / latest event", f"{p['earliest']}  /  {p['latest']}"],
                ["Hours seen (CST)", ", ".join(f"{h:02d}" for h in p["hours"])],
                ["Main device(s)", ", ".join(f"{k} {int(v)}" for k, v in p["devices"].items())],
                ["Network mix", ", ".join(f"{k} ({int(v)})" for k, v in p["networks"].items())],
                ["Localhost share", f"{p['localhost_share']:.0%}"],
            ],
            col_widths=[5.5, 11.5],
        )
        para(doc, "", space_after=8)

    heading(doc, "7. Accounts with no login in this window", 1)
    if inactive:
        rows = []
        for n in inactive:
            u = users.loc[users["name"] == n].iloc[0]
            created = pd.to_datetime(u["created_at"]).strftime("%Y-%m-%d")
            rows.append([n, str(u["role"]), created, "No events since recording began 23 Aug"])
        add_table(doc, ["Person", "Role", "Account created", "Note"], rows)
        para(
            doc,
            "Absence here does not mean the account is unused historically — only that it did not "
            "authenticate after logging started. Revisit in two weeks before treating these as dormant.",
            space_before=8,
        )
    para(
        doc,
        "benc and chy are admins with zero events in the window. If they still operate the site, "
        "they may be using a long-lived session from before recording began, or they simply did not "
        "open the app these four days.",
    )

    heading(doc, "8. Failures and security notes", 1)
    fails = hist.loc[~hist["success"]]
    fail_rows = []
    for _, r in fails.iterrows():
        fail_rows.append(
            [
                r["ts"].strftime("%Y-%m-%d %H:%M:%S"),
                r["who"],
                r["fail_reason"] or "—",
                r["device"],
                r["network"],
            ]
        )
    add_table(doc, ["Time (CST)", "Who", "Reason", "Device", "Network"], fail_rows)
    para(
        doc,
        "Nothing in this sample looks like credential stuffing. The two invalid_credentials rows "
        "are the same person, same network, 11 seconds apart, immediately followed by a success. "
        "The incomplete row is a documented self-test from a reserved TEST-NET address. "
        "Success rate among real people is 49/51 = 96%.",
        space_before=8,
    )

    heading(doc, "9. What to collect next", 1)
    para(
        doc,
        "This report will get much more useful after 10–15 weekdays. Until then, keep the overnight "
        "and 15:00 slots as defaults, and re-run the generator after Friday and the first weekend.",
    )
    para(
        doc,
        "Useful follow-ups: (1) persist session duration, not only login instant; (2) tag "
        "localhost vs production in the app so deploy impact is not mixed with local restarts; "
        "(3) record which page was opened after login if you want feature-level traffic, not just "
        "authentication; (4) re-score hours once Friday 15:00–17:00 and a full weekend exist.",
    )

    heading(doc, "10. Method", 1)
    para(
        doc,
        "Events were read from PostgreSQL through an SSH local-forward to 127.0.0.1:5432 on the "
        "application server, queried as logged_at AT TIME ZONE 'Asia/Shanghai'. "
        f"Visits collapse consecutive events for the same person with a gap of ≤ {BURST_MINUTES} minutes. "
        "Production = not localhost and not the self-test identifier. "
        "Hour risk uses pooled counts across the short window, so a single busy Tuesday afternoon "
        "has a large effect — that is stated rather than smoothed away.",
    )
    para(
        doc,
        "IP addresses are summarised into network groups in this document and are not listed in full.",
        color=MUTED,
        size=10,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))


def main() -> int:
    if sys.platform == "win32" and hasattr(sys.stdout, "buffer"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    configure_matplotlib()
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    users, hist = fetch()
    global ROLE_BY_NAME
    ROLE_BY_NAME = dict(zip(users["name"], users["role"]))
    people = hist.loc[~hist["is_selftest"]]
    sess_frames = []
    for name, g in people.groupby("who"):
        sess_frames.append(collapse_sessions(g))
    sessions = pd.concat(sess_frames, ignore_index=True) if sess_frames else pd.DataFrame()
    charts = make_charts(hist, sessions)
    build_doc(users, hist, charts, sessions)
    print(f"Wrote {REPORT_PATH}")
    print(f"Events={len(hist)} people={people['who'].nunique()} visits={len(sessions)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
