# -*- coding: utf-8 -*-
"""Weekly nginx + login Word report from reports/_week_traffic_raw."""
from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.font_manager import FontProperties, fontManager

ROOT = Path(__file__).resolve().parents[2]
RAW = ROOT / "reports" / "_week_traffic_raw"
OUT_DIR = ROOT / "reports"
CHART_DIR = OUT_DIR / "_week_traffic_charts"
REPORT_PATH = OUT_DIR / "Weekly_Traffic_Analysis_20260831-20260903.docx"

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_NAVY = "#1A365D"
C_TEAL = "#2B6CB0"
C_GOLD = "#C9A227"
C_GREEN = "#2F855A"
C_RED = "#C53030"
C_ORANGE = "#DD6B20"
C_GRAY = "#718096"
C_PURPLE = "#6B46C1"
C_LIGHT = "#E2E8F0"
PALETTE = [C_NAVY, C_TEAL, C_GOLD, C_GREEN, C_ORANGE, C_PURPLE, C_RED, C_GRAY]

_CN_FONT: FontProperties | None = None
TZ_LABEL = "Asia/Shanghai (UTC+8)"
WINDOW_LABEL = "Mon 31 Aug – Thu 3 Sep 2026 (through 16:12 CST)"


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    plt.rcParams["axes.edgecolor"] = "#CBD5E0"
    plt.rcParams["axes.grid"] = False
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.isfile(path):
            try:
                fontManager.addfont(path)
                _CN_FONT = FontProperties(fname=path)
                plt.rcParams["font.family"] = "sans-serif"
                plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "Microsoft YaHei", "SimHei"]
                return
            except Exception:
                continue
    _CN_FONT = FontProperties(family="Microsoft YaHei")


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.2
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=9, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "F7FAFC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            cell_text(
                table.rows[r_i + 1].cells[c_i],
                val,
                size=8,
                color=TEXT,
                align="center" if c_i else "left",
            )
            shade(table.rows[r_i + 1].cells[c_i], fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_picture(doc, path: Path, width_in=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    return p


def caption(doc, text: str):
    para(doc, text, size=8, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def style_ax(ax, title: str, xlabel: str, ylabel: str):
    ax.set_title(title, **fp(), fontsize=12, color=C_NAVY, pad=10)
    ax.set_xlabel(xlabel, **fp(), fontsize=9, color="#4A5568")
    ax.set_ylabel(ylabel, **fp(), fontsize=9, color="#4A5568")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(colors="#4A5568", labelsize=8)
    for lab in ax.get_xticklabels() + ax.get_yticklabels():
        lab.set_fontproperties(_CN_FONT)


def save_fig(fig, name: str) -> Path:
    path = CHART_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def fmt_int(n) -> str:
    return f"{int(n):,}"


def fmt_pct(n, d) -> str:
    if d == 0:
        return "—"
    return f"{100.0 * n / d:.1f}%"


def classify_login_device(ua: str | None) -> str:
    s = (ua or "").lower()
    if "iphone" in s or "android" in s or "mobile" in s or "huawei" in s:
        return "Mobile"
    if "edg/" in s:
        return "Edge"
    if "quark" in s:
        return "Quark"
    if "chrome" in s:
        return "Chrome"
    return "Other"


def classify_login_net(ip: str | None) -> str:
    ip = (ip or "").strip()
    if ip in ("::1", "127.0.0.1"):
        return "Localhost (dev)"
    if ip.startswith("116.237.193."):
        return "Office A"
    if ip.startswith("116.234.199."):
        return "Network B"
    if ip.startswith("116.234.86."):
        return "Network C"
    if ip.startswith("111.187."):
        return "Network D"
    return "Other"


def load() -> dict:
    summary = json.loads((RAW / "summary.json").read_text(encoding="utf-8"))
    daily = pd.read_csv(RAW / "daily.csv")
    hourly = pd.read_csv(RAW / "hourly.csv")
    paths = pd.read_csv(RAW / "top_paths.csv")
    pages = pd.read_csv(RAW / "top_pages.csv")
    status = pd.read_csv(RAW / "status.csv")
    devices = pd.read_csv(RAW / "devices.csv")
    networks = pd.read_csv(RAW / "networks.csv")
    methods = pd.read_csv(RAW / "methods.csv")
    ips = pd.read_csv(RAW / "ips.csv")
    login = pd.read_csv(RAW / "login.csv")
    users = pd.read_csv(RAW / "users.csv")
    login["ts"] = pd.to_datetime(login["ts"])
    login["success"] = login["success"].astype(str).str.lower().isin(["t", "true", "1"])
    login["who"] = login["who"].str.strip()
    login["who_norm"] = login["who"].str.lower()
    login["date"] = login["ts"].dt.date.astype(str)
    login["hour"] = login["ts"].dt.hour
    login["device"] = login["user_agent"].map(classify_login_device)
    login["network"] = login["ip"].map(classify_login_net)
    hourly["product"] = hourly[["page", "api", "admin", "login"]].sum(axis=1)
    daily["gb"] = daily["bytes"] / (1024 ** 3)
    return {
        "summary": summary,
        "daily": daily,
        "hourly": hourly,
        "paths": paths,
        "pages": pages,
        "status": status,
        "devices": devices,
        "networks": networks,
        "methods": methods,
        "ips": ips,
        "login": login,
        "users": users,
    }


def feature_rows(paths: pd.DataFrame) -> list[tuple[str, int, str]]:
    rules = [
        ("心跳 /api/presence", "/api/presence", "polling"),
        ("跟踪基金 API", "/ma/api/tracking-funds", "api"),
        ("CTP 行情 live", "/ma/api/ctp-market/live", "polling"),
        ("CTP 行情其它", "/ma/api/ctp-market", "api"),
        ("实时行情 API", "/ma/api/realtime-quotes", "api"),
        ("投资笔记 API", "/ma/api/investment-notes", "api"),
        ("私募产品 API", "/ma/api/private-funds", "api"),
        ("尽调表 API", "/ma/api/due-diligence-table", "api"),
        ("最近页面 API", "/ma/api/recent-pages", "polling"),
        ("运营/产品要素 API", "/ma/api/ops/", "api"),
        ("MOM API", "/ma/api/mom-analysis", "api"),
        ("FOF overview API", "/ma/api/:id/fof-overview", "api"),
        ("知识库 API", "/api/knowledge-base", "api"),
        ("鉴权 /api/auth/me", "/api/auth/me", "polling"),
        ("管理 API", "/api/admin/", "admin"),
        ("私募产品页", "/ma/dashboard/private-funds", "page"),
        ("MOM 分析页", "/ma/dashboard/mom-analysis", "page"),
        ("全天候", "/ma/dashboard/all-weather", "prefetch"),
        ("期货市场", "/ma/dashboard/futures-market", "prefetch"),
        ("期权市场", "/ma/dashboard/options-market", "prefetch"),
        ("宏观市场", "/ma/dashboard/macro-market", "prefetch"),
        ("AI 知识", "/ma/dashboard/ai-knowledge", "prefetch"),
        ("AI 投研", "/ma/dashboard/ai-researcher", "prefetch"),
        ("实时行情页", "/ma/dashboard/realtime-quotes", "prefetch"),
        ("股票市场", "/ma/dashboard/stock-market", "prefetch"),
        ("投研看板首页", "/ma/dashboard", "page"),
        ("登录页", "/login", "page"),
        ("Vercel insights 404", "/_vercel/insights", "noise"),
        ("图标 / 静态", "/icon", "static"),
        ("Next 静态资源", "/_next/", "static"),
    ]
    used = set()
    out = []
    for name, prefix, kind in rules:
        mask = paths["path"].astype(str).str.startswith(prefix)
        if prefix == "/ma/dashboard":
            mask = paths["path"].eq("/ma/dashboard") | paths["path"].eq("/ma/dashboard/:id")
        hits = int(paths.loc[mask & ~paths["path"].isin(used), "hits"].sum())
        for p in paths.loc[mask, "path"]:
            used.add(p)
        if hits:
            out.append((name, hits, kind))
    rest = int(paths.loc[~paths["path"].isin(used), "hits"].sum())
    if rest:
        out.append(("其它（未归入上表的路径）", rest, "other"))
    return out


def make_charts(d: dict) -> dict[str, Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    charts = {}
    daily = d["daily"]
    hourly = d["hourly"]
    login = d["login"]
    status = d["status"]
    networks = d["networks"]
    feats = feature_rows(d["paths"])

    kinds = ["page", "api", "heartbeat", "static", "probe", "admin", "login", "other"]
    colors = {
        "page": C_NAVY, "api": C_TEAL, "heartbeat": C_GOLD, "static": C_GRAY,
        "probe": C_RED, "admin": C_PURPLE, "login": C_GREEN, "other": C_ORANGE,
    }
    fig, ax = plt.subplots(figsize=(8.4, 3.8))
    x = np.arange(len(daily))
    bottom = np.zeros(len(daily))
    labels = ["Mon 8/31", "Tue 9/1", "Wed 9/2", "Thu 9/3*"]
    for k in kinds:
        vals = daily[k].to_numpy(dtype=float)
        ax.bar(x, vals, bottom=bottom, color=colors[k], width=0.62, label=k)
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels(labels, **fp())
    ax.legend(prop=_CN_FONT, fontsize=7, frameon=False, ncol=4, loc="upper right")
    style_ax(ax, "Daily nginx hits by request kind", "Date (CST)", "Requests")
    charts["daily_kind"] = save_fig(fig, "daily_kind.png")

    fig, ax = plt.subplots(figsize=(8.4, 3.6))
    product = daily["product"].to_numpy()
    rest = daily["all"].to_numpy() - product
    ax.bar(x - 0.18, product, width=0.36, color=C_NAVY, label="Product (page+API+admin+login)")
    ax.bar(x + 0.18, rest, width=0.36, color=C_LIGHT, label="Other (static / heartbeat / probe)")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, **fp())
    ax.legend(prop=_CN_FONT, fontsize=8, frameon=False)
    style_ax(ax, "Product traffic vs everything else", "Date (CST)", "Requests")
    charts["daily_product"] = save_fig(fig, "daily_product.png")

    hours = np.arange(24)
    prod_h = hourly.groupby("hour")["product"].sum().reindex(hours, fill_value=0)
    fig, ax = plt.subplots(figsize=(8.4, 3.5))
    bar_colors = [C_RED if v >= 8000 else (C_GOLD if v >= 4000 else C_TEAL) for v in prod_h]
    ax.bar(hours, prod_h.to_numpy(), color=bar_colors, width=0.72)
    ax.set_xticks(hours)
    style_ax(ax, "Product requests by hour of day (sum of four days)", "Hour (CST)", "Product requests")
    charts["hourly_product"] = save_fig(fig, "hourly_product.png")

    day_order = ["2026-08-31", "2026-09-01", "2026-09-02", "2026-09-03"]
    mat = (
        hourly.pivot_table(index="date", columns="hour", values="product", aggfunc="sum")
        .reindex(index=day_order, columns=hours, fill_value=0)
        .fillna(0)
    )
    fig, ax = plt.subplots(figsize=(8.6, 3.2))
    im = ax.imshow(mat.to_numpy(), aspect="auto", cmap="Blues")
    ax.set_xticks(hours)
    ax.set_yticks(range(len(day_order)))
    ax.set_yticklabels(["Mon 8/31", "Tue 9/1", "Wed 9/2", "Thu 9/3*"], **fp())
    ax.set_xlabel("Hour (CST)", **fp(), fontsize=9, color="#4A5568")
    ax.set_title("Product-request heatmap (page + API + admin + login)", **fp(), fontsize=12, color=C_NAVY, pad=10)
    ax.tick_params(colors="#4A5568", labelsize=8)
    cbar = fig.colorbar(im, ax=ax, fraction=0.03, pad=0.02)
    cbar.ax.tick_params(labelsize=7)
    cbar.set_label("Requests", **fp(), fontsize=8)
    fig.tight_layout()
    charts["heatmap"] = save_fig(fig, "heatmap.png")

    top = sorted(feats, key=lambda r: r[1], reverse=True)[:14]
    fig, ax = plt.subplots(figsize=(8.4, 4.6))
    names = [r[0] for r in top][::-1]
    vals = [r[1] for r in top][::-1]
    kinds_l = [r[2] for r in top][::-1]
    kind_color = {
        "polling": C_GOLD, "api": C_TEAL, "page": C_NAVY, "prefetch": C_PURPLE,
        "admin": C_ORANGE, "noise": C_RED, "static": C_GRAY, "other": C_ORANGE,
    }
    y = np.arange(len(names))
    ax.barh(y, vals, color=[kind_color.get(k, C_GRAY) for k in kinds_l], height=0.62)
    ax.set_yticks(y)
    ax.set_yticklabels(names, **fp())
    style_ax(ax, "Top path groups (colour = traffic type)", "", "Hits")
    charts["features"] = save_fig(fig, "features.png")

    st = status.sort_values("hits", ascending=False)
    fig, ax = plt.subplots(figsize=(8.4, 3.4))
    sc = [C_GREEN if s < 400 else (C_GOLD if s < 500 else C_RED) for s in st["status"]]
    ax.bar([str(s) for s in st["status"]], st["hits"], color=sc, width=0.7)
    style_ax(ax, "HTTP status counts", "Status", "Hits")
    charts["status"] = save_fig(fig, "status.png")

    fig, ax = plt.subplots(figsize=(8.4, 3.6))
    net = networks.sort_values("hits", ascending=True)
    y = np.arange(len(net))
    ax.barh(y, net["hits"], color=C_NAVY, height=0.55)
    ax.set_yticks(y)
    ax.set_yticklabels(list(net["network"]), **fp())
    style_ax(ax, "Hits by client network family", "", "Hits")
    charts["networks"] = save_fig(fig, "networks.png")

    daily_login = (
        login.assign(who=login["who"].str.replace("Musheng", "musheng", regex=False))
        .groupby(["date", "who"])
        .size()
        .unstack(fill_value=0)
        .sort_index()
    )
    fig, ax = plt.subplots(figsize=(8.4, 3.6))
    x2 = np.arange(len(daily_login))
    bottom = np.zeros(len(daily_login))
    for i, name in enumerate(daily_login.columns):
        vals = daily_login[name].to_numpy()
        ax.bar(x2, vals, bottom=bottom, color=PALETTE[i % len(PALETTE)], width=0.62, label=name)
        bottom += vals
    ax.set_xticks(x2)
    ax.set_xticklabels([pd.to_datetime(d).strftime("%a %m-%d") for d in daily_login.index], **fp())
    ax.legend(prop=_CN_FONT, fontsize=8, frameon=False, ncol=3, loc="upper left")
    style_ax(ax, "Login events by person", "Date (CST)", "Login events")
    charts["logins"] = save_fig(fig, "logins.png")
    return charts


def build_doc(d: dict, charts: dict[str, Path]) -> None:
    s = d["summary"]
    daily = d["daily"]
    hourly = d["hourly"]
    login = d["login"]
    users = d["users"]
    ips = d["ips"]
    status = d["status"]
    devices = d["devices"]
    methods = d["methods"]
    pages = d["pages"]
    feats = feature_rows(d["paths"])

    n = int(s["hits_in_week"])
    n_prod = int(s["product_hits"])
    n_page = int(s["page_hits"])
    n_ip = int(s["unique_ips"])
    gb = s["bytes_in_week"] / (1024 ** 3)
    kind = s["kind"]
    ok_2xx = int(status.loc[status["status"].between(200, 299), "hits"].sum())
    n_404 = int(status.loc[status["status"] == 404, "hits"].sum())
    n_499 = int(status.loc[status["status"] == 499, "hits"].sum())
    n_502 = int(status.loc[status["status"] == 502, "hits"].sum())
    n_500 = int(status.loc[status["status"] == 500, "hits"].sum())
    n_403 = int(status.loc[status["status"] == 403, "hits"].sum())
    n_probe = int(kind["probe"])
    n_hb = int(kind["heartbeat"])

    real_ips = ips[(ips["page"] > 0) | (ips["api"] > 0)]
    scanner_ips = ips[(ips["page"] == 0) & (ips["api"] == 0)]
    office_share = float(ips.loc[ips["ip"] == "116.237.193.158", "hits"].sum()) / n

    prefetch_pages = [
        "/ma/dashboard/all-weather", "/ma/dashboard/futures-market", "/ma/dashboard/options-market",
        "/ma/dashboard/macro-market", "/ma/dashboard/ai-knowledge", "/ma/dashboard/ai-researcher",
        "/ma/dashboard/realtime-quotes", "/ma/dashboard/stock-market",
    ]
    prefetch_hits = int(pages.loc[pages["path"].isin(prefetch_pages), "hits"].sum())
    private_page = int(pages.loc[pages["path"].str.startswith("/ma/dashboard/private-funds"), "hits"].sum())
    mom_page = int(pages.loc[pages["path"].str.startswith("/ma/dashboard/mom-analysis"), "hits"].sum())

    peak = hourly.loc[hourly["product"].idxmax()]
    peak_label = f"{peak['date']} {int(peak['hour']):02d}:00–{int(peak['hour']):02d}:59"

    login_ok = int(login["success"].sum())
    login_fail = int((~login["success"]).sum())
    people = sorted(login["who_norm"].unique())
    named = set(users["name"].str.lower())
    active = set(people)
    inactive = [n_ for n_ in users["name"] if n_.lower() not in active]
    new_users = users.loc[users["created"] >= "2026-08-31", "name"].tolist()

    by_who = (
        login.assign(who=np.where(login["who_norm"] == "musheng", "musheng", login["who"]))
        .groupby("who")
        .agg(events=("ts", "size"), ok=("success", "sum"), first=("ts", "min"), last=("ts", "max"))
        .sort_values("events", ascending=False)
    )

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    para(doc, "Market Dashboard  ·  Internal operations", size=10, color=GOLD, space_after=2)
    p = para(doc, "Weekly Website Traffic Analysis", size=22, bold=True, color=NAVY, space_after=4)
    p.runs[0].font.size = Pt(22)
    para(
        doc,
        f"Source: nginx access.log (+ rotated files) and public.auth_login_history  ·  "
        f"Timezone: {TZ_LABEL}  ·  Window: {WINDOW_LABEL}  ·  "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        size=9,
        color=MUTED,
        space_after=14,
    )

    heading(doc, "1. Executive summary", 1)
    para(
        doc,
        f"This is the first full weekday look at recorded HTTP traffic, not just logins. "
        f"From Monday 31 August 00:05 through Thursday 3 September 16:12 CST the site served "
        f"{fmt_int(n)} nginx requests ({gb:.2f} GB) from {fmt_int(n_ip)} unique client IPs. "
        f"{fmt_int(n_prod)} of those ({fmt_pct(n_prod, n)}) are product-facing "
        f"(pages, APIs, admin, login). Thursday is a partial day — do not read the drop versus "
        f"Monday as a demand collapse.",
    )
    para(
        doc,
        f"Almost {office_share:.0%} of all hits come from a single office NAT "
        f"(116.237.193.158). Unique IP is therefore not unique person. Login history is the "
        f"better headcount: {len(login)} events, {login_ok} success / {login_fail} failure, "
        f"{len(people)} accounts seen. Two new accounts were created this week (musheng, zhougang) "
        f"and both logged in on Tuesday evening.",
        space_after=8,
    )
    para(doc, "What the week actually looks like:", size=11, bold=True, color=NAVY, space_after=4)
    para(
        doc,
        "1) Office hours dominate. Product traffic is dense 09:00–17:30, with a lunch dip around "
        "12:00 and a second peak after 13:00. Overnight (00:00–07:00) is empty of real users and "
        "is mostly internet scanners.",
        space_after=3,
    )
    para(
        doc,
        f"2) Raw page counts overstate browsing. Eight market/AI sidebar routes each have "
        f"~3,040–3,100 hits — {fmt_int(prefetch_hits)} combined — a Next.js prefetch / nav-warmup "
        f"pattern, not eight equally loved pages. Intentional destinations this week are "
        f"私募产品 ({fmt_int(private_page)} page hits), MOM分析 ({fmt_int(mom_page)}), and the "
        f"dashboard home.",
        space_after=3,
    )
    para(
        doc,
        "3) The heaviest APIs are live polls, not clicks: /ma/api/ctp-market/live (~9.2k), "
        "tracking-funds (~9.7k), /api/presence (~11.5k). A quiet user with a quotes tab open "
        "still generates thousands of rows.",
        space_after=3,
    )
    para(
        doc,
        f"4) Reliability is acceptable but not invisible: {fmt_pct(ok_2xx, n)} of requests are 2xx. "
        f"There were {fmt_int(n_404)} 404s (Vercel insights script + scanners), {fmt_int(n_502)} 502s "
        f"(brief upstream unavailability, including a cluster around Thu 00:22), {fmt_int(n_500)} 500s, "
        f"and {fmt_int(n_499)} 499s (client closed the connection — usually a deploy or a slow page).",
        space_after=3,
    )
    para(
        doc,
        f"5) Internet noise is real: {fmt_int(n_probe)} classified probe hits plus several scanner "
        f"IPs (163.7.x, 34.24.x, 35.185.x) hitting /.env, /.git, wp-admin at 03:00–05:00. They did "
        f"not reach product APIs in material volume.",
        space_after=12,
    )

    heading(doc, "2. Data coverage and caveats", 1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["HTTP source", "/var/log/nginx/access.log + access.log.1 + .2.gz + .3.gz"],
            ["Login source", "public.auth_login_history"],
            ["Timezone", TZ_LABEL],
            ["Window start", "2026-08-31 00:00 CST (first hit 00:05:07)"],
            ["Window end", "2026-09-03 16:12:41 CST (partial Thursday)"],
            ["Lines scanned / in window", f"{fmt_int(s['lines_scanned'])} / {fmt_int(n)}"],
            ["Bytes served", f"{gb:.2f} GB"],
            ["Unique client IPs", fmt_int(n_ip)],
            ["IPs with page or API traffic", fmt_int(len(real_ips))],
            ["Product requests", f"{fmt_int(n_prod)} ({fmt_pct(n_prod, n)})"],
            ["Page / API / heartbeat / static / probe",
             f"{fmt_int(kind['page'])} / {fmt_int(kind['api'])} / {fmt_int(n_hb)} / "
             f"{fmt_int(kind['static'])} / {fmt_int(n_probe)}"],
            ["Login events (ok / fail)", f"{len(login)} ({login_ok} / {login_fail})"],
            ["Registered accounts / seen this week", f"{len(users)} / {len(people)}"],
            ["New accounts this week", ", ".join(new_users) if new_users else "—"],
        ],
        col_widths=[6.6, 10.4],
    )
    para(doc, "", space_after=6)
    para(
        doc,
        "A hit is one HTTP request, not a person and not a session. Presence heartbeats, "
        "CTP live polls, recent-pages writes, and Next.js route prefetch inflate volume. "
        "Office colleagues share 116.237.193.158, so IP uniqueness undercounts people at the "
        "desk and overcounts scanners on the public internet. Static files and /_next assets "
        "are kept in the totals so operators can see cache/304 behaviour, but they are not "
        "product usage. Friday–Sunday are not in this window.",
        size=10,
        color=MUTED,
        space_after=12,
    )

    heading(doc, "3. Volume by day", 1)
    para(
        doc,
        "Monday is the busiest full day (40.3k hits, 32.3k product). Tuesday is similar in "
        "product volume. Wednesday is lighter (24.8k / 16.9k product) even though evening "
        "onboarding of musheng and zhougang added login noise. Thursday through 16:12 is "
        "already 8.5k product requests — on pace for a normal weekday afternoon, not a holiday.",
    )
    add_table(
        doc,
        ["Date", "All hits", "Product", "Pages", "APIs", "Heartbeat", "Probe", "IPs", "GB"],
        [
            [
                row["date"],
                fmt_int(row["all"]),
                fmt_int(row["product"]),
                fmt_int(row["page"]),
                fmt_int(row["api"]),
                fmt_int(row["heartbeat"]),
                fmt_int(row["probe"]),
                fmt_int(row["unique_ips"]),
                f"{row['gb']:.2f}",
            ]
            for _, row in daily.iterrows()
        ],
        col_widths=[2.4, 1.7, 1.7, 1.5, 1.5, 1.7, 1.5, 1.3, 1.3],
    )
    para(doc, "", space_after=6)
    add_picture(doc, charts["daily_kind"], 6.4)
    caption(doc, "Figure 1. Daily nginx hits stacked by kind. Thursday is a partial day (to 16:12).")
    add_picture(doc, charts["daily_product"], 6.4)
    caption(doc, "Figure 2. Product-facing traffic versus static / heartbeat / scanner remainder.")

    heading(doc, "4. When the site is busy", 1)
    para(
        doc,
        f"The single busiest product hour was {peak_label} with {fmt_int(peak['product'])} "
        f"product requests ({fmt_int(peak['page'])} pages, {fmt_int(peak['api'])} APIs). "
        "Across the four days, 09:00, 13:00 and 15:00–16:00 are the densest bands. "
        "12:00 is a lunch trough every weekday. After 18:00 traffic is admin-led "
        "(cshen on Network B) plus a little mobile; it is not zero.",
    )
    para(
        doc,
        "Scanner bursts sit outside office hours and are easy to spot: Tue 03:00 (~895 hits, "
        "almost all probe/static/other, 0 pages), Wed 05:00 (~893, same shape), Thu 00:00 "
        "(~670, including a 502 cluster). Those hours should not be read as users.",
    )
    add_picture(doc, charts["hourly_product"], 6.4)
    caption(doc, "Figure 3. Product requests by clock hour, summed across Mon–Thu. Red ≥ 8,000; gold ≥ 4,000.")
    add_picture(doc, charts["heatmap"], 6.4)
    caption(doc, "Figure 4. Product-request heatmap. Empty overnight cells are real; scanner hours look busy in raw logs but not here.")

    heading(doc, "5. What people used", 1)
    para(
        doc,
        "Split the path list into four layers or the ranking will lie. Layer 1 is polling "
        "(presence, CTP live, auth/me, recent-pages). Layer 2 is feature APIs fired by a page "
        "(tracking-funds, private-funds, investment-notes, MOM). Layer 3 is HTML pages. "
        "Layer 4 is prefetch of sibling nav routes — eight market/AI pages with nearly identical "
        "counts, which is the signature of <Link prefetch>, not eight independent audiences.",
    )
    add_picture(doc, charts["features"], 6.4)
    caption(
        doc,
        "Figure 5. Top path groups. Navy = page, teal = API, gold = polling, purple = likely prefetch, red = noise.",
    )
    add_table(
        doc,
        ["Path group", "Hits", "Type"],
        [[name, fmt_int(hits), kind] for name, hits, kind in feats[:18]],
        col_widths=[7.0, 3.5, 6.5],
    )
    para(doc, "", space_after=8)
    para(doc, "Intentional page destinations (after collapsing :id):", size=11, bold=True, color=NAVY, space_after=4)
    page_rows = []
    for _, r in pages.head(16).iterrows():
        tag = "prefetch cluster" if r["path"] in prefetch_pages else "page"
        page_rows.append([r["path"], fmt_int(r["hits"]), tag])
    add_table(doc, ["Page", "Hits", "Note"], page_rows, col_widths=[9.5, 2.5, 5.0])
    para(doc, "", space_after=6)
    para(
        doc,
        "Reading that table as a product manager: 私募产品 is the working surface (list + "
        "product-code pages + APIs for notes, tags, due-diligence, share classes). MOM分析 "
        "is the second working surface (risk-report, trader-analysis, data-import, account-risk, "
        "carry-calc). Tracking-funds and CTP live are always-on widgets. Knowledge-base, "
        "reports, and mail-parse APIs are still a thin tail this week (knowledge-base 420, "
        "reports 75, email-parse 4).",
        space_after=12,
    )

    heading(doc, "6. Who was on the site", 1)
    para(
        doc,
        f"Login history is the person census. {len(login)} authentication events, "
        f"{login_ok} successful. One failure: Musheng (capital M) at 20:29 Tue from an iPhone "
        f"on 116.6.238.141 — invalid_credentials — then musheng (lowercase, the new account) "
        f"succeeded 13 seconds later. That is a first-login typo, not an attack.",
    )
    add_picture(doc, charts["logins"], 6.4)
    caption(doc, "Figure 6. Login events by person. cshen includes localhost bursts from local development.")
    who_rows = []
    for name, r in by_who.iterrows():
        who_rows.append([
            name,
            int(r.events),
            int(r.ok),
            r["first"].strftime("%m-%d %H:%M"),
            r["last"].strftime("%m-%d %H:%M"),
        ])
    add_table(
        doc,
        ["Person", "Events", "Success", "First (CST)", "Last (CST)"],
        who_rows,
        col_widths=[3.4, 2.4, 2.4, 4.4, 4.4],
    )
    para(doc, "", space_after=8)
    para(
        doc,
        "cshen (admin): weekday open is still ~09:07 from Office A, with localhost rows when "
        "the Next dev server is used at the desk, and Network B from ~19:40 on Mon/Tue/Wed. "
        "sunjie: Quark, Office A, mid-morning to mid-afternoon, no evening. luoshuang: Chrome "
        "from Network C, Mon–Wed mornings, last seen Wed 09:55. G.Wave: one Edge login Thu 13:29 "
        "from Network D (111.187.8.193). chenpeifeng: one Huawei mobile login Mon 21:33 on "
        "Network B. zhougang and musheng: created Tue 2 Sep; first seen that evening (20:08 "
        "Huawei on Network B; 20:29 iPhone).",
        space_after=8,
    )
    para(
        doc,
        "Accounts with no login this week: "
        + (", ".join(inactive) if inactive else "—")
        + ". Treat as dormant for this four-day window only — the previous login report already "
        "noted that some colleagues appear every other week.",
        space_after=8,
    )
    para(doc, "Network mix (nginx, all hits):", size=11, bold=True, color=NAVY, space_after=4)
    add_picture(doc, charts["networks"], 6.4)
    caption(doc, "Figure 7. Client network families. Office A is a shared NAT — several people behind one IP.")
    top_ip_rows = []
    for _, r in ips.head(12).iterrows():
        role = "office NAT / mixed users" if r.ip.endswith(".158") else (
            "evening / admin" if r.ip.endswith(".216") else (
                "scanner" if r.page == 0 and r.api == 0 else "user / mobile"
            )
        )
        top_ip_rows.append([
            r.ip, r.network, fmt_int(r.hits), fmt_int(r.page), fmt_int(r.api), role,
        ])
    add_table(
        doc,
        ["IP", "Network", "Hits", "Pages", "APIs", "Read as"],
        top_ip_rows,
        col_widths=[3.6, 4.2, 1.6, 1.6, 1.6, 4.4],
    )
    para(doc, "", space_after=6)
    para(
        doc,
        f"{fmt_int(len(real_ips))} IPs generated at least one page or API hit; "
        f"{fmt_int(len(scanner_ips))} of the top-listed IPs are scanner-only (no product path). "
        "Device mix of all hits: Edge 73.6k, Mobile 25.0k, Chrome 8.6k, Quark 2.3k, "
        "bot/script 2.7k. Mobile is real (chenpeifeng, musheng, some 39.144.* carrier IPs), "
        "not only user-agent spoofing.",
        space_after=12,
    )

    heading(doc, "7. Reliability and abuse", 1)
    para(
        doc,
        f"2xx share is {fmt_pct(ok_2xx, n)} ({fmt_int(ok_2xx)} of {fmt_int(n)}). "
        f"304 Not Modified is {fmt_int(int(status.loc[status['status']==304,'hits'].sum()))} — "
        "browsers are caching icons and some HTML. That is healthy. The errors that need a "
        "sentence each:",
    )
    add_picture(doc, charts["status"], 6.4)
    caption(doc, "Figure 8. HTTP status histogram. Green = 2xx/3xx, gold = 4xx, red = 5xx.")
    add_table(
        doc,
        ["Code", "Hits", "Share", "Reading"],
        [
            ["200", fmt_int(int(status.loc[status["status"] == 200, "hits"].sum())),
             fmt_pct(int(status.loc[status["status"] == 200, "hits"].sum()), n), "Normal success"],
            ["304", fmt_int(int(status.loc[status["status"] == 304, "hits"].sum())),
             fmt_pct(int(status.loc[status["status"] == 304, "hits"].sum()), n), "Cache hit"],
            ["404", fmt_int(n_404), fmt_pct(n_404, n),
             "Missing /_vercel/insights/script.js (~692) + scanner paths"],
            ["403", fmt_int(n_403), fmt_pct(n_403, n), "Auth / forbidden"],
            ["499", fmt_int(n_499), fmt_pct(n_499, n), "Client closed — deploy, reload, or slow handler"],
            ["500", fmt_int(n_500), fmt_pct(n_500, n), "Application errors — small, still worth a log grep"],
            ["502", fmt_int(n_502), fmt_pct(n_502, n), "nginx → Next upstream down (seen around Thu 00:22)"],
        ],
        col_widths=[2.0, 2.2, 2.2, 10.6],
    )
    para(doc, "", space_after=8)
    get_n = int(methods.loc[methods["method"] == "GET", "hits"].sum())
    post_n = int(methods.loc[methods["method"] == "POST", "hits"].sum())
    put_n = int(methods.loc[methods["method"] == "PUT", "hits"].sum())
    para(
        doc,
        f"Methods: GET {fmt_int(get_n)}, POST {fmt_int(post_n)}, PUT {fmt_int(put_n)}, "
        f"PATCH {fmt_int(int(methods.loc[methods['method']=='PATCH','hits'].sum()))}, "
        f"HEAD {fmt_int(int(methods.loc[methods['method']=='HEAD','hits'].sum()))}. "
        "Odd verbs (PRI, CONNECT) are HTTP/2 probes from scanners, not the app. "
        "PUT volume is largely /ma/api/recent-pages.",
        space_after=8,
    )
    para(
        doc,
        "Two hygiene items fall out of the 404 list. First, the front end still requests "
        "/_vercel/insights/script.js on this self-hosted box; every page load pays a 404. "
        "Remove the snippet. Second, the host is on the public internet, so /.env and "
        "WordPress probes will continue. They currently 404/502. Keep that as the outcome.",
        space_after=12,
    )

    heading(doc, "8. Deploy window (this week’s evidence)", 1)
    para(
        doc,
        "Use product-request density plus login clock, not unique IP. Ranked CST windows:",
        space_after=6,
    )
    para(doc, "1) Best: 00:30–07:00. No product users. Accept scanner 404s; do not confuse them with load.", space_after=3)
    para(doc, "2) Same-day low risk: 12:00–12:40 (lunch trough every weekday in this sample).", space_after=3)
    para(
        doc,
        "3) Acceptable if announced: 18:30–20:00 on days when cshen is not already on Network B. "
        "This week he often was, and Tue 20:00 also caught zhougang / musheng first logins.",
        space_after=3,
    )
    para(
        doc,
        "4) Do not deploy unannounced: 09:00–11:30 and 13:00–17:30. Those bands hold the "
        "heatmap core, live CTP polls, and overlapping colleagues (cshen + sunjie + luoshuang "
        "on Mon–Wed mornings; G.Wave Thursday afternoon).",
        space_after=12,
    )

    heading(doc, "9. Actions", 1)
    para(doc, "Product / ops", size=11, bold=True, color=NAVY, space_after=4)
    para(doc, "• 私募产品 and MOM分析 are the pages that matter this week; keep list APIs and detail pages fast during 09:00–17:30.", space_after=3)
    para(doc, "• CTP live and tracking-funds polling are the background load. A hung quotes tab will look like ‘heavy traffic’ on the deploy checker.", space_after=3)
    para(doc, "• Two colleagues onboarded Tuesday night on mobile. Confirm they can find 私募产品 without an admin sitting next to them.", space_after=8)
    para(doc, "Engineering hygiene", size=11, bold=True, color=NAVY, space_after=4)
    para(doc, "• Drop the Vercel insights script (recurring 404).", space_after=3)
    para(doc, "• Investigate the Thursday 00:22 502 burst — process restart or nginx upstream timeout.", space_after=3)
    para(doc, "• Optional: disable Next.js prefetch on the eight sibling market/AI nav links if those ~3k duplicate page hits are costing SSR.", space_after=3)
    para(doc, "• Keep using logged-in user presence, not unique IP, for deploy-readiness — this week’s office NAT would have hidden sunjie behind cshen.", space_after=8)
    para(doc, "Measurement", size=11, bold=True, color=NAVY, space_after=4)
    para(
        doc,
        "This report is four weekdays, cut at Thursday 16:12. Re-run Friday evening for a "
        "complete week, and keep the nginx rotation (14 days of access.log.*.gz) so next week’s "
        "file is still on disk. Login history remains the person-level source; nginx remains "
        "the feature-level source.",
        space_after=14,
    )
    para(
        doc,
        "Classification: internal. Traffic logs include client IPs and user-agents; do not forward this file outside the team.",
        size=8,
        color=MUTED,
    )
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))


def main() -> None:
    configure_matplotlib()
    d = load()
    charts = make_charts(d)
    build_doc(d, charts)
    print(f"Wrote {REPORT_PATH}")


if __name__ == "__main__":
    main()
