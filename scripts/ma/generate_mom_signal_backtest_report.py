# -*- coding: utf-8 -*-
"""Backtest MOM 量化 vs 主观 decision signals as a tradable overlay; write a Word report."""
from __future__ import annotations

import math
import os
import re
import traceback
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path

import matplotlib
import numpy as np
import pandas as pd

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties, fontManager
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "scripts" / "ma" / "_mom_signal_backtest_output"
CHART_DIR = OUT_DIR / "charts"
REPORT_PATH = ROOT / "MOM决策信号策略回测报告.docx"
REPORT_PATH_ASCII = ROOT / "MOM_signal_strategy_backtest.docx"

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _mom_20m_account import (  # noqa: E402
    START_EQUITY,
    COMM_RATE,
    SLIP_RATE,
    MAX_MARGIN_UTIL,
    account_stats,
    draw_account_charts,
    fmt_yuan,
    run_account,
)

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_NAVY = "#1A365D"
C_GOLD = "#C9A227"
C_RED = "#C53030"
C_GREEN = "#2F855A"
C_BLUE = "#2B6CB0"
C_ORANGE = "#DD6B20"
C_VIOLET = "#6D28D9"
C_AMBER = "#B45309"
C_TEAL = "#0F766E"
C_GRAY = "#718096"
C_SKY = "#0369A1"

QUANT_IDS = ("319", "324", "334", "339", "346", "350", "356")
CONSENSUS_MIN = 3.0
HEAVY_MIN = 8.0
LIGHT_MAX = 1.5
CROWD_SUM = 25.0
FLOW_MIN_ABS = 1_000_000.0
FLOW_REL = 0.05
FLOW_SCALE_FLOOR = 2_000_000.0
VOL_DAYS = 20
LOOKBACK_SPIKE = 40
TOP_N_PRODUCTS = 40
COST_BPS = 1.0  # one-way, per unit of |Δweight|
VOL_TARGET = 0.10

SECTOR_MAP = {
    "C": "农产", "CS": "农产", "WH": "农产", "PM": "农产", "RR": "农产", "RI": "农产", "JR": "农产", "LR": "农产",
    "A": "农产", "B": "农产", "M": "农产", "Y": "农产", "RM": "农产", "OI": "农产", "RS": "农产", "PK": "农产", "P": "农产",
    "SR": "农产", "CF": "农产", "CY": "农产", "LG": "农产", "SP": "农产", "OP": "农产",
    "AP": "生鲜", "CJ": "生鲜", "LH": "生鲜", "JD": "生鲜",
    "AU": "贵金属", "AG": "贵金属", "PT": "贵金属", "PD": "贵金属",
    "CU": "有色", "BC": "有色", "AL": "有色", "AO": "有色", "AD": "有色", "ZN": "有色", "PB": "有色", "NI": "有色", "SN": "有色",
    "LC": "新能源", "PS": "新能源", "SI": "新能源",
    "I": "黑色", "SF": "黑色", "SM": "黑色", "RB": "黑色", "HC": "黑色", "SS": "黑色", "WR": "黑色",
    "JM": "黑色", "J": "黑色", "ZC": "黑色", "FG": "黑色", "BB": "黑色", "FB": "黑色",
    "SC": "能源化工", "FU": "能源化工", "LU": "能源化工", "PG": "能源化工", "BU": "能源化工",
    "TA": "能源化工", "EG": "能源化工", "PF": "能源化工", "PR": "能源化工",
    "PL": "能源化工", "PP": "能源化工", "L": "能源化工",
    "BZ": "能源化工", "PX": "能源化工", "EB": "能源化工",
    "RU": "能源化工", "BR": "能源化工", "NR": "能源化工",
    "SA": "能源化工", "SH": "能源化工", "V": "能源化工",
    "UR": "能源化工", "MA": "能源化工",
    "EC": "航运",
    "IH": "股指", "IF": "股指", "IC": "股指", "IM": "股指", "MO": "股指",
    "TS": "国债", "TF": "国债", "T": "国债", "TL": "国债",
}
PROD_NAMES = {
    "C": "玉米", "CS": "淀粉", "WH": "强麦", "PM": "普麦", "RR": "粳米", "RI": "早籼稻", "JR": "粳稻", "LR": "晚籼稻",
    "A": "黄大豆1号", "B": "黄大豆2号", "M": "豆粕", "Y": "豆油", "RM": "菜籽粕", "OI": "菜籽油", "RS": "油菜籽", "PK": "花生", "P": "棕榈油",
    "SR": "白糖", "CF": "棉花", "CY": "棉纱", "LG": "原木", "SP": "纸浆", "OP": "双胶纸",
    "AP": "苹果", "CJ": "红枣", "LH": "生猪", "JD": "鸡蛋",
    "AU": "黄金", "AG": "白银", "PT": "铂", "PD": "钯",
    "CU": "沪铜", "BC": "国际铜", "AL": "沪铝", "AO": "氧化铝", "AD": "铝合金", "ZN": "沪锌", "PB": "沪铅", "NI": "沪镍", "SN": "沪锡",
    "LC": "碳酸锂", "PS": "多晶硅", "SI": "工业硅",
    "I": "铁矿石", "SF": "硅铁", "SM": "锰硅", "RB": "螺纹钢", "HC": "热卷", "SS": "不锈钢", "WR": "线材",
    "JM": "焦煤", "J": "焦炭", "ZC": "动力煤", "FG": "玻璃", "BB": "胶合板", "FB": "纤维板",
    "SC": "原油", "FU": "燃料油", "LU": "低硫燃料油", "PG": "液化石油气", "BU": "沥青",
    "TA": "PTA", "EG": "乙二醇", "PF": "短纤", "PR": "瓶片", "PL": "丙烯", "PP": "聚丙烯", "L": "塑料",
    "BZ": "纯苯", "PX": "对二甲苯", "EB": "苯乙烯",
    "RU": "天然橡胶", "BR": "丁二烯橡胶", "NR": "20号胶",
    "SA": "纯碱", "SH": "烧碱", "V": "PVC", "UR": "尿素", "MA": "甲醇",
    "EC": "航运指数",
    "IH": "上证50", "IF": "沪深300", "IC": "中证500", "IM": "中证1000", "MO": "中证1000期权",
    "TS": "2年期国债", "TF": "5年期国债", "T": "10年期国债", "TL": "30年期国债",
}
AKSHARE_CODE = {
    "A": "A0.DCE", "AD": "AD0.SHF", "AG": "AG0.SHF", "AL": "AL0.SHF", "AO": "AO0.SHF", "AP": "AP0.CZC",
    "AU": "AU0.SHF", "B": "B0.DCE", "BB": "BB0.DCE", "BC": "BCM.INE", "BR": "BR0.SHF", "BU": "BU0.SHF",
    "BZ": "BZ0.DCE", "C": "C0.DCE", "CF": "CF0.CZC", "CJ": "CJ0.CZC", "CS": "CS0.DCE", "CU": "CU0.SHF",
    "CY": "CY0.CZC", "EB": "EB0.DCE", "EC": "ECM.INE", "EG": "EG0.DCE", "FB": "FB0.DCE", "FG": "FG0.CZC",
    "FU": "FU0.SHF", "HC": "HC0.SHF", "I": "I0.DCE", "IC": "IC0.CFE", "IF": "IF0.CFE", "IH": "IH0.CFE",
    "IM": "IM0.CFE", "J": "J0.DCE", "JD": "JD0.DCE", "JM": "JM0.DCE", "JR": "JR0.CZC", "L": "L0.DCE",
    "LC": "LCM.GFE", "LG": "LG0.DCE", "LH": "LH0.DCE", "LR": "LR0.CZC", "LU": "LUM.INE", "M": "M0.DCE",
    "MA": "MA0.CZC", "NI": "NI0.SHF", "NR": "NRM.INE", "OI": "OI0.CZC", "OP": "OP0.SHF", "P": "P0.DCE",
    "PB": "PB0.SHF", "PD": "PDM.GFE", "PF": "PF0.CZC", "PG": "PG0.DCE", "PK": "PK0.CZC", "PL": "PL0.CZC",
    "PM": "PM0.CZC", "PP": "PP0.DCE", "PR": "PR0.CZC", "PS": "PSM.GFE", "PT": "PTM.GFE", "PX": "PX0.CZC",
    "RB": "RB0.SHF", "RI": "RI0.CZC", "RM": "RM0.CZC", "RR": "RR0.DCE", "RS": "RS0.CZC", "RU": "RU0.SHF",
    "SA": "SA0.CZC", "SC": "SCM.INE", "SF": "SF0.CZC", "SH": "SH0.CZC", "SI": "SIM.GFE", "SM": "SM0.CZC",
    "SN": "SN0.SHF", "SP": "SP0.SHF", "SR": "SR0.CZC", "SS": "SS0.SHF", "TA": "TA0.CZC", "T": "T0.CFE",
    "TF": "TF0.CFE", "TL": "TL0.CFE", "TS": "TS0.CFE", "UR": "UR0.CZC", "V": "V0.DCE", "WH": "WH0.CZC",
    "WR": "WR0.SHF", "Y": "Y0.DCE", "ZC": "ZC0.CZC", "ZN": "ZN0.SHF",
}
CODE_TO_PROD = {v: k for k, v in AKSHARE_CODE.items()}
ACTION_ORDER = ["加码", "暂缓加码", "减码准备", "控拥挤", "观望", "补风格"]
ACTION_COLORS = {
    "加码": C_RED,
    "暂缓加码": C_AMBER,
    "减码准备": C_ORANGE,
    "控拥挤": C_GOLD,
    "观望": C_VIOLET,
    "补风格": C_SKY,
}

_CN_FONT: FontProperties | None = None


def load_env() -> None:
    for fname in (".env.local", ".env"):
        for base in (Path("."), ROOT):
            f = base / fname
            if not f.is_file():
                continue
            for line in f.read_text(encoding="utf-8", errors="ignore").splitlines():
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                k, v = k.strip(), v.strip().strip('"').strip("'")
                if k and k not in os.environ:
                    os.environ[k] = v


def get_conn():
    import psycopg2

    url = os.environ.get("DATABASE_URL")
    if url and re.search(r"://[^/:@]+:[^@]+@", url):
        return psycopg2.connect(url)
    return psycopg2.connect(
        host=os.environ.get("DB_HOST", "127.0.0.1"),
        port=int(os.environ.get("DB_PORT", "5432")),
        dbname=os.environ.get("DB_NAME", "market_data"),
        user=os.environ.get("DB_USER", "market_user"),
        password=os.environ.get("DB_PASSWORD", ""),
    )


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    plt.rcParams["axes.edgecolor"] = "#CBD5E0"
    plt.rcParams["axes.grid"] = True
    plt.rcParams["grid.color"] = "#EDF2F7"
    plt.rcParams["grid.linewidth"] = 0.8
    for path in (r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\simsun.ttc"):
        if os.path.isfile(path):
            try:
                fontManager.addfont(path)
                _CN_FONT = FontProperties(fname=path)
                plt.rcParams["font.family"] = "sans-serif"
                plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "Microsoft YaHei", "SimHei"]
                return
            except Exception:
                continue
    _CN_FONT = FontProperties(family="Microsoft YaHei")


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def apply_font(ax) -> None:
    for lab in (*ax.get_xticklabels(), *ax.get_yticklabels()):
        if _CN_FONT is not None:
            lab.set_fontproperties(_CN_FONT)
    if ax.get_xlabel() and _CN_FONT is not None:
        ax.xaxis.label.set_fontproperties(_CN_FONT)
    if ax.get_ylabel() and _CN_FONT is not None:
        ax.yaxis.label.set_fontproperties(_CN_FONT)
    if ax.get_title() and _CN_FONT is not None:
        ax.title.set_fontproperties(_CN_FONT)
    legend = ax.get_legend()
    if legend is not None and _CN_FONT is not None:
        for t in legend.get_texts():
            t.set_fontproperties(_CN_FONT)


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None, italic=False):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or TEXT, italic=italic)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0, italic=False, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.28
    if first_line and align is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color or TEXT, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    size = {1: 16, 2: 13, 3: 12}.get(level, 11)
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=NAVY)
    return p


def shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color or TEXT)
    set_cell_border(cell)


def add_table(doc, headers, rows, col_widths=None, signed_cols=None):
    signed_cols = set(signed_cols or [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=9, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "F7FAFC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            color = TEXT
            s = "" if val is None else str(val)
            if c_i in signed_cols and s not in ("—", "", "-"):
                if s.startswith("-"):
                    color = GREEN
                elif s.startswith("+") or (s.endswith("%") and not s.startswith("0")):
                    color = RED
            cell_text(table.rows[r_i + 1].cells[c_i], s, size=8, color=color, align="left" if c_i == 0 else "center")
            shade(table.rows[r_i + 1].cells[c_i], fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def caption(doc, text):
    para(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_before=2, space_after=12)


def add_picture(doc, path: Path, width_cm=16.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    return p


def set_header_footer(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "MOM 每日风控  ·  决策信号策略回测", size=8, color=MUTED)
    fp_ = section.footer.paragraphs[0]
    fp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(fp_, "内部研究  ·  信号逻辑与 quant-vs-subjective-signals.ts 一致  ·  ", size=8, color=MUTED)
    run = fp_.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=8, color=MUTED)


def round1(n: float) -> float:
    return round(n * 10.0) / 10.0


def get_prefix(contract: str) -> str:
    if not contract:
        return ""
    m = re.match(r"^[A-Za-z]+", contract)
    if m:
        return m.group(0).upper()
    return contract.upper()


def get_sector(prefix: str) -> str:
    return SECTOR_MAP.get(prefix, "其他")


def prod_name(code: str) -> str:
    return PROD_NAMES.get(code, code)


def account_numeric_id(account: str) -> str:
    digits = re.sub(r"\D", "", str(account or ""))
    if not digits:
        return ""
    return digits.lstrip("0") or "0"


def is_quant(account: str) -> bool:
    return account_numeric_id(account) in QUANT_IDS


def zero_rollover_spikes(rets: np.ndarray) -> np.ndarray:
    out = rets.copy()
    n = len(rets)
    for i in range(LOOKBACK_SPIKE, n):
        win = np.abs(rets[i - LOOKBACK_SPIKE : i])
        med = float(np.median(win))
        mad = float(np.median(np.abs(win - med)))
        thr = max(0.06, med + 12.0 * mad * 1.4826)
        if abs(rets[i]) > thr:
            out[i] = 0.0
    return out


def lots_price(long_mv, short_mv, long_lots, short_lots) -> float:
    lots = abs(long_lots) + abs(short_lots)
    mv = abs(long_mv) + abs(short_mv)
    return mv / lots if lots > 1e-9 else 0.0


def decompose_net(prev, today):
    prev_net = prev["long_mv"] - prev["short_mv"]
    today_net = today["long_mv"] - today["short_mv"]
    d_lots = (today["long_lots"] - today["short_lots"]) - (prev["long_lots"] - prev["short_lots"])
    px = lots_price(today["long_mv"], today["short_mv"], today["long_lots"], today["short_lots"])
    if px == 0:
        px = lots_price(prev["long_mv"], prev["short_mv"], prev["long_lots"], prev["short_lots"])
    trade = d_lots * px
    return prev_net, today_net, trade, today_net - prev_net - trade


def flow_is_live(trade, prev_net, today_net) -> bool:
    base = max(abs(prev_net), abs(today_net), FLOW_SCALE_FLOOR)
    return abs(trade) >= max(FLOW_MIN_ABS, FLOW_REL * base)


def kind_from_trades(q_trade, s_trade, q_live, s_live, direction) -> str:
    if not q_live and not s_live:
        return "flat"
    if not q_live or not s_live:
        return "one_sided"
    if (q_trade > 0 and s_trade > 0) or (q_trade < 0 and s_trade < 0):
        return "both_add" if q_trade * direction > 0 else "both_cut"
    return "diverge"


def consensus_action_with_flow(crowded: bool, kind1d: str, kind5d: str) -> str:
    if crowded:
        return "控拥挤"
    if kind1d == "diverge":
        return "加码" if kind5d == "both_add" else "暂缓加码"
    if kind1d == "both_cut":
        return "加码" if kind5d == "both_add" else "减码准备"
    if kind1d == "both_add":
        return "加码"
    if kind5d == "diverge":
        return "暂缓加码"
    if kind5d == "both_cut":
        return "减码准备"
    return "加码"


def classify_row(q: float, s: float, kind1d: str | None, kind5d: str | None) -> tuple[str, str, float]:
    aq, as_ = abs(q), abs(s)
    same = (q > 0 and s > 0) or (q < 0 and s < 0)
    if same and aq >= CONSENSUS_MIN and as_ >= CONSENSUS_MIN:
        kind = "consensus_long" if q > 0 else "consensus_short"
        crowded = aq + as_ >= CROWD_SUM
        action = consensus_action_with_flow(crowded, kind1d or "flat", kind5d or "flat")
        if crowded:
            kind = "crowded"
        return kind, action, 1.0 if q > 0 else -1.0
    if (not same) and aq >= CONSENSUS_MIN and as_ >= CONSENSUS_MIN:
        return "divergence", "观望", 0.0
    if aq >= HEAVY_MIN and as_ < LIGHT_MAX:
        return "quant_only", "补风格", 1.0 if q > 0 else -1.0
    if as_ >= HEAVY_MIN and aq < LIGHT_MAX:
        return "subj_only", "补风格", 1.0 if s > 0 else -1.0
    return "neutral", "中性", 0.0


def empty_pos():
    return {"long_mv": 0.0, "short_mv": 0.0, "long_lots": 0.0, "short_lots": 0.0}


def load_data(conn):
    print("Loading positions…")
    quant_list = ",".join(f"'{x}'" for x in QUANT_IDS)
    option_re = r"[0-9][CP][0-9]"
    sql = f"""
        SELECT
          "交易日期"::date::text AS dt,
          CASE
            WHEN regexp_replace(TRIM("账户"), '[^0-9]', '', 'g') = '' THEN 'subjective'
            WHEN COALESCE(NULLIF(TRIM(LEADING '0' FROM regexp_replace(TRIM("账户"), '[^0-9]', '', 'g')), ''), '0')
                 IN ({quant_list}) THEN 'quant'
            ELSE 'subjective'
          END AS sleeve,
          UPPER(TRIM("合约")) AS contract,
          SUM(CASE WHEN COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("买持仓"::text, ''), ',', ''), ' ', ''), '')::numeric, 0) > 0
                   THEN COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("持仓市値"::text, ''), ',', ''), ' ', ''), '')::numeric, 0)
                   ELSE 0 END)::float8 AS long_mv,
          SUM(CASE WHEN COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("卖持仓"::text, ''), ',', ''), ' ', ''), '')::numeric, 0) > 0
                   THEN COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("持仓市値"::text, ''), ',', ''), ' ', ''), '')::numeric, 0)
                   ELSE 0 END)::float8 AS short_mv,
          SUM(COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("保证金"::text, ''), ',', ''), ' ', ''), '')::numeric, 0))::float8 AS margin,
          SUM(COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("买持仓"::text, ''), ',', ''), ' ', ''), '')::numeric, 0))::float8 AS long_lots,
          SUM(COALESCE(NULLIF(REPLACE(REPLACE(COALESCE("卖持仓"::text, ''), ',', ''), ' ', ''), '')::numeric, 0))::float8 AS short_lots
        FROM mom_position_details
        WHERE "交易日期" IS NOT NULL
          AND "合约" IS NOT NULL
          AND UPPER(TRIM("账户"::text)) NOT LIKE '%%GUOXIN%%'
          AND UPPER(TRIM("账户"::text)) NOT LIKE '%%GUOSEN%%'
          AND TRIM("账户"::text) NOT LIKE '%%国信%%'
          AND TRIM("账户"::text) <> '665300200077'
          AND UPPER(TRIM("合约")) !~ '{option_re}'
          AND TRIM("合约") NOT LIKE '%%-%%-%%'
        GROUP BY 1, 2, 3
        ORDER BY 1
    """
    pos = pd.read_sql(sql, conn)
    pos["dt"] = pos["dt"].astype(str).str.slice(0, 10)
    print(f"  position rows: {len(pos):,}")
    print("Loading futures returns…")
    px = pd.read_sql(
        """
        SELECT trade_date::text AS dt, code,
               pct_change::float8 AS pct,
               close::float8 AS close,
               clear::float8 AS settle
        FROM raw_akshare_futures_daily
        WHERE pct_change IS NOT NULL
        ORDER BY trade_date, code
        """,
        conn,
    )
    px["dt"] = px["dt"].astype(str).str.slice(0, 10)
    print(f"  price rows: {len(px):,}")
    return pos, px


def build_returns(px: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, list[str]]:
    px = px.copy()
    px["product"] = px["code"].map(CODE_TO_PROD)
    px = px.dropna(subset=["product"])
    px["ret"] = px["pct"] / 100.0
    px["px"] = px["close"].where(px["close"].notna() & (px["close"] > 0), px["settle"])
    wide = px.pivot_table(index="dt", columns="product", values="ret", aggfunc="last").sort_index()
    wide.index = wide.index.astype(str).str.slice(0, 10)
    close = px.pivot_table(index="dt", columns="product", values="px", aggfunc="last").sort_index()
    close.index = close.index.astype(str).str.slice(0, 10)
    clean = wide.copy()
    for col in clean.columns:
        vals = clean[col].fillna(0.0).to_numpy(dtype=float)
        clean[col] = zero_rollover_spikes(vals)
    dates = list(clean.index)
    return wide, clean, close, dates


def sigma_on(clean: pd.DataFrame, product: str, as_of: str) -> float:
    if product not in clean.columns:
        return 0.0
    idx = clean.index.searchsorted(as_of, side="right") - 1
    if idx < 2:
        return 0.0
    start = max(0, idx - VOL_DAYS)
    window = clean.iloc[start:idx][product].to_numpy(dtype=float)
    window = window[window != 0]
    if len(window) < 2:
        return 0.0
    return float(np.std(window, ddof=1))


def build_signals(pos: pd.DataFrame, clean: pd.DataFrame) -> pd.DataFrame:
    print("Building daily signals…")
    pos = pos.copy()
    pos["product"] = pos["contract"].map(get_prefix)
    pos = pos[pos["product"].str.len() > 0]
    pos["sector"] = pos["product"].map(get_sector)
    grouped = (
        pos.groupby(["dt", "sleeve", "product"], as_index=False)
        .agg(long_mv=("long_mv", "sum"), short_mv=("short_mv", "sum"),
             margin=("margin", "sum"), long_lots=("long_lots", "sum"), short_lots=("short_lots", "sum"))
    )
    dates = sorted(grouped["dt"].unique())
    products = sorted(grouped["product"].unique())
    pos_map: dict[tuple, dict] = {}
    for r in grouped.itertuples(index=False):
        pos_map[(r.dt, r.sleeve, r.product)] = {
            "long_mv": r.long_mv, "short_mv": r.short_mv,
            "long_lots": r.long_lots, "short_lots": r.short_lots, "margin": r.margin,
        }

    sigma_cache: dict[tuple[str, str], float] = {}
    median_sigma: dict[str, float] = {}
    rows = []
    date_index = {d: i for i, d in enumerate(dates)}

    for di, dt in enumerate(dates):
        if di == 0:
            continue
        prev = dates[di - 1]
        look5 = dates[max(0, di - 5)]
        day_prods = sorted({p for p in products if (dt, "quant", p) in pos_map or (dt, "subjective", p) in pos_map})
        raw_sig = []
        sig_today = {}
        for p in day_prods:
            key = (p, dt)
            if key not in sigma_cache:
                sigma_cache[key] = sigma_on(clean, p, dt)
            s = sigma_cache[key]
            sig_today[p] = s
            if s > 0:
                raw_sig.append(s)
        med = float(np.median(raw_sig)) if raw_sig else median_sigma.get(prev, 0.0)
        median_sigma[dt] = med

        q_abs = 0.0
        s_abs = 0.0
        nets = {}
        for p in day_prods:
            sig = sig_today[p] if sig_today[p] > 0 else med
            q = pos_map.get((dt, "quant", p), empty_pos())
            s = pos_map.get((dt, "subjective", p), empty_pos())
            q_net = q["long_mv"] - q["short_mv"]
            s_net = s["long_mv"] - s["short_mv"]
            q_risk = sig * q_net
            s_risk = sig * s_net
            nets[p] = (q_net, s_net, q_risk, s_risk, q, s, sig)
            q_abs += abs(q_risk)
            s_abs += abs(s_risk)

        scored = []
        for p, (q_net, s_net, q_risk, s_risk, qpos, spos, sig) in nets.items():
            q_pct = (q_risk / q_abs * 100.0) if q_abs > 0 else 0.0
            s_pct = (s_risk / s_abs * 100.0) if s_abs > 0 else 0.0
            scored.append((abs(q_pct) + abs(s_pct), p, q_pct, s_pct, qpos, spos, q_net, s_net))
        scored.sort(reverse=True)
        scored = scored[:TOP_N_PRODUCTS]

        for _, p, q_pct, s_pct, qpos, spos, q_net, s_net in scored:
            q_prev = pos_map.get((prev, "quant", p), empty_pos())
            s_prev = pos_map.get((prev, "subjective", p), empty_pos())
            q5 = pos_map.get((look5, "quant", p), empty_pos())
            s5 = pos_map.get((look5, "subjective", p), empty_pos())
            q_prev_net, q_today_net, q_trade, _ = decompose_net(q_prev, qpos)
            s_prev_net, s_today_net, s_trade, _ = decompose_net(s_prev, spos)
            q_px = lots_price(qpos["long_mv"], qpos["short_mv"], qpos["long_lots"], qpos["short_lots"]) or lots_price(q5["long_mv"], q5["short_mv"], q5["long_lots"], q5["short_lots"])
            s_px = lots_price(spos["long_mv"], spos["short_mv"], spos["long_lots"], spos["short_lots"]) or lots_price(s5["long_mv"], s5["short_mv"], s5["long_lots"], s5["short_lots"])
            q_trade5 = ((qpos["long_lots"] - qpos["short_lots"]) - (q5["long_lots"] - q5["short_lots"])) * q_px
            s_trade5 = ((spos["long_lots"] - spos["short_lots"]) - (s5["long_lots"] - s5["short_lots"])) * s_px
            q_live1 = flow_is_live(q_trade, q_prev_net, q_today_net)
            s_live1 = flow_is_live(s_trade, s_prev_net, s_today_net)
            q_live5 = flow_is_live(q_trade5, q_today_net - q_trade5, q_today_net)
            s_live5 = flow_is_live(s_trade5, s_today_net - s_trade5, s_today_net)
            direction = -1.0 if (q_pct + s_pct) < 0 else 1.0
            kind1d = kind_from_trades(q_trade, s_trade, q_live1, s_live1, direction)
            kind5d = kind_from_trades(q_trade5, s_trade5, q_live5, s_live5, direction)
            sig_kind, action, trade_dir = classify_row(round1(q_pct), round1(s_pct), kind1d, kind5d)
            if action == "中性":
                continue
            rows.append({
                "date": dt,
                "product": p,
                "name": prod_name(p),
                "sector": get_sector(p),
                "q_pct": round1(q_pct),
                "s_pct": round1(s_pct),
                "kind": sig_kind,
                "action": action,
                "trade_dir": trade_dir,
                "kind1d": kind1d,
                "kind5d": kind5d,
            })
        if di % 40 == 0:
            print(f"  {dt}  signals so far {len(rows):,}")

    sig = pd.DataFrame(rows)
    print(f"  product-level signal rows: {len(sig):,}")
    return sig


def fwd_return(wide: pd.DataFrame, product: str, dt: str, n: int) -> float | None:
    if product not in wide.columns:
        return None
    loc = int(wide.index.searchsorted(dt, side="right"))
    if loc >= len(wide.index):
        return None
    end = min(loc + n, len(wide.index))
    chunk = wide.iloc[loc:end][product].astype(float)
    if chunk.isna().all():
        return None
    return float((1.0 + chunk.fillna(0.0)).prod() - 1.0)


def next_day_return(wide: pd.DataFrame, product: str, dt: str) -> float | None:
    return fwd_return(wide, product, dt, 1)


def event_study(sig: pd.DataFrame, wide: pd.DataFrame) -> pd.DataFrame:
    print("Event study…")
    recs = []
    for r in sig.itertuples(index=False):
        r1 = next_day_return(wide, r.product, r.date)
        r5 = fwd_return(wide, r.product, r.date, 5)
        r10 = fwd_return(wide, r.product, r.date, 10)
        r20 = fwd_return(wide, r.product, r.date, 20)
        direction = r.trade_dir
        if r.action == "观望":
            q_sign = 1.0 if r.q_pct > 0 else -1.0
            aligned = None
        else:
            aligned = direction
        recs.append({
            "date": r.date,
            "product": r.product,
            "name": r.name,
            "sector": r.sector,
            "action": r.action,
            "kind": r.kind,
            "q_pct": r.q_pct,
            "s_pct": r.s_pct,
            "dir": direction,
            "r1": r1, "r5": r5, "r10": r10, "r20": r20,
            "signed_1": None if r1 is None or aligned in (None, 0) else aligned * r1,
            "signed_5": None if r5 is None or aligned in (None, 0) else aligned * r5,
            "signed_10": None if r10 is None or aligned in (None, 0) else aligned * r10,
            "signed_20": None if r20 is None or aligned in (None, 0) else aligned * r20,
            "quant_1": None if r1 is None else (1.0 if r.q_pct > 0 else -1.0) * r1,
            "subj_1": None if r1 is None else (1.0 if r.s_pct > 0 else -1.0) * r1,
        })
    return pd.DataFrame(recs)


def summarize_events(ev: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for action, g in ev.groupby("action"):
        def stats(col, signed=True):
            s = g[col].dropna()
            nn = len(s)
            if nn == 0:
                return 0, None, None, None, None
            hit = float((s > 0).mean()) if signed else None
            mu = float(s.mean())
            se = float(s.std(ddof=1) / math.sqrt(nn)) if nn > 1 else None
            t = mu / se if se and se > 0 else None
            return nn, mu, hit, t, float(s.median())

        n = len(g)
        n1, mu1, hit1, t1, med1 = stats("signed_1")
        _, mu5, hit5, t5, _ = stats("signed_5")
        _, mu10, hit10, t10, _ = stats("signed_10")
        _, mu20, hit20, t20, _ = stats("signed_20")
        if action == "观望":
            n1q, mu1q, hit1q, t1q, _ = stats("quant_1")
            n1s, mu1s, hit1s, t1s, _ = stats("subj_1")
        else:
            n1q = n1s = mu1q = mu1s = hit1q = hit1s = t1q = t1s = None
        rows.append({
            "action": action, "n": n,
            "n1": n1, "mu1": mu1, "hit1": hit1, "t1": t1, "med1": med1,
            "mu5": mu5, "hit5": hit5, "t5": t5,
            "mu10": mu10, "hit10": hit10, "t10": t10,
            "mu20": mu20, "hit20": hit20, "t20": t20,
            "mu1_quant": mu1q, "hit1_quant": hit1q, "t1_quant": t1q,
            "mu1_subj": mu1s, "hit1_subj": hit1s, "t1_subj": t1s,
        })
    out = pd.DataFrame(rows)
    out["ord"] = out["action"].map({a: i for i, a in enumerate(ACTION_ORDER)})
    return out.sort_values("ord").drop(columns=["ord"])


def weights_for_signal(action: str, kind: str, q_pct: float, s_pct: float, scheme: str) -> float:
    """Signed overlay weight before cross-sectional scaling. 0 = flat."""
    cons_dir = 1.0 if (q_pct + s_pct) >= 0 else -1.0
    if kind in ("consensus_long",):
        cons_dir = 1.0
    elif kind in ("consensus_short",):
        cons_dir = -1.0
    elif kind == "crowded":
        cons_dir = 1.0 if q_pct > 0 else -1.0
    heavy_dir = 1.0 if abs(q_pct) >= abs(s_pct) else (1.0 if s_pct > 0 else -1.0)
    if abs(q_pct) >= abs(s_pct):
        heavy_dir = 1.0 if q_pct > 0 else -1.0
    else:
        heavy_dir = 1.0 if s_pct > 0 else -1.0
    q_dir = 1.0 if q_pct > 0 else -1.0
    s_dir = 1.0 if s_pct > 0 else -1.0

    if scheme == "加码跟随":
        return cons_dir if action == "加码" else 0.0
    if scheme == "完整映射":
        if action == "加码":
            return cons_dir
        if action == "暂缓加码":
            return 0.5 * cons_dir
        if action == "减码准备":
            return 0.0
        if action == "控拥挤":
            return 0.3 * cons_dir
        if action == "观望":
            return 0.0
        if action == "补风格":
            return 0.5 * heavy_dir
        return 0.0
    if scheme == "加码+补风格":
        if action == "加码":
            return cons_dir
        if action == "补风格":
            return 0.5 * heavy_dir
        return 0.0
    if scheme == "分歧跟量化":
        if action == "加码":
            return cons_dir
        if action == "观望":
            return q_dir
        return 0.0
    if scheme == "分歧跟主观":
        if action == "加码":
            return cons_dir
        if action == "观望":
            return s_dir
        return 0.0
    if scheme == "量化方向":
        return q_dir if abs(q_pct) >= CONSENSUS_MIN else 0.0
    if scheme == "主观方向":
        return s_dir if abs(s_pct) >= CONSENSUS_MIN else 0.0
    return 0.0


def scale_weights(raw: dict[str, float]) -> dict[str, float]:
    gross = sum(abs(v) for v in raw.values())
    if gross <= 1e-12:
        return {k: 0.0 for k in raw}
    return {k: v / gross for k, v in raw.items()}


def portfolio_backtest(sig: pd.DataFrame, wide: pd.DataFrame, schemes: list[str], cost_bps: float, quiet=False) -> dict[str, pd.DataFrame]:
    if not quiet:
        print("Portfolio backtest…")
    by_date = {d: g for d, g in sig.groupby("date")}
    dates = sorted(by_date)
    # signal date t is applied to return of the next session after t
    out = {}
    for scheme in schemes:
        records = []
        prev_w: dict[str, float] = {}
        for dt in dates:
            g = by_date[dt]
            raw = {}
            for r in g.itertuples(index=False):
                w = weights_for_signal(r.action, r.kind, r.q_pct, r.s_pct, scheme)
                if w != 0:
                    raw[r.product] = raw.get(r.product, 0.0) + w
            w = scale_weights(raw)
            names = set(w) | set(prev_w)
            turnover = sum(abs(w.get(p, 0.0) - prev_w.get(p, 0.0)) for p in names)
            cost = turnover * (cost_bps / 1e4)
            loc = int(wide.index.searchsorted(dt, side="right"))
            if loc >= len(wide.index):
                break
            nxt = wide.index[loc]
            pnl = 0.0
            n_live = 0
            for p, wp in w.items():
                if wp == 0 or p not in wide.columns:
                    continue
                r = wide.at[nxt, p]
                if pd.isna(r):
                    continue
                pnl += wp * float(r)
                n_live += 1
            pnl -= cost
            records.append({
                "signal_date": dt,
                "return_date": nxt,
                "ret": pnl,
                "gross": sum(abs(v) for v in w.values()),
                "n": n_live,
                "turnover": turnover,
                "cost": cost,
            })
            prev_w = w
        df = pd.DataFrame(records)
        if df.empty:
            out[scheme] = df
            continue
        df["nav"] = (1.0 + df["ret"]).cumprod()
        df["dd"] = df["nav"] / df["nav"].cummax() - 1.0
        out[scheme] = df
    return out


def vol_target(df: pd.DataFrame, target=VOL_TARGET) -> pd.DataFrame:
    if df.empty or len(df) < 20:
        return df
    out = df.copy()
    roll = out["ret"].rolling(20, min_periods=10).std()
    scale = (target / math.sqrt(252.0)) / roll.replace(0, np.nan)
    scale = scale.shift(1).clip(upper=5.0).fillna(1.0)
    out["ret_vt"] = out["ret"] * scale
    out["nav_vt"] = (1.0 + out["ret_vt"]).cumprod()
    out["dd_vt"] = out["nav_vt"] / out["nav_vt"].cummax() - 1.0
    out["scale"] = scale
    return out


def perf_stats(ret: pd.Series, nav: pd.Series | None = None) -> dict:
    r = ret.dropna()
    if r.empty:
        return {k: None for k in ("n", "cagr", "vol", "sharpe", "maxdd", "calmar", "hit", "avg", "turnover")}
    n = len(r)
    years = n / 252.0
    nav = nav if nav is not None else (1.0 + r).cumprod()
    total = float(nav.iloc[-1] / nav.iloc[0]) if len(nav) else 1.0
    cagr = total ** (1.0 / years) - 1.0 if years > 0 and total > 0 else None
    vol = float(r.std(ddof=1) * math.sqrt(252)) if n > 1 else None
    sharpe = (float(r.mean()) / float(r.std(ddof=1)) * math.sqrt(252)) if n > 1 and r.std(ddof=1) > 0 else None
    maxdd = float((nav / nav.cummax() - 1.0).min())
    calmar = (cagr / abs(maxdd)) if cagr is not None and maxdd not in (0, None) else None
    hit = float((r > 0).mean())
    return {
        "n": n,
        "cagr": cagr,
        "vol": vol,
        "sharpe": sharpe,
        "maxdd": maxdd,
        "calmar": calmar,
        "hit": hit,
        "avg": float(r.mean()),
        "sum": float(r.sum()),
    }


def fmt_pct(v, digits=2, signed=True, already=False):
    if v is None or (isinstance(v, float) and (math.isnan(v) or math.isinf(v))):
        return "—"
    x = float(v) * 100.0 if not already else float(v)
    sign = "+" if signed and x > 0 else ""
    return f"{sign}{x:.{digits}f}%"


def fmt_num(v, digits=2):
    if v is None or (isinstance(v, float) and (math.isnan(v) or math.isinf(v))):
        return "—"
    return f"{float(v):.{digits}f}"


def fmt_t(v):
    if v is None or (isinstance(v, float) and (math.isnan(v) or math.isinf(v))):
        return "—"
    return f"{float(v):.2f}"


def save_fig(fig, name: str) -> Path:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def draw_equity(bt: dict[str, pd.DataFrame], use_vt=False) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 5.4), dpi=160)
    colors = [C_NAVY, C_RED, C_BLUE, C_ORANGE, C_VIOLET, C_TEAL, C_GOLD, C_GREEN]
    for i, (name, df) in enumerate(bt.items()):
        if df.empty:
            continue
        y = df["nav_vt"] if use_vt and "nav_vt" in df.columns else df["nav"]
        x = pd.to_datetime(df["return_date"])
        ax.plot(x, y, lw=1.7, color=colors[i % len(colors)], label=name)
    ax.axhline(1.0, color="#CBD5E0", lw=0.8)
    ax.set_title("波动目标 10% 后净值" if use_vt else "等权 overlay 累计净值（成本后）", fontsize=13, color=C_NAVY, **fp())
    ax.set_xlabel("收益归属日", **fp())
    ax.set_ylabel("净值（起始=1）", **fp())
    ax.legend(frameon=False, ncol=2, fontsize=8)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    apply_font(ax)
    return save_fig(fig, "equity_vt.png" if use_vt else "equity.png")


def draw_dd(bt: dict[str, pd.DataFrame], key="加码跟随") -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 3.8), dpi=160)
    df = bt.get(key)
    if df is not None and not df.empty:
        x = pd.to_datetime(df["return_date"])
        ax.fill_between(x, df["dd"] * 100, 0, color=C_RED, alpha=0.35, label=f"{key} 回撤")
        if "dd_vt" in df.columns:
            ax.plot(x, df["dd_vt"] * 100, color=C_NAVY, lw=1.2, label="波动目标后")
    ax.set_title(f"{key} 回撤", fontsize=13, color=C_NAVY, **fp())
    ax.set_xlabel("日期", **fp())
    ax.set_ylabel("回撤（%）", **fp())
    ax.legend(frameon=False, fontsize=8)
    apply_font(ax)
    return save_fig(fig, "drawdown.png")


def draw_event_bars(summary: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 5.2), dpi=160)
    actions = [a for a in ACTION_ORDER if a in set(summary["action"])]
    x = np.arange(len(actions))
    width = 0.22
    series = [("mu1", "次日"), ("mu5", "5日"), ("mu10", "10日"), ("mu20", "20日")]
    colors = [C_NAVY, C_BLUE, C_ORANGE, C_RED]
    for i, (col, lab) in enumerate(series):
        vals = []
        for a in actions:
            row = summary[summary["action"] == a].iloc[0]
            v = row[col]
            vals.append(0.0 if v is None or (isinstance(v, float) and math.isnan(v)) else v * 100)
        ax.bar(x + (i - 1.5) * width, vals, width, label=lab, color=colors[i])
    ax.axhline(0, color="#4A5568", lw=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels(actions, **fp())
    ax.set_ylabel("信号方向上的平均远期收益（%）", **fp())
    ax.set_title("各动作发出后，按信号方向对齐的平均收益", fontsize=13, color=C_NAVY, **fp())
    ax.legend(frameon=False, ncol=4, fontsize=8)
    apply_font(ax)
    return save_fig(fig, "event_bars.png")


def draw_hit(summary: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 4.6), dpi=160)
    actions = [a for a in ACTION_ORDER if a in set(summary["action"])]
    x = np.arange(len(actions))
    width = 0.35
    hit1, hit5 = [], []
    for a in actions:
        row = summary[summary["action"] == a].iloc[0]
        hit1.append((row["hit1"] or 0) * 100)
        hit5.append((row["hit5"] or 0) * 100)
    ax.bar(x - width / 2, hit1, width, label="次日胜率", color=C_NAVY)
    ax.bar(x + width / 2, hit5, width, label="5日胜率", color=C_ORANGE)
    ax.axhline(50, color=C_GRAY, ls="--", lw=1)
    ax.set_xticks(x)
    ax.set_xticklabels(actions, **fp())
    ax.set_ylabel("胜率（%）", **fp())
    ax.set_ylim(0, 100)
    ax.set_title("信号方向对齐后的胜率（虚线=50%）", fontsize=13, color=C_NAVY, **fp())
    ax.legend(frameon=False, fontsize=8)
    apply_font(ax)
    return save_fig(fig, "hit_rate.png")


def draw_signal_counts(sig: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 4.8), dpi=160)
    daily = sig.groupby(["date", "action"]).size().unstack(fill_value=0)
    daily.index = pd.to_datetime(daily.index)
    for a in ACTION_ORDER:
        if a in daily.columns:
            ax.plot(daily.index, daily[a].rolling(10, min_periods=1).mean(), lw=1.5, color=ACTION_COLORS.get(a, C_GRAY), label=a)
    ax.set_title("每日信号条数（10日均线）", fontsize=13, color=C_NAVY, **fp())
    ax.set_xlabel("日期", **fp())
    ax.set_ylabel("条数", **fp())
    ax.legend(frameon=False, ncol=3, fontsize=8)
    apply_font(ax)
    return save_fig(fig, "signal_counts.png")


def draw_monthly(df: pd.DataFrame, title: str, fname: str) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 4.6), dpi=160)
    if df.empty:
        ax.text(0.5, 0.5, "无数据", ha="center", **fp())
        return save_fig(fig, fname)
    s = df.set_index(pd.to_datetime(df["return_date"]))["ret"].resample("ME").apply(lambda x: (1 + x).prod() - 1)
    colors = [C_RED if v >= 0 else C_GREEN for v in s.values]
    ax.bar(s.index, s.values * 100, width=20, color=colors, align="center")
    ax.axhline(0, color="#4A5568", lw=0.8)
    ax.set_title(title, fontsize=13, color=C_NAVY, **fp())
    ax.set_ylabel("月收益（%）", **fp())
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    apply_font(ax)
    return save_fig(fig, fname)


def draw_sector_contrib(ev: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 5.0), dpi=160)
    g = ev[ev["action"] == "加码"].dropna(subset=["signed_1"])
    if g.empty:
        ax.text(0.5, 0.5, "无加码样本", ha="center", **fp())
        return save_fig(fig, "sector_contrib.png")
    sec = g.groupby("sector")["signed_1"].agg(["mean", "count"])
    sec = sec[sec["count"] >= 8].sort_values("mean")
    colors = [C_RED if v >= 0 else C_GREEN for v in sec["mean"]]
    ax.barh(sec.index, sec["mean"] * 100, color=colors)
    ax.axvline(0, color="#4A5568", lw=0.8)
    ax.set_xlabel("加码信号次日、按方向对齐的平均收益（%）", **fp())
    ax.set_title("加码信号：分板块次日收益", fontsize=13, color=C_NAVY, **fp())
    apply_font(ax)
    return save_fig(fig, "sector_contrib.png")


def draw_product_table_chart(ev: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 6.0), dpi=160)
    g = ev[ev["action"] == "加码"].dropna(subset=["signed_1"])
    top = g.groupby(["product", "name"]).agg(mu=("signed_1", "mean"), n=("signed_1", "count"), hit=("signed_1", lambda s: (s > 0).mean()))
    top = top[top["n"] >= 12].sort_values("mu", ascending=False).head(12)
    if top.empty:
        ax.text(0.5, 0.5, "样本不足", ha="center", **fp())
        return save_fig(fig, "product_top.png")
    labels = [f"{n}({p})" for p, n in top.index]
    colors = [C_RED if v >= 0 else C_GREEN for v in top["mu"]]
    ax.barh(labels[::-1], (top["mu"] * 100).iloc[::-1], color=colors[::-1])
    ax.axvline(0, color="#4A5568", lw=0.8)
    ax.set_xlabel("平均次日对齐收益（%）", **fp())
    ax.set_title("加码信号样本≥12 的品种：次日收益最高的 12 个", fontsize=13, color=C_NAVY, **fp())
    apply_font(ax)
    return save_fig(fig, "product_top.png")


def draw_turnover(df: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 3.8), dpi=160)
    if not df.empty:
        x = pd.to_datetime(df["return_date"])
        ax.plot(x, df["turnover"].rolling(10, min_periods=1).mean(), color=C_NAVY, lw=1.4)
        ax.set_ylabel("换手（|Δw| 之和）", **fp())
    ax.set_title("加码跟随：10 日平均换手", fontsize=13, color=C_NAVY, **fp())
    apply_font(ax)
    return save_fig(fig, "turnover.png")


def draw_cost_sensitivity(sig, wide, scheme="加码跟随") -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 4.6), dpi=160)
    bps_list = [0, 1, 2, 5, 10]
    sharpes = []
    cagrs = []
    for bps in bps_list:
        bt = portfolio_backtest(sig, wide, [scheme], bps, quiet=True)[scheme]
        st = perf_stats(bt["ret"], bt["nav"]) if not bt.empty else {}
        sharpes.append(st.get("sharpe") or 0)
        cagrs.append((st.get("cagr") or 0) * 100)
    ax.plot(bps_list, sharpes, marker="o", color=C_NAVY, lw=1.6, label="夏普")
    ax.set_xlabel("单边成本（bps）", **fp())
    ax.set_ylabel("夏普比率", **fp())
    ax2 = ax.twinx()
    ax2.plot(bps_list, cagrs, marker="s", color=C_ORANGE, lw=1.4, label="年化%")
    ax2.set_ylabel("年化收益（%）", **fp())
    ax.set_title("加码跟随：交易成本敏感性", fontsize=13, color=C_NAVY, **fp())
    lines, labels = ax.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax.legend(lines + lines2, labels + labels2, frameon=False, fontsize=8)
    apply_font(ax)
    apply_font(ax2)
    return save_fig(fig, "cost_sens.png")


def draw_split(bt_df: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(11.2, 4.4), dpi=160)
    if bt_df.empty:
        return save_fig(fig, "split.png")
    df = bt_df.copy()
    df["year"] = pd.to_datetime(df["return_date"]).dt.year
    years = sorted(df["year"].unique())
    cagr, sharpe = [], []
    for y in years:
        sub = df[df["year"] == y]
        st = perf_stats(sub["ret"], (1 + sub["ret"]).cumprod())
        cagr.append((st["cagr"] or 0) * 100)
        sharpe.append(st["sharpe"] or 0)
    axes[0].bar([str(y) for y in years], cagr, color=C_NAVY)
    axes[0].axhline(0, color="#4A5568", lw=0.8)
    axes[0].set_title("分年年化收益", color=C_NAVY, **fp())
    axes[0].set_ylabel("%", **fp())
    axes[1].bar([str(y) for y in years], sharpe, color=C_BLUE)
    axes[1].axhline(0, color="#4A5568", lw=0.8)
    axes[1].set_title("分年夏普", color=C_NAVY, **fp())
    for ax in axes:
        apply_font(ax)
    return save_fig(fig, "yearly.png")


def kind_label(k: str) -> str:
    return {
        "consensus_long": "共识做多",
        "consensus_short": "共识做空",
        "divergence": "方向分歧",
        "quant_only": "仅量化",
        "subj_only": "仅主观",
        "crowded": "共识但拥挤",
    }.get(k, k)


def add_account_chapter(doc, ctx: dict):
    acct = ctx["acct"]
    acct0 = ctx["acct_core"]
    daily = acct["daily"]
    holds = acct["holds"]
    trades = acct["trades"]
    st = ctx["acct_stats"]
    st0 = ctx["acct_core_stats"]
    ch = ctx["acct_charts"]

    heading(doc, "二、2,000 万账户：按信号下单会赚多少", 1)
    para(
        doc,
        f"假设单独开一个期货账户，初始权益 {fmt_yuan(START_EQUITY)}。"
        "每个交易日收盘后读取当日 MOM 决策信号，下一交易日按规则调仓并承担当日涨跌。"
        "加码：与量化/主观共识同向开仓；控拥挤：与共识反向开仓；观望、暂缓加码、减码准备：不开仓。"
        "补风格：跟随重仓一侧的方向（仅量化重仓则跟量化，仅主观重仓则跟主观）。"
        "仓位按当日权益的约 2.2 倍名义、最多 8 个品种等权分配，单品种一日 1σ 亏损不超过权益的 1.2%，"
        f"保证金占用不超过权益的 {MAX_MARGIN_UTIL:.0%}（券商保证金系数 1.1）。手数取整。"
        f"费用按成交名义本金单边手续费 {COMM_RATE * 1e4:.1f}bp + 滑点 {SLIP_RATE * 1e4:.1f}bp，"
        "且每手手续费不低于 3 元，开平都收，用来贴近国内期货公司的全部摩擦。",
    )

    add_table(
        doc,
        ["信号", "本账户怎么做"],
        [
            ["加码", "与量化/主观存量同向开仓（共识做多则买、共识做空则卖）"],
            ["控拥挤", "与量化/主观共识反向开仓（两边都已很重，账户做对冲）"],
            ["观望", "不开仓"],
            ["暂缓加码", "不开仓（存量同向但边际已背离，不再加）"],
            ["减码准备", "不开仓；若昨有仓则平掉"],
            ["补风格", "跟重仓一侧同向（覆盖提示被译成交易：复制已有方向）"],
        ],
        col_widths=[3.2, 13.2],
    )
    caption(doc, "表 A-1  2,000 万账户的下单规则。信号消失或改为观望/暂缓/减码时，下一交易日平仓。")

    if st:
        para(
            doc,
            f"样本内（{daily['return_date'].iloc[0]} 至 {daily['return_date'].iloc[-1]}，{st['n']} 个交易日）"
            f"期末权益 {fmt_yuan(st['end'])}，净盈亏 {fmt_yuan(st['pnl'], signed=True)}，"
            f"相对起始资金 {fmt_pct((st['end'] - START_EQUITY) / START_EQUITY)}。"
            f"年化 {fmt_pct(st.get('cagr'))}，波动 {fmt_pct(st.get('vol'))}，夏普 {fmt_num(st.get('sharpe'))}，"
            f"最大回撤 {fmt_pct(st.get('maxdd'))}（约 {fmt_yuan(st.get('maxdd_yuan'))}），"
            f"日胜率 {fmt_pct(st.get('hit'), already=False)}。"
            f"期间手续费 {fmt_yuan(st.get('total_comm'))}、滑点 {fmt_yuan(st.get('total_slip'))}，"
            f"费用合计 {fmt_yuan(st.get('total_cost'))}，吃掉毛利润 {fmt_yuan(st.get('total_gross'), signed=True)} 中的一部分。"
            f"日均持仓 {fmt_num(st.get('avg_names'), 1)} 个品种，日均保证金 {fmt_yuan(st.get('avg_margin'))}，"
            f"占用率 {fmt_pct(st.get('avg_util'))}，日均名义 {fmt_yuan(st.get('avg_notional'))}。",
        )
    if st0:
        para(
            doc,
            f"若严格按你举的三个例子、不做补风格（只做加码同向 + 控拥挤反向，其余空仓），"
            f"期末权益 {fmt_yuan(st0.get('end'))}，净盈亏 {fmt_yuan(st0.get('pnl'), signed=True)}，"
            f"年化 {fmt_pct(st0.get('cagr'))}，夏普 {fmt_num(st0.get('sharpe'))}，回撤 {fmt_pct(st0.get('maxdd'))}。"
            "补风格条数最多，把它做成跟重仓一侧会显著改变路径；两套结果都列在下面，避免只报好看的那条。",
        )

    if "acct_equity" in ch:
        add_picture(doc, ch["acct_equity"])
        caption(doc, "图 A-1  含补风格规则下，2,000 万账户扣费后权益。")
    if "acct_dd" in ch:
        add_picture(doc, ch["acct_dd"])
        caption(doc, "图 A-2  账户回撤。")
    if "acct_cost" in ch:
        add_picture(doc, ch["acct_cost"])
        caption(doc, "图 A-3  毛盈亏、交易费用与净盈亏。费用按成交名义逐笔计提，不是事后统一扣一个点数。")
    if "acct_margin" in ch:
        add_picture(doc, ch["acct_margin"])
        caption(doc, "图 A-4  保证金占用与名义本金。名义随信号个数与权益复利变化。")
    if "acct_n" in ch:
        add_picture(doc, ch["acct_n"])
        caption(doc, "图 A-5  每日持仓品种数（上限 8）。")
    if "acct_monthly" in ch:
        add_picture(doc, ch["acct_monthly"])
        caption(doc, "图 A-6  月度净盈亏（万元，红赚绿亏）。")
    if "acct_by_action" in ch:
        add_picture(doc, ch["acct_by_action"])
        caption(doc, "图 A-7  持仓毛盈亏按信号动作拆开。可直接看加码、控拥挤、补风格各自贡献。")
    if "acct_by_prod" in ch:
        add_picture(doc, ch["acct_by_prod"])
        caption(doc, "图 A-8  品种层累计毛盈亏两端。")

    heading(doc, "2.1 账户绩效与费用", 2)
    add_table(
        doc,
        ["口径", "期末权益", "净盈亏", "年化", "波动", "夏普", "最大回撤", "费用合计"],
        [
            [
                "加码同向 + 控拥挤反向 + 补风格跟重仓",
                fmt_yuan(st.get("end")),
                fmt_yuan(st.get("pnl"), signed=True),
                fmt_pct(st.get("cagr")),
                fmt_pct(st.get("vol")),
                fmt_num(st.get("sharpe")),
                fmt_pct(st.get("maxdd")),
                fmt_yuan(st.get("total_cost")),
            ],
            [
                "只做加码同向 + 控拥挤反向（不做补风格）",
                fmt_yuan(st0.get("end")),
                fmt_yuan(st0.get("pnl"), signed=True),
                fmt_pct(st0.get("cagr")),
                fmt_pct(st0.get("vol")),
                fmt_num(st0.get("sharpe")),
                fmt_pct(st0.get("maxdd")),
                fmt_yuan(st0.get("total_cost")),
            ],
        ],
        signed_cols={2, 3, 5, 6},
    )
    caption(doc, "表 A-2  两套下单口径。推荐记账用「含补风格」对照「不含」，不要只看一条。")

    add_table(
        doc,
        ["费用项", "金额", "说明"],
        [
            ["手续费", fmt_yuan(st.get("total_comm")), f"成交名义 × {COMM_RATE * 1e4:.1f}bp，每手不低于 3 元，开平都收"],
            ["滑点", fmt_yuan(st.get("total_slip")), f"成交名义 × {SLIP_RATE * 1e4:.1f}bp，约 1～2 个跳价"],
            ["费用合计", fmt_yuan(st.get("total_cost")), "上面两项之和"],
            ["毛盈亏", fmt_yuan(st.get("total_gross"), signed=True), "持仓名义 × 次日涨跌，未扣费"],
            ["净盈亏", fmt_yuan(st.get("pnl"), signed=True), "毛盈亏 − 费用"],
        ],
        signed_cols={1},
    )
    caption(doc, "表 A-3  含补风格口径的费用拆解。主力连续合约没有真实换月移仓，换月成本未另计。")

    if not daily.empty:
        m = daily.copy()
        m["ym"] = pd.to_datetime(m["return_date"]).dt.strftime("%Y-%m")
        mg = m.groupby("ym").agg(pnl=("pnl_net", "sum"), cost=("cost", "sum"), n=("n", "mean"), eq=("equity", "last"))
        rows = []
        for ym, r in mg.iterrows():
            rows.append([
                ym,
                fmt_yuan(r["pnl"], signed=True),
                fmt_yuan(r["cost"]),
                fmt_num(r["n"], 1),
                fmt_yuan(r["eq"]),
            ])
        add_table(doc, ["月份", "净盈亏", "费用", "日均品种数", "月末权益"], rows, signed_cols={1})
        caption(doc, "表 A-4  含补风格口径的分月账户结果。")

    heading(doc, "2.2 持仓：期末与 2026-09-01", 2)
    para(
        doc,
        "持仓在信号日收盘后生成，下一交易日生效。下表是账户在样本最后一天、以及与风控页同一天（2026-09-01）的实际合约手数。"
        "方向已经按规则处理：加码与共识同向，控拥挤与共识反向。",
    )

    def hold_table(df: pd.DataFrame, title: str):
        if df is None or df.empty:
            para(doc, f"{title}：当日无持仓。", first_line=False)
            return
        rows = []
        for r in df.sort_values("notional", ascending=False).itertuples(index=False):
            rows.append([
                f"{r.name}({r.product})",
                r.sector,
                r.action,
                r.dir,
                str(int(r.lots)),
                f"{r.price:,.2f}" if pd.notna(r.price) else "—",
                fmt_yuan(r.notional),
                fmt_yuan(r.margin),
                fmt_yuan(r.pnl, signed=True),
            ])
        tot_n = df["notional"].sum()
        tot_m = df["margin"].sum()
        tot_p = df["pnl"].sum()
        rows.append(["合计", "", "", "", str(int(df["lots"].abs().sum())), "", fmt_yuan(tot_n), fmt_yuan(tot_m), fmt_yuan(tot_p, signed=True)])
        add_table(doc, ["品种", "板块", "信号", "方向", "手数", "价格", "名义", "保证金", "当日盈亏"], rows, signed_cols={8})
        caption(doc, title)

    if not holds.empty:
        last_d = holds["hold_date"].max()
        hold_table(holds[holds["hold_date"] == last_d], f"表 A-5  样本末日 {last_d} 持仓（含补风格）。")
        snap = holds[holds["hold_date"] == "2026-09-01"]
        if snap.empty:
            snap = holds[holds["signal_date"] == "2026-09-01"]
        if not snap.empty:
            hold_table(snap, "表 A-6  2026-09-01 持仓（与风控页同一天）。")
        elif last_d != "2026-09-01":
            para(doc, "2026-09-01 当天按规则没有可开的仓（只有观望/暂缓/减码，或价格缺失下不了单）。", first_line=False)

        heading(doc, "2.3 全程持仓统计", 2)
        hs = holds.groupby(["product", "name", "sector"]).agg(
            days=("hold_date", "nunique"),
            lots=("lots", lambda s: s.abs().mean()),
            notional=("notional", "mean"),
            pnl=("pnl", "sum"),
            long_days=("dir", lambda s: int((s == "多").sum())),
            short_days=("dir", lambda s: int((s == "空").sum())),
        ).sort_values("pnl", ascending=False)
        rows = []
        for (p, nme, sec), r in hs.iterrows():
            rows.append([
                f"{nme}({p})",
                sec,
                f"{int(r['days'])}",
                f"{int(r['long_days'])}/{int(r['short_days'])}",
                f"{r['lots']:.1f}",
                fmt_yuan(r["notional"]),
                fmt_yuan(r["pnl"], signed=True),
            ])
        add_table(doc, ["品种", "板块", "持仓天数", "多/空天数", "平均|手数|", "平均名义", "累计毛盈亏"], rows[:40], signed_cols={6})
        caption(doc, "表 A-7  全程出现过的持仓（按累计毛盈亏排序，最多 40 行）。多/空天数是持仓日计数。")

        if not trades.empty:
            heading(doc, "2.4 成交与换手", 2)
            t20 = trades.sort_values("notional", ascending=False).head(18)
            trows = []
            for r in t20.itertuples(index=False):
                trows.append([
                    r.trade_date,
                    f"{r.name}({r.product})",
                    r.action,
                    r.side,
                    f"{int(r.old_lots)}→{int(r.new_lots)}",
                    fmt_yuan(r.notional),
                    fmt_yuan(r.cost),
                ])
            add_table(doc, ["成交日", "品种", "信号", "动作", "手数", "成交名义", "费用"], trows)
            caption(doc, "表 A-8  名义最大的 18 笔调仓。费用=该笔手续费+滑点。")
            para(
                doc,
                f"全样本成交 {len(trades):,} 笔，成交名义合计 {fmt_yuan(trades['notional'].sum())}，"
                f"费用合计 {fmt_yuan(trades['cost'].sum())}，平均每笔 {fmt_yuan(trades['cost'].mean())}。"
                "名单日度变化大时，换手会明显高于「抱着共识品种不动」。",
            )

    para(
        doc,
        "这不是券商结算单：合约用主力连续价，没有指定月换月盈亏；保证金率是品种近似值而非当日交易所公布值；"
        "没有涨跌停、强平、资金划转。但手数、保证金上限、开平双边费用和滑点已经按可交易账户来做，"
        "权益数字是「若用 2,000 万跟信号」的数量级，而不是单位净值 overlay。",
    )


def build_report(ctx: dict) -> Path:
    configure_matplotlib()
    sig: pd.DataFrame = ctx["sig"]
    ev: pd.DataFrame = ctx["ev"]
    summary: pd.DataFrame = ctx["summary"]
    bt: dict = ctx["bt"]
    bt_vt = {k: vol_target(v) for k, v in bt.items()}
    charts = ctx["charts"]
    start, end = ctx["start"], ctx["end"]
    n_days = ctx["n_days"]

    core = bt.get("加码跟随", pd.DataFrame())
    core_vt = bt_vt.get("加码跟随", pd.DataFrame())
    st_core = perf_stats(core["ret"], core["nav"]) if not core.empty else {}
    st_vt = perf_stats(core_vt["ret_vt"], core_vt["nav_vt"]) if not core_vt.empty and "ret_vt" in core_vt.columns else {}
    add_row = summary[summary["action"] == "加码"]
    div_row = summary[summary["action"] == "观望"]
    mu1 = float(add_row["mu1"].iloc[0]) if not add_row.empty and pd.notna(add_row["mu1"].iloc[0]) else None
    hit1 = float(add_row["hit1"].iloc[0]) if not add_row.empty and pd.notna(add_row["hit1"].iloc[0]) else None
    t1 = float(add_row["t1"].iloc[0]) if not add_row.empty and pd.notna(add_row["t1"].iloc[0]) else None
    n_add = int(add_row["n"].iloc[0]) if not add_row.empty else 0

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    set_header_footer(doc)

    para(doc, "MOM 每日风控", size=12, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=4)
    para(doc, "量化 vs 主观  ·  决策信号策略回测报告", size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=8)
    para(doc, "If we trade the MOM decision signals, how do they perform?", size=11, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=14)
    para(
        doc,
        f"样本：{start} 至 {end}　　交易日 {n_days}　　产品级信号 {len(sig):,} 条　　模拟账户 {fmt_yuan(START_EQUITY)}　　手续费+滑点单边 {(COMM_RATE + SLIP_RATE) * 1e4:.1f}bp",
        size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=16,
    )

    para(
        doc,
        "本报告把风控页「量化 vs 主观」里的 MOM 决策信号，当成一套可交易的品种 overlay 来回测。"
        "信号本身是给 MOM 做资金分配用的（加码 / 暂缓 / 减码 / 观望 / 补风格 / 控拥挤），不是预测模型。"
        "问题是：如果在信号日收盘后，按信号方向在对应期货品种上开仓，次日及之后表现如何。"
        "实现与页面相同：风险% 阈值、1 日/5 日主动调仓、换月跳空剔除，均复刻 lib/ma/quant-vs-subjective-signals.ts。",
        space_after=10,
    )

    heading(doc, "一、结论先行", 1)
    ast = ctx.get("acct_stats") or {}
    ast0 = ctx.get("acct_core_stats") or {}
    if ast:
        para(
            doc,
            f"用 {fmt_yuan(START_EQUITY)} 跟信号做期货：加码与量化/主观同向、控拥挤反向、观望/暂缓/减码空仓、补风格跟重仓一侧。"
            f"扣掉单边手续费 {COMM_RATE * 1e4:.1f}bp 与滑点 {SLIP_RATE * 1e4:.1f}bp、手数取整并受保证金 50% 上限约束后，"
            f"期末权益 {fmt_yuan(ast.get('end'))}，净盈亏 {fmt_yuan(ast.get('pnl'), signed=True)}，"
            f"年化 {fmt_pct(ast.get('cagr'))}，夏普 {fmt_num(ast.get('sharpe'))}，最大回撤 {fmt_pct(ast.get('maxdd'))}。"
            f"费用合计 {fmt_yuan(ast.get('total_cost'))}。若不做补风格、只做加码同向+控拥挤反向，"
            f"期末 {fmt_yuan(ast0.get('end'))}，净盈亏 {fmt_yuan(ast0.get('pnl'), signed=True)}，夏普 {fmt_num(ast0.get('sharpe'))}。"
            "详细持仓、成交和分月盈亏见第二章。后面各章仍是单位净值 overlay，用来解释哪种信号有方向，不是账户结算。",
        )
    fade = bt.get("分歧跟主观", pd.DataFrame())
    fade_q = bt.get("分歧跟量化", pd.DataFrame())
    st_fade = perf_stats(fade["ret"], fade["nav"]) if not fade.empty else {}
    st_fade_q = perf_stats(fade_q["ret"], fade_q["nav"]) if not fade_q.empty else {}

    if mu1 is not None:
        lead = (
            f"「加码」是六种动作里方向最干净的一条：{n_add:,} 条产品级样本，方向对齐次日收益均值 {fmt_pct(mu1)}，"
            f"胜率 {fmt_pct(hit1, already=False) if hit1 is not None else '—'}，t={fmt_t(t1)}。"
            "t 在 1.6 附近，是弱但同号的证据，不是很强的预测。"
        )
    else:
        lead = "样本期内未能对「加码」信号估计出稳定的次日收益。"
    if st_core.get("sharpe") is not None:
        lead += (
            f" 只交易加码、做成等权多空 overlay（单边 {COST_BPS:.0f}bp 后），年化 {fmt_pct(st_core.get('cagr'))}、"
            f"波动 {fmt_pct(st_core.get('vol'))}、夏普 {fmt_num(st_core.get('sharpe'))}、最大回撤 {fmt_pct(st_core.get('maxdd'))}。"
        )
    if st_fade.get("sharpe") is not None:
        lead += (
            f" 样本内最好的可交易组合是「加码跟共识，观望日改跟主观」：年化 {fmt_pct(st_fade.get('cagr'))}、"
            f"夏普 {fmt_num(st_fade.get('sharpe'))}、回撤 {fmt_pct(st_fade.get('maxdd'))}。"
        )
        if st_fade_q.get("sharpe") is not None:
            lead += (
                f" 对称地，观望日改跟量化则年化 {fmt_pct(st_fade_q.get('cagr'))}、夏普 {fmt_num(st_fade_q.get('sharpe'))}。"
            )
        lead += " 这一年主观侧在分歧品种上占优，但不能外推成「永远站主观」——单看主观方向对照策略本身接近零。"
    para(doc, lead)

    para(
        doc,
        "动作必须拆开看，不能六种信号捆成一篮子。「暂缓加码 / 减码准备」存量仍同向，但主动调仓已经不再同加，"
        "次日对齐收益为负或接近零，支持页面把它们从「共识加码」KPI 里拆出去。"
        "「补风格」条数最多，次日对齐收益略负，说明它更像覆盖/引进投顾的提示，不适合用期货去复制重仓一侧。"
        "「完整映射」把暂缓、拥挤、补风格一并下单，夏普差于只做加码。"
        "观望本身没有方向；组合层面「跟主观」好、事件研究次日 t 只有约 0.7，所以更稳妥的用法仍是：分歧品种先不加 beta，而不是自动站队。",
    )

    kpi_rows = []
    for _, r in summary.iterrows():
        kpi_rows.append([
            r["action"],
            f"{int(r['n']):,}",
            fmt_pct(r["mu1"]),
            fmt_pct(r["hit1"], already=False) if r["hit1"] is not None else "—",
            fmt_t(r["t1"]),
            fmt_pct(r["mu5"]),
            fmt_pct(r["hit5"], already=False) if r["hit5"] is not None else "—",
            fmt_pct(r["mu20"]),
        ])
    add_table(
        doc,
        ["动作", "样本", "次日均值", "次日胜率", "t", "5日均值", "5日胜率", "20日均值"],
        kpi_rows,
        signed_cols={2, 3, 5, 6, 7},
    )
    caption(doc, "表 1  产品级信号的事件研究。收益按信号方向对齐（做多共识看多、做空共识看空）；观望无方向，次日均值留空，详见第五节。")

    add_account_chapter(doc, ctx)

    heading(doc, "三、单位净值 overlay 怎么做", 1)
    heading(doc, "3.1 数据与信号", 2)
    para(
        doc,
        f"持仓来自 public.mom_position_details（只读），量化账户为 {', '.join(QUANT_IDS)}，剔除国信/国盛及指定账户，期权合约不计入。"
        "品种收益用 raw_akshare_futures_daily 的主力连续合约日涨跌幅，换月跳空按与线上一致的 MAD 规则置零。"
        f"风险% = 品种净市值 × 近 {VOL_DAYS} 日波动 / 该风格全部品种 |净风险| 之和。波动只用信号日之前的收益，不含当日。"
        f"每日只保留 |量化风险%|+|主观风险%| 最大的 {TOP_N_PRODUCTS} 个品种，与页面名单构建方式一致；回测不截断「前 18 条」。",
    )
    para(
        doc,
        "信号阈值：两边同向且 |风险%| 都 ≥ 3% 为共识；都 ≥ 8% 为重仓共识；合计 ≥ 25% 改记控拥挤。"
        "反向且两边都 ≥ 3% 为观望。一侧 ≥ 8% 且另一侧 < 1.5% 为补风格。"
        "主动调仓 = Δ手数 × 当日价（不含涨跌），需大约 ≥ 100 万且 ≥ 仓位的 5%。"
        "存量同向但 1 日两边调仓反向 → 暂缓加码（若 5 日仍同加则维持加码）；1 日两边都在减 → 减码准备。",
    )
    heading(doc, "3.2 Overlay 交易规则", 2)
    para(
        doc,
        "时点：用 t 日收盘后的持仓与信号，赚取 t 之后第一个交易日的涨跌幅，避免用当日持仓解释当日行情。"
        "组合每日对有信号的品种等权，再把多空权重的绝对值之和缩放为 1（当日无仓则收益为 0）。"
        f"换手按 |Δw| 之和计，单边成本 {COST_BPS:.0f}bp。未做滑点、涨跌停和保证金约束。",
    )
    add_table(
        doc,
        ["策略", "如何把信号变成仓位"],
        [
            ["加码跟随", "只交易「加码」：共识做多+1、共识做空−1；其余动作空仓"],
            ["完整映射", "加码 1.0；暂缓加码 0.5；控拥挤 0.3；补风格 0.5（跟重仓一侧）；减码/观望 0"],
            ["加码+补风格", "加码跟共识，补风格跟重仓一侧，其余空仓"],
            ["分歧跟量化", "加码跟共识；观望日改跟量化方向"],
            ["分歧跟主观", "加码跟共识；观望日改跟主观方向"],
            ["量化方向", "对照：凡 |量化风险%|≥3% 即跟量化净方向（忽略主观）"],
            ["主观方向", "对照：凡 |主观风险%|≥3% 即跟主观净方向（忽略量化）"],
        ],
        col_widths=[3.4, 13.0],
    )
    caption(doc, "表 2  策略定义。前五条是「用信号交易」；后两条是风格仓位对照，用来判断信号是否比单看一侧更有用。")

    heading(doc, "四、Overlay 组合表现", 1)
    add_picture(doc, charts["equity"])
    caption(doc, "图 1  各策略成本后累计净值（等权 overlay，毛暴露=1）。")
    add_picture(doc, charts["equity_vt"])
    caption(doc, "图 2  同一套日收益，滚动 20 日波动目标 10% 年化后的净值，便于比较夏普。")

    perf_rows = []
    for name, df in bt.items():
        if df.empty:
            continue
        st = perf_stats(df["ret"], df["nav"])
        dfv = bt_vt[name]
        stv = perf_stats(dfv["ret_vt"], dfv["nav_vt"]) if "ret_vt" in dfv.columns else {}
        to = float(df["turnover"].mean()) if "turnover" in df.columns else None
        perf_rows.append([
            name,
            fmt_pct(st.get("cagr")),
            fmt_pct(st.get("vol")),
            fmt_num(st.get("sharpe")),
            fmt_pct(st.get("maxdd")),
            fmt_num(st.get("calmar")),
            fmt_pct(st.get("hit"), already=False),
            fmt_num(to, 2),
            fmt_num(stv.get("sharpe")),
        ])
    add_table(
        doc,
        ["策略", "年化", "波动", "夏普", "最大回撤", "Calmar", "日胜率", "日均换手", "10%目标夏普"],
        perf_rows,
        signed_cols={1, 3, 4, 5, 6, 8},
    )
    caption(doc, "表 3  全样本绩效。换手为每日 |Δ权重| 之和的均值；最后一列为波动目标后的夏普。")

    para(
        doc,
        "读表时注意三件事。第一，等权 overlay 的波动取决于当日有多少个信号、品种本身有多吵，所以波动目标后的夏普更适合横向比较。"
        "第二，「量化方向 / 主观方向」不是 MOM 实盘，只是把该侧净暴露当成交易信号。本样本里两条对照都接近零或为负，"
        "说明「永远跟一侧」没有优势；加码跟随好于单侧，是因为只在两边同时确认时才下单。"
        "第三，完整映射把暂缓、拥挤、补风格一并下单，夏普差于只做加码，和事件研究一致。"
        "第四，「分歧跟主观」夏普最高，来自加码日加上观望日站主观；这是样本内最优，但观望次日事件 t 不强，只宜当研究线索。",
    )

    add_picture(doc, charts["drawdown"])
    caption(doc, "图 3  加码跟随的回撤路径。")
    add_picture(doc, charts["monthly"])
    caption(doc, "图 4  加码跟随月度收益（红涨绿跌）。")
    add_picture(doc, charts["yearly"])
    caption(doc, "图 5  加码跟随分年年化与夏普。样本短的年份只作参考。")
    add_picture(doc, charts["turnover"])
    caption(doc, "图 6  加码跟随换手。信号名单日度变化大时，成本会明显侵蚀收益。")

    heading(doc, "五、事件研究：每种动作发出之后", 1)
    para(
        doc,
        "组合回测把同一天的多个品种捆在一起，会掩盖「哪种动作真正有方向」。"
        "事件研究改成：每条产品级信号单独看，把随后 1/5/10/20 个交易日的品种收益乘以信号方向。"
        "加码、暂缓、减码、控拥挤、补风格都有明确方向；观望没有，因此表 1 的对齐收益对观望为空白，下面单独拆重量化侧与主观侧。",
    )
    add_picture(doc, charts["event_bars"])
    caption(doc, "图 7  各动作在信号方向上的平均远期收益。")
    add_picture(doc, charts["hit"])
    caption(doc, "图 8  方向对齐后的胜率。")
    add_picture(doc, charts["signal_counts"])
    caption(doc, "图 9  各类信号出现频率随时间的变化（10 日均线）。")

    if not div_row.empty:
        r = div_row.iloc[0]
        para(
            doc,
            f"观望（方向分歧）共 {int(r['n']):,} 条。若在分歧日改跟量化，次日均值 {fmt_pct(r['mu1_quant'])}、"
            f"胜率 {fmt_pct(r['hit1_quant'], already=False) if r['hit1_quant'] is not None else '—'}、t={fmt_t(r['t1_quant'])}；"
            f"改跟主观则为 {fmt_pct(r['mu1_subj'])}、胜率 {fmt_pct(r['hit1_subj'], already=False) if r['hit1_subj'] is not None else '—'}、t={fmt_t(r['t1_subj'])}。"
            "两边数字恰好互为相反数（分歧日量化与主观方向相反）。次日 t 绝对值不到 1，事件研究单独看并不显著；"
            "组合回测里「跟主观」却明显好于「跟量化」，说明优势更多来自把分歧品种加进 overlay 之后的权重缩放和持有路径，而不是每条信号都稳赢。"
            "实盘更稳妥的读法仍是页面原意：分歧品种先不加该 beta，而不是自动站主观。",
        )

    kind_rows = []
    for kind, g in ev.groupby("kind"):
        s = g["signed_1"].dropna() if kind != "divergence" else g["quant_1"].dropna()
        if s.empty:
            continue
        se = s.std(ddof=1) / math.sqrt(len(s)) if len(s) > 1 else np.nan
        tstat = float(s.mean() / se) if se and se > 0 else None
        kind_rows.append([
            kind_label(kind),
            f"{len(g):,}",
            fmt_pct(float(s.mean())) if kind != "divergence" else fmt_pct(float(g["quant_1"].dropna().mean())) + "（跟量化）",
            fmt_pct(float((s > 0).mean()), already=False),
            fmt_t(tstat),
        ])
    if kind_rows:
        add_table(doc, ["信号类型", "样本", "次日对齐收益", "胜率", "t"], kind_rows, signed_cols={2, 3})
        caption(doc, "表 4  按信号类型（共识多/空、拥挤、仅一侧、分歧）的次日表现。分歧列的收益是「跟量化」。")

    heading(doc, "六、哪些板块和品种在贡献", 1)
    add_picture(doc, charts["sector"])
    caption(doc, "图 10  加码信号按板块的次日对齐收益（样本≥8）。")
    add_picture(doc, charts["product"])
    caption(doc, "图 11  加码信号样本≥12 的品种中，次日对齐收益最高的 12 个。")

    sec_rows = []
    g_add = ev[ev["action"] == "加码"].dropna(subset=["signed_1"])
    if not g_add.empty:
        sec = g_add.groupby("sector").agg(n=("signed_1", "count"), mu=("signed_1", "mean"), hit=("signed_1", lambda s: (s > 0).mean()))
        sec = sec.sort_values("n", ascending=False)
        for sec_name, r in sec.iterrows():
            se = g_add[g_add["sector"] == sec_name]["signed_1"].std(ddof=1) / math.sqrt(r["n"]) if r["n"] > 1 else np.nan
            tstat = float(r["mu"] / se) if se and se > 0 else None
            sec_rows.append([sec_name, f"{int(r['n']):,}", fmt_pct(float(r["mu"])), fmt_pct(float(r["hit"]), already=False), fmt_t(tstat)])
        add_table(doc, ["板块", "加码样本", "次日均值", "胜率", "t"], sec_rows, signed_cols={2, 3})
        caption(doc, "表 5  加码信号的板块分解。样本很少的板块 t 值不可靠。")

    prod_rows = []
    if not g_add.empty:
        pg = g_add.groupby(["product", "name"]).agg(n=("signed_1", "count"), mu=("signed_1", "mean"), hit=("signed_1", lambda s: (s > 0).mean()))
        pg = pg[pg["n"] >= 10].sort_values("mu", ascending=False)
        for (p, nme), r in list(pg.head(15).iterrows()) + list(pg.tail(8).iterrows() if len(pg) > 15 else []):
            prod_rows.append([f"{nme}({p})", get_sector(p), f"{int(r['n']):,}", fmt_pct(float(r["mu"])), fmt_pct(float(r["hit"]), already=False)])
        if prod_rows:
            add_table(doc, ["品种", "板块", "样本", "次日均值", "胜率"], prod_rows, signed_cols={3, 4})
            caption(doc, "表 6  加码信号样本≥10 的品种：头部 15 个与尾部若干，用于看集中度，而不是推荐交易名单。")

    heading(doc, "七、稳健性", 1)
    add_picture(doc, charts["cost"])
    caption(doc, "图 12  加码跟随对单边成本的敏感性。若 5–10bp 后夏普仍明显为正，结论对摩擦更稳健；若很快掉到零附近，则 overlay 更适合作为低换手的配置提示，而不是高频调仓。")

    # half sample
    if not core.empty:
        mid = core["return_date"].iloc[len(core) // 2]
        a = core[core["return_date"] <= mid]
        b = core[core["return_date"] > mid]
        sa, sb = perf_stats(a["ret"], (1 + a["ret"]).cumprod()), perf_stats(b["ret"], (1 + b["ret"]).cumprod())
        add_table(
            doc,
            ["子样本", "区间", "年化", "夏普", "最大回撤", "日胜率"],
            [
                ["前半", f"{a['return_date'].iloc[0]} ~ {a['return_date'].iloc[-1]}" if len(a) else "—", fmt_pct(sa.get("cagr")), fmt_num(sa.get("sharpe")), fmt_pct(sa.get("maxdd")), fmt_pct(sa.get("hit"), already=False)],
                ["后半", f"{b['return_date'].iloc[0]} ~ {b['return_date'].iloc[-1]}" if len(b) else "—", fmt_pct(sb.get("cagr")), fmt_num(sb.get("sharpe")), fmt_pct(sb.get("maxdd")), fmt_pct(sb.get("hit"), already=False)],
            ],
            signed_cols={2, 3, 4, 5},
        )
        caption(doc, "表 7  加码跟随前后半样本。若半段反号，则全样本结论不能外推。")

    para(
        doc,
        "5 日、10 日、20 日远期收益有重叠，t 值会被高估，只能当方向参考，不能当严格显著性。"
        "次日收益重叠少，是本报告最看重的一列。另外，信号用的是已实现持仓，属于「盘后看见两边仓位再交易」，"
        "不是开盘前可交易的预测；实盘若要跟，最早也是次日开盘，还会遇到隔夜跳空。",
    )

    heading(doc, "八、怎么用、不要怎么用", 1)
    para(
        doc,
        "建议把回测当成对信号设计的检验，而不是承诺一组期货策略的收益。页面上的加码，含义是「两种风格已经站在同一边，MOM 可以考虑增加该 beta」；"
        "回测显示这一条在品种收益上往往最干净。暂缓加码和减码准备被拆出来是对的：存量同向并不等于现在还该加。"
        "观望首先是风险预算红灯。本样本里「观望日站主观」组合最好，但事件研究不显著、且主观对照策略本身没有超额，"
        "因此不能写成交易规则。若要试验，应单独记账、限制名义、并准备在下一年度失效。"
        "补风格是覆盖问题（要不要引进另一侧投顾），用期货去「补」只会复制重仓一侧。"
        "控拥挤即使方向对，也不该按满仓 overlay 去加。",
    )
    para(
        doc,
        "2,000 万账户一章已经计入双边手续费、滑点、手数取整和保证金上限，但合约仍是主力连续、保证金率是品种近似值，没有换月与涨跌停。"
        "量化账户名单固定为当前七户；每日前 40 名品种会随截面变化。"
        "本报告只读 MOM 持仓与行情，不改 ETL，也不写入 mom_* 表。",
    )

    heading(doc, "附录  样本与计数", 1)
    act_counts = sig.groupby("action").size().reindex(ACTION_ORDER).dropna()
    add_table(
        doc,
        ["动作", "产品级信号条数", "占总信号%"],
        [[a, f"{int(n):,}", fmt_pct(n / len(sig), already=False)] for a, n in act_counts.items()],
    )
    caption(doc, "表 8  全样本信号条数。一条=某一交易日某一个品种触发该动作。")
    para(
        doc,
        f"持仓交易日 {ctx['pos_days']}，行情 overlapping 后用于回测的信号日 {n_days}。"
        f"报告生成于 {date.today().isoformat()}。图表原文件在 scripts/ma/_mom_signal_backtest_output/charts/。",
        size=10, color=MUTED, first_line=False,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))
    try:
        doc.save(str(REPORT_PATH_ASCII))
    except Exception:
        pass
    return REPORT_PATH


def main():
    load_env()
    configure_matplotlib()
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    conn = get_conn()
    try:
        pos, px = load_data(conn)
    finally:
        conn.close()
    if pos.empty:
        raise SystemExit("mom_position_details is empty")
    if px.empty:
        raise SystemExit("raw_akshare_futures_daily is empty")

    wide, clean, close, _ = build_returns(px)
    sig = build_signals(pos, clean)
    if sig.empty:
        raise SystemExit("no signals reconstructed")
    ev = event_study(sig, wide)
    summary = summarize_events(ev)
    schemes = ["加码跟随", "完整映射", "加码+补风格", "分歧跟量化", "分歧跟主观", "量化方向", "主观方向"]
    bt = portfolio_backtest(sig, wide, schemes, COST_BPS)
    bt_vt_core = vol_target(bt["加码跟随"]) if "加码跟随" in bt else pd.DataFrame()

    print("20M account backtest…")
    acct = run_account(sig, close, wide, clean, include_bufengge=True)
    acct_core = run_account(sig, close, wide, clean, include_bufengge=False)
    acct_stats_ = account_stats(acct["daily"])
    acct_core_stats = account_stats(acct_core["daily"])
    acct_charts = draw_account_charts(acct, CHART_DIR, fp, apply_font, save_fig)
    if not acct["daily"].empty:
        acct["daily"].to_csv(OUT_DIR / "acct_daily.csv", index=False, encoding="utf-8-sig")
    if not acct["holds"].empty:
        acct["holds"].to_csv(OUT_DIR / "acct_holdings.csv", index=False, encoding="utf-8-sig")
    if not acct["trades"].empty:
        acct["trades"].to_csv(OUT_DIR / "acct_trades.csv", index=False, encoding="utf-8-sig")
    acct_core["daily"].to_csv(OUT_DIR / "acct_core_daily.csv", index=False, encoding="utf-8-sig")

    print("Drawing charts…")
    charts = {
        "equity": draw_equity(bt, use_vt=False),
        "equity_vt": draw_equity({k: vol_target(v) for k, v in bt.items()}, use_vt=True),
        "drawdown": draw_dd({"加码跟随": bt_vt_core if not bt_vt_core.empty else bt.get("加码跟随", pd.DataFrame())}),
        "event_bars": draw_event_bars(summary),
        "hit": draw_hit(summary),
        "signal_counts": draw_signal_counts(sig),
        "monthly": draw_monthly(bt.get("加码跟随", pd.DataFrame()), "加码跟随：月度收益", "monthly.png"),
        "yearly": draw_split(bt.get("加码跟随", pd.DataFrame())),
        "sector": draw_sector_contrib(ev),
        "product": draw_product_table_chart(ev),
        "turnover": draw_turnover(bt.get("加码跟随", pd.DataFrame())),
        "cost": draw_cost_sensitivity(sig, wide),
    }

    sig.to_csv(OUT_DIR / "signals.csv", index=False, encoding="utf-8-sig")
    ev.to_csv(OUT_DIR / "events.csv", index=False, encoding="utf-8-sig")
    summary.to_csv(OUT_DIR / "event_summary.csv", index=False, encoding="utf-8-sig")
    for name, df in bt.items():
        safe = re.sub(r"[^\w]+", "_", name)
        df.to_csv(OUT_DIR / f"bt_{safe}.csv", index=False, encoding="utf-8-sig")

    ctx = {
        "sig": sig,
        "ev": ev,
        "summary": summary,
        "bt": bt,
        "charts": charts,
        "start": str(sig["date"].min()),
        "end": str(sig["date"].max()),
        "n_days": int(sig["date"].nunique()),
        "pos_days": int(pos["dt"].nunique()),
        "acct": acct,
        "acct_core": acct_core,
        "acct_stats": acct_stats_,
        "acct_core_stats": acct_core_stats,
        "acct_charts": acct_charts,
    }
    print("Writing Word report…")
    path = build_report(ctx)
    print(f"Wrote {path}")
    print(f"Also {REPORT_PATH_ASCII}")


if __name__ == "__main__":
    try:
        main()
    except Exception:
        traceback.print_exc()
        raise
