# -*- coding: utf-8 -*-
"""Generate Word report for running 在管产品 (as of 2026-08-17)."""
from __future__ import annotations

import json
import os
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties, fontManager
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
DATA_PATH = BASE / "_cta_lowvol_report_data.json"
OUT_DIR = BASE / "_cta_lowvol_report_output"
CHART_DIR = OUT_DIR / "charts"
REPORT_PATH = Path(__file__).resolve().parents[2] / "在管产品分析报告_20260817.docx"

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_NAVY = "#1A365D"
C_GOLD = "#C9A227"
C_RED = "#C53030"
C_GREEN = "#2F855A"
C_TEAL = "#2B6CB0"
C_ORANGE = "#DD6B20"
C_GRAY = "#718096"
PALETTE = ["#1A365D", "#C53030", "#2B6CB0", "#C9A227", "#2F855A", "#805AD5", "#DD6B20", "#319795", "#B83280", "#4A5568", "#9B2C2C"]

_CN_FONT: FontProperties | None = None


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.isfile(path):
            try:
                fontManager.addfont(path)
                _CN_FONT = FontProperties(fname=path)
                plt.rcParams["font.family"] = "sans-serif"
                plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "Microsoft YaHei", "SimHei"]
                return
            except Exception:
                continue
    _CN_FONT = FontProperties(family="Microsoft YaHei")


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_table(doc, headers, rows, col_widths=None, pct_cols=None):
    pct_cols = set(pct_cols or [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=9, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "F7FAFC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            color = TEXT
            if c_i in pct_cols and isinstance(val, str) and val not in ("—", "", "-"):
                if val.startswith("+") or (val.endswith("%") and not val.startswith("-") and val != "0.00%"):
                    if val.startswith("-"):
                        color = GREEN
                    elif val.startswith("+"):
                        color = RED
                elif val.startswith("-"):
                    color = GREEN
            cell_text(table.rows[r_i + 1].cells[c_i], val, size=8, color=color, align="center" if c_i else "left")
            shade(table.rows[r_i + 1].cells[c_i], fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def fmt_pct(v, digits=2, signed=True):
    if v is None:
        return "—"
    x = float(v) * 100 if abs(float(v)) <= 2 else float(v)
    sign = "+" if signed and x > 0 else ""
    return f"{sign}{x:.{digits}f}%"


def fmt_num(v, digits=4):
    if v is None:
        return "—"
    return f"{float(v):.{digits}f}"


def fmt_money(v):
    if v is None:
        return "—"
    n = float(v)
    if abs(n) >= 1e8:
        return f"{n / 1e8:.2f}亿"
    if abs(n) >= 1e4:
        return f"{n / 1e4:,.0f}万"
    return f"{n:,.0f}"


def fmt_yi(v):
    if v is None:
        return "—"
    return f"{float(v) / 1e8:.2f}"


def short_mgr(name: str | None) -> str:
    if not name:
        return "未登记"
    n = name.replace("私募证券基金管理有限公司", "").replace("私募基金管理有限公司", "")
    n = n.replace("资产管理有限公司", "").replace("有限公司", "")
    return n


BUCKETS = [
    "强势股/量化多头",
    "短周期/打板",
    "主观CTA/趋势",
    "期限结构/套利CTA",
    "商品/期现套利",
    "市场中性/对冲",
    "宏观/配置",
    "指数增强",
    "股票成长/价值",
    "现金/未识别",
]


def sleeve_of(p: dict) -> str:
    name = p.get("short_name") or ""
    l1 = p.get("company_strategy_l1") or ""
    if "恒盈2号" in name or (l1 == "期货策略"):
        return "主观多板块CTA"
    if "海宸" in name:
        return "期限结构CTA"
    if name.startswith("衡颐"):
        return "衡颐低波FOF"
    if name.startswith("荣熙"):
        return "荣熙FOF"
    if "抱朴" in name:
        return "抱朴多资产FOF"
    if name.startswith("金舆"):
        return "金舆FOF系列"
    if l1:
        return l1
    return "其他"


def classify_holding(name: str, asset_class: str | None = None, product_name: str = "") -> str:
    n = clean_name(name).replace(" ", "")
    if "银河期货" in n or "保证金" in n or "备付金" in n or "银行存款" in n:
        return "IGNORE"
    if "海宸" in product_name:
        return "期限结构/套利CTA"
    if "恒盈2号" in product_name:
        return "主观CTA/趋势"
    if re.search(r"豪鑫|滚雪粒", n):
        return "短周期/打板"
    if re.search(r"CTA|恒盈|九木|敦和芝诺|天戈|博衍九溪|汇融|浦江CTA|林健", n):
        return "主观CTA/趋势"
    if re.search(r"套利|如川|鑫安|稳博鹏瑞", n):
        return "商品/期现套利"
    if re.search(r"对冲|中性|鲲鹏", n):
        return "市场中性/对冲"
    if re.search(r"宏观|配置|宁苑|交睿", n):
        return "宏观/配置"
    if re.search(r"指增|指数增强|中证1000|添增2000", n):
        return "指数增强"
    if re.search(r"成长|价值|静瑞|铨景|墨雪", n):
        return "股票成长/价值"
    if re.search(
        r"九紫|木盛|瀛岳|青钱|棕榈滩|星阔|郁金香|务扬|古曲|立心|笃熙|贞元|常晋|铭跃|多璨|大风|守正|百奕|奇盾|英旷|熙典|留芳|国优|乾上|众量|澜熙|俊丹|添禄|峰云|准星|芬德|正源|金时信|博牛|坤望|天鹭|藤创|金麦穗|臻财|常行|自然红|臻选|图南|骐骥|桫罗",
        n,
    ):
        return "强势股/量化多头"
    ac = asset_class or ""
    if ac in ("商品期货", "国债期货", "股指期货"):
        return "主观CTA/趋势"
    return "其他"


def advisor_key(name: str) -> str:
    n = clean_name(name)
    n = re.sub(r"[ABC类]$", "", n)
    for prefix in ("六妙星", "棕榈滩", "瀛岳", "宁苑", "青钱", "荣熙恒盈", "荣熙如川", "交睿", "诚奇", "众量", "汇融", "博衍", "铨景", "墨雪", "锡和鑫安", "锡和骐骥", "澜熙", "俊丹"):
        if prefix in n:
            return prefix
    n = re.sub(r"[0-9一二三四五六七八九十]+号.*$", "", n)
    return n[:10]


def product_mix(p: dict) -> dict[str, float]:
    nav = float(p.get("net_asset_value") or 0)
    mix: dict[str, float] = defaultdict(float)
    pname = p.get("short_name") or ""
    holds = p.get("fof_holdings") or []
    if holds and nav > 0:
        for h in holds:
            b = classify_holding(h.get("name") or "", None, pname)
            if b == "IGNORE":
                continue
            mix[b] += abs(h.get("market_value") or 0)
        identified = sum(mix.values())
        mix["现金/未识别"] += max(0.0, nav - identified)
        return {k: mix.get(k, 0) / nav for k in BUCKETS}
    total = 0.0
    for h in p.get("valuation_holdings") or []:
        b = classify_holding(h.get("name") or "", h.get("asset_class"), pname)
        if b == "IGNORE":
            continue
        mv = abs(h.get("market_value") or 0)
        mix[b] += mv
        total += mv
    if total <= 0:
        return {k: 0.0 for k in BUCKETS}
    return {k: mix.get(k, 0) / total for k in BUCKETS}


def holding_weights(p: dict) -> dict[str, float]:
    nav = float(p.get("net_asset_value") or 1) or 1
    out: dict[str, float] = {}
    for h in p.get("fof_holdings") or []:
        key = advisor_key(h.get("name") or "")
        out[key] = out.get(key, 0) + abs(h.get("market_value") or 0) / nav
    return out


def overlap_matrix(products: list) -> list[list[float]]:
    sets = [holding_weights(p) for p in products]
    n = len(products)
    mat = [[0.0] * n for _ in range(n)]
    for i in range(n):
        for j in range(n):
            if i == j:
                mat[i][j] = 1.0
                continue
            common = set(sets[i]) & set(sets[j])
            mat[i][j] = sum(min(sets[i][k], sets[j][k]) for k in common)
    return mat


def role_of(p: dict) -> tuple[str, str]:
    """Infer mandate / goal from name + look-through mix."""
    name = p.get("short_name") or ""
    mix = product_mix(p)
    if "恒盈2号" in name:
        return "收益引擎", "主观多板块CTA，用商品+利率+股指赚取与股票低相关的弹性收益。"
    if "海宸" in name:
        return "低相关对冲仓", "商品期限结构/跨期套利CTA：毛敞口高、净方向低，用来压组合波动而非博趋势。"
    if "海泰" in name:
        return "低波权益旗舰", "强势股/量化多头FOF（周报口径），目标稳健权益beta，几乎不配CTA。"
    if "承和" in name:
        return "波动压制FOF", "权益约六成 + CTA约两成 + 对冲约一成半，把波动压到2%附近。"
    if "共赢" in name:
        return "内部再配置FOF", "把荣熙自有CTA与外部强势股/打板/配置打包，给不能直接持有期货的资金。"
    if "抱朴" in name or "祥和" in name:
        return "全天候多策略", "股票、打板、CTA、宏观、对冲、指增同时在篮子里，做长周期客户主账户。"
    if "基石" in name:
        return "金舆核心仓", "最分散的新发FOF，对冲+多头+CTA+指增，仍有约三分之一现金在建仓。"
    if "追风" in name:
        return "进攻权益卫星", "底层约七成强势股/短周期，与海泰同源投顾但更集中、更进攻。"
    if "守安" in name:
        return "防守增强", "权益多头为主，叠如川/鑫安套利与成长，名字即回撤约束。"
    if "锡泰" in name:
        return "宏观+套利（建仓中）", "交睿宏观与商品套利权重明显高于追风；约一半净值尚未投出。"
    if "稳健增长" in name:
        return "命名约束FOF", "成长约三分之一、强势股约三分之一、套利约两成，按名字做股债/套利平衡。"
    top = max(mix, key=lambda k: mix.get(k, 0))
    return "未命名桶", f"当前主导底层为{top}。"


def clean_name(name: str) -> str:
    n = name or ""
    n = re.sub(r"^场外[_／/].*?成本[．.]", "", n)
    n = n.replace("成本.", "").replace("私募证券投资基金", "").replace("私募投资基金", "")
    n = re.sub(r"^[场外已上市开放式私募成本．.]+", "", n)
    return n.strip() or name


def vol_bucket(ann_vol: float | None, n_obs: int) -> str:
    if ann_vol is None:
        return "样本不足"
    if n_obs < 30:
        return "样本较短"
    pct = ann_vol * 100
    if pct < 6:
        return "低波动"
    if pct < 12:
        return "中低波动"
    if pct < 20:
        return "中波动"
    return "中高波动"


def save_fig(name: str):
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DIR / name
    plt.tight_layout()
    plt.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def add_chart(doc, path: Path, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))


def caption(doc, text: str):
    para(doc, text, size=8, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


# ── charts ──────────────────────────────────────────────────────────────────

def chart_aum_bar(products):
    fig, ax = plt.subplots(figsize=(9.2, 4.6))
    names = [p["short_name"] for p in products]
    vals = [ (p["net_asset_value"] or 0) / 1e8 for p in products]
    colors = [PALETTE[i % len(PALETTE)] for i in range(len(names))]
    y = np.arange(len(names))
    ax.barh(y, vals, color=colors, height=0.62)
    ax.set_yticks(y)
    ax.set_yticklabels(names, **fp())
    ax.invert_yaxis()
    ax.set_xlabel("资产净值（亿元）", **fp())
    ax.set_title("在管产品资产净值分布", **fp())
    ax.grid(axis="x", linestyle=":", alpha=0.4)
    for i, v in enumerate(vals):
        ax.text(v + 0.02, i, f"{v:.2f}", va="center", fontsize=8, **fp())
    return save_fig("aum_bar.png")


def chart_mgr_pie(products):
    fig, ax = plt.subplots(figsize=(6.4, 4.4))
    buckets = defaultdict(float)
    for p in products:
        buckets[short_mgr(p.get("manager"))] += p.get("net_asset_value") or 0
    labels = list(buckets)
    sizes = [buckets[k] / 1e8 for k in labels]
    ax.pie(sizes, labels=[f"{k}\n{v:.2f}亿" for k, v in zip(labels, sizes)], colors=PALETTE[: len(labels)],
           startangle=90, textprops={"fontsize": 8, **fp()})
    ax.set_title("管理人规模占比", **fp())
    return save_fig("mgr_pie.png")


def chart_returns(products):
    fig, ax = plt.subplots(figsize=(9.4, 4.8))
    labels = [p["short_name"].replace("私募证券投资基金", "")[:10] for p in products]
    x = np.arange(len(labels))
    width = 0.18
    series = [("近一周", "ret_1w"), ("近一月", "ret_1m"), ("近三月", "ret_3m"), ("近六月", "ret_6m")]
    colors = [C_NAVY, C_TEAL, C_GOLD, C_ORANGE]
    for i, (lab, key) in enumerate(series):
        vals = [ (p.get(key) or 0) * 100 for p in products]
        ax.bar(x + (i - 1.5) * width, vals, width, label=lab, color=colors[i])
    ax.axhline(0, color="#CBD5E0", linewidth=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=28, ha="right", **fp())
    ax.set_ylabel("区间收益率（%）", **fp())
    ax.set_title("区间收益对比", **fp())
    ax.legend(prop=_CN_FONT, fontsize=8, frameon=False, ncol=4)
    ax.grid(axis="y", linestyle=":", alpha=0.4)
    return save_fig("returns.png")


def chart_risk_scatter(products):
    fig, ax = plt.subplots(figsize=(7.6, 4.8))
    for i, p in enumerate(products):
        r = p.get("risk") or {}
        vol = (r.get("ann_vol") or 0) * 100
        dd = abs(r.get("max_dd") or 0) * 100
        ret = (p.get("ret_3m") or 0) * 100
        ax.scatter(vol, dd, s=max(40, (p.get("net_asset_value") or 0) / 4e5), color=PALETTE[i % len(PALETTE)], alpha=0.85, label=p["short_name"][:8])
        ax.annotate(p["short_name"][:6], (vol, dd), textcoords="offset points", xytext=(4, 4), fontsize=7, **fp())
    ax.set_xlabel("年化波动率（%，净值序列）", **fp())
    ax.set_ylabel("最大回撤绝对值（%）", **fp())
    ax.set_title("波动—回撤散点（气泡大小=资产净值）", **fp())
    ax.grid(True, linestyle=":", alpha=0.4)
    return save_fig("risk_scatter.png")


def chart_nav_overlay(products):
    fig, ax = plt.subplots(figsize=(9.4, 4.8))
    for i, p in enumerate(products):
        pts = p.get("nav_chart") or []
        if len(pts) < 5:
            continue
        xs = [datetime.strptime(x["date"], "%Y-%m-%d") for x in pts]
        base = pts[0]["nav"] or 1
        ys = [x["nav"] / base for x in pts]
        ax.plot(xs, ys, color=PALETTE[i % len(PALETTE)], linewidth=1.3, label=p["short_name"][:8])
    ax.set_title("净值归一化走势（起点=1）", **fp())
    ax.set_ylabel("归一化净值", **fp())
    ax.legend(prop=_CN_FONT, fontsize=7, ncol=3, frameon=False, loc="upper left")
    ax.grid(True, linestyle=":", alpha=0.35)
    fig.autofmt_xdate()
    return save_fig("nav_overlay.png")


def chart_lookthrough(bucket_mv: dict):
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    items = sorted(bucket_mv.items(), key=lambda x: x[1], reverse=True)
    labels = [k for k, _ in items]
    vals = [v / 1e8 for _, v in items]
    ax.barh(labels, vals, color=PALETTE[: len(labels)])
    ax.invert_yaxis()
    ax.set_xlabel("底层持仓市值（亿元，未穿透重复）", **fp())
    ax.set_title("FOF底层策略桶市值", **fp())
    ax.grid(axis="x", linestyle=":", alpha=0.4)
    return save_fig("lookthrough.png")


def chart_hengying_fut(holdings):
    fig, ax = plt.subplots(figsize=(6.4, 4.2))
    buckets = defaultdict(float)
    for h in holdings:
        cls = h.get("asset_class") or "其他"
        buckets[cls] += abs(h.get("market_value") or 0)
    labels = list(buckets)
    sizes = [buckets[k] for k in labels]
    if not sizes:
        plt.close()
        return None
    ax.pie(sizes, labels=labels, colors=PALETTE[: len(labels)], startangle=90, textprops={"fontsize": 8, **fp()})
    ax.set_title("荣熙恒盈2号期货名义敞口（绝对值）", **fp())
    return save_fig("hengying_fut.png")


def chart_sleeve_aum(products):
    fig, ax = plt.subplots(figsize=(6.6, 4.2))
    buckets = defaultdict(float)
    for p in products:
        buckets[sleeve_of(p)] += p.get("net_asset_value") or 0
    labels = list(buckets)
    sizes = [buckets[k] / 1e8 for k in labels]
    ax.pie(sizes, labels=[f"{k}\n{v:.2f}亿" for k, v in zip(labels, sizes)], colors=PALETTE[: len(labels)],
           startangle=90, textprops={"fontsize": 8, **fp()})
    ax.set_title("策略桶规模", **fp())
    return save_fig("sleeve_pie.png")


def chart_strategy_stack(products):
    fig, ax = plt.subplots(figsize=(9.4, 5.0))
    labels = [p["short_name"][:8] for p in products]
    mixes = [product_mix(p) for p in products]
    keys = [k for k in BUCKETS if any(m.get(k, 0) > 0.02 for m in mixes)]
    x = np.arange(len(labels))
    bottom = np.zeros(len(labels))
    colors = PALETTE + ["#A0AEC0", "#E2E8F0"]
    for i, k in enumerate(keys):
        vals = np.array([m.get(k, 0) * 100 for m in mixes])
        ax.bar(x, vals, bottom=bottom, label=k, color=colors[i % len(colors)], width=0.72)
        bottom = bottom + vals
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=28, ha="right", **fp())
    ax.set_ylabel("占净值或占已识别名义敞口（%）", **fp())
    ax.set_title("分产品底层策略结构", **fp())
    ax.set_ylim(0, 100)
    ax.legend(prop=_CN_FONT, fontsize=7, ncol=3, frameon=False, loc="upper center", bbox_to_anchor=(0.5, -0.22))
    ax.grid(axis="y", linestyle=":", alpha=0.35)
    return save_fig("strategy_stack.png")


def chart_overlap(products):
    mat = overlap_matrix(products)
    n = len(products)
    fig, ax = plt.subplots(figsize=(8.4, 6.6))
    arr = np.array(mat) * 100
    np.fill_diagonal(arr, np.nan)
    im = ax.imshow(arr, cmap="YlOrRd", vmin=0, vmax=50)
    ax.set_xticks(range(n))
    ax.set_yticks(range(n))
    names = [p["short_name"][:8] for p in products]
    ax.set_xticklabels(names, rotation=40, ha="right", **fp())
    ax.set_yticklabels(names, **fp())
    for i in range(n):
        for j in range(n):
            if i == j:
                continue
            ax.text(j, i, f"{arr[i, j]:.0f}", ha="center", va="center", fontsize=7, color="#1A202C")
    ax.set_title("FOF投顾重叠（共同投顾权重取较小值，%）", **fp())
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    return save_fig("overlap.png")


# ── report ──────────────────────────────────────────────────────────────────

def build():
    configure_matplotlib()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    products = data["products"]
    kb_prod = data.get("kb_by_product") or {}
    kb_strat = data.get("strategy_kb") or []
    as_of = data.get("as_of") or "2026-08-17"
    total_nav = data.get("total_nav") or sum(p.get("net_asset_value") or 0 for p in products)

    look_names = defaultdict(float)
    look_bucket = defaultdict(float)
    look_count = defaultdict(int)
    for p in products:
        for h in p.get("fof_holdings") or []:
            nm = clean_name(h.get("name") or "")
            mv = h.get("market_value") or 0
            look_names[nm] += mv
            bkt = classify_holding(h.get("name") or "", None, p.get("short_name") or "")
            if bkt != "IGNORE":
                look_bucket[bkt] += mv
            look_count[nm] += 1

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # cover
    para(doc, "内部资料 · 请勿外传", size=9, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    para(doc, "在管产品分析报告", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, "投资分析 · 运行中全量组合", size=14, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, f"数据截止日  {as_of}", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"产品只数  {len(products)}    合计资产净值  {total_nav/1e8:.2f} 亿元", size=12, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(
        doc,
        "数据来源：PostgreSQL（ops_managed_products_list_cache、净值明细缓存、估值表/FOF底层、private_fund_info）及 AI 知识库 kb_chunks（内部/外部尽调资料）。",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    para(
        doc,
        "口径说明：列表收益与单位净值与投资分析「在管产品」界面一致；年化波动与最大回撤由明细净值序列计算（日收益×√252）。FOF底层权重按持仓市值/产品资产净值重算。",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )

    # 1 executive
    heading(doc, "一、摘要与结论", 1)
    para(
        doc,
        f"截至 {as_of}，在管运行中产品共 {len(products)} 只，合计资产净值 {total_nav/1e8:.2f} 亿元。"
        "这不是11只独立策略，而是四家管理人各自的「产品货架」：用不同底层配比，把同一套投顾池切成进攻、稳健、对冲、全天候等不同风险预算，再卖给/配置给不同资金。"
        "直投期货只有两只（荣熙恒盈2号主观趋势、衡颐海宸期限结构），其余九只是FOF。",
        size=11,
        space_after=8,
    )
    bullets = [
        "创设逻辑：同一管理人很少做重复产品。衡颐用海泰吃权益低波、承和用CTA+对冲把波动再压一档、海宸单独做跨期套利；荣熙用恒盈2号做引擎、共赢做「CTA+股票」包装；金舆五只是2026年中一次性上架的风险分层货架（基石/追风/守安/锡泰/稳健增长）。",
        "底层证据：海泰与追风的投顾重叠约46%（六妙星、棕榈滩、瀛岳、特夫），差别在浓度——海泰更均衡低波，追风把强势股提到约七成。守安/锡泰/稳健增长重叠20%–27%，共用铨景、如川、交睿，差别在现金比例与套利/宏观权重。",
        "组合后果：六妙星、棕榈滩、荣熙（恒盈/如川）被5–6只FOF同时持有。表面上11条净值曲线，穿透后是少数投顾的杠杆叠加。荣熙恒盈2号近三月-5.9%会同时打击共赢、抱朴、基石。",
        "波动分层：恒盈2号年化波动约21%是唯一中高波引擎；海宸虽是期货但净方向低，实现波动约8.6%；承和FOF仅2.4%。金舆新品低波动不可外推——锡泰约一半净值仍是现金。",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, b, size=11, color=TEXT)

    # 2 overview
    heading(doc, "二、组合概览", 1)
    heading(doc, "2.1 产品一览（与界面字段对齐）", 2)
    headers = ["产品", "备案号", "管理人", "策略", "净值日", "单位净值", "日涨跌", "资产净值", "近一周", "近一月", "近三月", "近六月"]
    rows = []
    for p in products:
        strat = " / ".join([x for x in [p.get("company_strategy_l1"), p.get("company_strategy_l2")] if x]) or sleeve_of(p)
        rows.append([
            p["short_name"],
            p.get("beian_hao") or "—",
            short_mgr(p.get("manager")),
            strat,
            (p.get("nav_date") or "—")[:10],
            fmt_num(p.get("unit_nav"), 4),
            fmt_pct(p.get("return_pct")),
            fmt_money(p.get("net_asset_value")),
            fmt_pct(p.get("ret_1w")),
            fmt_pct(p.get("ret_1m")),
            fmt_pct(p.get("ret_3m")),
            fmt_pct(p.get("ret_6m")),
        ])
    add_table(doc, headers, rows, pct_cols={6, 8, 9, 10, 11})
    caption(doc, "表1  在管产品列表指标。日涨跌与区间收益为小数口径×100后展示，红涨绿跌。")

    heading(doc, "2.2 规模结构", 2)
    add_chart(doc, chart_aum_bar(products))
    caption(doc, "图1  单只资产净值。金舆基石、荣熙恒盈2号、荣熙共赢、抱朴祥和构成规模第一梯队。")
    add_chart(doc, chart_mgr_pie(products), width=5.4)
    caption(doc, "图2  按管理人汇总。金舆系列合计约 2.02 亿元，为最大委托人；荣熙两只合计约 1.40 亿元。")
    add_chart(doc, chart_sleeve_aum(products), width=5.4)
    caption(doc, "图3  按策略桶划分。FOF 仍是规模主力；直投期货现为两只：主观趋势（恒盈2号）与期限结构（海宸）。")

    mgr_aum = defaultdict(float)
    for p in products:
        mgr_aum[short_mgr(p.get("manager"))] += p.get("net_asset_value") or 0
    mgr_rows = [[k, f"{len([p for p in products if short_mgr(p.get('manager'))==k])}只", fmt_money(v), f"{v/total_nav*100:.1f}%"] for k, v in sorted(mgr_aum.items(), key=lambda x: -x[1])]
    add_table(doc, ["管理人", "产品数", "资产净值", "占比"], mgr_rows)
    caption(doc, "表2  管理人集中度。前两家（金舆+荣熙）合计超过 78%。")

    # 3 performance & vol
    heading(doc, "三、收益与波动", 1)
    heading(doc, "3.1 区间收益", 2)
    para(
        doc,
        "近一周组合整体偏暖，11只中多数收涨；近一月则普遍走弱或接近持平，显示7月下旬至8月上旬的商品/权益脉冲对低波FOF与CTA均有扰动。"
        "近三月分化最大：衡颐海泰（强势股低波FOF）录得约 +4.0%，荣熙恒盈2号约 -5.9%，抱朴祥和约 -3.5%。",
        size=11,
    )
    add_chart(doc, chart_returns(products))
    caption(doc, "图4  近一周/一月/三月/六月收益。金舆新品成立不足两月，近三月与近六月基本重合，解读时需贴现历史长度。")

    heading(doc, "3.2 波动等级与回撤", 2)
    para(
        doc,
        "波动等级按净值序列年化波动划分：<6% 低波动，6%–12% 中低波动，12%–20% 中波动，≥20% 中高波动。"
        "金舆守安/锡泰/稳健增长成立仅数周，波动被压低，表中单独标注「样本较短」。",
        size=11,
    )
    risk_headers = ["产品", "成立/序列起点", "样本天数", "年化波动", "波动等级", "最大回撤", "夏普(序列)", "近一年夏普(列表)"]
    risk_rows = []
    for p in products:
        r = p.get("risk") or {}
        n = int(r.get("n") or 0)
        vol = r.get("ann_vol")
        risk_rows.append([
            p["short_name"],
            (p.get("inception_date") or r.get("start_date") or "—")[:10],
            str(n),
            fmt_pct(vol, signed=False),
            vol_bucket(vol, n),
            fmt_pct(r.get("max_dd")),
            fmt_num(r.get("sharpe"), 2),
            fmt_num(p.get("sharpe_1y"), 2),
        ])
    add_table(doc, risk_headers, risk_rows)
    caption(doc, "表3  风险指标。夏普无风险利率按 2% 计。列表夏普来自 ops_managed_products_list_cache（近一年窗口，样本不足则为空）。")
    add_chart(doc, chart_risk_scatter(products), width=5.8)
    caption(doc, "图5  波动—回撤。右上角为荣熙恒盈2号；左下角为金舆新发FOF。气泡越大规模越大。")
    add_chart(doc, chart_nav_overlay(products))
    caption(doc, "图6  归一化净值。抱朴祥和历史最长（2021年成立、缓存序列自2024-11）；荣熙恒盈2号弹性显著高于FOF群。")

    # 4 why these products
    heading(doc, "四、为什么这样排：目标、差异与底层证据", 1)
    para(
        doc,
        "团队策略字段几乎是空的，不能从标签读出意图。下面全部用估值表权重反推：每个产品在「强势股 / 打板 / 主观CTA / 期限结构 / 套利 / 对冲 / 宏观 / 指增 / 成长」上的真实配比，以及FOF之间共用了哪些投顾。"
        "推断对象有两层——（1）各管理人为什么要做这一系列产品；（2）在管账户为什么同时拿下这11只。",
        size=11,
    )
    heading(doc, "4.1 一句话角色表", 2)
    role_rows = []
    for p in products:
        role, goal = role_of(p)
        mix = product_mix(p)
        top3 = sorted(((k, v) for k, v in mix.items() if v >= 0.03), key=lambda x: -x[1])[:3]
        mix_s = "；".join(f"{k} {v*100:.0f}%" for k, v in top3) or "—"
        role_rows.append([p["short_name"], short_mgr(p.get("manager")), role, goal, mix_s])
    add_table(doc, ["产品", "管理人", "推断角色", "目标（由持仓反推）", "主导底层"], role_rows)
    caption(doc, "表A  产品角色不是宣传口径，而是看完底层百分比后的工作假设。金舆锡泰、基石的「现金/未识别」高，说明仍在建仓，角色会随投放变化。")

    heading(doc, "4.2 底层策略结构（百分比）", 2)
    short_b = {
        "强势股/量化多头": "强势股",
        "短周期/打板": "打板",
        "主观CTA/趋势": "趋势CTA",
        "期限结构/套利CTA": "期限CTA",
        "商品/期现套利": "套利",
        "市场中性/对冲": "对冲",
        "宏观/配置": "宏观",
        "指数增强": "指增",
        "股票成长/价值": "成长",
        "现金/未识别": "现金",
    }
    mix_headers = ["产品"] + [short_b[k] for k in BUCKETS]
    mix_rows = []
    for p in products:
        mix = product_mix(p)
        mix_rows.append(
            [p["short_name"][:10]]
            + [fmt_pct(mix.get(k, 0), 0, signed=False) if mix.get(k, 0) >= 0.015 else "—" for k in BUCKETS]
        )
    add_table(doc, mix_headers, mix_rows)
    caption(doc, "表B  FOF为占资产净值比；两只直投期货为占已识别期货名义敞口比。现金=净值−已识别底层市值（含建仓未完成与解析遗漏）。")
    add_chart(doc, chart_strategy_stack(products))
    caption(doc, "图A  堆叠后差异立刻可见：海泰/追风几乎是纯权益多头；承和明显多了CTA和对冲；稳健增长把成长与套利抬到与多头并列；锡泰/基石现金柱很高。")

    heading(doc, "4.3 投顾重叠：看起来11只，穿透后是一套池子", 2)
    para(
        doc,
        "定义重叠度为：两个FOF共同投顾上，取双方权重的较小值再加总。重叠46%意味着将近一半的风险预算压在同一批投顾上，净值相关性会被显著抬高。",
        size=11,
    )
    add_chart(doc, chart_overlap(products), width=5.8)
    caption(doc, "图B  热力越高越「不是独立产品」。海泰–追风、抱朴–共赢、共赢–追风是最红的三对；两只直投期货没有FOF代码重叠（它们是被FOF持有的底层，而不是对等FOF）。")
    para(
        doc,
        "重复投顾的含义：六妙星（九紫=强势股、豪鑫=打板）出现在海泰、抱朴、共赢、基石、追风、守安；棕榈滩出现在海泰、抱朴、共赢、基石、追风、锡泰；荣熙（恒盈2号或如川套利）出现在抱朴、共赢、守安、锡泰、稳健增长。"
        "在管组合等于给这三条投顾链加了杠杆。若打板拥挤或商品趋势反转，受伤的不会只是名字里带CTA的产品。",
        size=11,
    )

    heading(doc, "4.4 各管理人的产品货架——为什么要做成不一样的", 2)
    heading(doc, "衡颐：同一投研平台，切出三种风险预算", 3)
    para(
        doc,
        "海泰（约0.08亿）底层约四分之三是强势股/量化多头：六妙星九紫约19%、赢仕木盛约17%、棕榈滩约14%、瀛岳约13%、特夫郁金香约12%。几乎没有CTA。知识库与周报都把它写成「低波动·稳健运作·强势股」。这是衡颐给客户看的权益旗舰，目标是平滑的股票多头，不是期货账户。",
        size=11,
    )
    para(
        doc,
        "承和FOF（约0.12亿）把权益降到约62%（瀛岳单独就约26%），加入博衍九溪CTA约13% + 博衍留芳等，CTA合计约21%，再配金和和善/乾上泉等对冲约17%。结果是组合里波动最低的产品（年化约2.4%、回撤约-1.1%）。名字里的「和」就是相关性对冲：股票不好时让CTA/中性顶住。海泰负责赚钱弹性，承和负责睡得着。",
        size=11,
    )
    para(
        doc,
        "海宸（约0.10亿）不是FOF。估值表是银河期货账户，持仓以跨期配对为主：聚氯乙烯2609多 vs 2701空、纸浆2611多 vs 2609空、白糖701多 vs 705空，沥青、焦煤同样有反向月份。毛名义杠杆高（单腿可超过净值30%），但多空对冲后净方向小，所以实现波动约8.6%、回撤约-4.5%，远低于恒盈2号的21%/-17%。"
        "衡颐把它单独备案，是因为期限结构CTA的收益来源（展期、供需错配）与海泰的股票动量、承和的FOF配置都不一样，混在一个净值里会说不清归因。",
        size=11,
    )

    heading(doc, "荣熙：引擎 + 包装", 3)
    para(
        doc,
        "恒盈2号（约0.77亿）是货架上的发动机：团队策略明确写「期货/主观/多板块」。估值表同时出现豆粕、黄金、沪铜、棕榈油、焦煤、碳酸锂、国债（10年/30年）和IM股指，既有单边也有跨月对锁。目标是商品与利率上的主观相对价值+趋势，给组合提供与A股强势股低相关的收益。近一年夏普约1.9，但近三月-5.9%、历史回撤约-17%，弹性与回撤是同一枚硬币。",
        size=11,
    )
    para(
        doc,
        "共赢（约0.63亿）是把发动机装进FOF外壳。底层：六妙星九紫约20% + 豪鑫约6%（短周期），宁苑沛华约14%（配置），恒盈2号C类约11%，其余瀛岳、赢仕、古曲、务扬等强势股。股票多头约45%、打板约11%、宏观配置约14%、自家CTA约11%。"
        "为什么需要共赢：不少资金不能或不愿直接持有期货CTA；共赢用约一成净值去买恒盈2号，其余买股票投顾，波动降到约9.5%。代价是与海泰/追风/抱朴在六妙星上高度重叠（与海泰重叠约29%，与追风约35%，与抱朴约44%），「包装」并没有换一套投顾。",
        size=11,
    )

    heading(doc, "抱朴：长周期全天候主账户", 3)
    para(
        doc,
        "祥和一号（约0.62亿）是全书最老的产品（2021年成立，单位净值1.913）。登记为多资产/FOF。底层几乎把策略地图铺满：宁苑配置约16%、恒盈2号A约14%、六妙星豪鑫约9%+九紫相关、青钱约7%、静瑞价值约7%、棕榈滩约7%、百奕约6%，再加上九木/敦和/天戈等CTA合计约7%、锡和鑫安套利约3%、磐松指增约4%、嘉鸿对冲约3%。"
        "没有单一腿超过16%。这是「客户把钱放进来就不用再选策略」的全天候账户，目标是长期复利而不是风格极致。近一周+1.23%领涨、近三月-3.5%，涨时股票与CTA一起贡献，跌时两边也会一起吐——全天候不是低回撤保证。",
        size=11,
    )

    heading(doc, "金舆：一次上架五只，按名字做风险分层", 3)
    para(
        doc,
        "五只都在2026年6–7月成立，合计约2.02亿，是当前在管最大委托人。持仓高度同源（铨景、如川、交睿、众量、澜熙、棕榈滩、六妙星），差别主要是权重和现金，而不是投顾宇宙。这是标准的「一个投顾池、多条份额/多只产品」做法：用产品名告诉客户这只更进攻还是更防守，方便不同风险偏好下单，而不用真的建五套投研。",
        size=11,
    )
    para(
        doc,
        "基石（约0.75亿，最大只）：28条底层，对冲（诚奇约6.7%）+ 强势股（六妙星、棕榈滩、青钱、古曲）+ 宏观（交睿约2.8%）+ 指增 + 少量CTA。现金/未识别约34%——核心仓还在铺。名字「基石」= 金舆希望客户把这只当底仓，所以最分散、对冲最多。",
        size=11,
    )
    para(
        doc,
        "追风（约0.16亿）：强势股约73%（六妙星九紫一号+二号合计约22%，棕榈滩六号约9%+四号约6%，再加瀛岳、青钱、特夫、笃熙、俊丹、澜熙）。与海泰重叠约46%，是全书最高。名字「追风」= 追动量。目标不是低波，而是用更干净的强势股篮子放大海泰那套权益逻辑。规模小，适合当卫星。",
        size=11,
    )
    para(
        doc,
        "守安（约0.30亿）：强势股约49%，套利约10%（如川约6.7%），成长约10%（铨景约10%），现金约28%。「守」= 少用短周期打板、多用不吃方向的套利和成长均衡，回撤约束写进名字。实现波动约2%是建仓期假象，但配比方向确实比追风防守。",
        size=11,
    )
    para(
        doc,
        "锡泰（约0.52亿）：现金约50%，已投部分里交睿宏观约9.7%、如川约5.8%、澜熙约5.8%、铨景约5.8%、锡和鑫安约2.9%、棕榈滩约3.7%。宏观+商品套利的相对权重在金舆五只里最高，很像「锡和鑫安 + 海泰式权益」的缩写。规模已经不小但一半没出去，当前净值几乎不波动（0.65%）不能当风险特征。",
        size=11,
    )
    para(
        doc,
        "稳健增长FOF（约0.30亿）：墨雪约16.7% + 铨景约16.6%（成长合计约33%），强势股约33%，如川约10% + 鑫安约6.7%（套利约20%），汇融CTA约6.7%。这是五只里最「按名字施工」的：股票成长与动量对半，再用套利和CTA降回撤。目标收益来自权益，目标体验来自套利缓冲。",
        size=11,
    )

    heading(doc, "4.5 在管账户为什么同时拿这11只", 2)
    para(
        doc,
        "站在配置人而不是产品发行人的角度，这本在管账更像一个有意的 barbell，而不是11次互不相关的尽调结论。",
        size=11,
    )
    alloc = [
        "核心规模放在「可讲故事的FOF主账户」：金舆基石、荣熙共赢、抱朴祥和、荣熙恒盈2号四只就占了约2.77亿（64%）。前三只都是多策略包装，第四只是真正的CTA引擎。",
        "低波体验用衡颐三只小产品来做实验田：合计仅0.31亿。海泰验证强势股低波是否成立，承和验证「股票+CTA+对冲」能否把波动打到2%，海宸验证期限结构能否提供第三条不相关收益。小规模是因为这是策略实验室，不是目前的规模载体。",
        "金舆五只是同一委托人的风险分层下单：基石底仓、追风卫星、守安/稳健增长给要「名字里有稳健」的资金、锡泰给宏观套利偏好。在管把五只都接进来，等于承接金舆整条产品线，而不是精选一只。",
        "相关收益来源其实只有四条：①A股强势股/打板（六妙星、棕榈滩、瀛岳、青钱）；②主观商品CTA（恒盈2号）；③期限结构/套利（海宸、如川、鑫安）；④股票对冲与宏观（诚奇、交睿、宁苑）。11只净值是这四条链的不同配比，不是11条链。",
        "因此配置目标可以概括为：用FOF给客户提供可理解的产品名和波动分层；用两只直投期货获取不在股票里的收益；用小规模衡颐产品试新策略。代价是投顾拥挤——要管理的是六妙星/棕榈滩/恒盈2号的合计穿透，而不是产品只数。",
    ]
    for b in alloc:
        bp = doc.add_paragraph(style="List Bullet")
        add_text(bp, b, size=11, color=TEXT)

    heading(doc, "4.6 差异对照（同一问题的不同答案）", 2)
    diff_rows = [
        ["谁来赚股票的钱？", "海泰、追风、共赢、抱朴、基石", "追风最纯（约73%多头）；海泰更均衡；共赢/抱朴被CTA稀释"],
        ["谁来压波动？", "承和、守安、稳健增长、海宸", "承和用对冲+CTA；稳健增长用套利；海宸用跨期对锁；守安用现金+套利"],
        ["谁来提供商品收益？", "恒盈2号、海宸、如川/鑫安（嵌在FOF里）", "恒盈=主观多板块；海宸=期限结构；如川/鑫安=套利，方向暴露更低"],
        ["谁是客户主账户？", "抱朴祥和、金舆基石、荣熙共赢", "持仓条数多、单一投顾<16%、故事好讲"],
        ["谁还在建仓？", "锡泰、基石、守安", "现金34%–50%，当前波动不能当产品风险标签"],
        ["谁和其他人最像？", "追风 ≈ 海泰的进攻版", "重叠约46%，差别是浓度不是投顾名单"],
        ["谁和其他人最不像？", "恒盈2号、海宸", "直投期货，FOF重叠矩阵为0；二者彼此也不一样（趋势 vs 跨期）"],
    ]
    add_table(doc, ["配置问题", "主要承担者", "差异要点"], diff_rows)
    caption(doc, "表C  用「问题—产品」而不是「产品—介绍」来看货架，才能看出为什么不能把11只等权当分散。")

    heading(doc, "4.7 恒盈2号期货结构（引擎细节）", 2)
    hy = next(p for p in products if p["beian_hao"] == "SBAH99")
    para(
        doc,
        f"荣熙恒盈2号成立于 {hy.get('inception_date')}，最新单位净值 {fmt_num(hy.get('unit_nav'))}（{hy.get('nav_date')}），"
        f"资产净值 {fmt_money(hy.get('net_asset_value'))}。"
        f"年化波动约 {fmt_pct(hy['risk']['ann_vol'], signed=False)}，最大回撤约 {fmt_pct(hy['risk']['max_dd'])}，近一年列表夏普 {fmt_num(hy.get('sharpe_1y'), 2)}。"
        "它同时被抱朴、共赢、基石持有，所以既是在管直投，也是FOF底层——引擎故障会传导三次。",
        size=11,
    )
    fut_path = chart_hengying_fut(hy.get("valuation_holdings") or [])
    if fut_path:
        add_chart(doc, fut_path, width=5.2)
        caption(doc, "图7  荣熙恒盈2号期货名义敞口绝对值（已剔除保证金/备付金）。商品为主，国债与股指为卫星。")
    fut_rows = []
    for h in (hy.get("valuation_holdings") or [])[:12]:
        fut_rows.append([h.get("name"), h.get("asset_class") or "—", fmt_money(h.get("market_value"))])
    if fut_rows:
        add_table(doc, ["合约/科目", "资产类别", "市值（名义）"], fut_rows)
        caption(doc, "表4  主要合约。豆粕、锰硅等跨月多空并存，说明不是纯单边趋势，含期限结构对冲。")

    if kb_strat:
        para(doc, "知识库摘录（定性印证，不替代净值）：", size=11, bold=True)
        for item in kb_strat[:6]:
            src = (item.get("source") or "").replace("\\", "/")
            win = (item.get("windows") or ["（无摘录）"])[0]
            kp = doc.add_paragraph(style="List Bullet")
            add_text(kp, f"{src.split('/')[-1][:48]}：", size=9, bold=True, color=NAVY)
            add_text(kp, win[:220], size=9, color=TEXT)

    # 5 look through
    heading(doc, "五、底层资产与穿透", 1)
    para(
        doc,
        "以下按 FOF 估值表底层汇总。同一投顾产品若被多只在管FOF同时持有，市值为简单加总（未按在管产品之间的交叉持股去重），用于揭示拥挤度。",
        size=11,
    )
    add_chart(doc, chart_lookthrough(look_bucket), width=5.6)
    caption(doc, "图8  底层策略桶。股票主观与短周期权重大，CTA/商品与市场中性构成风险分散层。")

    top_u = sorted(look_names.items(), key=lambda x: -x[1])[:18]
    u_rows = []
    for name, mv in top_u:
        u_rows.append([name[:22], look_count[name], fmt_money(mv), f"{mv/total_nav*100:.1f}%"])
    add_table(doc, ["底层产品（清洗后）", "出现只数", "合计市值", "占在管总净值"], u_rows)
    caption(doc, "表5  穿透后最大底层。出现只数≥3 的标的（如六妙星、棕榈滩、瀛岳、如川套利）构成组合拥挤核心。")

    heading(doc, "5.1 分产品底层（前8）", 2)
    for p in products:
        holds = p.get("fof_holdings") or []
        if not holds:
            continue
        heading(doc, f"{p['short_name']}（{p.get('beian_hao')}）", 3)
        hrows = []
        for h in holds[:8]:
            hrows.append([
                clean_name(h.get("name") or "")[:24],
                h.get("code") or "—",
                classify_holding(h.get("name") or "", None, p.get("short_name") or ""),
                fmt_money(h.get("market_value")),
                fmt_pct(h.get("market_weight"), signed=False),
            ])
        add_table(doc, ["底层", "代码", "策略桶", "市值", "权重"], hrows)
        date = holds[0].get("valuation_date") or ""
        caption(doc, f"估值日 {date}。权重=市值/产品资产净值。")

    # 6 product cards
    heading(doc, "六、分产品要点", 1)
    for p in products:
        heading(doc, f"{p['short_name']}  {p.get('beian_hao')}", 2)
        r = p.get("risk") or {}
        lines = [
            f"管理人：{p.get('manager') or '—'}；成立日：{(p.get('inception_date') or '—')[:10]}。",
            f"团队策略：{p.get('company_strategy_l1') or '未配置'} / {p.get('company_strategy_l2') or '—'}；平台策略：{p.get('platform_strategy_l1') or '—'} / {p.get('platform_strategy_l2') or '—'}。",
            f"最新净值 {fmt_num(p.get('unit_nav'))}（{(p.get('nav_date') or '')[:10]}），日涨跌 {fmt_pct(p.get('return_pct'))}，资产净值 {fmt_money(p.get('net_asset_value'))}，托管账户余额 {fmt_money(p.get('custody_balance'))}。",
            f"区间收益：1周 {fmt_pct(p.get('ret_1w'))}，1月 {fmt_pct(p.get('ret_1m'))}，3月 {fmt_pct(p.get('ret_3m'))}，6月 {fmt_pct(p.get('ret_6m'))}。",
            f"风险：年化波动 {fmt_pct(r.get('ann_vol'), signed=False)}（{vol_bucket(r.get('ann_vol'), int(r.get('n') or 0))}），最大回撤 {fmt_pct(r.get('max_dd'))}，序列夏普 {fmt_num(r.get('sharpe'), 2)}。",
        ]
        for line in lines:
            para(doc, line, size=11, space_after=4)

        role, goal = role_of(p)
        para(doc, f"推断角色：{role}。{goal}", size=11)
        mix = product_mix(p)
        top = sorted(((k, v) for k, v in mix.items() if v >= 0.04), key=lambda x: -x[1])[:4]
        if top:
            para(doc, "底层结构：" + "；".join(f"{k} {v*100:.0f}%" for k, v in top) + "。", size=11)

        hits = kb_prod.get(p.get("beian_hao") or "", [])
        if hits:
            para(doc, "知识库摘录：", size=10, bold=True, color=NAVY, space_after=2)
            for hit in hits[:2]:
                src = (hit.get("source") or "").replace("\\", "/").split("/")[-1][:56]
                win = (hit.get("windows") or [""])[0][:240]
                para(doc, f"{src} — {win}", size=9, color=TEXT, space_after=4)

    # 7 risks
    heading(doc, "七、风险与后续跟踪", 1)
    risks = [
        "策略拥挤：六妙星、棕榈滩、瀛岳、荣熙恒盈2号、如川套利被多只FOF同时持有。海泰与追风重叠约46%，抱朴与共赢约44%。应按投顾合并穿透上限，而不是按产品只数看分散。",
        "波动误读：金舆锡泰现金约50%、基石约34%、守安约28%仍在建仓。当前0.2%–2%的实现波动不是产品风险特征，应用投顾历史波动做预算。",
        "引擎传导：荣熙恒盈2号近三月 -5.9%、历史最大回撤约 -17%，且被抱朴、共赢、基石重复持有，是组合尾部风险的主要放大器。海宸是另一类期货（跨期），不要和恒盈2号当成同一只CTA。",
        "货架重复：金舆五只共享投顾池，差异主要是权重。配置上应把五只合成一个金舆风险预算，再决定总规模，而不是五次独立加仓。",
        "数据缺口：多数产品团队策略/标签未写入缓存；九紫等底层存在份额类别映射（SBPC20 / ABCX2）。海宸已确认为期货账户而非FOF解析错误。",
    ]
    for b in risks:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, b, size=11)

    heading(doc, "八、附录：方法与来源", 1)
    para(
        doc,
        "1）产品全集取自 managed_products ⋈ ops_managed_products_list_cache，运行中定义为资产净值为空或大于 0，共 11 只，与界面「共 11 条」一致。"
        "2）区间收益、单位净值、托管余额来自列表缓存（与邮件净值/估值表夜间ETL一致）。"
        "3）年化波动、最大回撤、序列夏普来自 ops_private_fund_detail_nav_cache 净值序列。"
        "4）FOF底层来自 ops_managed_fof_underlying；荣熙恒盈2号与衡颐海宸改用 ops_email_valuation_holdings（直投期货）。"
        "5）策略桶由底层名称规则映射（九紫/棕榈滩/瀛岳→强势股，豪鑫→打板，恒盈/九木/汇融→CTA，如川/鑫安→套利，诚奇→对冲，交睿/宁苑→宏观）。重叠度=共同投顾权重取min后加总。"
        "6）产品角色与配置意图为基于持仓百分比的推断，不是管理人官方说明书。"
        "7）定性描述另检索 kb_chunks。本报告为投研内部分析，不构成业绩承诺。",
        size=10,
        color=TEXT,
    )
    para(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}（本地）", size=9, color=MUTED, space_before=12)

    doc.save(str(REPORT_PATH))
    ascii_copy = Path(__file__).resolve().parents[2] / "managed_products_report_20260817.docx"
    ascii_copy.write_bytes(REPORT_PATH.read_bytes())
    print(f"wrote {REPORT_PATH}")
    print(f"wrote {ascii_copy}")
    return str(REPORT_PATH)


if __name__ == "__main__":
    build()
