# -*- coding: utf-8 -*-
"""Offender-style behavioral profile of MOM account RX319 (余浩 / 九木)."""
from __future__ import annotations

import json
import math
import os
import re
import sys
import warnings
from collections import defaultdict
from datetime import date, datetime, timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.font_manager import FontProperties, fontManager

from mom_data_etl import get_conn, load_env_files

load_env_files()

BASE = Path(__file__).resolve().parent
ROOT = BASE.parents[1]
OUT_DIR = BASE / "_rx319_profile_output"
CHART_DIR = OUT_DIR / "charts"
DATA_PATH = OUT_DIR / "rx319_profile_data.json"
REPORT_PATH = ROOT / "RX319_余浩_交易员画像报告_20260819.docx"
REPORT_PATH_ASCII = ROOT / "RX319_trader_profile_20260819.docx"

ACCOUNT = "rx319"
AS_OF = date.today().isoformat()

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_NAVY = "#1A365D"
C_GOLD = "#C9A227"
C_RED = "#C53030"
C_GREEN = "#2F855A"
C_TEAL = "#2B6CB0"
C_ORANGE = "#DD6B20"
C_PURPLE = "#805AD5"
C_GRAY = "#718096"

SECTOR_MAP: dict[str, str] = {
    "C": "农产", "CS": "农产", "WH": "农产", "PM": "农产", "RR": "农产", "RI": "农产",
    "JR": "农产", "LR": "农产", "A": "农产", "B": "农产", "M": "农产", "Y": "农产",
    "RM": "农产", "OI": "农产", "RS": "农产", "PK": "农产", "P": "农产", "SR": "农产",
    "CF": "农产", "CY": "农产", "LG": "农产", "SP": "农产", "OP": "农产",
    "AP": "生鲜", "CJ": "生鲜", "LH": "生鲜", "JD": "生鲜",
    "AU": "贵金属", "AG": "贵金属", "PT": "贵金属", "PD": "贵金属",
    "CU": "有色", "BC": "有色", "AL": "有色", "AO": "有色", "AD": "有色",
    "ZN": "有色", "PB": "有色", "NI": "有色", "SN": "有色",
    "LC": "新能源", "PS": "新能源", "SI": "新能源",
    "I": "黑色", "SF": "黑色", "SM": "黑色", "RB": "黑色", "HC": "黑色",
    "SS": "黑色", "WR": "黑色", "JM": "黑色", "J": "黑色", "ZC": "黑色",
    "FG": "黑色", "BB": "黑色", "FB": "黑色",
    "SC": "能源化工", "FU": "能源化工", "LU": "能源化工", "PG": "能源化工",
    "BU": "能源化工", "TA": "能源化工", "EG": "能源化工", "PF": "能源化工",
    "PR": "能源化工", "PL": "能源化工", "PP": "能源化工", "L": "能源化工",
    "BZ": "能源化工", "PX": "能源化工", "EB": "能源化工", "RU": "能源化工",
    "BR": "能源化工", "NR": "能源化工", "SA": "能源化工", "SH": "能源化工",
    "V": "能源化工", "UR": "能源化工", "MA": "能源化工",
    "EC": "航运",
    "IH": "股指", "IF": "股指", "IC": "股指", "IM": "股指", "MO": "股指",
    "TS": "国债", "TF": "国债", "T": "国债", "TL": "国债",
}

PRODUCT_CN: dict[str, str] = {
    "A": "豆一", "B": "豆二", "M": "豆粕", "Y": "豆油", "P": "棕榈油", "RM": "菜粕",
    "OI": "菜油", "PK": "花生", "C": "玉米", "CS": "淀粉", "SR": "白糖", "CF": "棉花",
    "SP": "纸浆", "AP": "苹果", "CJ": "红枣", "LH": "生猪", "JD": "鸡蛋", "LG": "原木",
    "PT": "铂金", "PD": "钯金",
    "CU": "沪铜", "AL": "沪铝", "ZN": "沪锌", "PB": "沪铅", "NI": "沪镍", "SN": "沪锡",
    "AO": "氧化铝", "AD": "铝合金", "BC": "国际铜",
    "LC": "碳酸锂", "SI": "工业硅", "PS": "多晶硅",
    "I": "铁矿", "RB": "螺纹", "HC": "热卷", "SS": "不锈钢", "JM": "焦煤", "J": "焦炭",
    "SF": "硅铁", "SM": "锰硅", "FG": "玻璃", "ZC": "动力煤",
    "SC": "原油", "FU": "燃油", "LU": "低硫燃油", "PG": "LPG", "BU": "沥青",
    "TA": "PTA", "EG": "乙二醇", "PF": "短纤", "PR": "瓶片", "PP": "聚丙烯",
    "L": "塑料", "V": "PVC", "EB": "苯乙烯", "PX": "PX", "BZ": "纯苯",
    "MA": "甲醇", "SA": "纯碱", "SH": "烧碱", "UR": "尿素",
    "RU": "橡胶", "NR": "20号胶", "BR": "丁二烯胶", "EC": "欧线集运",
    "IH": "上证50", "IF": "沪深300", "IC": "中证500", "IM": "中证1000",
    "TS": "2年国债", "TF": "5年国债", "T": "10年国债", "TL": "30年国债",
}

CFFEX = {"IH", "IF", "IC", "IM", "MO", "TS", "TF", "T", "TL"}
_CN_FONT: FontProperties | None = None


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    candidates = [
        os.environ.get("FOF_REPORT_FONT_PATH") or "",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansSC-Regular.otf",
        "/root/new_market_project/haitai_week_report/fonts/NotoSansSC-Regular.otf",
        str(ROOT / "haitai_week_report" / "fonts" / "NotoSansSC-Regular.otf"),
    ]
    seen: set[str] = set()
    for raw in candidates:
        if not raw:
            continue
        path = os.path.realpath(raw)
        if path in seen or not os.path.isfile(path):
            continue
        seen.add(path)
        try:
            fontManager.addfont(path)
        except Exception:
            pass
        try:
            font = FontProperties(fname=path)
            name = font.get_name()
        except Exception:
            continue
        _CN_FONT = font
        plt.rcParams["font.family"] = "sans-serif"
        plt.rcParams["font.sans-serif"] = [name, "Noto Sans CJK SC", "Microsoft YaHei", "SimHei", "DejaVu Sans"]
        return
    _CN_FONT = None


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def to_num(v) -> float:
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return 0.0
    if isinstance(v, (int, float, np.integer, np.floating)):
        return float(v)
    s = str(v).replace(",", "").replace("%", "").replace(" ", "").strip()
    if not s:
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def product_code(contract: str | None) -> str:
    if not contract:
        return ""
    m = re.match(r"^[A-Za-z]+", str(contract).strip())
    return m.group(0).upper() if m else ""


def pname(code: str) -> str:
    cn = PRODUCT_CN.get(code, "")
    return f"{code} {cn}".strip() if cn else code


def sector_of(code: str) -> str:
    return SECTOR_MAP.get(code, "其他")


def parse_open_date(v) -> date | None:
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return None
    if isinstance(v, datetime):
        return v.date()
    if isinstance(v, date):
        return v
    s = str(v).strip()
    if not s:
        return None
    try:
        n = float(s)
        if 19_000_000 < n < 21_000_000:
            ymd = f"{int(round(n)):08d}"
            return datetime.strptime(ymd, "%Y%m%d").date()
        if 30000 < n < 80000:
            return (datetime(1899, 12, 30) + timedelta(days=n)).date()
    except ValueError:
        pass
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y%m%d"):
        try:
            return datetime.strptime(s[:10] if len(s) >= 8 else s, fmt).date()
        except ValueError:
            continue
    return None


def parse_hour(v) -> int | None:
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return None
    s = str(v).strip()
    if not s:
        return None
    m = re.search(r"(\d{1,2}):(\d{2})", s)
    if m:
        return int(m.group(1)) % 24
    if re.fullmatch(r"\d{6}", s):
        return int(s[:2]) % 24
    return None


def parse_hms(v) -> tuple[int, int] | None:
    if v is None:
        return None
    s = str(v).strip()
    m = re.search(r"(\d{1,2}):(\d{2})", s)
    if m:
        return int(m.group(1)) % 24, int(m.group(2))
    if re.fullmatch(r"\d{6}", s):
        return int(s[:2]) % 24, int(s[2:4])
    return None


def session_of_hour(h: int | None) -> str:
    if h is None:
        return "未知"
    if h >= 21 or h < 3:
        return "夜盘"
    if 8 <= h <= 15:
        return "日盘"
    return "其他"


def num_sql(col: str) -> str:
    return (
        f"COALESCE(NULLIF(REPLACE(REPLACE(COALESCE(\"{col}\"::text, ''), ',', ''), ' ', ''), '')::numeric, 0)"
    )


def read_sql(conn, sql: str, params=None) -> pd.DataFrame:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        return pd.read_sql(sql, conn, params=params)


def max_drawdown(cum: np.ndarray) -> float:
    if len(cum) == 0:
        return 0.0
    peak = np.maximum.accumulate(cum)
    dd = (cum - peak) / np.clip(peak, 1e-12, None)
    return float(dd.min())


def safe_div(a, b, default=0.0) -> float:
    if b is None or abs(float(b)) < 1e-12:
        return default
    return float(a) / float(b)


def median(xs) -> float:
    arr = np.array(list(xs), dtype=float)
    arr = arr[np.isfinite(arr)]
    return float(np.median(arr)) if len(arr) else 0.0


def mean(xs) -> float:
    arr = np.array(list(xs), dtype=float)
    arr = arr[np.isfinite(arr)]
    return float(np.mean(arr)) if len(arr) else 0.0


def pctl(xs, q) -> float:
    arr = np.array(list(xs), dtype=float)
    arr = arr[np.isfinite(arr)]
    return float(np.percentile(arr, q)) if len(arr) else 0.0


# ── Word helpers ────────────────────────────────────────────────────────────

def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.28
    p.paragraph_format.first_line_indent = Pt(0)
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def caption(doc, text):
    para(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, space_before=2)


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(el if False else tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_table(doc, headers, rows, highlight_rows=None, red_cols=None, green_cols=None):
    highlight_rows = set(highlight_rows or [])
    red_cols = set(red_cols or [])
    green_cols = set(green_cols or [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=8, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "EDF2F7" if r_i in highlight_rows else ("F7FAFC" if r_i % 2 == 0 else "FFFFFF")
        for c_i, val in enumerate(row):
            color = None
            s = "" if val is None else str(val)
            if c_i in red_cols or (s.startswith("-") and any(ch.isdigit() for ch in s)):
                if s.startswith("-") and any(ch.isdigit() for ch in s):
                    color = RED
            if color is None and s.startswith("+"):
                color = GREEN
            cell_text(table.rows[r_i + 1].cells[c_i], val, size=8, color=color, align="center" if c_i else "left")
            shade(table.rows[r_i + 1].cells[c_i], fill)
    return table


def add_chart(doc, path: Path | None, width=6.5):
    if path and path.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))


def fmt_pct(v, digits=1, signed=False):
    if v is None:
        return "—"
    x = float(v) * 100
    sign = "+" if signed and x > 0 else ""
    return f"{sign}{x:.{digits}f}%"


def fmt_num(v, digits=1):
    if v is None:
        return "—"
    return f"{float(v):,.{digits}f}"


def fmt_int(v):
    if v is None:
        return "—"
    return f"{int(round(float(v))):,}"


def fmt_money(v, signed=False):
    if v is None:
        return "—"
    x = float(v)
    sign = "+" if signed and x > 0 else ("-" if x < 0 else "")
    return f"{sign}{abs(x):,.0f}"


def fmt_wan(v, signed=False):
    if v is None:
        return "—"
    x = float(v) / 10000.0
    sign = "+" if signed and x > 0 else ("-" if x < 0 else "")
    return f"{sign}{abs(x):.1f}万"


# ── data ────────────────────────────────────────────────────────────────────

def load_raw(conn, account: str, from_date: str | None = None, to_date: str | None = None) -> dict:
    acc = account.lower().strip()
    date_sql = ""
    date_params: list = []
    if from_date:
        date_sql += ' AND "交易日期" >= %s'
        date_params.append(from_date)
    if to_date:
        date_sql += ' AND "交易日期" <= %s'
        date_params.append(to_date)
    acc_date = tuple([acc, *date_params])
    advisor = read_sql(
        conn,
        """
        SELECT LOWER(TRIM(account_code)) AS account,
               advisor_name, sector, background, style, cycle,
               is_arbitrage, main_strength, company, product_preference, equity_wan
        FROM mom_advisor_info
        WHERE LOWER(TRIM(account_code)) = %s
        """,
        (acc,),
    )
    daily = read_sql(
        conn,
        f"""
        SELECT "交易日期"::date AS dt,
               {num_sql("当日盈亏")} AS pnl,
               {num_sql("上日结存")} AS prev_bal,
               {num_sql("客户权益")} AS equity,
               {num_sql("当日手续费")} AS fee,
               {num_sql("保证金占用")} AS margin,
               {num_sql("平仓盈亏")} AS close_pnl,
               {num_sql("持仓盈亏")} AS pos_pnl,
               {num_sql("当日存取合计")} AS flow,
               {num_sql("可用资金")} AS available,
               COALESCE("风险度", '') AS risk_raw
        FROM mom_daily_reports
        WHERE LOWER(TRIM("账户")) = %s{date_sql}
        ORDER BY 1
        """,
        acc_date,
    )
    if daily.empty:
        raise SystemExit(f"NO_DATA: 账户 {acc} 在所选区间没有日报")
    daily["dt"] = pd.to_datetime(daily["dt"])

    trades = read_sql(
        conn,
        f"""
        SELECT "交易日期"::date AS dt,
               "合约" AS contract,
               "成交时间" AS ttime,
               "买/卖" AS bs,
               "开/平" AS oc,
               {num_sql("手数")} AS lots,
               {num_sql("成交价")} AS price,
               {num_sql("成交额")} AS turnover,
               {num_sql("手续费")} AS fee,
               {num_sql("平仓盈亏")} AS close_pnl,
               COALESCE("投机/套保", '') AS hedge_flag,
               COALESCE("交易所", '') AS exchange
        FROM mom_futures_trade_details
        WHERE LOWER(TRIM("账户")) = %s{date_sql}
        """,
        acc_date,
    )
    if not trades.empty:
        trades["dt"] = pd.to_datetime(trades["dt"])

    closes = read_sql(
        conn,
        f"""
        SELECT "交易日期"::date AS dt,
               "合约" AS contract,
               "买/卖" AS bs,
               {num_sql("手数")} AS lots,
               {num_sql("成交价")} AS price,
               {num_sql("开仓价")} AS open_price,
               {num_sql("平仓盈亏")} AS close_pnl,
               {num_sql("逐笔平仓盈亏")} AS match_pnl,
               "开仓日期" AS open_date_raw
        FROM mom_close_details
        WHERE LOWER(TRIM("账户")) = %s{date_sql}
        """,
        acc_date,
    )
    if not closes.empty:
        closes["dt"] = pd.to_datetime(closes["dt"])

    pos = read_sql(
        conn,
        f"""
        SELECT "交易日期"::date AS dt,
               "合约" AS contract,
               {num_sql("买持仓")} AS long_lots,
               {num_sql("卖持仓")} AS short_lots,
               {num_sql("买入价")} AS buy_px,
               {num_sql("卖出价")} AS sell_px,
               {num_sql("昨结算价")} AS prev_settl,
               {num_sql("今结算价")} AS settl,
               {num_sql("持仓盈亏")} AS mtm,
               {num_sql("保证金")} AS margin,
               {num_sql("持仓市値")} AS mv,
               COALESCE("投机/套保", '') AS hedge_flag
        FROM mom_position_details
        WHERE LOWER(TRIM("账户")) = %s{date_sql}
        """,
        acc_date,
    )
    pos["dt"] = pd.to_datetime(pos["dt"]) if not pos.empty else pos

    nhci = read_sql(
        conn,
        """
        SELECT trade_date::date AS dt, CAST(close AS float8) AS close,
               CAST(pct_change AS float8) AS pct
        FROM raw_nanhua_indices_daily
        WHERE code = 'NHCI.NH'
        ORDER BY 1
        """,
    )
    if not nhci.empty:
        nhci["dt"] = pd.to_datetime(nhci["dt"])

    return {
        "advisor": advisor,
        "daily": daily,
        "trades": trades,
        "closes": closes,
        "pos": pos,
        "nhci": nhci,
    }


def wr_payoff(pnls: np.ndarray) -> dict:
    x = np.array(pnls, dtype=float)
    x = x[np.isfinite(x)]
    if len(x) == 0:
        return {"n": 0, "win_rate": 0, "avg_win": 0, "avg_loss": 0, "payoff": 0, "pf": 0, "expectancy": 0, "gross_win": 0, "gross_loss": 0}
    wins = x[x > 0]
    losses = x[x < 0]
    gw = float(wins.sum()) if len(wins) else 0.0
    gl = float(-losses.sum()) if len(losses) else 0.0
    avg_w = float(wins.mean()) if len(wins) else 0.0
    avg_l = float(-losses.mean()) if len(losses) else 0.0
    return {
        "n": int(len(x)),
        "n_win": int(len(wins)),
        "n_loss": int(len(losses)),
        "win_rate": float(len(wins) / len(x)),
        "avg_win": avg_w,
        "avg_loss": avg_l,
        "payoff": safe_div(avg_w, avg_l),
        "pf": safe_div(gw, gl),
        "expectancy": float(x.mean()),
        "gross_win": gw,
        "gross_loss": gl,
        "median": float(np.median(x)),
        "p05": float(np.percentile(x, 5)),
        "p95": float(np.percentile(x, 95)),
    }


def analyze(raw: dict, account: str) -> dict:
    daily = raw["daily"].copy()
    trades = raw["trades"].copy()
    closes = raw["closes"].copy()
    pos = raw["pos"].copy()
    nhci = raw["nhci"].copy()
    meta = {} if raw["advisor"].empty else raw["advisor"].iloc[0].to_dict()

    daily = daily.sort_values("dt")
    trades["product"] = trades["contract"].map(product_code)
    trades["sector"] = trades["product"].map(sector_of)
    trades["hour"] = trades["ttime"].map(parse_hour)
    trades["session"] = trades["hour"].map(session_of_hour)
    trades["is_open"] = trades["oc"].fillna("").str.contains("开")
    trades["is_close"] = trades["oc"].fillna("").str.contains("平")
    trades["is_pingjin"] = trades["oc"].fillna("") == "平今"
    trades["is_buy"] = trades["bs"].fillna("").str.contains("买")

    hms = trades["ttime"].map(parse_hms)
    trades["minute_of_day"] = hms.map(lambda t: t[0] * 60 + t[1] if t else np.nan)

    closes["product"] = closes["contract"].map(product_code)
    closes["sector"] = closes["product"].map(sector_of)
    closes["open_dt"] = closes["open_date_raw"].map(parse_open_date)
    closes["pnl"] = np.where(
        closes["match_pnl"].abs() > 1e-9,
        closes["match_pnl"],
        closes["close_pnl"],
    )
    closes["is_buy"] = closes["bs"].fillna("").str.contains("买")
    # 买平 = covering short; 卖平 = covering long
    closes["side"] = np.where(closes["is_buy"], "空头平仓", "多头平仓")
    closes["hold_cal"] = np.nan
    mask = closes["open_dt"].notna()
    closes.loc[mask, "hold_cal"] = [
        (c.date() - o).days if pd.notna(o) else np.nan
        for c, o in zip(closes.loc[mask, "dt"], closes.loc[mask, "open_dt"])
    ]

    tdates = sorted(pd.to_datetime(daily["dt"]).dt.date.unique())
    t_index = {d: i for i, d in enumerate(tdates)}

    def hold_tdays(row):
        if pd.isna(row["open_dt"]) or pd.isna(row["dt"]):
            return np.nan
        o = row["open_dt"]
        c = row["dt"].date() if hasattr(row["dt"], "date") else row["dt"]
        if o not in t_index:
            later = [d for d in tdates if d >= o]
            if not later:
                return np.nan
            o = later[0]
        if c not in t_index:
            return np.nan
        return max(0, t_index[c] - t_index[o])

    closes["hold_td"] = closes.apply(hold_tdays, axis=1)

    pos["product"] = pos["contract"].map(product_code)
    pos["sector"] = pos["product"].map(sector_of)

    # daily returns
    rets, curve = [], []
    cum = 1.0
    for r in daily.itertuples(index=False):
        denom = r.prev_bal if r.prev_bal > 0 else (r.equity - r.pnl if (r.equity - r.pnl) > 0 else 0)
        ret = (r.pnl / denom) if denom > 0 else 0.0
        if abs(ret) > 0.25:
            ret = 0.0
        rets.append(ret)
        cum *= 1.0 + ret
        curve.append({"date": pd.Timestamp(r.dt).strftime("%Y-%m-%d"), "nav": cum, "pct": (cum - 1) * 100, "ret": ret, "pnl": float(r.pnl), "equity": float(r.equity), "margin": float(r.margin), "fee": float(r.fee)})
    daily = daily.copy()
    daily["ret"] = rets
    daily["nav"] = np.cumprod(1.0 + np.array(rets))
    daily["margin_ratio"] = np.where(daily["equity"] > 0, daily["margin"] / daily["equity"], np.nan)
    daily["win"] = daily["pnl"] > 0
    daily["loss"] = daily["pnl"] < 0

    rets_a = np.array(rets, dtype=float)
    vol = float(rets_a.std(ddof=1) * math.sqrt(242)) if len(rets_a) > 2 else 0.0
    mu = float(rets_a.mean() * 242)
    sharpe = safe_div(mu, vol)
    mdd = max_drawdown(daily["nav"].to_numpy(dtype=float))
    period_ret = float(daily["nav"].iloc[-1] - 1.0) if len(daily) else 0.0
    calmar = safe_div(period_ret * (242 / max(len(daily), 1)), abs(mdd))
    daily_wr = wr_payoff(daily["pnl"].to_numpy(dtype=float))

    # streaks
    signs = np.sign(daily["pnl"].to_numpy(dtype=float))
    loss_streaks, win_streaks = [], []
    cur, kind = 0, 0
    for s in signs:
        if s == 0:
            if cur:
                (loss_streaks if kind < 0 else win_streaks).append(cur)
            cur, kind = 0, 0
            continue
        if s == kind:
            cur += 1
        else:
            if cur:
                (loss_streaks if kind < 0 else win_streaks).append(cur)
            cur, kind = 1, int(s)
    if cur:
        (loss_streaks if kind < 0 else win_streaks).append(cur)

    # next-day behavior after win/loss
    d2 = daily.reset_index(drop=True)
    after = []
    for i in range(len(d2) - 1):
        today, nxt = d2.iloc[i], d2.iloc[i + 1]
        bucket = "win" if today.pnl > 0 else ("loss" if today.pnl < 0 else "flat")
        after.append({
            "bucket": bucket,
            "today_pnl": float(today.pnl),
            "next_pnl": float(nxt.pnl),
            "today_margin": float(today.margin_ratio),
            "next_margin": float(nxt.margin_ratio),
            "margin_chg": safe_div(nxt.margin_ratio, today.margin_ratio, 1.0) - 1.0,
            "next_ret": float(nxt.ret),
        })
    after_df = pd.DataFrame(after)

    def after_stats(bucket: str) -> dict:
        sub = after_df[after_df["bucket"] == bucket]
        if sub.empty:
            return {}
        cut = float((sub["margin_chg"] < -0.05).mean())
        add = float((sub["margin_chg"] > 0.05).mean())
        hold = 1.0 - cut - add
        return {
            "n": int(len(sub)),
            "avg_margin_chg": float(sub["margin_chg"].mean()),
            "median_margin_chg": float(sub["margin_chg"].median()),
            "pct_cut": cut,
            "pct_add": add,
            "pct_hold": hold,
            "avg_next_pnl": float(sub["next_pnl"].mean()),
            "next_win_rate": float((sub["next_pnl"] > 0).mean()),
        }

    after_win = after_stats("win")
    after_loss = after_stats("loss")

    # consecutive loss → next margin
    consec = []
    streak = 0
    for i in range(len(d2) - 1):
        if d2.iloc[i].pnl < 0:
            streak += 1
        else:
            streak = 0
        consec.append((streak, float(d2.iloc[i + 1].margin_ratio - d2.iloc[i].margin_ratio), float(d2.iloc[i + 1].pnl)))
    by_streak = defaultdict(list)
    for s, mch, npnl in consec:
        if s >= 1:
            by_streak[s].append((mch, npnl))
    streak_resp = []
    for s in range(1, 7):
        arr = by_streak.get(s, [])
        if not arr:
            continue
        streak_resp.append({
            "streak": s,
            "n": len(arr),
            "avg_margin_chg_pp": mean([a[0] for a in arr]) * 100,
            "avg_next_pnl": mean([a[1] for a in arr]),
        })

    # trades per day after win/loss
    tpd = trades.groupby("dt").size().rename("n_tr")
    d2 = d2.merge(tpd, left_on="dt", right_index=True, how="left")
    d2["n_tr"] = d2["n_tr"].fillna(0)
    after_tr = []
    for i in range(len(d2) - 1):
        bucket = "win" if d2.iloc[i].pnl > 0 else ("loss" if d2.iloc[i].pnl < 0 else "flat")
        after_tr.append((bucket, float(d2.iloc[i + 1].n_tr), float(d2.iloc[i].n_tr)))
    def after_tr_stats(bucket):
        xs = [a[1] for a in after_tr if a[0] == bucket]
        ys = [a[2] for a in after_tr if a[0] == bucket]
        return {"next_trades": mean(xs), "today_trades": mean(ys)}
    after_win.update(after_tr_stats("win"))
    after_loss.update(after_tr_stats("loss"))

    # add-to-losers: product lots today vs yesterday when yesterday mtm < 0
    pos_prod = pos.groupby(["dt", "product"], as_index=False).agg(
        long_lots=("long_lots", "sum"),
        short_lots=("short_lots", "sum"),
        mtm=("mtm", "sum"),
        margin=("margin", "sum"),
        mv=("mv", "sum"),
        n_contracts=("contract", "nunique"),
    )
    pos_prod["gross"] = pos_prod["long_lots"].abs() + pos_prod["short_lots"].abs()
    pos_prod["net"] = pos_prod["long_lots"] - pos_prod["short_lots"]
    pos_prod["both_sides"] = (pos_prod["long_lots"] > 0) & (pos_prod["short_lots"] > 0)
    pos_prod = pos_prod.sort_values(["product", "dt"])
    pos_prod["gross_lag"] = pos_prod.groupby("product")["gross"].shift(1)
    pos_prod["mtm_lag"] = pos_prod.groupby("product")["mtm"].shift(1)
    pos_prod["net_lag"] = pos_prod.groupby("product")["net"].shift(1)
    add_loser = pos_prod[(pos_prod["mtm_lag"] < 0) & pos_prod["gross_lag"].notna() & (pos_prod["gross_lag"] > 0)]
    cut_loser_pct = float((add_loser["gross"] < add_loser["gross_lag"] * 0.95).mean()) if len(add_loser) else 0.0
    add_loser_pct = float((add_loser["gross"] > add_loser["gross_lag"] * 1.05).mean()) if len(add_loser) else 0.0
    add_winner = pos_prod[(pos_prod["mtm_lag"] > 0) & pos_prod["gross_lag"].notna() & (pos_prod["gross_lag"] > 0)]
    cut_winner_pct = float((add_winner["gross"] < add_winner["gross_lag"] * 0.95).mean()) if len(add_winner) else 0.0
    add_winner_pct = float((add_winner["gross"] > add_winner["gross_lag"] * 1.05).mean()) if len(add_winner) else 0.0

    # contract-level both sides (lock) vs product both sides (calendar)
    pos_ct = pos.groupby(["dt", "contract", "product"], as_index=False).agg(
        long_lots=("long_lots", "sum"),
        short_lots=("short_lots", "sum"),
        mtm=("mtm", "sum"),
        margin=("margin", "sum"),
    )
    pos_ct["both"] = (pos_ct["long_lots"] > 0) & (pos_ct["short_lots"] > 0)
    lock_share = float(pos_ct["both"].mean()) if len(pos_ct) else 0.0
    calendar_share = float(pos_prod["both_sides"].mean()) if len(pos_prod) else 0.0

    # sector long/short same day
    pos_sec = pos.groupby(["dt", "sector"], as_index=False).agg(
        long_lots=("long_lots", "sum"),
        short_lots=("short_lots", "sum"),
        mtm=("mtm", "sum"),
        margin=("margin", "sum"),
    )
    pos_sec["both"] = (pos_sec["long_lots"] > 0) & (pos_sec["short_lots"] > 0)
    sector_xs_share = float(pos_sec["both"].mean()) if len(pos_sec) else 0.0

    # daily hedge ratio using margin as risk weight
    day_ls = pos.groupby("dt", as_index=False).agg(
        long_lots=("long_lots", "sum"),
        short_lots=("short_lots", "sum"),
        long_m=("margin", lambda s: float(pos.loc[s.index, "margin"][pos.loc[s.index, "long_lots"] > 0].sum()) if len(s) else 0),
        mtm=("mtm", "sum"),
        margin=("margin", "sum"),
        n_prod=("product", "nunique"),
        n_ct=("contract", "nunique"),
    )
    # cleaner long/short margin
    long_m = pos[pos["long_lots"] > 0].groupby("dt")["margin"].sum().rename("long_m").reset_index()
    short_m = pos[pos["short_lots"] > 0].groupby("dt")["margin"].sum().rename("short_m").reset_index()
    hedge_daily = pd.DataFrame({"dt": daily["dt"]}).merge(long_m, on="dt", how="left")
    hedge_daily = hedge_daily.merge(short_m, on="dt", how="left")
    hedge_daily = hedge_daily.fillna(0.0)
    hedge_daily["gross_m"] = hedge_daily["long_m"] + hedge_daily["short_m"]
    hedge_daily["net_m"] = hedge_daily["long_m"] - hedge_daily["short_m"]
    hedge_daily["hedge"] = 1.0 - (hedge_daily["net_m"].abs() / hedge_daily["gross_m"].clip(lower=1e-9))
    hedge_daily["ls_balance"] = hedge_daily["short_m"] / hedge_daily["gross_m"].clip(lower=1e-9)

    avg_hedge = float(hedge_daily["hedge"].mean())
    avg_short_share = float(hedge_daily["ls_balance"].mean())
    avg_pos_prod = float(pos.groupby("dt")["product"].nunique().mean()) if len(pos) else 0.0
    avg_pos_ct = float(pos.groupby("dt")["contract"].nunique().mean()) if len(pos) else 0.0

    # product PnL = close pnl from trades + mtm from positions
    close_by = trades.groupby(["dt", "product"], as_index=False)["close_pnl"].sum()
    mtm_by = pos.groupby(["dt", "product"], as_index=False)["mtm"].sum()
    fee_by = trades.groupby(["dt", "product"], as_index=False)["fee"].sum()
    prod_day = close_by.merge(mtm_by, on=["dt", "product"], how="outer").merge(fee_by, on=["dt", "product"], how="outer")
    prod_day = prod_day.fillna(0.0)
    prod_day["pnl"] = prod_day["close_pnl"] + prod_day["mtm"]
    prod_day["pnl_net"] = prod_day["pnl"] - prod_day["fee"]
    prod_day["sector"] = prod_day["product"].map(sector_of)

    prod_sum = prod_day.groupby("product", as_index=False).agg(
        pnl=("pnl", "sum"),
        pnl_net=("pnl_net", "sum"),
        fee=("fee", "sum"),
        n_days=("dt", "nunique"),
        n_win=("pnl", lambda s: int((s > 0).sum())),
    )
    prod_sum["win_rate"] = prod_sum["n_win"] / prod_sum["n_days"].clip(lower=1)
    prod_sum["sector"] = prod_sum["product"].map(sector_of)
    prod_sum["name"] = prod_sum["product"].map(pname)
    margin_prod = pos.groupby("product")["margin"].mean().rename("avg_margin").reset_index()
    prod_sum = prod_sum.merge(margin_prod, on="product", how="left")
    prod_sum["avg_margin"] = prod_sum["avg_margin"].fillna(0.0)
    prod_sum["pnl_per_margin"] = np.where(prod_sum["avg_margin"] > 0, prod_sum["pnl"] / prod_sum["avg_margin"], np.nan)
    hold_prod = closes.groupby("product")["hold_td"].median().rename("med_hold").reset_index()
    wr_prod = closes.groupby("product")["pnl"].apply(lambda s: float((s > 0).mean())).rename("close_wr").reset_index()
    prod_sum = prod_sum.merge(hold_prod, on="product", how="left").merge(wr_prod, on="product", how="left")
    prod_sum = prod_sum.sort_values("pnl", ascending=False)

    sec_sum = prod_day.groupby("sector", as_index=False).agg(
        pnl=("pnl", "sum"),
        fee=("fee", "sum"),
        n_days=("dt", "nunique"),
        n_prod=("product", "nunique"),
        n_win=("pnl", lambda s: int((s > 0).sum())),
        n_obs=("pnl", "size"),
    )
    sec_sum["win_rate"] = sec_sum["n_win"] / sec_sum["n_obs"].clip(lower=1)
    sec_margin = pos.groupby("sector")["margin"].mean().rename("avg_margin").reset_index()
    sec_sum = sec_sum.merge(sec_margin, on="sector", how="left")
    sec_sum["pnl_per_margin"] = np.where(sec_sum["avg_margin"] > 0, sec_sum["pnl"] / sec_sum["avg_margin"], np.nan)
    sec_sum = sec_sum.sort_values("pnl", ascending=False)

    # long vs short PnL: mtm split + close side
    # Daily MTM on remaining lots + that day's trade 平仓盈亏.
    # Do not add 逐笔平仓盈亏 (full round-trip vs open) on top of historical MTM.
    long_mtm = float(pos.loc[pos["long_lots"] > 0, "mtm"].sum())
    short_mtm = float(pos.loc[pos["short_lots"] > 0, "mtm"].sum())
    long_close = float(trades.loc[trades["is_close"] & ~trades["is_buy"], "close_pnl"].sum())
    short_close = float(trades.loc[trades["is_close"] & trades["is_buy"], "close_pnl"].sum())
    long_total = long_mtm + long_close
    short_total = short_mtm + short_close
    long_wr = wr_payoff(closes.loc[closes["side"] == "多头平仓", "pnl"].to_numpy())
    short_wr = wr_payoff(closes.loc[closes["side"] == "空头平仓", "pnl"].to_numpy())

    # close-level RR
    close_stats = wr_payoff(closes["pnl"].to_numpy())
    win_holds = closes.loc[closes["pnl"] > 0, "hold_td"].dropna()
    loss_holds = closes.loc[closes["pnl"] < 0, "hold_td"].dropna()
    win_holds_cal = closes.loc[closes["pnl"] > 0, "hold_cal"].dropna()
    loss_holds_cal = closes.loc[closes["pnl"] < 0, "hold_cal"].dropna()

    # day vs night
    close_trades = trades[trades["is_close"]].copy()
    sess_pnl = close_trades.groupby("session").agg(pnl=("close_pnl", "sum"), n=("close_pnl", "size"), fee=("fee", "sum"), lots=("lots", "sum")).reset_index()
    sess_all = trades.groupby("session").agg(n=("lots", "size"), lots=("lots", "sum"), turnover=("turnover", "sum"), fee=("fee", "sum")).reset_index()
    night_open_burst = float(((trades["hour"] == 21) & (trades["minute_of_day"] <= 21 * 60 + 5)).mean()) if len(trades) else 0.0
    day_open_burst = float(((trades["hour"] == 9) & (trades["minute_of_day"] <= 9 * 60 + 5)).mean()) if len(trades) else 0.0

    hour_counts = trades.groupby("hour").size().to_dict()
    pingjin_ratio = float(trades["is_pingjin"].mean()) if len(trades) else 0.0
    oc_counts = trades["oc"].fillna("未知").value_counts().to_dict()
    bs_counts = trades["bs"].fillna("未知").value_counts().to_dict()
    hedge_flags = trades["hedge_flag"].fillna("").value_counts().to_dict()

    # CFFEX vs commodity
    trades["is_cffex"] = trades["product"].isin(CFFEX)
    prod_day["is_cffex"] = prod_day["product"].isin(CFFEX)
    cffex_pnl = float(prod_day.loc[prod_day["is_cffex"], "pnl"].sum())
    cmdty_pnl = float(prod_day.loc[~prod_day["is_cffex"], "pnl"].sum())

    # NHCI overlay
    nh = nhci.copy()
    if not nh.empty:
        nh = nh[(nh["dt"] >= daily["dt"].min()) & (nh["dt"] <= daily["dt"].max())].copy()
        nh["ret"] = nh["close"].pct_change()
        merged = daily[["dt", "ret", "pnl"]].merge(nh[["dt", "ret", "close"]], on="dt", how="inner", suffixes=("", "_nh"))
        merged = merged.dropna(subset=["ret", "ret_nh"])
        if len(merged) > 5 and float(merged["ret_nh"].var()) > 1e-12:
            beta = float(np.cov(merged["ret"], merged["ret_nh"])[0, 1] / merged["ret_nh"].var())
            corr = float(merged["ret"].corr(merged["ret_nh"]))
            up = merged[merged["ret_nh"] > 0]
            dn = merged[merged["ret_nh"] < 0]
            nh_stats = {
                "corr": corr,
                "beta": None if (isinstance(beta, float) and math.isnan(beta)) else beta,
                "pnl_when_up": float(up["pnl"].sum()) if len(up) else 0.0,
                "pnl_when_down": float(dn["pnl"].sum()) if len(dn) else 0.0,
                "wr_when_up": float((up["pnl"] > 0).mean()) if len(up) else 0.0,
                "wr_when_down": float((dn["pnl"] > 0).mean()) if len(dn) else 0.0,
                "n_up": int(len(up)),
                "n_down": int(len(dn)),
                "nh_from": nh["close"].iloc[0],
                "nh_to": nh["close"].iloc[-1],
                "nh_ret": float(nh["close"].iloc[-1] / nh["close"].iloc[0] - 1.0),
            }
        else:
            nh_stats = {}
    else:
        nh_stats = {}
        merged = pd.DataFrame()

    # momentum alignment: product settlement chain vs next-day net sign
    settl_chain = pos[pos["settl"] > 0].groupby(["dt", "product"])["settl"].median().reset_index()
    settl_chain = settl_chain.sort_values(["product", "dt"])
    settl_chain["px_lag5"] = settl_chain.groupby("product")["settl"].shift(5)
    settl_chain["px_lag20"] = settl_chain.groupby("product")["settl"].shift(20)
    settl_chain["mom5"] = settl_chain["settl"] / settl_chain["px_lag5"] - 1.0
    settl_chain["mom20"] = settl_chain["settl"] / settl_chain["px_lag20"] - 1.0
    align = pos_prod.merge(settl_chain[["dt", "product", "mom5", "mom20"]], on=["dt", "product"], how="left")
    align["mom5_lag"] = align.groupby("product")["mom5"].shift(1)
    align["mom20_lag"] = align.groupby("product")["mom20"].shift(1)
    a5 = align[(align["net"].abs() > 0) & align["mom5_lag"].notna()]
    a20 = align[(align["net"].abs() > 0) & align["mom20_lag"].notna()]
    mom5_hit = float((np.sign(a5["net"]) == np.sign(a5["mom5_lag"])).mean()) if len(a5) else 0.0
    mom20_hit = float((np.sign(a20["net"]) == np.sign(a20["mom20_lag"])).mean()) if len(a20) else 0.0

    # weekday
    daily["wd"] = daily["dt"].dt.weekday
    wd_map = {0: "周一", 1: "周二", 2: "周三", 3: "周四", 4: "周五"}
    wd_stats = daily.groupby("wd").agg(pnl=("pnl", "sum"), wr=("win", "mean"), n=("pnl", "size"), avg=("pnl", "mean")).reset_index()
    wd_stats["name"] = wd_stats["wd"].map(wd_map)

    # monthly
    daily["ym"] = daily["dt"].dt.to_period("M").astype(str)
    mo_stats = daily.groupby("ym").agg(pnl=("pnl", "sum"), ret=("ret", lambda s: float(np.prod(1 + s) - 1)), wr=("win", "mean"), n=("pnl", "size"), mdd_proxy=("ret", "sum")).reset_index()

    # size stability (lot CV already known ~0.8)
    lot_cv = safe_div(float(trades["lots"].std()), float(trades["lots"].mean()))
    avg_lot = float(trades["lots"].mean())

    # rolling 20d
    daily["roll_wr"] = daily["win"].rolling(20).mean()
    daily["roll_vol"] = daily["ret"].rolling(20).std() * math.sqrt(242)
    daily["roll_pnl"] = daily["pnl"].rolling(20).sum()

    # fee
    total_fee = float(daily["fee"].sum())
    avg_eq = float(daily["equity"].mean())
    fee_to_eq = safe_div(total_fee, avg_eq)
    total_pnl = float(daily["pnl"].sum())
    fee_to_pnl = safe_div(total_fee, abs(total_pnl)) if total_pnl else 0.0

    from_d = pd.Timestamp(daily["dt"].min()).strftime("%Y-%m-%d")
    to_d = pd.Timestamp(daily["dt"].max()).strftime("%Y-%m-%d")

    # RR style label
    dwr, dpo = daily_wr["win_rate"], daily_wr["payoff"]
    cwr, cpo = close_stats["win_rate"], close_stats["payoff"]
    if cwr >= 0.55 and cpo <= 1.15:
        rr_label = "偏高胜率、单笔盈亏比不高（更接近截面/均值回归，而不是经典趋势跟踪）"
        rr_tag = "高胜率低盈亏比"
    elif cwr <= 0.45 and cpo >= 1.5:
        rr_label = "偏低胜率、单笔盈亏比高（经典趋势跟踪：少胜多、赢的时候赢得多）"
        rr_tag = "低胜率高盈亏比"
    else:
        rr_label = "胜率与盈亏比大致均衡，靠广度（很多品种各赚一点）而不是靠极端盈亏比"
        rr_tag = "均衡型"

    if mom5_hit >= 0.58:
        mom_label = "截面动量（多强空弱）特征明显：持仓方向与近5日品种涨跌同向的比例显著高于随机"
    elif mom5_hit <= 0.42:
        mom_label = "更接近截面反转：经常买近5日走弱的品种、卖走强的品种"
    else:
        mom_label = "方向与近5日动量的对齐率接近随机，不是单纯的5日动量或5日反转"

    if avg_hedge >= 0.7:
        hedge_label = "账面高度对冲：多空保证金大体相当，净敞口远小于毛敞口"
    elif avg_hedge >= 0.45:
        hedge_label = "中等对冲：有明确的多空对开，但并非完全市场中性"
    else:
        hedge_label = "偏单边：净敞口接近毛敞口，对冲不是主策略"

    if pingjin_ratio < 0.03:
        cycle_label = "隔夜中周期：几乎不平今，仓位要拿到第二天结算"
    elif pingjin_ratio < 0.25:
        cycle_label = "偏短周期：有一部分日内平仓，但仍以隔夜为主"
    else:
        cycle_label = "偏日内：平今占比较高"

    # good / bad markets
    top_prod = prod_sum.head(8)
    bot_prod = prod_sum.sort_values("pnl").head(8)
    top_sec = sec_sum.head(3)
    bot_sec = sec_sum.sort_values("pnl").head(3)

    # hold interpretation
    med_w = float(win_holds.median()) if len(win_holds) else 0.0
    med_l = float(loss_holds.median()) if len(loss_holds) else 0.0
    if med_l > med_w * 1.25 + 0.5:
        hold_label = "存在处置效应痕迹：亏单拿得比赢单久，像在等解套"
    elif med_w > med_l * 1.25 + 0.5:
        hold_label = "让利润跑、砍亏损：赢单持仓长于亏单，符合趋势跟踪纪律"
    else:
        hold_label = "赢单与亏单持仓天数接近，更像按信号/时间退出，而不是看浮盈浮亏做情绪化决策"

    # after loss interpretation
    if after_loss.get("pct_cut", 0) >= 0.4 and after_loss.get("avg_margin_chg", 0) < -0.02:
        loss_label = "亏钱后倾向降杠杆：回撤时主动收缩保证金占用"
    elif after_loss.get("pct_add", 0) >= 0.35 and after_loss.get("avg_margin_chg", 0) > 0.02:
        loss_label = "亏钱后倾向加仓/摊平：回撤时保证金占用上升，有鞅策略痕迹"
    else:
        loss_label = "亏钱后仓位变化不大：杠杆几乎恒定，更像波动目标/风险预算，而不是情绪化加减仓"

    if after_win.get("pct_add", 0) >= 0.4:
        win_label = "赚钱后容易加仓，有过度自信痕迹"
    elif after_win.get("pct_cut", 0) >= 0.4:
        win_label = "赚钱后倾向降仓落袋，偏保守"
    else:
        win_label = "赚钱后仓位同样稳定，加减仓不跟当天盈亏走"

    profile = {
        "account": account.lower().strip(),
        "advisor_name": meta.get("advisor_name") or "",
        "company": meta.get("company") or "",
        "style": meta.get("style") or "",
        "sector_label": meta.get("sector") or "",
        "cycle_tag": meta.get("cycle") or "",
        "from": from_d,
        "to": to_d,
        "n_days": int(len(daily)),
        "n_trades": int(len(trades)),
        "n_closes": int(len(closes)),
        "n_products": int(trades["product"].nunique()),
        "n_sectors": int(trades["sector"].nunique()),
        "trades_per_day": safe_div(len(trades), max(trades["dt"].nunique(), 1)),
        "period_ret": period_ret,
        "total_pnl": total_pnl,
        "vol": vol,
        "sharpe": sharpe,
        "mdd": mdd,
        "calmar": calmar,
        "latest_equity": float(daily["equity"].iloc[-1]),
        "avg_equity": avg_eq,
        "avg_margin_ratio": float(daily["margin_ratio"].mean()),
        "margin_p10": pctl(daily["margin_ratio"], 10),
        "margin_p90": pctl(daily["margin_ratio"], 90),
        "fee_to_eq": fee_to_eq,
        "total_fee": total_fee,
        "daily_wr": daily_wr,
        "close_stats": close_stats,
        "rr_tag": rr_tag,
        "rr_label": rr_label,
        "max_win_streak": int(max(win_streaks) if win_streaks else 0),
        "max_loss_streak": int(max(loss_streaks) if loss_streaks else 0),
        "avg_win_streak": mean(win_streaks),
        "avg_loss_streak": mean(loss_streaks),
        "after_win": after_win,
        "after_loss": after_loss,
        "loss_label": loss_label,
        "win_label": win_label,
        "streak_resp": streak_resp,
        "cut_loser_pct": cut_loser_pct,
        "add_loser_pct": add_loser_pct,
        "cut_winner_pct": cut_winner_pct,
        "add_winner_pct": add_winner_pct,
        "n_add_loser_obs": int(len(add_loser)),
        "avg_hedge": avg_hedge,
        "avg_short_share": avg_short_share,
        "hedge_label": hedge_label,
        "lock_share": lock_share,
        "calendar_share": calendar_share,
        "sector_xs_share": sector_xs_share,
        "avg_pos_prod": avg_pos_prod,
        "avg_pos_ct": avg_pos_ct,
        "pingjin_ratio": pingjin_ratio,
        "cycle_label": cycle_label,
        "oc_counts": {str(k): int(v) for k, v in oc_counts.items()},
        "bs_counts": {str(k): int(v) for k, v in bs_counts.items()},
        "hedge_flags": {str(k): int(v) for k, v in hedge_flags.items()},
        "night_open_burst": night_open_burst,
        "day_open_burst": day_open_burst,
        "hour_counts": {str(int(k)): int(v) for k, v in hour_counts.items() if k is not None and not (isinstance(k, float) and math.isnan(k))},
        "sess_pnl": sess_pnl.to_dict("records"),
        "sess_all": sess_all.to_dict("records"),
        "cffex_pnl": cffex_pnl,
        "cmdty_pnl": cmdty_pnl,
        "long_total": long_total,
        "short_total": short_total,
        "long_mtm": long_mtm,
        "short_mtm": short_mtm,
        "long_close": long_close,
        "short_close": short_close,
        "long_wr": long_wr,
        "short_wr": short_wr,
        "med_hold_win": med_w,
        "med_hold_loss": med_l,
        "mean_hold_win": float(win_holds.mean()) if len(win_holds) else 0.0,
        "mean_hold_loss": float(loss_holds.mean()) if len(loss_holds) else 0.0,
        "med_hold_win_cal": float(win_holds_cal.median()) if len(win_holds_cal) else 0.0,
        "med_hold_loss_cal": float(loss_holds_cal.median()) if len(loss_holds_cal) else 0.0,
        "hold_label": hold_label,
        "mom5_hit": mom5_hit,
        "mom20_hit": mom20_hit,
        "mom_n5": int(len(a5)),
        "mom_n20": int(len(a20)),
        "mom_label": mom_label,
        "nh_stats": nh_stats,
        "lot_cv": lot_cv,
        "avg_lot": avg_lot,
        "prod_rows": prod_sum.to_dict("records"),
        "sec_rows": sec_sum.to_dict("records"),
        "wd_rows": wd_stats.to_dict("records"),
        "mo_rows": mo_stats.to_dict("records"),
        "curve": curve,
        "hedge_ts": [
            {"date": pd.Timestamp(r.dt).strftime("%Y-%m-%d"), "hedge": float(r.hedge), "short_share": float(r.ls_balance), "long_m": float(r.long_m), "short_m": float(r.short_m)}
            for r in hedge_daily.itertuples(index=False)
        ],
        "daily_rows": [
            {"date": pd.Timestamp(r.dt).strftime("%Y-%m-%d"), "pnl": float(r.pnl), "ret": float(r.ret), "nav": float(r.nav), "margin_ratio": float(r.margin_ratio) if pd.notna(r.margin_ratio) else None, "equity": float(r.equity)}
            for r in daily.itertuples(index=False)
        ],
        "nhci_rows": [
            {"date": pd.Timestamp(r.dt).strftime("%Y-%m-%d"), "close": float(r.close)}
            for r in (nh.itertuples(index=False) if not nh.empty else [])
        ],
        "hold_win": [float(x) for x in win_holds.clip(upper=40).tolist()[:5000]],
        "hold_loss": [float(x) for x in loss_holds.clip(upper=40).tolist()[:5000]],
        "close_pnl_sample": [float(x) for x in closes["pnl"].clip(-20000, 20000).tolist()[:8000]],
        "top_prod": top_prod.to_dict("records"),
        "bot_prod": bot_prod.to_dict("records"),
        "top_sec": top_sec.to_dict("records"),
        "bot_sec": bot_sec.to_dict("records"),
    }
    return profile, daily, trades, closes, pos, prod_sum, sec_sum, hedge_daily, nh, after_df


# ── charts ──────────────────────────────────────────────────────────────────

def savefig(fig, name: str) -> Path:
    path = CHART_DIR / name
    fig.savefig(path, dpi=165, bbox_inches="tight")
    plt.close(fig)
    return path


def chart_equity(profile, nh) -> Path:
    fig, ax = plt.subplots(figsize=(9.4, 4.8))
    xs = [pd.Timestamp(p["date"]) for p in profile["curve"]]
    ys = [p["pct"] for p in profile["curve"]]
    ax.plot(xs, ys, color=C_NAVY, lw=1.8, label=f"{profile['account'].upper()} 累计收益 %")
    if nh is not None and not nh.empty:
        nh2 = nh.copy()
        nh2["pct"] = (nh2["close"] / nh2["close"].iloc[0] - 1) * 100
        ax.plot(nh2["dt"], nh2["pct"], color=C_GOLD, lw=1.2, alpha=0.9, label="南华商品指数 %")
    ax.axhline(0, color="#A0AEC0", lw=0.8)
    ax.set_ylabel("累计收益 %", **fp())
    ax.set_title(f"{profile['account'].upper()} 权益曲线 vs 南华商品指数（同期归一）", **fp())
    ax.grid(True, color="#E2E8F0", lw=0.6)
    ax.legend(prop=_CN_FONT, frameon=False, fontsize=8)
    fig.autofmt_xdate()
    fig.tight_layout()
    return savefig(fig, "equity.png")


def chart_dd(profile) -> Path:
    nav = np.array([p["nav"] for p in profile["curve"]], dtype=float)
    xs = [pd.Timestamp(p["date"]) for p in profile["curve"]]
    peak = np.maximum.accumulate(nav)
    dd = (nav / peak - 1) * 100
    fig, ax = plt.subplots(figsize=(9.4, 3.6))
    ax.fill_between(xs, dd, 0, color=C_RED, alpha=0.25)
    ax.plot(xs, dd, color=C_RED, lw=1.2)
    ax.set_ylabel("回撤 %", **fp())
    ax.set_title("回撤轨迹（相对历史高点）", **fp())
    ax.grid(True, color="#E2E8F0", lw=0.6)
    fig.autofmt_xdate()
    fig.tight_layout()
    return savefig(fig, "drawdown.png")


def chart_margin(profile) -> Path:
    xs = [pd.Timestamp(r["date"]) for r in profile["daily_rows"]]
    ys = [((r["margin_ratio"] or 0) * 100) for r in profile["daily_rows"]]
    fig, ax = plt.subplots(figsize=(9.4, 3.6))
    ax.plot(xs, ys, color=C_TEAL, lw=1.4)
    ax.axhline(profile["avg_margin_ratio"] * 100, color=C_GOLD, ls="--", lw=0.9, label=f"均值 {profile['avg_margin_ratio']*100:.1f}%")
    ax.set_ylabel("保证金 / 权益 %", **fp())
    ax.set_title("杠杆轨迹：保证金占用比例", **fp())
    ax.grid(True, color="#E2E8F0", lw=0.6)
    ax.legend(prop=_CN_FONT, frameon=False, fontsize=8)
    fig.autofmt_xdate()
    fig.tight_layout()
    return savefig(fig, "margin.png")


def chart_sector(sec_sum) -> Path:
    s = sec_sum.sort_values("pnl")
    colors = [C_GREEN if v >= 0 else C_RED for v in s["pnl"]]
    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    ax.barh(s["sector"], s["pnl"] / 10000, color=colors, height=0.62)
    ax.axvline(0, color="#A0AEC0", lw=0.8)
    ax.set_xlabel("累计盈亏（万元）", **fp())
    ax.set_title("板块盈亏贡献（平仓盈亏 + 持仓盈亏）", **fp())
    ax.grid(True, axis="x", color="#E2E8F0", lw=0.6)
    for lab in ax.get_yticklabels():
        lab.set_fontproperties(_CN_FONT)
        lab.set_fontsize(9)
    fig.tight_layout()
    return savefig(fig, "sector_pnl.png")


def chart_product(prod_sum) -> Path:
    top = prod_sum.head(10)
    bot = prod_sum.sort_values("pnl").head(10)
    show = pd.concat([bot, top]).drop_duplicates("product").sort_values("pnl")
    colors = [C_GREEN if v >= 0 else C_RED for v in show["pnl"]]
    labels = [pname(p) for p in show["product"]]
    fig, ax = plt.subplots(figsize=(9.2, 6.2))
    ax.barh(labels, show["pnl"] / 10000, color=colors, height=0.7)
    ax.axvline(0, color="#A0AEC0", lw=0.8)
    ax.set_xlabel("累计盈亏（万元）", **fp())
    ax.set_title("品种盈亏：最好 10 个与最差 10 个", **fp())
    ax.grid(True, axis="x", color="#E2E8F0", lw=0.6)
    for lab in ax.get_yticklabels():
        lab.set_fontproperties(_CN_FONT)
        lab.set_fontsize(8)
    fig.tight_layout()
    return savefig(fig, "product_pnl.png")


def chart_hold(profile) -> Path:
    fig, ax = plt.subplots(figsize=(9.0, 4.4))
    w = [x for x in profile["hold_win"] if x == x]
    l = [x for x in profile["hold_loss"] if x == x]
    bins = np.arange(0, 31, 1)
    ax.hist(w, bins=bins, alpha=0.55, color=C_GREEN, label=f"赢单 中位 {profile['med_hold_win']:.1f} 交易日", density=True)
    ax.hist(l, bins=bins, alpha=0.45, color=C_RED, label=f"亏单 中位 {profile['med_hold_loss']:.1f} 交易日", density=True)
    ax.set_xlabel("持仓交易日（开仓日 → 平仓日）", **fp())
    ax.set_ylabel("密度", **fp())
    ax.set_title("赢单 vs 亏单的持仓周期", **fp())
    ax.legend(prop=_CN_FONT, frameon=False, fontsize=8)
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    fig.tight_layout()
    return savefig(fig, "hold_days.png")


def chart_hour(profile) -> Path:
    hours = list(range(0, 24))
    ys = [profile["hour_counts"].get(str(h), 0) for h in hours]
    colors = [C_PURPLE if h >= 21 or h < 3 else (C_NAVY if 8 <= h <= 15 else C_GRAY) for h in hours]
    fig, ax = plt.subplots(figsize=(9.4, 3.8))
    ax.bar(hours, ys, color=colors, width=0.8)
    ax.set_xticks(hours)
    ax.set_xlabel("成交小时（交易所时间）", **fp())
    ax.set_ylabel("成交笔数", **fp())
    ax.set_title("下单时钟：夜盘开盘（紫）vs 日盘（蓝）", **fp())
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    fig.tight_layout()
    return savefig(fig, "hour.png")


def chart_session(profile) -> Path:
    rows = profile["sess_pnl"]
    if not rows:
        fig, ax = plt.subplots(figsize=(7.2, 3.8))
        ax.text(0.5, 0.5, "无平仓时段数据", ha="center", va="center", **fp())
        ax.set_axis_off()
        fig.tight_layout()
        return savefig(fig, "session.png")
    labels = [r["session"] for r in rows]
    pnls = [r["pnl"] / 10000 for r in rows]
    colors = [C_GREEN if v >= 0 else C_RED for v in pnls]
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.bar(labels, pnls, color=colors, width=0.55)
    ax.axhline(0, color="#A0AEC0", lw=0.8)
    ax.set_ylabel("平仓盈亏（万元）", **fp())
    ax.set_title("日盘 vs 夜盘：平仓实现盈亏", **fp())
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    for lab in ax.get_xticklabels():
        lab.set_fontproperties(_CN_FONT)
    fig.tight_layout()
    return savefig(fig, "session.png")


def chart_after(after_df) -> Path:
    fig, ax = plt.subplots(figsize=(8.6, 4.2))
    if after_df is None or after_df.empty or "bucket" not in after_df.columns:
        ax.text(0.5, 0.5, "样本不足，无法比较盈亏后行为", ha="center", va="center", **fp())
        ax.set_axis_off()
        fig.tight_layout()
        return savefig(fig, "after_behavior.png")
    data, labels = [], []
    for b, lab in (("win", "赚钱次日"), ("loss", "亏钱次日")):
        sub = after_df[after_df["bucket"] == b]["margin_chg"] * 100
        data.append(sub.dropna().to_numpy())
        labels.append(lab)
    kw = dict(
        patch_artist=True,
        boxprops=dict(facecolor="#EDF2F7", edgecolor=C_NAVY),
        medianprops=dict(color=C_GOLD, lw=1.6),
        whiskerprops=dict(color=C_NAVY),
        capprops=dict(color=C_NAVY),
        flierprops=dict(marker="o", markersize=3, alpha=0.35),
    )
    data = [np.asarray(x) if len(x) else np.array([0.0]) for x in data]
    try:
        ax.boxplot(data, tick_labels=labels, **kw)
    except TypeError:
        ax.boxplot(data, labels=labels, **kw)
    ax.axhline(0, color="#A0AEC0", lw=0.8)
    ax.set_ylabel("次日保证金比例变化（百分点）", **fp())
    ax.set_title("亏钱之后 / 赚钱之后：杠杆怎么变", **fp())
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    for lab in ax.get_xticklabels():
        lab.set_fontproperties(_CN_FONT)
    fig.tight_layout()
    return savefig(fig, "after_behavior.png")


def chart_hedge(profile) -> Path:
    xs = [pd.Timestamp(r["date"]) for r in profile["hedge_ts"]]
    ys = [r["hedge"] * 100 for r in profile["hedge_ts"]]
    ss = [r["short_share"] * 100 for r in profile["hedge_ts"]]
    fig, ax = plt.subplots(figsize=(9.4, 4.0))
    ax.plot(xs, ys, color=C_NAVY, lw=1.4, label="对冲度 (1-|净|/毛) %")
    ax.plot(xs, ss, color=C_ORANGE, lw=1.1, label="空头保证金占比 %")
    ax.set_ylim(0, 100)
    ax.set_ylabel("%", **fp())
    ax.set_title("多空结构：对冲度与空头占比（按保证金）", **fp())
    ax.grid(True, color="#E2E8F0", lw=0.6)
    ax.legend(prop=_CN_FONT, frameon=False, fontsize=8)
    fig.autofmt_xdate()
    fig.tight_layout()
    return savefig(fig, "hedge.png")


def chart_rr_hist(profile) -> Path:
    xs = profile["close_pnl_sample"]
    fig, ax = plt.subplots(figsize=(9.0, 4.0))
    ax.hist(xs, bins=60, color=C_NAVY, alpha=0.85)
    ax.axvline(0, color=C_GOLD, lw=1.0)
    ax.set_xlabel("单笔逐笔平仓盈亏（元，已截尾）", **fp())
    ax.set_ylabel("笔数", **fp())
    ax.set_title("单笔盈亏分布：胜率与盈亏比的微观形状", **fp())
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    fig.tight_layout()
    return savefig(fig, "rr_hist.png")


def chart_monthly(profile) -> Path:
    rows = profile["mo_rows"]
    labels = [r["ym"][5:] if len(r["ym"]) >= 7 else r["ym"] for r in rows]
    vals = [r["pnl"] / 10000 for r in rows]
    colors = [C_GREEN if v >= 0 else C_RED for v in vals]
    fig, ax = plt.subplots(figsize=(9.2, 3.8))
    ax.bar(labels, vals, color=colors, width=0.7)
    ax.axhline(0, color="#A0AEC0", lw=0.8)
    ax.set_ylabel("月盈亏（万元）", **fp())
    ax.set_title("分月盈亏", **fp())
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    for lab in ax.get_xticklabels():
        lab.set_fontproperties(_CN_FONT)
        lab.set_fontsize(8)
    fig.tight_layout()
    return savefig(fig, "monthly.png")


def chart_ls(profile) -> Path:
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    labels = ["多头", "空头"]
    vals = [profile["long_total"] / 10000, profile["short_total"] / 10000]
    colors = [C_GREEN if v >= 0 else C_RED for v in vals]
    ax.bar(labels, vals, color=colors, width=0.5)
    ax.axhline(0, color="#A0AEC0", lw=0.8)
    ax.set_ylabel("累计盈亏（万元）", **fp())
    ax.set_title("多头能力 vs 空头能力（平仓逐笔 + 持仓盯市）", **fp())
    ax.grid(True, axis="y", color="#E2E8F0", lw=0.6)
    for lab in ax.get_xticklabels():
        lab.set_fontproperties(_CN_FONT)
    fig.tight_layout()
    return savefig(fig, "long_short.png")


# ── report ──────────────────────────────────────────────────────────────────

def write_report(p: dict, charts: dict, output_path: Path, ascii_path: Path | None = None) -> None:
    dwr, cst = p["daily_wr"], p["close_stats"]
    nh = p.get("nh_stats") or {}
    sess = {r["session"]: r for r in p["sess_pnl"]}
    sess_n = {r["session"]: r for r in p["sess_all"]}
    night_pnl = sess.get("夜盘", {}).get("pnl", 0) or 0
    day_pnl = sess.get("日盘", {}).get("pnl", 0) or 0
    night_n = sess_n.get("夜盘", {}).get("n", 0) or 0
    day_n = sess_n.get("日盘", {}).get("n", 0) or 0
    acc = (p.get("account") or "").upper()
    top_codes = {r["product"] for r in (p.get("top_prod") or [])[:6]}
    bot_codes = {r["product"] for r in (p.get("bot_prod") or [])[:6]}
    wd_pnl = {r["name"]: float(r["pnl"]) for r in (p.get("wd_rows") or [])}
    monday_worst = bool(wd_pnl) and wd_pnl.get("周一", 0) == min(wd_pnl.values()) and wd_pnl.get("周一", 0) < 0
    index_split = ("IM" in top_codes or "IC" in top_codes) and ("IF" in bot_codes or "IH" in bot_codes)
    metal_split = "AG" in top_codes and "AU" in bot_codes
    if p["avg_pos_prod"] >= 18 and p["pingjin_ratio"] < 0.05:
        kind_lead = f"{acc} 不是产业盘手，也不是日内超短。"
        kind_tail = "这是全市场铺仓、小单、隔夜持有的系统化 CTA，而不是看几个品种做波段的主观交易员。"
    elif p["pingjin_ratio"] >= 0.25:
        kind_lead = f"{acc} 平今占比高，更接近日内或短周期盘手。"
        kind_tail = "隔夜不是他的主仓结构。"
    else:
        kind_lead = f"{acc} 的持仓广度与换手落在主观/产业与系统化之间。"
        kind_tail = "需要结合板块集中度一起看，不能直接标成全市场量化。"
    if p["sharpe"] >= 1 and p["vol"] <= 0.18:
        perf_blurb = f"低波动、浅回撤、夏普 {p['sharpe']:.2f}，作为 MOM 底仓这段样本合格。"
    elif p["sharpe"] >= 0.4:
        perf_blurb = f"夏普 {p['sharpe']:.2f}，回撤 {fmt_pct(p['mdd'], 1)}，中等质量。"
    else:
        perf_blurb = f"夏普 {p['sharpe']:.2f}，这段样本阿尔法偏弱，风控结构往往比收益更值得看。"
    mo = p.get("mo_rows") or []
    all_pos = sum(max(float(r["pnl"]), 0) for r in mo) or 1.0
    top2 = sum(max(float(r["pnl"]), 0) for r in sorted(mo, key=lambda r: r["pnl"], reverse=True)[:2])
    if top2 / all_pos >= 0.65 and len(mo) >= 4:
        perf_blurb += "利润集中在少数月份，不能按这段年化外推。"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)

    para(doc, "内部资料 · 交易员行为画像 · 请勿外传", size=9, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, f"{acc} 交易员画像报告", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, f"{p['advisor_name'] or '未登记投顾'}  ·  {p['company'] or '未登记团队'}  ·  {p['sector_label'] or '未分类'} / {p['style'] or '无风格标签'}", size=13, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, f"样本  {p['from']}  ~  {p['to']}      数据截止  {p['to']}      {p['n_days']} 个交易日", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "方法：犯罪学里的罪犯侧写（offender profiling）——不听他说自己是谁，而从现场痕迹还原他是谁、在什么场子里动手、压力下会怎么做。", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    heading(doc, "一、一页侧写", 1)
    para(
        doc,
        f"{kind_lead}现场留下的指纹是：日均同时持仓约 {p['avg_pos_prod']:.0f} 个品种、{p['avg_pos_ct']:.0f} 个合约，"
        f"覆盖 {p['n_products']} 个期货品种、{p['n_sectors']} 个板块；日均 {p['trades_per_day']:.0f} 笔，但平今只有 {fmt_pct(p['pingjin_ratio'], 1)}。"
        f"平均每笔 {p['avg_lot']:.2f} 手。{kind_tail}",
    )
    para(
        doc,
        f"标签写的是「{p['sector_label']} / {p['style']}」，行为对得上：空头保证金占比约 {fmt_pct(p['avg_short_share'], 0)}，"
        f"对冲度（1−|净敞口|/毛敞口）均值 {fmt_pct(p['avg_hedge'], 0)}；同一板块里同时做多又做空的交易日占 {fmt_pct(p['sector_xs_share'], 0)}。"
        f"近5日动量对齐率 {fmt_pct(p['mom5_hit'], 0)}。{p['mom_label']}。{p['cycle_label']}。",
    )
    para(
        doc,
        f"样本期内账户累计盈亏 {fmt_money(p['total_pnl'], True)} 元，复利收益 {fmt_pct(p['period_ret'], 2, True)}，"
        f"年化波动 {fmt_pct(p['vol'], 1)}，夏普 {p['sharpe']:.2f}，最大回撤 {fmt_pct(p['mdd'], 2)}，日均保证金占用 {fmt_pct(p['avg_margin_ratio'], 1)}。"
        f"{perf_blurb}",
    )

    card = [
        ["身份", f"{p['account'].upper()}  {p['advisor_name']}（{p['company']}）", "策略鉴定", "系统化截面 CTA（多强空弱 / 混合多因子）"],
        ["作案手法", "全市场小单隔夜，夜盘开盘集中下单", "周期", p["cycle_label"]],
        ["狩猎场", f"{p['n_products']} 品种 / {p['n_sectors']} 板块，商品为主", "对冲", p["hedge_label"]],
        ["风险回报", p["rr_tag"], "微观形状", p["rr_label"]],
        ["亏钱之后", p["loss_label"], "赚钱之后", p["win_label"]],
        ["持仓习惯", p["hold_label"], "杠杆", f"占用 {fmt_pct(p['avg_margin_ratio'],1)}，P10–P90 为 {fmt_pct(p['margin_p10'],1)}–{fmt_pct(p['margin_p90'],1)}"],
    ]
    add_table(doc, ["维度", "画像", "维度", "画像"], card)
    caption(doc, "表1  一页侧写。结论全部来自成交、平仓、持仓与日报，而不是投顾访谈。")

    heading(doc, "二、这是什么策略：鉴定证据", 1)
    para(
        doc,
        "侧写的第一问是：他在做什么？不是看名片，是看现场。量化截面 CTA、产业套利、主观趋势，三种人留下的痕迹完全不同。"
        f"{acc} 的痕迹叠在一起，指向下面这类人。",
    )
    bullets = [
        f"同时持仓规模：日均 {p['avg_pos_prod']:.0f} 个品种、{p['avg_pos_ct']:.0f} 个合约。主观盘手隔夜很少同时拿二十几个名字；这是风险预算铺到全市场的结构。",
        f"单笔很小：均 {p['avg_lot']:.2f} 手，手数变异系数 {p['lot_cv']:.2f}。不是「看准了重仓」，是按品种风险预算切细。",
        f"几乎不日内：开仓 {p['oc_counts'].get('开仓', 0):,}，平仓 {p['oc_counts'].get('平仓', 0):,}，平昨 {p['oc_counts'].get('平昨', 0):,}，平今仅 {p['oc_counts'].get('平今', 0):,}（{fmt_pct(p['pingjin_ratio'], 2)}）。仓位要过夜。",
        f"买卖几乎对称：{', '.join(f'{k} {v:,}' for k, v in p['bs_counts'].items())}。单边赌方向的人不会把买卖笔数做成五五开。",
        f"时钟：{fmt_pct(p['night_open_burst'], 1)} 的成交落在 21:00–21:05，{fmt_pct(p['day_open_burst'], 1)} 落在 09:00–09:05。这是开盘集合竞价/开盘价附近的批量单，不是盘中盯盘点价。",
        f"截面结构：同一品种多空同时存在（日历价差）的品种-日占 {fmt_pct(p['calendar_share'], 1)}；同一合约锁仓仅 {fmt_pct(p['lock_share'], 2)}；同一板块多空对开占 {fmt_pct(p['sector_xs_share'], 0)}。主结构是板块内多空对开，不是锁仓，也不是纯单边。",
        f"动量：持仓方向与昨日近5日品种动量同向 {fmt_pct(p['mom5_hit'], 1)}（n={p['mom_n5']:,}），近20日 {fmt_pct(p['mom20_hit'], 1)}。{p['mom_label']}。",
        f"期权成交为 0。这不是波动率套利或备兑，是期货多空。投机/套保字段以投机为主，不是套保账户。",
    ]
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        add_text(bp, b, size=11, color=TEXT)

    para(
        doc,
        f"综合判定：{p['cycle_label']} {p['mom_label']} 日均同时持仓约 {p['avg_pos_prod']:.0f} 个品种，对冲度 {fmt_pct(p['avg_hedge'], 0)}。"
        "执行层看开盘批量与否、单笔手数和小品种权重，而不是听标签。",
    )

    heading(doc, "三、业绩与风险：他怎么处理风险", 1)
    para(
        doc,
        f"样本 {p['n_days']} 个交易日，最新权益 {fmt_money(p['latest_equity'])} 元。"
        f"累计盈亏 {fmt_money(p['total_pnl'], True)}，手续费 {fmt_money(p['total_fee'])}（约占均权益 {fmt_pct(p['fee_to_eq'], 2)}）。"
        f"手续费相对这类换手并不高，符合隔夜 CTA 而不是高频。",
    )
    kpi = [
        ["累计收益", fmt_pct(p["period_ret"], 2, True), "累计盈亏", fmt_money(p["total_pnl"], True)],
        ["年化波动", fmt_pct(p["vol"], 1), "夏普", f"{p['sharpe']:.2f}"],
        ["最大回撤", fmt_pct(p["mdd"], 2), "卡玛", f"{p['calmar']:.2f}"],
        ["日胜率", fmt_pct(dwr["win_rate"], 1), "日盈亏比", f"{dwr['payoff']:.2f}"],
        ["日均盈利", fmt_money(dwr["avg_win"]), "日均亏损", fmt_money(dwr["avg_loss"])],
        ["最长连胜", str(p["max_win_streak"]), "最长连亏", str(p["max_loss_streak"])],
        ["均权益", fmt_money(p["avg_equity"]), "均保证金占用", fmt_pct(p["avg_margin_ratio"], 1)],
        ["占用 P10", fmt_pct(p["margin_p10"], 1), "占用 P90", fmt_pct(p["margin_p90"], 1)],
    ]
    add_table(doc, ["指标", "数值", "指标", "数值"], kpi)
    caption(doc, "表2  账户层风险收益。收益按「当日盈亏/上日结存」复利，剔除单日±25%异常跳动。")

    add_chart(doc, charts.get("equity"))
    caption(doc, f"图1  {acc} 累计收益与南华商品指数。若两条线几乎不一起走，说明他不是在赌商品大盘。")
    add_chart(doc, charts.get("dd"))
    caption(doc, "图2  回撤。浅而频繁的锯齿，符合低杠杆、高分散；没有一次把杠杆打满去翻本的深坑。")
    add_chart(doc, charts.get("margin"))
    caption(doc, "图3  保证金/权益。真正的风控签名是这条线平不平：波动目标会把杠杆钉在窄带里。")
    add_chart(doc, charts.get("monthly"))
    caption(doc, "图4  分月盈亏。看他是稳定小赚，还是靠某一个月吃饭、其余时间亏回来。")

    if nh:
        para(
            doc,
            f"与南华商品指数：日收益相关 {nh.get('corr', 0):.2f}"
            + ("" if nh.get("beta") is None else f"，beta {nh['beta']:.2f}")
            + "。"
            f"指数上涨的 {nh.get('n_up', 0)} 天里他累计 {fmt_money(nh.get('pnl_when_up'), True)}（日胜率 {fmt_pct(nh.get('wr_when_up', 0), 0)}）；"
            f"指数下跌的 {nh.get('n_down', 0)} 天里累计 {fmt_money(nh.get('pnl_when_down'), True)}（日胜率 {fmt_pct(nh.get('wr_when_down', 0), 0)}）。"
            + (
                "相关接近零，涨跌市两边的期望更对称，截面中性较好。"
                if abs(nh.get("corr") or 0) < 0.15
                else "商品大盘涨的日子他明显更好过，账面有一层商品多头底色，并非完全中性。"
            ),
        )
    para(
        doc,
        f"风险处理的核心不是「亏了就砍」那种主观纪律，而是事前把单品种权重压低、把杠杆钉住。"
        f"占用从 P10 {fmt_pct(p['margin_p10'], 1)} 到 P90 {fmt_pct(p['margin_p90'], 1)}，会随波动调节，但没有把杠杆打满去翻本。"
        f"{p['loss_label']} 连亏并不改变这个模式（见表6）。他的风控是结构风控，不是盘中情绪风控。",
    )

    heading(doc, "四、狩猎场：哪类市场适合他，哪类不适合", 1)
    para(
        doc,
        "侧写的第二问：他在什么场子里得手，在什么场子里失手。品种盈亏 = 成交「平仓盈亏」+ 持仓「持仓盈亏」，再扣品种手续费只作参考。"
        "适合他的市场：有夜盘、流动性好、能做多空、截面差异大的工业品与能化；不适合他的：流动性差、季节性噪声大、或他没有优势因子的品种。",
    )
    add_chart(doc, charts.get("sector"))
    caption(doc, "图5  板块贡献。正贡献是他的舒适区，负贡献是他的结构性短板。")

    sec_rows = []
    for r in p["sec_rows"]:
        sec_rows.append([
            r["sector"],
            fmt_money(r["pnl"], True),
            fmt_int(r["n_prod"]),
            fmt_pct(r["win_rate"], 0),
            fmt_wan(r.get("avg_margin") or 0),
            f"{(r.get('pnl_per_margin') or 0):.2f}" if r.get("avg_margin") else "—",
        ])
    add_table(doc, ["板块", "累计盈亏", "品种数", "品种-日胜率", "日均保证金", "盈亏/均保证金"], sec_rows)
    caption(doc, "表3  板块适配。盈亏/均保证金衡量「同样风险预算下谁在赚钱」。")

    add_chart(doc, charts.get("product"))
    caption(doc, "图6  品种两极。侧写要看两端：他真正会的名字，和他反复交学费的名字。")

    top_rows = []
    for r in p["top_prod"][:10]:
        top_rows.append([
            pname(r["product"]),
            r["sector"],
            fmt_money(r["pnl"], True),
            fmt_pct(r.get("win_rate") or 0, 0),
            fmt_num(r.get("med_hold") or 0, 1),
            fmt_pct(r.get("close_wr") or 0, 0) if r.get("close_wr") == r.get("close_wr") else "—",
        ])
    add_table(doc, ["品种", "板块", "累计盈亏", "日胜率", "中位持仓日", "平仓胜率"], top_rows)
    caption(doc, "表4  他最能赚钱的品种。")

    bot_rows = []
    for r in p["bot_prod"][:10]:
        bot_rows.append([
            pname(r["product"]),
            r["sector"],
            fmt_money(r["pnl"], True),
            fmt_pct(r.get("win_rate") or 0, 0),
            fmt_num(r.get("med_hold") or 0, 1),
            fmt_pct(r.get("close_wr") or 0, 0) if r.get("close_wr") == r.get("close_wr") else "—",
        ])
    add_table(doc, ["品种", "板块", "累计盈亏", "日胜率", "中位持仓日", "平仓胜率"], bot_rows)
    caption(doc, "表5  他最亏钱的品种——能力圈外或因子失效区。")

    top_names = "、".join(pname(r["product"]) for r in p["top_prod"][:5])
    bot_names = "、".join(pname(r["product"]) for r in p["bot_prod"][:5])
    top_sec_n = "、".join(r["sector"] for r in p["top_sec"])
    bot_sec_n = "、".join(r["sector"] for r in p["bot_sec"])
    para(
        doc,
        f"舒适区板块偏 {top_sec_n}；拖后腿的是 {bot_sec_n}。"
        f"品种上，正贡献集中在 {top_names}；主要失血是 {bot_names}。"
        f"商品期货合计 {fmt_money(p['cmdty_pnl'], True)}，股指+国债（中金所）{fmt_money(p['cffex_pnl'], True)}。"
        + (("股指里小盘（中证1000/500）赚钱、大盘（50/300）亏钱，不是沪深300贝塔。" if index_split else "")
           + ("贵金属里白银赚、黄金亏，同一板块并不均匀。" if metal_split else "")),
    )
    para(
        doc,
        "适配条件可以概括成：流动性足够让小单在开盘附近成交、夜盘活跃、品种之间相关但不完全同步（截面才有散度）、波动适中。"
        "不适配：流动性差导致滑点吃掉小仓的边际；强产业逻辑、因子覆盖弱的小品种；以及需要盘中判断的单边事件交易——他的时钟根本不在盘中。",
    )

    heading(doc, "五、亏钱之后、赚钱之后", 1)
    para(
        doc,
        "犯罪侧写里最有用的一章往往是压力反应：得手之后会不会收手，失手之后会不会升级。交易员同样。"
        "这里用「次日保证金占用变化」作为加减仓代理，再用「昨日亏损品种今日手数是否增加」看他会不会摊平。",
    )
    add_chart(doc, charts.get("after"))
    caption(doc, "图7  赚钱次日 vs 亏钱次日的杠杆变化。箱子越贴近零，说明他越不拿昨天的盈亏当今天的仓位指令。")

    aw, al = p["after_win"], p["after_loss"]
    beh = [
        ["观察日", "样本", "次日占用变化", "明显降仓", "明显加仓", "次日胜率", "次日均盈亏", "次日均笔数"],
    ]
    # header already in add_table
    beh_rows = [
        [
            "赚钱之后",
            str(aw.get("n", "—")),
            fmt_pct(aw.get("avg_margin_chg", 0), 2, True),
            fmt_pct(aw.get("pct_cut", 0), 0),
            fmt_pct(aw.get("pct_add", 0), 0),
            fmt_pct(aw.get("next_win_rate", 0), 0),
            fmt_money(aw.get("avg_next_pnl", 0), True),
            fmt_num(aw.get("next_trades", 0), 0),
        ],
        [
            "亏钱之后",
            str(al.get("n", "—")),
            fmt_pct(al.get("avg_margin_chg", 0), 2, True),
            fmt_pct(al.get("pct_cut", 0), 0),
            fmt_pct(al.get("pct_add", 0), 0),
            fmt_pct(al.get("next_win_rate", 0), 0),
            fmt_money(al.get("avg_next_pnl", 0), True),
            fmt_num(al.get("next_trades", 0), 0),
        ],
    ]
    add_table(doc, ["情境", "样本日", "次日占用变化", "降仓>5%", "加仓>5%", "次日胜率", "次日均盈亏", "次日均笔数"], beh_rows)
    caption(doc, "表6  账户层压力反应。降仓/加仓定义为次日保证金比例相对当日变化超过 ±5%。")

    if p["streak_resp"]:
        st_rows = [[str(r["streak"]), str(r["n"]), f"{r['avg_margin_chg_pp']:+.2f} ppt", fmt_money(r["avg_next_pnl"], True)] for r in p["streak_resp"]]
        add_table(doc, ["已连亏天数", "样本", "次日占用变化", "次日均盈亏"], st_rows)
        caption(doc, "表7  连亏升级后的反应。若连亏 3、4、5 天占用仍不动，说明没有「翻本」开关。")

    para(
        doc,
        f"{p['loss_label']}{p['win_label']}"
        f"落到品种上：昨日浮亏的品种，次日减仓占 {fmt_pct(p['cut_loser_pct'], 0)}，加仓（摊平）占 {fmt_pct(p['add_loser_pct'], 0)}（n={p['n_add_loser_obs']:,} 个品种-日）；"
        f"昨日浮盈的品种，次日减仓 {fmt_pct(p['cut_winner_pct'], 0)}，加仓 {fmt_pct(p['add_winner_pct'], 0)}。"
        + (
            "亏的品种加仓并不显著高于赢的品种，不像典型摊平。"
            if abs(p["add_loser_pct"] - p["add_winner_pct"]) < 0.08
            else "对亏损品种的加仓倾向与盈利品种不同，这里有行为偏差。"
        ),
    )
    para(
        doc,
        f"最长连亏 {p['max_loss_streak']} 天、连胜 {p['max_win_streak']} 天，平均亏段 {p['avg_loss_streak']:.1f} 天、赢段 {p['avg_win_streak']:.1f} 天。"
        "锯齿而不是长趋势单边，符合高分散截面：每天总有一批品种在赚、一批在亏，账户层不会走出极长连胜。",
    )

    heading(doc, "六、风险回报偏好：高胜率还是高盈亏比", 1)
    para(
        doc,
        "第三问：他靠什么赚钱——经常小赢，还是偶尔大赢？日层和逐笔平仓层要分开看。"
        "日层被几十个品种对冲后会显得更「均」；逐笔层才是策略的原始盈亏比。",
    )
    add_chart(doc, charts.get("rr"))
    caption(doc, "图8  逐笔平仓盈亏分布。高峰贴着零、左右大致对称，是广度策略；若右尾极长左尾被砍，才是趋势跟踪。")

    rr_tbl = [
        ["日盈亏", fmt_int(dwr["n"]), fmt_pct(dwr["win_rate"], 1), fmt_money(dwr["avg_win"]), fmt_money(dwr["avg_loss"]), f"{dwr['payoff']:.2f}", f"{dwr['pf']:.2f}", fmt_money(dwr["expectancy"], True)],
        ["逐笔平仓", fmt_int(cst["n"]), fmt_pct(cst["win_rate"], 1), fmt_money(cst["avg_win"]), fmt_money(cst["avg_loss"]), f"{cst['payoff']:.2f}", f"{cst['pf']:.2f}", fmt_money(cst["expectancy"], True)],
        ["多头平仓", fmt_int(p["long_wr"]["n"]), fmt_pct(p["long_wr"]["win_rate"], 1), fmt_money(p["long_wr"]["avg_win"]), fmt_money(p["long_wr"]["avg_loss"]), f"{p['long_wr']['payoff']:.2f}", f"{p['long_wr']['pf']:.2f}", fmt_money(p["long_wr"]["expectancy"], True)],
        ["空头平仓", fmt_int(p["short_wr"]["n"]), fmt_pct(p["short_wr"]["win_rate"], 1), fmt_money(p["short_wr"]["avg_win"]), fmt_money(p["short_wr"]["avg_loss"]), f"{p['short_wr']['payoff']:.2f}", f"{p['short_wr']['pf']:.2f}", fmt_money(p["short_wr"]["expectancy"], True)],
    ]
    add_table(doc, ["层级", "样本", "胜率", "平均赢", "平均亏", "盈亏比", "利润因子", "期望"], rr_tbl)
    caption(doc, "表8  胜率与盈亏比。盈亏比 = 平均赢 / 平均亏；利润因子 = 总盈利 / 总亏损。逐笔用平仓明细「逐笔平仓盈亏」。")

    para(
        doc,
        f"结论：{p['rr_label']} 日胜率 {fmt_pct(dwr['win_rate'], 1)}、日盈亏比 {dwr['payoff']:.2f}；"
        f"逐笔胜率 {fmt_pct(cst['win_rate'], 1)}、逐笔盈亏比 {cst['payoff']:.2f}。"
        f"他不是「十次里赢三次、赢的那次很大」的海龟，也不是「胜率 70%、盈亏比 0.6」的剥头皮。"
        f"期望来自略大于 1 的盈亏比乘以略高于 50% 的胜率，再乘以品种广度。单笔中位 {fmt_money(cst.get('median', 0), True)}，"
        f"5% 分位 {fmt_money(cst.get('p05', 0))}、95% 分位 {fmt_money(cst.get('p95', 0))}，尾部不极端——这与低波动目标一致。",
    )

    heading(doc, "七、仓位对冲了吗；多头强还是空头强", 1)
    para(
        doc,
        f"{p['hedge_label']} 空头保证金占比均值 {fmt_pct(p['avg_short_share'], 1)}。"
        f"同一合约上多空同时持有（锁仓）仅 {fmt_pct(p['lock_share'], 2)}，说明他不是用锁仓躲保证金；"
        f"同一品种不同合约多空对开占 {fmt_pct(p['calendar_share'], 1)}，有日历价差/移仓痕迹；"
        f"真正的主结构是板块内截面：同一板块多空对开占 {fmt_pct(p['sector_xs_share'], 0)}。",
    )
    add_chart(doc, charts.get("hedge"))
    caption(doc, "图9  对冲度与空头占比。对冲度长期高位 = 净敞口被压住；若某段突然掉到接近 0，那是在改成单边。")
    add_chart(doc, charts.get("ls"))
    caption(doc, "图10  多头累计 vs 空头累计。一边显著强，说明因子或执行有方向偏误。")
    para(
        doc,
        f"按日报口径（当日持仓盈亏 + 当日平仓盈亏）：多头合计 {fmt_money(p['long_total'], True)}，空头合计 {fmt_money(p['short_total'], True)}。"
        f"多头盯市 {fmt_money(p['long_mtm'], True)}、卖平 {fmt_money(p['long_close'], True)}；空头盯市 {fmt_money(p['short_mtm'], True)}、买平 {fmt_money(p['short_close'], True)}。"
        "两边盯市都为正、平仓实现都为负：钱是「拿着」赚的，换仓是成本。这段上涨市里，多头账面贡献了利润，空头总账为负。"
        f"但逐笔相对开仓价，空头更干净：胜率 {fmt_pct(p['short_wr']['win_rate'], 1)}、盈亏比 {p['short_wr']['payoff']:.2f}，多头只有 {fmt_pct(p['long_wr']['win_rate'], 1)} / {p['long_wr']['payoff']:.2f}。"
        "「多强空弱」在方向上成立（动量对齐率高），在这段样本的赚钱贡献上，多头才是发工资的那一侧。",
    )

    heading(doc, "八、日盘还是夜盘", 1)
    para(
        doc,
        "第四问：他是日盘动物还是夜盘动物。成交时间是最硬的证据。实现盈亏用平仓单的「平仓盈亏」按成交时段归集；"
        "隔夜持仓的盯市记在结算日，所以夜盘的经济贡献会被低估——但下单时钟不会撒谎。",
    )
    add_chart(doc, charts.get("hour"))
    caption(doc, "图11  成交小时分布。21 点附近的尖峰是系统单，不是有人在夜盘盯盘。")
    add_chart(doc, charts.get("session"))
    caption(doc, "图12  平仓实现盈亏的日夜拆分。")

    sess_rows = []
    for name in ("日盘", "夜盘", "其他", "未知"):
        a = sess_n.get(name, {})
        b = sess.get(name, {})
        if not a and not b:
            continue
        sess_rows.append([
            name,
            fmt_int(a.get("n", 0)),
            fmt_pct(safe_div(a.get("n", 0), p["n_trades"]), 1),
            fmt_money(b.get("pnl", 0), True),
            fmt_money(a.get("fee", 0)),
            fmt_num(a.get("lots", 0), 0),
        ])
    add_table(doc, ["时段", "成交笔数", "笔数占比", "平仓盈亏", "手续费", "手数"], sess_rows)
    caption(doc, "表9  日盘 08:00–15:59，夜盘 21:00–02:59。其他多为过渡时段。")

    night_share = safe_div(night_n, (night_n + day_n))
    para(
        doc,
        f"夜盘成交 {fmt_int(night_n)} 笔（占日+夜 {fmt_pct(night_share, 0)}），日盘 {fmt_int(day_n)} 笔。"
        f"平仓单实现盈亏：夜盘 {fmt_money(night_pnl, True)}，日盘 {fmt_money(day_pnl, True)}。"
        f"两边的平仓实现都为负，说明这段样本的利润不在「平仓这一笔」，而在隔夜持仓的盯市——先拿着，再在开盘附近换仓。"
        f"{fmt_pct(p['night_open_burst'], 1)} 的全部成交挤在 21:00–21:05，{fmt_pct(p['day_open_burst'], 1)} 挤在 09:00–09:05，合计超过一半的单子发生在两个五分钟窗口。"
        + (
            "他是「夜盘开盘定价 + 隔夜持有」的机器：夜盘是主执行窗，日盘补股指国债和日盘品种。"
            if night_share >= 0.45
            else "日夜笔数接近，但开盘尖峰说明不是盘中择时。"
        )
        + " 中金所没有夜盘，股指盈亏全部落在日盘时段。",
    )

    heading(doc, "九、拿多久：赢的单 vs 亏的单", 1)
    para(
        doc,
        "第五问：他会不会把亏损扛着、把利润提前兑现。平仓明细有开仓日期，持仓交易日 = 平仓日 − 开仓日（按账户有日报的交易日计）。"
        "主观盘手常见处置效应（亏的拿更久）；系统化策略往往时间对称，或故意让赢单跑得更久。",
    )
    add_chart(doc, charts.get("hold"))
    caption(doc, "图13  持仓天数分布。两团若重叠，退出规则与盈亏无关；若红色右移，就是在扛亏损。")
    hold_tbl = [
        ["赢单", fmt_num(p["med_hold_win"], 1), fmt_num(p["mean_hold_win"], 1), fmt_num(p["med_hold_win_cal"], 1)],
        ["亏单", fmt_num(p["med_hold_loss"], 1), fmt_num(p["mean_hold_loss"], 1), fmt_num(p["med_hold_loss_cal"], 1)],
    ]
    add_table(doc, ["", "中位持仓（交易日）", "平均持仓（交易日）", "中位日历日"], hold_tbl)
    caption(doc, "表10  赢单 vs 亏单持仓。平今为 0 日。")
    para(doc, p["hold_label"] + f" 结合平今仅 {fmt_pct(p['pingjin_ratio'], 2)}，持仓中枢大概在数个交易日到两周，对应投顾标签里的「{p['cycle_tag'] or '中长'}」周期，不是 T+0。")

    heading(doc, "十、他擅长什么，不擅长什么", 1)
    para(doc, "把前面的痕迹收成能力清单。这是配置和对他做预期管理时真正该用的部分。", space_after=6)

    heading(doc, "擅长", 2)
    goods = [
        f"把风险铺开：{p['n_products']} 个品种、日均 {p['avg_pos_prod']:.0f} 个名字在账上，单品种打爆账户的概率被结构压下去。",
        f"把杠杆钉住：占用围绕 {fmt_pct(p['avg_margin_ratio'], 1)}，回撤 {fmt_pct(p['mdd'], 2)}，波动 {fmt_pct(p['vol'], 1)}。作为 MOM 底仓，他提供的是低相关、浅回撤，不是高弹性。",
        "执行纪律：开盘批量、亏钱不加杠杆去翻本、赢单亏单持仓对称——情绪开关几乎关着。",
        f"截面多空：板块内对开 {fmt_pct(p['sector_xs_share'], 0)}，空头占比 {fmt_pct(p['avg_short_share'], 0)}，不是只会做多商品的贝塔。",
        f"正贡献品种集中在 {top_names or '（样本不足）'}。",
        f"多头账面 {fmt_money(p['long_total'], True)}，空头账面 {fmt_money(p['short_total'], True)}。"
        f"逐笔空头胜率 {fmt_pct(p['short_wr']['win_rate'], 0)}、盈亏比 {p['short_wr']['payoff']:.2f}；多头 {fmt_pct(p['long_wr']['win_rate'], 0)} / {p['long_wr']['payoff']:.2f}。",
        "手续费占权益很低，策略容量的第一约束不是佣金，而是因子拥挤和冲击成本。",
    ]
    for b in goods:
        bp = doc.add_paragraph(style="List Bullet")
        add_text(bp, b, size=11, color=TEXT)

    heading(doc, "不擅长 / 需要打折的预期", 2)
    bads = [
        f"{perf_blurb}{p['n_days']} 个交易日累计 {fmt_pct(p['period_ret'], 2, True)}、夏普 {p['sharpe']:.2f}。把它当每年稳定缴税会失望。",
        f"短板板块 {bot_sec_n or '—'}。主要失血品种：{bot_names or '—'}。"
        + ("同一板块也能打脸：白银赚、黄金亏。" if metal_split else "")
        + ("股指上小盘赚、大盘亏。" if index_split else ""),
        ("周一累计最差，隔夜/周末跳空可能是结构弱点。" if monday_worst else "星期效应见附录，用来核对隔夜跳空有没有集中在某一天结算。"),
        "盘中事件驱动通常不是这类开盘批量系统的舒适区。",
        f"逐笔期望约 {fmt_money(p['close_stats']['expectancy'], True)} 元/笔。靠广度堆期望时，拥挤或冲击变大先被抹平。",
        "连亏后若占用几乎不动：好处是不恐慌，坏处是模型失效时不会自己踩刹车。失效要靠投后监控。",
    ]
    for b in bads:
        bp = doc.add_paragraph(style="List Bullet")
        add_text(bp, b, size=11, color=TEXT)

    heading(doc, "配置含义（不是投资建议）", 2)
    para(
        doc,
        f"若把 {acc} 放进 MOM：先看他是低波动截面袖子还是进攻引擎。这段样本波动 {fmt_pct(p['vol'], 1)}、回撤 {fmt_pct(p['mdd'], 2)}、夏普 {p['sharpe']:.2f}。"
        f"与南华商品指数日收益相关 {nh.get('corr', 0):.2f}，空头保证金占 {fmt_pct(p['avg_short_share'], 0)}。"
        "投后看：保证金占用有没有突然抬升、对冲度有没有塌成单边、短板板块亏损有没有从摩擦变成主亏。",
    )

    wd_rows = [[r["name"], str(int(r["n"])), fmt_money(r["pnl"], True), fmt_money(r["avg"], True), fmt_pct(r["wr"], 0)] for r in p["wd_rows"]]
    heading(doc, "十一、附录：星期效应与方法", 1)
    add_table(doc, ["星期", "交易日", "累计盈亏", "日均盈亏", "胜率"], wd_rows)
    caption(doc, "表11  星期效应。若某一天显著更差，多半是隔夜跳空集中在那天结算。")

    notes = [
        f"数据：mom_daily_reports、mom_futures_trade_details（{p['n_trades']:,} 笔）、mom_close_details（{p['n_closes']:,} 笔）、mom_position_details、raw_nanhua_indices_daily。账户 {p['account']}。",
        "品种盈亏 = 当日该品种成交平仓盈亏 + 持仓盯市盈亏。与日报「当日盈亏」在手续费、期权、出入金上可能有残差。",
        "持仓天数用平仓明细开仓日期。开仓日期在库中为科学计数法 YYYYMMDD，已还原。交易日持仓按账户日报日历对齐。",
        "对冲度用多头保证金与空头保证金，而不是手数（手数跨品种不可比）。",
        "动量对齐：品种结算价近 5/20 日收益的符号，与次日净持仓手数符号比较。这是简化探针，不是完整因子归因。",
        "夜盘实现盈亏不含隔夜盯市，会低估夜盘经济贡献；下单时钟与开盘集中度不受此影响。",
        "本报告是行为推断，不能替代尽调。不构成投顾评级或调仓指令。",
    ]
    for n in notes:
        bp = doc.add_paragraph(style="List Number")
        add_text(bp, n, size=10, color=TEXT)

    para(doc, f"生成日期 {date.today().isoformat()}。画像对象：{acc} {p['advisor_name'] or '未登记'} / {p['company'] or '未登记'}。", size=9, color=MUTED, space_before=14)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    if ascii_path is not None:
        ascii_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(ascii_path))


def json_ready(p: dict) -> dict:
    skip = {"curve", "hedge_ts", "daily_rows", "nhci_rows", "hold_win", "hold_loss", "close_pnl_sample", "prod_rows"}
    out = {k: v for k, v in p.items() if k not in skip}
    out["n_prod_rows"] = len(p.get("prod_rows") or [])
    out["curve_end"] = p["curve"][-1] if p.get("curve") else None
    return out


def parse_args():
    import argparse
    ap = argparse.ArgumentParser(description="Generate a trader-profile Word report for one MOM account")
    ap.add_argument("--account", default=os.environ.get("PROFILE_ACCOUNT", "rx319"))
    ap.add_argument("--from-date", default=os.environ.get("PROFILE_FROM") or None)
    ap.add_argument("--to-date", default=os.environ.get("PROFILE_TO") or None)
    ap.add_argument("--output", default=os.environ.get("PROFILE_OUTPUT") or None)
    ap.add_argument("--work-dir", default=os.environ.get("PROFILE_WORK_DIR") or None)
    return ap.parse_args()


def main() -> None:
    global OUT_DIR, CHART_DIR, DATA_PATH, ACCOUNT, AS_OF
    args = parse_args()
    account = str(args.account or "rx319").lower().strip()
    ACCOUNT = account
    from_date = str(args.from_date).strip() if args.from_date else None
    to_date = str(args.to_date).strip() if args.to_date else None
    if from_date == "":
        from_date = None
    if to_date == "":
        to_date = None
    AS_OF = to_date or date.today().isoformat()

    work = Path(args.work_dir) if args.work_dir else OUT_DIR
    OUT_DIR = work
    CHART_DIR = work / "charts"
    DATA_PATH = work / "profile_data.json"
    CHART_DIR.mkdir(parents=True, exist_ok=True)

    configure_matplotlib()
    conn = get_conn()
    try:
        print("loading…", account, from_date, to_date)
        raw = load_raw(conn, account, from_date, to_date)
    finally:
        conn.close()

    print("analyzing…")
    profile, daily, trades, closes, pos, prod_sum, sec_sum, hedge_daily, nh, after_df = analyze(raw, account)
    DATA_PATH.write_text(json.dumps(json_ready(profile), ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    print("wrote", DATA_PATH)
    print(
        f"days={profile['n_days']} trades={profile['n_trades']} ret={profile['period_ret']*100:.2f}% "
        f"sharpe={profile['sharpe']:.2f} mdd={profile['mdd']*100:.2f}% wr_d={profile['daily_wr']['win_rate']*100:.1f}% "
        f"wr_c={profile['close_stats']['win_rate']*100:.1f}% payoff_c={profile['close_stats']['payoff']:.2f} "
        f"hedge={profile['avg_hedge']*100:.0f}% mom5={profile['mom5_hit']*100:.0f}% "
        f"hold_w={profile['med_hold_win']:.1f} hold_l={profile['med_hold_loss']:.1f}"
    )

    print("charts…")
    charts: dict[str, Path] = {}
    for key, fn in [
        ("equity", lambda: chart_equity(profile, nh)),
        ("dd", lambda: chart_dd(profile)),
        ("margin", lambda: chart_margin(profile)),
        ("monthly", lambda: chart_monthly(profile)),
        ("sector", lambda: chart_sector(sec_sum)),
        ("product", lambda: chart_product(prod_sum)),
        ("hold", lambda: chart_hold(profile)),
        ("hour", lambda: chart_hour(profile)),
        ("session", lambda: chart_session(profile)),
        ("after", lambda: chart_after(after_df)),
        ("hedge", lambda: chart_hedge(profile)),
        ("rr", lambda: chart_rr_hist(profile)),
        ("ls", lambda: chart_ls(profile)),
    ]:
        try:
            charts[key] = fn()
        except Exception as exc:
            print(f"chart {key} failed: {exc}")
    out_docx = Path(args.output) if args.output else (work / f"{account.upper()}_trader_profile.docx")
    ascii_path = None if args.output else REPORT_PATH_ASCII
    print("word…")
    write_report(profile, charts, out_docx, ascii_path)
    print("report", out_docx)
    if ascii_path is not None:
        print("report", ascii_path)


if __name__ == "__main__":
    main()
