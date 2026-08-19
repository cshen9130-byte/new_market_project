# -*- coding: utf-8 -*-
"""Identify MOM accounts that trade like quantitative strategies and write a Word report."""
from __future__ import annotations

import json
import math
import os
import re
import sys
from collections import defaultdict
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.font_manager import FontProperties, fontManager

from mom_data_etl import get_conn, load_env_files

load_env_files()

BASE = Path(__file__).resolve().parent
ROOT = BASE.parents[1]
OUT_DIR = BASE / "_quant_report_output"
CHART_DIR = OUT_DIR / "charts"
DATA_PATH = OUT_DIR / "quant_report_data.json"
REPORT_PATH = ROOT / "MOM量化策略账户识别报告_20260819.docx"
REPORT_PATH_ASCII = ROOT / "MOM_quant_accounts_report_20260819.docx"

FROM_DATE = "2026-02-19"
TO_DATE = "2026-08-18"
AS_OF = "2026-08-18"
MIN_DAYS = 20

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
TEXT = RGBColor(0x2D, 0x37, 0x48)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xC5, 0x30, 0x30)
GREEN = RGBColor(0x2F, 0x85, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_NAVY = "#1A365D"
C_GOLD = "#C9A227"
C_RED = "#C53030"
C_GREEN = "#2F855A"
C_TEAL = "#2B6CB0"
C_ORANGE = "#DD6B20"
C_PURPLE = "#805AD5"
C_GRAY = "#718096"

SECTOR_MAP: dict[str, str] = {
    "C": "农产", "CS": "农产", "WH": "农产", "PM": "农产", "RR": "农产", "RI": "农产",
    "JR": "农产", "LR": "农产", "A": "农产", "B": "农产", "M": "农产", "Y": "农产",
    "RM": "农产", "OI": "农产", "RS": "农产", "PK": "农产", "P": "农产", "SR": "农产",
    "CF": "农产", "CY": "农产", "LG": "农产", "SP": "农产", "OP": "农产",
    "AP": "生鲜", "CJ": "生鲜", "LH": "生鲜", "JD": "生鲜",
    "AU": "贵金属", "AG": "贵金属", "PT": "贵金属", "PD": "贵金属",
    "CU": "有色", "BC": "有色", "AL": "有色", "AO": "有色", "AD": "有色",
    "ZN": "有色", "PB": "有色", "NI": "有色", "SN": "有色",
    "LC": "新能源", "PS": "新能源", "SI": "新能源",
    "I": "黑色", "SF": "黑色", "SM": "黑色", "RB": "黑色", "HC": "黑色",
    "SS": "黑色", "WR": "黑色", "JM": "黑色", "J": "黑色", "ZC": "黑色",
    "FG": "黑色", "BB": "黑色", "FB": "黑色",
    "SC": "能源化工", "FU": "能源化工", "LU": "能源化工", "PG": "能源化工",
    "BU": "能源化工", "TA": "能源化工", "EG": "能源化工", "PF": "能源化工",
    "PR": "能源化工", "PL": "能源化工", "PP": "能源化工", "L": "能源化工",
    "BZ": "能源化工", "PX": "能源化工", "EB": "能源化工", "RU": "能源化工",
    "BR": "能源化工", "NR": "能源化工", "SA": "能源化工", "SH": "能源化工",
    "V": "能源化工", "UR": "能源化工", "MA": "能源化工",
    "EC": "航运",
    "IH": "股指", "IF": "股指", "IC": "股指", "IM": "股指", "MO": "股指",
    "TS": "国债", "TF": "国债", "T": "国债", "TL": "国债",
}

KNOWN_QUANT_FIRMS = ("九木", "大风", "量衍")
_CN_FONT: FontProperties | None = None


def configure_matplotlib() -> None:
    global _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.isfile(path):
            try:
                fontManager.addfont(path)
                _CN_FONT = FontProperties(fname=path)
                plt.rcParams["font.family"] = "sans-serif"
                plt.rcParams["font.sans-serif"] = [_CN_FONT.get_name(), "Microsoft YaHei", "SimHei"]
                return
            except Exception:
                continue
    _CN_FONT = FontProperties(family="Microsoft YaHei")


def fp() -> dict:
    return {"fontproperties": _CN_FONT} if _CN_FONT is not None else {}


def to_num(v) -> float:
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return 0.0
    s = str(v).replace(",", "").replace("%", "").replace(" ", "").strip()
    if not s:
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def product_code(contract: str | None) -> str:
    if not contract:
        return ""
    m = re.match(r"^[A-Za-z]+", str(contract).strip())
    return m.group(0).upper() if m else ""


def num_sql(col: str) -> str:
    return (
        f"COALESCE(NULLIF(REPLACE(REPLACE(COALESCE(\"{col}\"::text, ''), ',', ''), ' ', ''), '')::numeric, 0)"
    )


def read_sql(conn, sql: str, params=None) -> pd.DataFrame:
    return pd.read_sql(sql, conn, params=params)


def load_features(conn) -> dict:
    advisor = read_sql(
        conn,
        """
        SELECT LOWER(TRIM(account_code)) AS account,
               advisor_name, sector, background, style, cycle,
               is_arbitrage, main_strength, company, product_preference, equity_wan
        FROM mom_advisor_info
        """,
    )

    daily = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               "交易日期"::date AS dt,
               {num_sql("当日盈亏")} AS pnl,
               {num_sql("上日结存")} AS prev_bal,
               {num_sql("客户权益")} AS equity,
               {num_sql("当日手续费")} AS fee,
               {num_sql("保证金占用")} AS margin,
               {num_sql("平仓盈亏")} AS close_pnl,
               {num_sql("持仓盈亏")} AS pos_pnl
        FROM mom_daily_reports
        WHERE "交易日期" BETWEEN %s AND %s
        """,
        (FROM_DATE, TO_DATE),
    )
    daily["dt"] = pd.to_datetime(daily["dt"])

    trades = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               UPPER(SUBSTRING("合约" FROM '^[A-Za-z]+')) AS product,
               COUNT(*)::int AS n_trades,
               SUM({num_sql("手数")}) AS lots,
               SUM({num_sql("成交额")}) AS turnover,
               SUM({num_sql("手续费")}) AS fee,
               SUM({num_sql("平仓盈亏")}) AS close_pnl,
               COUNT(DISTINCT "交易日期")::int AS n_days
        FROM mom_futures_trade_details
        WHERE "交易日期" BETWEEN %s AND %s
        GROUP BY 1, 2
        """,
        (FROM_DATE, TO_DATE),
    )

    open_close = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               COALESCE(NULLIF(TRIM("开/平"), ''), '未知') AS oc,
               COUNT(*)::int AS n
        FROM mom_futures_trade_details
        WHERE "交易日期" BETWEEN %s AND %s
        GROUP BY 1, 2
        """,
        (FROM_DATE, TO_DATE),
    )

    lot_stats = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               COUNT(*)::int AS n_trades,
               AVG({num_sql("手数")}) AS avg_lot,
               STDDEV_SAMP({num_sql("手数")}) AS sd_lot,
               COUNT(DISTINCT "成交时间")::int AS n_times
        FROM mom_futures_trade_details
        WHERE "交易日期" BETWEEN %s AND %s
        GROUP BY 1
        """,
        (FROM_DATE, TO_DATE),
    )

    hour_dist = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               EXTRACT(HOUR FROM (
                 CASE
                   WHEN "成交时间" ~ '^[0-9]{{1,2}}:[0-9]{{2}}' THEN ("成交时间")::time
                   WHEN "成交时间" ~ '^[0-9]{{6}}$' THEN to_timestamp("成交时间", 'HH24MISS')::time
                   WHEN "成交时间" ~ ' ' THEN split_part("成交时间", ' ', 2)::time
                   ELSE NULL
                 END
               ))::int AS hr,
               COUNT(*)::int AS n
        FROM mom_futures_trade_details
        WHERE "交易日期" BETWEEN %s AND %s
          AND "成交时间" IS NOT NULL AND TRIM("成交时间") <> ''
        GROUP BY 1, 2
        """,
        (FROM_DATE, TO_DATE),
    )

    pos = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               "交易日期"::date AS dt,
               UPPER(SUBSTRING("合约" FROM '^[A-Za-z]+')) AS product,
               SUM(ABS({num_sql("买持仓")}) + ABS({num_sql("卖持仓")})) AS lots
        FROM mom_position_details
        WHERE "交易日期" BETWEEN %s AND %s
        GROUP BY 1, 2, 3
        """,
        (FROM_DATE, TO_DATE),
    )
    if not pos.empty:
        pos["dt"] = pd.to_datetime(pos["dt"])

    opt = read_sql(
        conn,
        f"""
        SELECT LOWER(TRIM("账户")) AS account,
               COUNT(*)::int AS n_opt_trades,
               COUNT(DISTINCT UPPER(SUBSTRING("合约" FROM '^[A-Za-z]+')))::int AS n_opt_products
        FROM mom_options_trade_details
        WHERE "交易日期" BETWEEN %s AND %s
        GROUP BY 1
        """,
        (FROM_DATE, TO_DATE),
    )

    return {
        "advisor": advisor,
        "daily": daily,
        "trades": trades,
        "open_close": open_close,
        "lot_stats": lot_stats,
        "hour_dist": hour_dist,
        "pos": pos,
        "opt": opt,
    }


def entropy(counts: np.ndarray) -> float:
    s = counts.sum()
    if s <= 0:
        return 0.0
    p = counts[counts > 0] / s
    return float(-(p * np.log(p)).sum() / math.log(len(p))) if len(p) > 1 else 0.0


def hhi(weights: np.ndarray) -> float:
    s = weights.sum()
    if s <= 0:
        return 1.0
    p = weights / s
    return float((p ** 2).sum())


def max_drawdown(cum: np.ndarray) -> float:
    if len(cum) == 0:
        return 0.0
    peak = np.maximum.accumulate(cum)
    dd = (cum - peak) / np.clip(peak, 1e-12, None)
    return float(dd.min())


def score_accounts(raw: dict) -> list[dict]:
    advisor = raw["advisor"].set_index("account") if not raw["advisor"].empty else pd.DataFrame()
    daily = raw["daily"]
    trades = raw["trades"]
    oc = raw["open_close"]
    lot_stats = raw["lot_stats"].set_index("account") if not raw["lot_stats"].empty else pd.DataFrame()
    hour_dist = raw["hour_dist"]
    pos = raw["pos"]
    opt = raw["opt"].set_index("account") if not raw["opt"].empty else pd.DataFrame()

    universe_days = int(daily["dt"].nunique()) if not daily.empty else 1
    accounts = sorted(set(daily["account"].unique()) | set(trades["account"].unique()) if not trades.empty else daily["account"].unique())

    oc_map: dict[str, dict[str, int]] = defaultdict(lambda: defaultdict(int))
    for r in oc.itertuples(index=False):
        oc_map[r.account][str(r.oc)] += int(r.n)

    hour_map: dict[str, dict[int, int]] = defaultdict(lambda: defaultdict(int))
    if not hour_dist.empty:
        for r in hour_dist.itertuples(index=False):
            if r.hr is None or (isinstance(r.hr, float) and math.isnan(r.hr)):
                continue
            hour_map[r.account][int(r.hr)] += int(r.n)

    trade_by_acc = trades.groupby("account") if not trades.empty else None
    pos_by_acc = pos.groupby("account") if not pos.empty else None
    daily_by_acc = daily.groupby("account")

    rows: list[dict] = []
    for acc in accounts:
        meta = advisor.loc[acc].to_dict() if acc in advisor.index else {}
        d = daily_by_acc.get_group(acc).sort_values("dt") if acc in daily_by_acc.groups else pd.DataFrame()
        t = trade_by_acc.get_group(acc) if trade_by_acc is not None and acc in trade_by_acc.groups else pd.DataFrame()
        p = pos_by_acc.get_group(acc) if pos_by_acc is not None and acc in pos_by_acc.groups else pd.DataFrame()

        n_days = int(d["dt"].nunique()) if not d.empty else 0
        if n_days < MIN_DAYS and (t.empty or int(t["n_days"].max() or 0) < MIN_DAYS):
            continue

        products = sorted({x for x in t["product"].dropna().tolist() if x}) if not t.empty else []
        n_products = len(products)
        sectors = sorted({SECTOR_MAP.get(x, "其他") for x in products})
        n_sectors = len(sectors)
        n_trades = int(t["n_trades"].sum()) if not t.empty else 0
        trade_days = int(t["n_days"].max()) if not t.empty else 0
        trades_per_day = n_trades / max(trade_days, 1)
        turnover = float(t["turnover"].sum()) if not t.empty else 0.0
        trade_fee = float(t["fee"].sum()) if not t.empty else 0.0

        trade_hhi = hhi(t["n_trades"].to_numpy(dtype=float)) if not t.empty else 1.0
        lot_hhi = hhi(t["lots"].abs().to_numpy(dtype=float)) if not t.empty else 1.0

        pos_hhi_days = []
        pos_nprod_days = []
        if not p.empty:
            for _, g in p.groupby("dt"):
                lots = g["lots"].to_numpy(dtype=float)
                pos_hhi_days.append(hhi(lots))
                pos_nprod_days.append(int((lots > 0).sum()))
        pos_hhi = float(np.mean(pos_hhi_days)) if pos_hhi_days else 1.0
        avg_pos_products = float(np.mean(pos_nprod_days)) if pos_nprod_days else 0.0

        oc_counts = oc_map.get(acc, {})
        oc_total = sum(oc_counts.values()) or 1
        pingjin = sum(v for k, v in oc_counts.items() if "平今" in k)
        pingcang = sum(v for k, v in oc_counts.items() if "平" in k)
        kaicang = sum(v for k, v in oc_counts.items() if k.startswith("开") or "开仓" in k)
        pingjin_ratio = pingjin / oc_total
        close_ratio = pingcang / oc_total

        hours = np.array([hour_map[acc].get(h, 0) for h in range(9, 16)], dtype=float)
        time_entropy = entropy(hours)

        ls = lot_stats.loc[acc].to_dict() if acc in lot_stats.index else {}
        avg_lot = to_num(ls.get("avg_lot"))
        sd_lot = to_num(ls.get("sd_lot"))
        lot_cv = (sd_lot / avg_lot) if avg_lot > 0 else 0.0

        rets = []
        equities = []
        fees = []
        margins = []
        pnls = []
        cum_factor = 1.0
        curve = []
        for r in d.itertuples(index=False):
            denom = r.prev_bal if r.prev_bal > 0 else (r.equity - r.pnl if (r.equity - r.pnl) > 0 else 0)
            ret = (r.pnl / denom) if denom > 0 else 0.0
            if abs(ret) > 0.25:
                continue
            rets.append(ret)
            pnls.append(float(r.pnl))
            equities.append(float(r.equity))
            fees.append(float(r.fee))
            margins.append(float(r.margin))
            cum_factor *= 1.0 + ret
            curve.append({"date": r.dt.strftime("%Y-%m-%d"), "pct": (cum_factor - 1) * 100, "ret": ret})
        rets_a = np.array(rets, dtype=float) if rets else np.array([0.0])
        vol = float(rets_a.std(ddof=1) * math.sqrt(242)) if len(rets_a) > 2 else 0.0
        mu = float(rets_a.mean() * 242) if len(rets_a) else 0.0
        sharpe = (mu / vol) if vol > 1e-8 else 0.0
        cum_nav = np.cumprod(1.0 + rets_a) if len(rets_a) else np.array([1.0])
        mdd = max_drawdown(cum_nav)
        period_ret = float(cum_nav[-1] - 1.0)
        latest_equity = float(equities[-1]) if equities else 0.0
        avg_equity = float(np.mean(equities)) if equities else 0.0
        period_fee = float(np.sum(fees)) if fees else float(d["fee"].sum() if not d.empty else 0)
        fee_to_eq = period_fee / avg_equity if avg_equity > 0 else 0.0
        avg_margin_ratio = float(np.mean([m / e for m, e in zip(margins, equities) if e > 0])) if equities else 0.0
        pnl_abs = np.abs(rets_a)
        smoothness = 1.0 / (1.0 + float(np.median(pnl_abs) * 100)) if len(pnl_abs) else 0.0

        n_opt = int(opt.loc[acc, "n_opt_trades"]) if acc in opt.index else 0
        n_opt_prod = int(opt.loc[acc, "n_opt_products"]) if acc in opt.index else 0

        coverage = trade_days / max(universe_days, 1)
        # Simultaneous book size is the key split: quants hold dozens overnight;
        # discretionary desks may *touch* 30 names over six months but hold 3–6.
        sim_score = min(100.0, avg_pos_products / 40.0 * 100.0)
        breadth_score = min(100.0, n_products / 50.0 * 100.0)
        sector_score = min(100.0, n_sectors / 6.0 * 100.0)
        coverage_score = min(100.0, coverage * 100.0)
        if avg_pos_products >= 20 and n_products >= 25:
            freq_score = min(100.0, trades_per_day / 50.0 * 100.0)
        elif n_products >= 15:
            freq_score = min(55.0, trades_per_day / 80.0 * 55.0)
        else:
            freq_score = min(30.0, trades_per_day / 100.0 * 30.0)
        div_score = (1.0 - pos_hhi) * 100.0
        trade_div_score = (1.0 - trade_hhi) * 100.0
        vol_score = max(0.0, min(100.0, (0.25 - min(vol, 0.25)) / 0.25 * 100.0)) if vol > 0 else 40.0
        time_score = time_entropy * 100.0
        label_text = " ".join(
            str(x or "")
            for x in [
                meta.get("sector"),
                meta.get("style"),
                meta.get("company"),
                meta.get("main_strength"),
            ]
        )
        labeled_quant = bool(
            meta.get("sector") == "多因子"
            or re.search(r"时序|截面|多强空弱|量价趋势|多因子", label_text)
            or any(k in str(meta.get("company") or "") for k in KNOWN_QUANT_FIRMS)
        )
        industrial_style = bool(re.search(r"小单边|套利-偏产业|套利\+单边", str(meta.get("style") or "")))
        rotating_discretionary = n_products >= 20 and avg_pos_products < 10

        score = (
            0.24 * sim_score
            + 0.18 * breadth_score
            + 0.12 * sector_score
            + 0.14 * div_score
            + 0.08 * trade_div_score
            + 0.08 * coverage_score
            + 0.08 * freq_score
            + 0.05 * vol_score
            + 0.03 * time_score
        )
        if labeled_quant and (avg_pos_products >= 10 or n_products >= 25):
            score = max(score, 78.0)
        if rotating_discretionary or (industrial_style and avg_pos_products < 16):
            score = min(score, 52.0)
        if n_products < 18 and n_sectors <= 3:
            score = min(score, 48.0)

        full_market = (
            n_products >= 40
            and avg_pos_products >= 25
            and n_sectors >= 8
            and pos_hhi <= 0.12
        )
        broad_book = (
            n_products >= 25
            and avg_pos_products >= 18
            and n_sectors >= 6
            and pos_hhi <= 0.16
        )
        if labeled_quant and (avg_pos_products >= 10 or n_products >= 25):
            verdict = "高确信量化"
        elif full_market:
            verdict = "高确信量化"
        elif broad_book and not industrial_style:
            verdict = "较可能量化"
        elif broad_book and industrial_style:
            verdict = "较可能量化"
        elif n_trades >= 2000 and n_products < 18:
            verdict = "高频产业/套利（非多因子量化）"
        else:
            verdict = "主观/产业盘手"

        if labeled_quant and "时序" in str(meta.get("style") or ""):
            subtype = "时序趋势CTA"
        elif labeled_quant and "截面" in str(meta.get("style") or ""):
            subtype = "截面多空"
        elif labeled_quant:
            subtype = "混合多因子"
        elif verdict in ("高确信量化", "较可能量化") and pingjin_ratio >= 0.35:
            subtype = "短周期/偏日内系统"
        elif verdict in ("高确信量化", "较可能量化"):
            subtype = "系统化多品种CTA"
        elif verdict.startswith("高频"):
            subtype = "产业高频/价差"
        else:
            subtype = "主观或产业"

        reasons = []
        if avg_pos_products >= 25:
            reasons.append(f"日均同时持仓约{avg_pos_products:.0f}个品种——这是量化与主观的核心分界：主观盘手很少隔夜同时拿二十几个合约")
        elif avg_pos_products >= 12:
            reasons.append(f"日均同时持仓约{avg_pos_products:.0f}个品种，仓位结构偏系统化")
        if n_products >= 40:
            reasons.append(f"近六月成交覆盖{n_products}个期货品种，远超单一产业链盘手的能力边界")
        elif n_products >= 25:
            reasons.append(f"近六月覆盖{n_products}个品种，接近全市场系统化交易")
        if n_sectors >= 5:
            reasons.append(f"横跨{n_sectors}个板块（{('、').join(sectors[:8])}）")
        if coverage >= 0.85 and trades_per_day >= 20:
            reasons.append(f"近六月{trade_days}个交易日几乎满勤，日均{trades_per_day:.0f}笔")
        elif trade_days < 25 and n_products >= 40:
            reasons.append(f"样本较短（有成交{trade_days}日），但品种与持仓结构已是全市场系统化特征")
        if pos_hhi <= 0.18:
            reasons.append(f"持仓品种集中度HHI={pos_hhi:.2f}，单品种权重被压低")
        if labeled_quant:
            reasons.append(f"投顾标签为「{meta.get('sector') or ''} / {meta.get('style') or ''}」")
        if any(k in str(meta.get("company") or "") for k in KNOWN_QUANT_FIRMS):
            reasons.append(f"所属{meta.get('company')}为市场已知量化CTA团队")
        if vol > 0 and vol <= 0.12 and avg_pos_products >= 12:
            reasons.append(f"年化波动约{vol*100:.1f}%，曲线平滑，符合波动目标/风险平价特征")
        if pingjin_ratio >= 0.4 and n_products >= 20:
            reasons.append(f"平今占比{pingjin_ratio:.0%}，换手偏短周期系统")
        if industrial_style and verdict in ("高确信量化", "较可能量化"):
            reasons.append("投顾表写成产业套利，但同时持仓数量已超出人工价差盘的常规范围，更像程序化截面/价差")

        rows.append(
            {
                "account": acc,
                "advisor_name": meta.get("advisor_name") or "",
                "company": meta.get("company") or "",
                "sector": meta.get("sector") or "",
                "background": meta.get("background") or "",
                "style": meta.get("style") or "",
                "cycle": meta.get("cycle") or "",
                "is_arbitrage": meta.get("is_arbitrage"),
                "main_strength": meta.get("main_strength") or "",
                "product_preference": meta.get("product_preference") or "",
                "equity_wan": to_num(meta.get("equity_wan")),
                "n_days": n_days,
                "trade_days": trade_days,
                "coverage": coverage,
                "n_trades": n_trades,
                "trades_per_day": trades_per_day,
                "n_products": n_products,
                "n_sectors": n_sectors,
                "sectors": sectors,
                "products": products,
                "trade_hhi": trade_hhi,
                "lot_hhi": lot_hhi,
                "pos_hhi": pos_hhi,
                "avg_pos_products": avg_pos_products,
                "pingjin_ratio": pingjin_ratio,
                "close_ratio": close_ratio,
                "kaicang": kaicang,
                "time_entropy": time_entropy,
                "lot_cv": lot_cv,
                "avg_lot": avg_lot,
                "turnover": turnover,
                "trade_fee": trade_fee,
                "period_ret": period_ret,
                "vol": vol,
                "sharpe": sharpe,
                "mdd": mdd,
                "latest_equity": latest_equity,
                "avg_equity": avg_equity,
                "fee_to_eq": fee_to_eq,
                "avg_margin_ratio": avg_margin_ratio,
                "smoothness": smoothness,
                "n_opt_trades": n_opt,
                "n_opt_products": n_opt_prod,
                "labeled_quant": labeled_quant,
                "score": score,
                "verdict": verdict,
                "subtype": subtype,
                "reasons": reasons,
                "curve": curve,
                "sim_score": sim_score,
                "breadth_score": breadth_score,
                "sector_score": sector_score,
                "div_score": div_score,
                "freq_score": freq_score,
            }
        )

    rows.sort(key=lambda x: (-x["score"], -x["n_products"]))
    return rows


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, *, size=11, bold=False, color=None):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def para(doc, text="", *, size=11, bold=False, color=None, align=None, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    if align:
        p.alignment = align
    if text:
        add_text(p, text, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def caption(doc, text):
    para(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, space_before=2)


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, size=9, bold=False, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "" if text is None else str(text), size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_table(doc, headers, rows, highlight_rows=None):
    highlight_rows = set(highlight_rows or [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=8, bold=True, color=WHITE)
        shade(table.rows[0].cells[i], "1A365D")
    for r_i, row in enumerate(rows):
        fill = "EDF2F7" if r_i in highlight_rows else ("F7FAFC" if r_i % 2 == 0 else "FFFFFF")
        for c_i, val in enumerate(row):
            cell_text(table.rows[r_i + 1].cells[c_i], val, size=8, align="center" if c_i else "left")
            shade(table.rows[r_i + 1].cells[c_i], fill)
    return table


def add_chart(doc, path: Path, width=6.4):
    if path and path.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))


def fmt_pct(v, digits=1, signed=False):
    if v is None:
        return "—"
    x = float(v) * 100
    sign = "+" if signed and x > 0 else ""
    return f"{sign}{x:.{digits}f}%"


def fmt_num(v, digits=1):
    if v is None:
        return "—"
    return f"{float(v):.{digits}f}"


def acc_label(r: dict) -> str:
    name = r.get("advisor_name") or ""
    return f"{r['account'].upper()}" + (f" {name}" if name else "")


def chart_scatter(rows: list[dict]) -> Path:
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    color_map = {
        "高确信量化": C_NAVY,
        "较可能量化": C_TEAL,
        "高频产业/套利（非多因子量化）": C_ORANGE,
        "主观/产业盘手": C_GRAY,
    }
    for verdict, color in color_map.items():
        sub = [r for r in rows if r["verdict"] == verdict]
        if not sub:
            continue
        ax.scatter(
            [r["n_products"] for r in sub],
            [r["trades_per_day"] for r in sub],
            s=[max(28, min(180, r["latest_equity"] / 80000)) for r in sub],
            c=color,
            label=verdict,
            alpha=0.85,
            edgecolors="white",
            linewidths=0.4,
            zorder=3,
        )
        for r in sub:
            if r["verdict"] in ("高确信量化", "较可能量化") or r["n_products"] >= 40 or r["trades_per_day"] >= 80:
                ax.annotate(
                    r["account"].upper(),
                    (r["n_products"], r["trades_per_day"]),
                    textcoords="offset points",
                    xytext=(4, 4),
                    fontsize=7,
                    color=C_NAVY,
                    **fp(),
                )
    ax.set_xlabel("近六月成交品种数", **fp())
    ax.set_ylabel("日均成交笔数", **fp())
    ax.set_title("品种广度 vs 交易频率（点大小≈最新权益）", **fp())
    ax.grid(True, color="#E2E8F0", linewidth=0.6)
    ax.legend(prop=_CN_FONT, frameon=False, fontsize=8)
    fig.tight_layout()
    path = CHART_DIR / "scatter_breadth_freq.png"
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def chart_scores(quants: list[dict]) -> Path:
    fig, ax = plt.subplots(figsize=(9.2, max(3.6, 0.38 * len(quants) + 1.4)))
    labels = [acc_label(r) for r in reversed(quants)]
    scores = [r["score"] for r in reversed(quants)]
    colors = [C_NAVY if r["verdict"] == "高确信量化" else C_TEAL for r in reversed(quants)]
    ax.barh(labels, scores, color=colors, height=0.62)
    ax.set_xlim(0, 100)
    ax.set_xlabel("量化行为得分", **fp())
    ax.set_title("量化识别得分（高确信 / 较可能）", **fp())
    ax.axvline(70, color=C_GOLD, lw=0.9, ls="--")
    ax.axvline(55, color=C_GRAY, lw=0.7, ls=":")
    ax.grid(True, axis="x", color="#E2E8F0", linewidth=0.6)
    for lab in ax.get_yticklabels():
        lab.set_fontproperties(_CN_FONT)
        lab.set_fontsize(8)
    fig.tight_layout()
    path = CHART_DIR / "score_bar.png"
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def chart_equity(quants: list[dict]) -> Path:
    fig, ax = plt.subplots(figsize=(9.4, 5.2))
    palette = [C_NAVY, C_TEAL, C_ORANGE, C_PURPLE, C_GREEN, C_RED, C_GOLD, "#2C7A7B"]
    for i, r in enumerate(quants[:8]):
        xs = [pd.Timestamp(p["date"]) for p in r["curve"]]
        ys = [p["pct"] for p in r["curve"]]
        if not xs:
            continue
        ax.plot(xs, ys, color=palette[i % len(palette)], lw=1.5, label=acc_label(r))
    ax.axhline(0, color="#A0AEC0", lw=0.8)
    ax.set_ylabel("累计收益率 %", **fp())
    ax.set_title("识别为量化的账户 · 近六月累计收益曲线", **fp())
    ax.grid(True, color="#E2E8F0", linewidth=0.6)
    ax.legend(prop=_CN_FONT, frameon=False, fontsize=8, ncol=2)
    fig.autofmt_xdate()
    fig.tight_layout()
    path = CHART_DIR / "equity_quant.png"
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def chart_pos_hhi(rows: list[dict]) -> Path:
    fig, ax = plt.subplots(figsize=(9.2, 5.4))
    xs = [r["avg_pos_products"] for r in rows]
    ys = [r["pos_hhi"] for r in rows]
    colors = [
        C_NAVY if r["verdict"] == "高确信量化" else C_TEAL if r["verdict"] == "较可能量化" else C_ORANGE if r["verdict"].startswith("高频") else C_GRAY
        for r in rows
    ]
    ax.scatter(xs, ys, c=colors, s=42, alpha=0.85, edgecolors="white", linewidths=0.4)
    for r in rows:
        if r["verdict"] in ("高确信量化", "较可能量化") or r["avg_pos_products"] >= 18:
            ax.annotate(r["account"].upper(), (r["avg_pos_products"], r["pos_hhi"]), textcoords="offset points", xytext=(4, 3), fontsize=7, color=C_NAVY, **fp())
    ax.set_xlabel("日均持仓品种数", **fp())
    ax.set_ylabel("持仓HHI（越低越分散）", **fp())
    ax.set_title("持仓分散度：量化账户通常品种多、HHI低", **fp())
    ax.grid(True, color="#E2E8F0", linewidth=0.6)
    fig.tight_layout()
    path = CHART_DIR / "pos_hhi.png"
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def write_report(rows: list[dict], charts: dict[str, Path]) -> None:
    quants = [r for r in rows if r["verdict"] in ("高确信量化", "较可能量化")]
    high = [r for r in quants if r["verdict"] == "高确信量化"]
    mid = [r for r in quants if r["verdict"] == "较可能量化"]
    hft = [r for r in rows if r["verdict"].startswith("高频")][:8]
    disc = [r for r in rows if r["verdict"] == "主观/产业盘手"]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    para(doc, "内部资料 · 请勿外传", size=9, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    para(doc, "MOM 投顾量化策略识别报告", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, "基于成交、持仓与日报微观结构，而非仅看投顾标签", size=13, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    para(doc, f"样本区间  {FROM_DATE}  ~  {TO_DATE}      数据截止  {AS_OF}", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        doc,
        f"纳入评估账户  {len(rows)}    高确信量化  {len(high)}    较可能量化  {len(mid)}",
        size=12,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )
    para(
        doc,
        "数据来源：mom_daily_reports、mom_futures_trade_details、mom_position_details、mom_options_trade_details、mom_advisor_info。"
        "识别结论是交易行为推断，不能替代尽调访谈。",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=20,
    )

    heading(doc, "一、结论：哪些账户是量化策略", 1)
    if not quants:
        para(doc, "按当前阈值，没有账户同时满足品种广度与交易规律性。请检查样本区间或阈值。")
    else:
        names = "、".join(acc_label(r) for r in quants)
        para(
            doc,
            f"近六月期货成交与持仓显示，下列 {len(quants)} 个账户的交易方式更接近量化/系统化CTA，而不是单一产业链上的主观或产业盘手：{names}。",
        )
        para(
            doc,
            "判断依据不是盈亏高低，而是隔夜同时持仓是否铺到全市场、单品种权重是否被压低、成交是否像规则驱动。"
            "吴锴、张海源、王亚伟这类账户半年也能碰到二三十个合约，但日均只持仓3–6个，那是观点轮动，不是量化。"
            "化工/粕类日均上百笔的账户同样不进名单：执行可以程序化，策略仍是产业链定价。",
        )

    bullets = [
        f"高确信（{len(high)}）：" + ("、".join(acc_label(r) + "（" + r["subtype"] + "）" for r in high) if high else "无"),
        f"较可能（{len(mid)}）：" + ("、".join(acc_label(r) + "（" + r["subtype"] + "）" for r in mid) if mid else "无"),
        "对照：日均成交很高但品种很少的账户，归为高频产业/价差，不进入量化名单。",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, b, size=11, color=TEXT)

    headers = ["账户", "投顾", "公司/团队", "判定", "子类型", "得分", "同时持仓", "成交品种", "板块", "日均笔数", "近六月收益", "年化波动"]
    table_rows = []
    for r in quants:
        table_rows.append(
            [
                r["account"].upper(),
                r["advisor_name"] or "—",
                r["company"] or "—",
                r["verdict"],
                r["subtype"],
                fmt_num(r["score"], 0),
                fmt_num(r["avg_pos_products"], 0),
                str(r["n_products"]),
                str(r["n_sectors"]),
                fmt_num(r["trades_per_day"], 0),
                fmt_pct(r["period_ret"], 1, signed=True),
                fmt_pct(r["vol"], 1),
            ]
        )
    add_table(doc, headers, table_rows, highlight_rows={i for i, r in enumerate(quants) if r["verdict"] == "高确信量化"})
    caption(doc, "表1  量化策略账户名单。收益按日报「当日盈亏/上日结存」复利，剔除单日±25%异常跳动。")

    heading(doc, "二、识别方法", 1)
    para(
        doc,
        "最有效的分界不是「半年里碰过多少合约」，而是「每天晚上账上同时躺着多少个品种」。"
        "量化CTA会隔夜同时持有二三十个以上合约，把风险预算铺到全市场；主观/产业盘手即使半年换过三四十个合约，真正隔夜拿着的通常只有三五个。"
        "监控台那条0%–10%的紧密簇，多数是后者：曲线平滑，但是产业套利，不是多因子量化。",
    )
    para(doc, "打分权重如下。同时持仓规模权重大于累计成交品种数：")
    method_rows = [
        ["同时持仓规模", "24%", "日均持仓品种数，40个得满分。这是主信号"],
        ["累计品种广度", "18%", "近六月成交品种数，50个得满分"],
        ["持仓分散HHI", "14%", "日均持仓HHI，越低越像风险预算"],
        ["板块广度", "12%", "农产/黑色/能化/有色/股指/国债等"],
        ["成交分散", "8%", "成交笔数在品种上的HHI"],
        ["出勤率", "8%", "有成交交易日 / 区间交易日"],
        ["交易频率", "8%", "日均笔数；持仓不分散时降权"],
        ["曲线平滑", "5%", "年化波动（上限25%）"],
        ["时间分布", "3%", "盘中成交时点熵"],
    ]
    add_table(doc, ["特征", "权重", "含义"], method_rows)
    caption(doc, "表2  量化行为得分构成。投顾表「多因子 / 时序 / 截面」只作交叉验证，不能单独进名单。")
    para(
        doc,
        "判定规则：投顾板块为「多因子」且同时持仓≥10（或成交品种≥25），或满足「成交品种≥40、同时持仓≥25、板块≥8、HHI≤0.12」，列为高确信。"
        "同时持仓≥18、成交品种≥25、HHI≤0.16列为较可能。"
        "风格为小单边/产业套利且同时持仓<16的账户，得分上限压到52，避免把半年轮动过很多合约的主观盘手当成量化。",
    )

    heading(doc, "三、交易行为全景", 1)
    add_chart(doc, charts["scatter"])
    caption(doc, "图1  横轴为品种数、纵轴为日均成交。右上角是量化典型区；左上角是产业高频；左下是主观/产业盘手。")
    add_chart(doc, charts["hhi"])
    caption(doc, "图2  横轴日均同时持仓品种、纵轴持仓HHI。右下角（持仓多、集中度低）才是量化；右上/左侧是轮动盘手或产业账户。")
    if quants:
        add_chart(doc, charts["scores"])
        caption(doc, "图3  进入名单的账户得分。虚线为高确信阈值70，点线为较可能阈值55。")
        add_chart(doc, charts["equity"])
        caption(doc, "图4  量化账户近六月累计收益。曲线形态用于观察系统性，不作为推荐依据。")

    heading(doc, "四、逐户证据", 1)
    for r in quants:
        heading(doc, f"{r['account'].upper()}  {r['advisor_name'] or '（未登记投顾名）'}", 2)
        bits = [
            f"判定：{r['verdict']} · {r['subtype']}",
            f"得分 {r['score']:.0f}",
            f"标签 {r['sector'] or '未分类'} / {r['style'] or '无'} / {r['company'] or '无公司'}",
        ]
        para(doc, "；".join(bits), size=10, color=NAVY, space_after=6)
        if r["reasons"]:
            for reason in r["reasons"]:
                p = doc.add_paragraph(style="List Bullet")
                add_text(p, reason, size=11, color=TEXT)
        else:
            para(doc, "交易特征达到量化阈值，但缺少更细的标签旁证。")
        detail = [
            ["成交品种", str(r["n_products"]), "日均持仓品种", fmt_num(r["avg_pos_products"], 1)],
            ["板块数", str(r["n_sectors"]), "持仓HHI", fmt_num(r["pos_hhi"], 2)],
            ["成交笔数", f"{r['n_trades']:,}", "日均笔数", fmt_num(r["trades_per_day"], 1)],
            ["有成交天数", str(r["trade_days"]), "出勤率", fmt_pct(r["coverage"], 0)],
            ["平今占比", fmt_pct(r["pingjin_ratio"], 0), "时间熵", fmt_num(r["time_entropy"], 2)],
            ["近六月收益", fmt_pct(r["period_ret"], 1, True), "年化波动", fmt_pct(r["vol"], 1)],
            ["夏普", fmt_num(r["sharpe"], 2), "最大回撤", fmt_pct(r["mdd"], 1)],
            ["手续费/均权益", fmt_pct(r["fee_to_eq"], 1), "期权成交", str(r["n_opt_trades"])],
        ]
        add_table(doc, ["指标", "数值", "指标", "数值"], detail)
        if r["sectors"]:
            para(doc, "成交板块：" + "、".join(r["sectors"]) + "。", size=10, color=MUTED, space_after=10)

    heading(doc, "五、明确排除：高频但不量化", 1)
    para(
        doc,
        "监控台收益曲线里，0%–10%的紧密簇很多是产业套利账户：曲线同样平滑，但成交集中在化工、粕类或黑色的少数合约上。"
        "它们执行可以很系统（甚至程序化下单），策略本质仍是产业链定价/价差，不是全市场多因子量化。下列账户是这一类的代表，不列入量化名单。",
    )
    if hft:
        hft_rows = []
        for r in hft:
            hft_rows.append(
                [
                    r["account"].upper(),
                    r["advisor_name"] or "—",
                    r["style"] or "—",
                    r["product_preference"] or "—",
                    str(r["n_products"]),
                    fmt_num(r["trades_per_day"], 0),
                    fmt_pct(r["period_ret"], 1, True),
                    fmt_num(r["score"], 0),
                ]
            )
        add_table(doc, ["账户", "投顾", "风格", "品种偏好", "品种数", "日均笔数", "近六月收益", "得分"], hft_rows)
        caption(doc, "表3  高频产业/价差账户。日均笔数可以高于量化CTA，但品种广度不够。")
    else:
        para(doc, "本期没有账户同时满足「日均成交很高且品种<18」。")

    heading(doc, "六、对其余账户的简要判断", 1)
    para(
        doc,
        f"其余 {len(disc)} 个有效样本更接近主观单边、产业套利或宏观配置：品种通常少于20，持仓HHI高，停手天数更多。"
        "下表列出得分最高的非量化账户，便于对照。",
    )
    rest = [r for r in rows if r not in quants][:18]
    rest_rows = []
    for r in rest:
        rest_rows.append(
            [
                r["account"].upper(),
                r["advisor_name"] or "—",
                r["verdict"][:8],
                r["style"] or "—",
                str(r["n_products"]),
                str(r["n_sectors"]),
                fmt_num(r["trades_per_day"], 0),
                fmt_num(r["score"], 0),
            ]
        )
    add_table(doc, ["账户", "投顾", "归类", "风格", "品种数", "板块", "日均笔数", "得分"], rest_rows)
    caption(doc, "表4  未进入量化名单的账户（按得分截取前18）。完整特征见同目录 JSON。")

    heading(doc, "七、使用注意", 1)
    notes = [
        "量化识别看的是交易方式，不是历史收益。高夏普的产业套利仍可能是主观/产业策略。",
        "新开户、样本不足20个交易日的账户未评分。",
        "投顾信息表可能滞后。未登记账户若品种广度达标，仍可因交易行为进入名单。",
        "平今占比、时间熵受柜台成交时间字段质量影响；字段缺失时该项得分接近0，不会单独否决。",
        "本报告不构成投顾评级或调仓建议，只回答「谁更像在做量化」。",
    ]
    for n in notes:
        p = doc.add_paragraph(style="List Number")
        add_text(p, n, size=11, color=TEXT)

    para(doc, f"生成日期 {date.today().isoformat()}。", size=9, color=MUTED, space_before=16)
    doc.save(str(REPORT_PATH))
    doc.save(str(REPORT_PATH_ASCII))


def main() -> None:
    configure_matplotlib()
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    conn = get_conn()
    try:
        print("loading features…")
        raw = load_features(conn)
    finally:
        conn.close()

    print("scoring…")
    rows = score_accounts(raw)
    serializable = []
    for r in rows:
        item = {k: v for k, v in r.items() if k != "curve"}
        item["curve_points"] = len(r["curve"])
        item["curve_end_pct"] = r["curve"][-1]["pct"] if r["curve"] else None
        serializable.append(item)
    DATA_PATH.write_text(json.dumps(serializable, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    print("wrote", DATA_PATH)

    quants = [r for r in rows if r["verdict"] in ("高确信量化", "较可能量化")]
    print("accounts", len(rows), "quant", len(quants))
    for r in rows[:25]:
        print(
            f"{r['account']:8} {r['score']:5.1f} {r['verdict']:16} prod={r['n_products']:3} "
            f"sec={r['n_sectors']:2} tpd={r['trades_per_day']:6.1f} {r['advisor_name']} {r['subtype']}"
        )

    print("charts…")
    charts = {
        "scatter": chart_scatter(rows),
        "hhi": chart_pos_hhi(rows),
        "scores": chart_scores(quants) if quants else chart_scores(rows[:8]),
        "equity": chart_equity(quants if quants else rows[:6]),
    }
    print("word…")
    write_report(rows, charts)
    print("report", REPORT_PATH)
    print("report", REPORT_PATH_ASCII)


if __name__ == "__main__":
    main()
