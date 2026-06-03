"""国信期货交易策略分析报告 — 从 PostgreSQL 数据库读取数据版本

用法：
    python generate_guoxin_word_report_db.py

环境变量（优先级高于默认值）：
    DATABASE_URL  postgresql://user:pass@host:port/dbname
"""
from __future__ import annotations

import io
import os
import sys
import traceback as _traceback
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

# Force UTF-8 stdout/stderr on Windows so Chinese chars in tracebacks don't
# cause a secondary UnicodeEncodeError that swallows the real error message.
if sys.platform == "win32":
    if hasattr(sys.stdout, "buffer"):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    if hasattr(sys.stderr, "buffer"):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)

# Patch requests to enforce a 30-second timeout on all akshare HTTP calls.
# Without this, a single hung HTTP connection stalls the whole script.
import pickle
import time

# Suppress tqdm progress bars so they don't pollute stdout when running
# headless (e.g., spawned by Node.js child_process).
import os as _os
_os.environ.setdefault("TQDM_DISABLE", "1")
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    import psycopg2
except ImportError:
    sys.exit("请先安装 psycopg2: pip install psycopg2-binary")

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "report_output"
CHART_DIR = OUTPUT_DIR / "charts"
REPORT_PATH = OUTPUT_DIR / "国信期货交易策略分析报告.docx"

# ---------------------------------------------------------------------------
# Constants (same as CSV version)
# ---------------------------------------------------------------------------

PRODUCT_MARKET = {
    # GFEX
    "碳酸锂": "GFEX",
    "多晶硅": "GFEX",
    # CZCE
    "玻璃": "CZCE",
    "花生仁": "CZCE",
    # DCE
    "鲜鸡蛋": "DCE",
    "生猪": "DCE",
    "聚氯乙烯": "DCE",
    # SHFE
    "铝": "SHFE",
    "不锈钢": "SHFE",
    "石油沥青": "SHFE",
}

PRODUCT_NAME_ORDER = ["碳酸锂", "铝", "不锈钢", "鲜鸡蛋", "玻璃"]

DISPLAY_NAME = {
    "碳酸锂": "碳酸锂",
    "铝": "铝",
    "不锈钢": "不锈钢",
    "鲜鸡蛋": "鲜鸡蛋",
    "玻璃": "玻璃",
}

SPREAD_DEFS = {
    "玻璃 FG609-FG605": ("FG609", "FG605"),
    "玻璃 FG609-FG701": ("FG609", "FG701"),
    "鸡蛋 JD2607-JD2606": ("JD2607", "JD2606"),
    "铝 AL2605-AL2606": ("AL2605", "AL2606"),
    "不锈钢 SS2606-SS2607": ("SS2606", "SS2607"),
    "碳酸锂 LC2607-LC2609": ("LC2607", "LC2609"),
    "碳酸锂 LC2608-LC2609": ("LC2608", "LC2609"),
    "碳酸锂 LC2610-LC2609": ("LC2610", "LC2609"),
}

PAIR_EVENT_MAP = {
    "玻璃 FG609-FG605": {"entry": "2026-04-17", "exit": "2026-04-20"},
    "玻璃 FG609-FG701": {"entry": "2026-04-21", "exit": None},
    "鸡蛋 JD2607-JD2606": {"entry": "2026-04-21", "exit": "2026-04-22"},
    "铝 AL2605-AL2606": {"entry": "2026-04-23", "exit": None},
    "不锈钢 SS2606-SS2607": {"entry": "2026-04-23", "exit": None},
    "碳酸锂 LC2607-LC2609": {"entry": "2026-04-17", "exit": None},
    "碳酸锂 LC2608-LC2609": {"entry": "2026-04-17", "exit": None},
    "碳酸锂 LC2610-LC2609": {"entry": "2026-04-17", "exit": "2026-04-17"},
}


@dataclass
class AnalysisResult:
    account: pd.DataFrame
    trades: pd.DataFrame
    positions: pd.DataFrame
    closed: pd.DataFrame
    market_hist: pd.DataFrame
    trade_mkt: pd.DataFrame
    equity_stats: dict
    turnover_stats: pd.DataFrame
    product_netting: pd.DataFrame
    execution_stats: pd.DataFrame
    spread_stats: pd.DataFrame
    trade_clusters: dict
    close_clusters: dict
    chart_paths: dict[str, Path]


def ensure_output_dirs() -> None:
    CHART_DIR.mkdir(parents=True, exist_ok=True)


def configure_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Noto Sans CJK SC",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


# ---------------------------------------------------------------------------
# Load data from PostgreSQL (replaces load_data from CSV version)
# ---------------------------------------------------------------------------

def load_data_from_db() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    db_url = os.environ.get(
        "DATABASE_URL",
        "postgresql://market_user:2026SmartDashboard!@127.0.0.1:5433/market_data",
    )
    print(f"[DB] Connecting to database…")
    conn = psycopg2.connect(db_url, connect_timeout=8)

    def query_df(sql: str) -> pd.DataFrame:
        with conn.cursor() as cur:
            cur.execute(sql)
            cols = [desc[0] for desc in cur.description]
            return pd.DataFrame(cur.fetchall(), columns=cols)

    account = query_df("SELECT * FROM guosen_account_summary ORDER BY trade_date")
    trades = query_df(
        "SELECT * FROM guosen_transaction_records ORDER BY trade_date, settlement_date, row_num"
    )
    positions = query_df(
        "SELECT * FROM guosen_position_summary ORDER BY settlement_date, product, instrument"
    )
    closed = query_df(
        "SELECT * FROM guosen_position_closed ORDER BY settlement_date, product, instrument"
    )
    conn.close()
    print(f"[DB] Loaded: {len(account)} account rows, {len(trades)} trade rows, "
          f"{len(positions)} position rows, {len(closed)} closed rows")

    account["trade_date"] = pd.to_datetime(account["trade_date"])
    trades["trade_date"] = pd.to_datetime(trades["trade_date"])
    trades["settlement_date"] = pd.to_datetime(trades["settlement_date"])
    positions["settlement_date"] = pd.to_datetime(positions["settlement_date"])
    closed["settlement_date"] = pd.to_datetime(closed["settlement_date"])

    # Ensure premium_rp column exists (may be absent in some DB schemas)
    if "premium_rp" not in trades.columns:
        trades["premium_rp"] = 0.0

    numeric_sets = [
        (
            account,
            [
                "balance_bf", "realized_pl", "mtm_pl", "commission",
                "exercise_fee", "delivery_fee", "client_equity",
                "margin_occupied", "fund_avail", "risk_degree", "balance_cf",
            ],
        ),
        (trades, ["price", "lots", "turnover", "fee", "realized_pl", "premium_rp"]),
        (
            positions,
            [
                "long_pos", "short_pos", "avg_buy_price", "avg_sell_price",
                "prev_settl", "settl_today", "mtm_pl", "margin_occupied",
            ],
        ),
        (closed, ["lots", "pos_open_price", "prev_settl", "trans_price", "realized_pl"]),
    ]
    for frame, cols in numeric_sets:
        for col in cols:
            if col in frame.columns:
                frame[col] = pd.to_numeric(frame[col], errors="coerce")

    return account, trades, positions, closed


# ---------------------------------------------------------------------------
# Market-data cache helpers
# ---------------------------------------------------------------------------

_CACHE_PATH = BASE_DIR / "market_data_cache.pkl"
_CACHE_MAX_AGE_HOURS = 12


def _load_market_cache(start_date: str, end_date: str) -> pd.DataFrame | None:
    if not _CACHE_PATH.exists():
        return None
    try:
        with open(_CACHE_PATH, "rb") as _f:
            cached = pickle.load(_f)
        age_h = (time.time() - cached.get("ts", 0)) / 3600
        if age_h > _CACHE_MAX_AGE_HOURS:
            print(f"[CACHE] Cache expired ({age_h:.1f}h old), will refresh.", flush=True)
            return None
        if cached.get("start_date") == start_date and cached.get("end_date") == end_date:
            print(f"[CACHE] Using cached market data ({age_h:.1f}h old).", flush=True)
            return cached["data"]
    except Exception as _e:
        print(f"[CACHE] Could not load cache: {_e}", flush=True)
    return None


def _save_market_cache(df: pd.DataFrame, start_date: str, end_date: str) -> None:
    try:
        with open(_CACHE_PATH, "wb") as _f:
            pickle.dump({"ts": time.time(), "start_date": start_date, "end_date": end_date, "data": df}, _f)
        print("[CACHE] Market data saved to disk cache.", flush=True)
    except Exception as _e:
        print(f"[CACHE] Could not save cache: {_e}", flush=True)


# ---------------------------------------------------------------------------
# Remaining functions identical to the CSV version
# ---------------------------------------------------------------------------

import re as _re
_CZCE_3DIGIT = _re.compile(r'^([A-Z]+)([0-9])(\d{2})$')


def _to_canonical(sym: str) -> str:
    """Normalize CZCE 3-digit contract codes to 4-digit format used in DB.

    CZCE traditionally encodes contracts as e.g. FG605 (product=FG, year=6, month=05).
    The database stores them as FG2605.  All other exchanges already use 4-digit
    year format (AL2605, JD2606, LC2606 …) so they pass through unchanged.
    """
    s = sym.upper()
    m = _CZCE_3DIGIT.match(s)
    if m:
        product, year_digit, month = m.groups()
        # 2-digit year: digit 5-9 → "2X" (2025-2029), digit 0-4 → "3X" (2030-2034)
        decade = "2" if int(year_digit) >= 5 else "3"
        return f"{product}{decade}{year_digit}{month}"
    return s


def fetch_market_history(trades: pd.DataFrame) -> pd.DataFrame:
    # Exclude options contracts (product names containing '期权', or instrument
    # codes containing '-P-' / '-C-' for puts/calls).
    is_option = (
        trades["product"].str.contains("期权", na=False)
        | trades["instrument"].str.contains(r"-[PC]-", na=False, regex=True)
    )
    futures_trades = trades[~is_option]
    relevant = futures_trades[["product", "instrument"]].drop_duplicates().copy()
    # Canonicalize: convert CZCE 3-digit codes (FG605) to 4-digit (FG2605) so
    # they match the DB column raw_futures_contracts_daily.contract.
    relevant["symbol"] = relevant["instrument"].apply(_to_canonical)

    start_date = (trades["trade_date"].min() - pd.Timedelta(days=20)).strftime("%Y%m%d")
    end_date = trades["trade_date"].max().strftime("%Y%m%d")

    cached = _load_market_cache(start_date, end_date)
    if cached is not None:
        return cached

    needed_symbols = relevant["symbol"].unique().tolist()
    start_dt = pd.to_datetime(start_date).date()
    end_dt = pd.to_datetime(end_date).date()

    print(f"[DB] Fetching market data for {len(needed_symbols)} symbols from raw_futures_contracts_daily…", flush=True)

    db_url = os.environ.get(
        "DATABASE_URL",
        "postgresql://market_user:2026SmartDashboard!@127.0.0.1:5433/market_data",
    )
    conn = psycopg2.connect(db_url, connect_timeout=8)
    try:
        with conn.cursor() as cur:
            # DISTINCT ON deduplicates (trade_date, symbol) keeping the row with
            # an exchange suffix (e.g. AL2606.SHF over AL2606) when both exist.
            cur.execute(
                """
                SELECT DISTINCT ON (trade_date, sym)
                    trade_date,
                    UPPER(SPLIT_PART(contract, '.', 1)) AS sym,
                    open, high, low, close,
                    clear        AS settle,
                    preclear     AS pre_settle,
                    volume,
                    hqoi         AS open_interest
                FROM raw_futures_contracts_daily
                WHERE trade_date BETWEEN %s AND %s
                  AND UPPER(SPLIT_PART(contract, '.', 1)) = ANY(%s)
                  AND clear IS NOT NULL
                ORDER BY trade_date, sym,
                    CASE WHEN contract LIKE '%%.%%' THEN 0 ELSE 1 END,
                    contract
                """,
                (start_dt, end_dt, needed_symbols),
            )
            cols = [desc[0] for desc in cur.description]
            rows = cur.fetchall()
    finally:
        conn.close()

    if not rows:
        print("[DB] No market data found in DB.", flush=True)
        return pd.DataFrame()

    df = pd.DataFrame(rows, columns=cols)
    df = df.rename(columns={"sym": "symbol", "trade_date": "date"})
    df["date"] = pd.to_datetime(df["date"])
    df["symbol"] = df["symbol"].astype(str)
    for col in ["open", "high", "low", "close", "settle", "pre_settle", "volume", "open_interest"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    found_symbols = set(df["symbol"].unique())
    missing = set(needed_symbols) - found_symbols
    if missing:
        print(f"[DB] Symbols not found in DB (settle data unavailable): {sorted(missing)}", flush=True)
    else:
        print(f"[DB] All {len(needed_symbols)} symbols found in DB ({len(df)} rows).", flush=True)

    market_hist = df.sort_values(["symbol", "date"]).reset_index(drop=True)
    _save_market_cache(market_hist, start_date, end_date)
    return market_hist


def build_trade_market_table(trades: pd.DataFrame, market_hist: pd.DataFrame) -> pd.DataFrame:
    merged = trades.copy()
    # Use canonical 4-digit format so CZCE contracts like FG605 → FG2605 match
    # the symbols stored in market_hist (which come from the DB in 4-digit form).
    merged["symbol"] = merged["instrument"].apply(_to_canonical)
    merged = merged.merge(
        market_hist[["date", "symbol", "open", "high", "low", "close", "settle", "pre_settle", "volume", "open_interest"]],
        left_on=["trade_date", "symbol"],
        right_on=["date", "symbol"],
        how="left",
    )
    merged["range_pct"] = np.where(
        (merged["high"] - merged["low"]).abs() > 0,
        (merged["price"] - merged["low"]) / (merged["high"] - merged["low"]),
        np.nan,
    )
    merged["good_exec_vs_settle"] = np.where(
        merged["bs"] == "买",
        merged["settle"] - merged["price"],
        merged["price"] - merged["settle"],
    )
    merged["good_exec_vs_close"] = np.where(
        merged["bs"] == "买",
        merged["close"] - merged["price"],
        merged["price"] - merged["close"],
    )
    return merged


def summarize_equity(account: pd.DataFrame) -> dict:
    ordered = account.sort_values("trade_date").reset_index(drop=True)
    first = ordered.iloc[0]
    last = ordered.iloc[-1]
    fees_total = ordered[["commission", "exercise_fee", "delivery_fee"]].fillna(0).sum().sum()
    return {
        "start_date": first["trade_date"],
        "end_date": last["trade_date"],
        "equity_start": float(first["client_equity"]),
        "equity_end": float(last["client_equity"]),
        "net_change": float(last["client_equity"] - first["client_equity"]),
        "return_pct": float((last["client_equity"] - first["client_equity"]) / first["client_equity"] * 100),
        "realized_total": float(ordered["realized_pl"].sum()),
        "mtm_total": float(ordered["mtm_pl"].sum()),
        "fees_total": float(fees_total),
        "max_risk_degree": float(ordered["risk_degree"].max()),
        "max_margin_occupied": float(ordered["margin_occupied"].max()),
    }


def summarize_turnover(trades: pd.DataFrame) -> pd.DataFrame:
    # Use row count for "trades" column (compatible with both id and row_num schemas)
    count_col = "id" if "id" in trades.columns else "row_num" if "row_num" in trades.columns else "lots"
    turnover = trades.groupby("product", as_index=False).agg(
        turnover=("turnover", "sum"),
        lots=("lots", "sum"),
        trades=(count_col, "count"),
    )
    turnover["share_pct"] = turnover["turnover"] / turnover["turnover"].sum() * 100
    turnover["product"] = turnover["product"].map(lambda value: DISPLAY_NAME.get(value, value))
    return turnover.sort_values("turnover", ascending=False).reset_index(drop=True)


def summarize_product_netting(positions: pd.DataFrame) -> pd.DataFrame:
    netting = positions.groupby(["settlement_date", "product"], as_index=False).agg(
        long_lots=("long_pos", "sum"),
        short_lots=("short_pos", "sum"),
        mtm=("mtm_pl", "sum"),
        margin=("margin_occupied", "sum"),
    )
    netting["net_lots"] = netting["long_lots"] - netting["short_lots"]
    netting["product"] = netting["product"].map(lambda value: DISPLAY_NAME.get(value, value))
    return netting.sort_values(["settlement_date", "product"]).reset_index(drop=True)


def summarize_execution(trade_mkt: pd.DataFrame) -> pd.DataFrame:
    grouped = trade_mkt.groupby("product").apply(
        lambda group: pd.Series(
            {
                "交易笔数": len(group),
                "成交手数": int(group["lots"].sum()),
                "买入均值区间位置": group.loc[group["bs"] == "买", "range_pct"].mean(),
                "卖出均值区间位置": group.loc[group["bs"] == "卖", "range_pct"].mean(),
                "相对结算价执行优势": group["good_exec_vs_settle"].mean(),
                "相对收盘价执行优势": group["good_exec_vs_close"].mean(),
            }
        )
    ).reset_index()
    grouped["product"] = grouped["product"].map(lambda value: DISPLAY_NAME.get(value, value))
    return grouped.sort_values("交易笔数", ascending=False).reset_index(drop=True)


def summarize_spreads(market_hist: pd.DataFrame) -> pd.DataFrame:
    pivot = market_hist.pivot_table(index="date", columns="symbol", values="settle", aggfunc="last").sort_index()
    rows = []
    for name, (leg_a, leg_b) in SPREAD_DEFS.items():
        if leg_a not in pivot.columns or leg_b not in pivot.columns:
            continue
        spread = pivot[leg_a] - pivot[leg_b]
        z20 = (spread - spread.rolling(20, min_periods=5).mean()) / spread.rolling(20, min_periods=5).std()
        event = PAIR_EVENT_MAP.get(name, {})
        entry_date = pd.Timestamp(event["entry"]) if event.get("entry") else None
        exit_date = pd.Timestamp(event["exit"]) if event.get("exit") else None
        rows.append(
            {
                "spread_name": name,
                "entry_date": entry_date,
                "exit_date": exit_date,
                "entry_spread": float(spread.loc[entry_date]) if entry_date in spread.index else np.nan,
                "entry_z20": float(z20.loc[entry_date]) if entry_date in z20.index else np.nan,
                "exit_spread": float(spread.loc[exit_date]) if exit_date and exit_date in spread.index else np.nan,
                "exit_z20": float(z20.loc[exit_date]) if exit_date and exit_date in z20.index else np.nan,
                "spread_mean": float(spread.mean()),
                "spread_std": float(spread.std()),
                "spread_min": float(spread.min()),
                "spread_max": float(spread.max()),
            }
        )
    return pd.DataFrame(rows)


def build_trade_clusters(trade_mkt: pd.DataFrame) -> dict:
    clusters: dict = {}
    sort_col = "row_num" if "row_num" in trade_mkt.columns else "trade_date"
    open_trades = trade_mkt[trade_mkt["oc"] == "开"].sort_values(["trade_date", "product", sort_col])
    for (trade_date, product), group in open_trades.groupby(["trade_date", "product"]):
        agg = group.groupby(["instrument", "bs"], as_index=False).agg(
            lots=("lots", "sum"),
            avg_price=("price", "mean"),
            avg_range_pct=("range_pct", "mean"),
            avg_exec_vs_settle=("good_exec_vs_settle", "mean"),
        )
        clusters[(trade_date, product)] = {
            "summary": agg,
            "turnover": float(group["turnover"].sum()),
            "fees": float(group["fee"].sum()),
        }
    return clusters


def build_close_clusters(closed: pd.DataFrame) -> dict:
    clusters: dict = {}
    for (settlement_date, product), group in closed.groupby(["settlement_date", "product"]):
        agg = group.groupby(["instrument", "bs"], as_index=False).agg(
            lots=("lots", "sum"),
            realized=("realized_pl", "sum"),
        )
        clusters[(settlement_date, product)] = {
            "summary": agg,
            "realized": float(group["realized_pl"].sum()),
            "lots": float(group["lots"].sum()),
        }
    return clusters


def format_percent(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}%"


def format_number(value: float, digits: int = 2) -> str:
    return f"{value:,.{digits}f}"


def save_chart(fig: plt.Figure, name: str) -> Path:
    path = CHART_DIR / name
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_charts(
    account: pd.DataFrame,
    turnover_stats: pd.DataFrame,
    product_netting: pd.DataFrame,
    spread_stats: pd.DataFrame,
    market_hist: pd.DataFrame,
) -> dict[str, Path]:
    chart_paths: dict[str, Path] = {}

    ordered_account = account.sort_values("trade_date")
    fig, ax1 = plt.subplots(figsize=(10, 5))
    ax1.plot(ordered_account["trade_date"], ordered_account["client_equity"], marker="o", linewidth=2, color="#0f5c5e")
    ax1.set_title("账户权益与风险度变化")
    ax1.set_ylabel("客户权益")
    ax1.grid(alpha=0.25)
    ax2 = ax1.twinx()
    ax2.plot(ordered_account["trade_date"], ordered_account["risk_degree"], marker="s", linestyle="--", color="#c84c09")
    ax2.set_ylabel("风险度(%)")
    chart_paths["equity"] = save_chart(fig, "equity_risk.png")

    fig, ax = plt.subplots(figsize=(9, 5))
    colors = ["#204e4a", "#2f6f6d", "#4e9689", "#9fbf7b", "#d5a021"]
    ax.bar(turnover_stats["product"], turnover_stats["share_pct"], color=colors[: len(turnover_stats)])
    ax.set_title("各品种成交额占比")
    ax.set_ylabel("占比(%)")
    ax.grid(axis="y", alpha=0.2)
    chart_paths["turnover"] = save_chart(fig, "turnover_share.png")

    margin_by_date = product_netting.groupby("settlement_date", as_index=False)["margin"].sum()
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(margin_by_date["settlement_date"], margin_by_date["margin"], color="#46637f")
    ax.set_title("每日保证金占用")
    ax.set_ylabel("保证金")
    ax.grid(axis="y", alpha=0.2)
    chart_paths["margin"] = save_chart(fig, "margin_usage.png")

    pivot = market_hist.pivot_table(index="date", columns="symbol", values="settle", aggfunc="last").sort_index()
    key_spreads = ["碳酸锂 LC2607-LC2609", "鸡蛋 JD2607-JD2606", "铝 AL2605-AL2606"]
    for spread_name in key_spreads:
        leg_a, leg_b = SPREAD_DEFS[spread_name]
        if leg_a not in pivot.columns or leg_b not in pivot.columns:
            continue
        spread = pivot[leg_a] - pivot[leg_b]
        z20 = (spread - spread.rolling(20, min_periods=5).mean()) / spread.rolling(20, min_periods=5).std()
        fig, axes = plt.subplots(2, 1, figsize=(10, 7), sharex=True)
        axes[0].plot(spread.index, spread, color="#0f5c5e", linewidth=2)
        axes[0].set_title(f"{spread_name} 价差与20日Z分数")
        axes[0].set_ylabel("价差")
        axes[0].grid(alpha=0.2)
        axes[1].plot(z20.index, z20, color="#c84c09", linewidth=1.8)
        axes[1].axhline(-1.5, color="#666666", linestyle="--", linewidth=1)
        axes[1].axhline(0, color="#666666", linestyle=":", linewidth=1)
        axes[1].set_ylabel("Z20")
        axes[1].grid(alpha=0.2)
        event = PAIR_EVENT_MAP[spread_name]
        entry_dt = pd.Timestamp(event["entry"])
        for axis in axes:
            if entry_dt in spread.index:
                axis.axvline(entry_dt, color="#1d7a52", linestyle="--", linewidth=1.2)
            if event["exit"]:
                exit_dt = pd.Timestamp(event["exit"])
                if exit_dt in spread.index:
                    axis.axvline(exit_dt, color="#ad2e24", linestyle="--", linewidth=1.2)
        chart_paths[f"spread_{leg_a}_{leg_b}"] = save_chart(fig, f"spread_{leg_a}_{leg_b}.png")

    return chart_paths


def compose_analysis(
    account: pd.DataFrame,
    trades: pd.DataFrame,
    positions: pd.DataFrame,
    closed: pd.DataFrame,
    market_hist: pd.DataFrame,
    trade_mkt: pd.DataFrame,
) -> AnalysisResult:
    equity_stats = summarize_equity(account)
    turnover_stats = summarize_turnover(trades)
    product_netting = summarize_product_netting(positions)
    execution_stats = summarize_execution(trade_mkt)
    spread_stats = summarize_spreads(market_hist)
    trade_clusters = build_trade_clusters(trade_mkt)
    close_clusters = build_close_clusters(closed)
    chart_paths = build_charts(account, turnover_stats, product_netting, spread_stats, market_hist)
    return AnalysisResult(
        account=account,
        trades=trades,
        positions=positions,
        closed=closed,
        market_hist=market_hist,
        trade_mkt=trade_mkt,
        equity_stats=equity_stats,
        turnover_stats=turnover_stats,
        product_netting=product_netting,
        execution_stats=execution_stats,
        spread_stats=spread_stats,
        trade_clusters=trade_clusters,
        close_clusters=close_clusters,
        chart_paths=chart_paths,
    )


def set_document_style(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(document: Document, title: str, subtitle: str) -> None:
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(subtitle)
    run.font.size = Pt(10.5)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, dataframe: pd.DataFrame, columns: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, (header, _) in enumerate(columns):
        hdr[idx].text = header
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, (_, key) in enumerate(columns):
            value = row[key]
            if isinstance(value, pd.Timestamp):
                cells[idx].text = value.strftime("%Y-%m-%d")
            elif pd.isna(value):
                cells[idx].text = ""
            elif isinstance(value, (float, np.floating)):
                cells[idx].text = f"{value:,.2f}"
            else:
                cells[idx].text = str(value)


def write_report(result: AnalysisResult) -> None:
    document = Document()
    set_document_style(document)
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    add_title(
        document,
        "国信期货账户交易策略分析报告",
        f"样本区间：{result.equity_stats['start_date']:%Y-%m-%d} 至 {result.equity_stats['end_date']:%Y-%m-%d}",
    )

    document.add_heading("一、结论摘要", level=1)
    add_bullets(
        document,
        [
            "该账户的核心策略属于商品期货跨期价差相对价值交易，整体不是单边趋势型 CTA。",
            "所有样本日的品种级净持仓均为 0，说明交易者严格保持品种内多空对冲，主要赚取期限结构错配回归。",
            f"碳酸锂是主导仓位，成交额占比 {result.turnover_stats.loc[result.turnover_stats['product'] == '碳酸锂', 'share_pct'].iloc[0]:.2f}%，同时也是主要波动与保证金来源。",
            "碳酸锂结构并非普通两腿跨期，而是以 lc2609 为核心空头、对冲 lc2607/lc2608/lc2610 的曲线形状交易，可视为 short belly / long wings 风格。",
            "玻璃、鸡蛋、铝、不锈钢则更接近标准日历价差，入场点普遍出现在 20 日价差 Z 分数的负偏离区域，符合均值回复交易特征。",
        ],
    )

    document.add_heading("二、账户收益与风险特征", level=1)
    eq = result.equity_stats
    add_bullets(
        document,
        [
            f"期初权益：{format_number(eq['equity_start'])}；期末权益：{format_number(eq['equity_end'])}；净变动：{format_number(eq['net_change'])}。",
            f"样本收益率：{format_percent(eq['return_pct'])}；累计平仓盈亏：{format_number(eq['realized_total'])}；累计盯市盈亏：{format_number(eq['mtm_total'])}。",
            f"累计手续费与交割相关费用：{format_number(eq['fees_total'])}；最大风险度：{format_percent(eq['max_risk_degree'])}；最大保证金占用：{format_number(eq['max_margin_occupied'])}。",
            "收益表现为小幅正收益，但资金占用较高，说明该策略更像低方向暴露、靠结构性错配获取稳定边际收益的资金效率型策略。",
        ],
    )
    document.add_picture(str(result.chart_paths["equity"]), width=Inches(6.7))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_picture(str(result.chart_paths["margin"]), width=Inches(6.7))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("三、品种结构与资金集中度", level=1)
    add_bullets(
        document,
        [
            "成交额集中在碳酸锂、铝和不锈钢，其中碳酸锂一项就超过总成交额的一半，是账户的风险锚点。",
            "逐日持仓显示所有品种 long lots = short lots，说明交易目标不是方向性行情，而是月间价差和曲线结构。",
            "玻璃与鸡蛋的持仓规模较小，更像短周期战术仓位；铝和不锈钢则偏向标准近远月均值回复；碳酸锂则属于主策略仓位。",
        ],
    )
    add_table(
        document,
        result.turnover_stats,
        [("品种", "product"), ("成交额", "turnover"), ("成交手数", "lots"), ("成交笔数", "trades"), ("占比(%)", "share_pct")],
    )
    document.add_paragraph()
    document.add_picture(str(result.chart_paths["turnover"]), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("四、交易策略类别判断", level=1)
    add_bullets(
        document,
        [
            "主类别：商品期货跨期价差交易（calendar spread / term structure relative value）。",
            "细分类别：均值回复型相对价值策略，而非趋势跟踪或事件驱动。",
            "碳酸锂子策略：short belly / long wings 的期限结构曲线交易，实质上是在做 2609 合约相对周边月份偏贵时的回落。",
            "执行层面偏向离散批次建仓、日内择价，而非高频撮合；交易记录中同日双腿开仓明显，说明其下单逻辑以配对完成为核心。",
        ],
    )

    document.add_heading("五、入场与离场逻辑推断", level=1)
    add_bullets(
        document,
        [
            "入场逻辑大概率采用价差偏离度筛选。多个已知开仓点都落在 20 日价差 Z 分数显著偏离区域，尤其铝 AL2605-AL2606 在 -2.35 左右开仓，鸡蛋 JD2607-JD2606 在 -1.53 左右开仓。",
            "离场逻辑更像部分均值回归止盈，而非等待完全回到历史均值。玻璃 FG609-FG605 从 103 扩大到 115 后即在 Z 分数由 -1.28 收敛到 -0.59 时平仓，鸡蛋价差也表现出类似行为。",
            "碳酸锂仓位存在加仓痕迹。2026-04-21 又增加 lc2607 多头和 lc2609 空头，说明当核心偏离持续存在时，交易者愿意顺着原有结构加码，而不是立即止损。",
            "从执行区间位置看，多数买单在当日价格区间中下部、卖单在中上部，说明交易者并非无条件吃单，而是对配对腿的成交质量有明确要求。",
        ],
    )

    document.add_heading("六、关键价差统计", level=1)
    spread_table = result.spread_stats.copy()
    spread_table["entry_date"] = spread_table["entry_date"].dt.strftime("%Y-%m-%d")
    spread_table["exit_date"] = spread_table["exit_date"].dt.strftime("%Y-%m-%d")
    add_table(
        document,
        spread_table,
        [
            ("价差", "spread_name"),
            ("开仓日", "entry_date"),
            ("开仓价差", "entry_spread"),
            ("开仓Z20", "entry_z20"),
            ("平仓日", "exit_date"),
            ("平仓价差", "exit_spread"),
            ("平仓Z20", "exit_z20"),
        ],
    )
    for chart_key in ["spread_LC2607_LC2609", "spread_JD2607_JD2606", "spread_AL2605_AL2606"]:
        if chart_key in result.chart_paths:
            document.add_paragraph()
            document.add_picture(str(result.chart_paths[chart_key]), width=Inches(6.5))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("七、开平仓批次解读", level=1)
    for (trade_date, product), cluster in sorted(result.trade_clusters.items(), key=lambda item: (item[0][0], item[0][1])):
        product_name = DISPLAY_NAME.get(product, product)
        document.add_heading(f"{trade_date:%Y-%m-%d} {product_name} 开仓批次", level=2)
        lines = []
        for _, row in cluster["summary"].iterrows():
            side = row["bs"]
            lines.append(
                f"{row['instrument']} {side} {int(row['lots'])} 手，均价 {row['avg_price']:.2f}，"
                f"日内区间位置 {row['avg_range_pct']:.3f}，相对结算价执行优势 {row['avg_exec_vs_settle']:.2f}。"
            )
        lines.append(f"该批次合计成交额 {format_number(cluster['turnover'])}，手续费 {format_number(cluster['fees'])}。")
        add_bullets(document, lines)

    for (settlement_date, product), cluster in sorted(result.close_clusters.items(), key=lambda item: (item[0][0], item[0][1])):
        product_name = DISPLAY_NAME.get(product, product)
        document.add_heading(f"{settlement_date:%Y-%m-%d} {product_name} 平仓批次", level=2)
        lines = [f"批次实现盈亏 {format_number(cluster['realized'])}，平仓手数 {cluster['lots']:.0f}。"]
        for _, row in cluster["summary"].iterrows():
            lines.append(f"{row['instrument']} {row['bs']} {int(row['lots'])} 手，实现盈亏 {format_number(row['realized'])}。")
        add_bullets(document, lines)

    document.add_heading("八、策略优缺点与后续建议", level=1)
    add_bullets(
        document,
        [
            "优点：方向暴露低、持仓结构清晰、入场点普遍落在历史偏离区域，具备较强的相对价值纪律。",
            "风险：碳酸锂仓位集中度高，且中间月份短空可能在曲线扭转时带来较大盯市波动；收益率相对保证金占用偏低。",
            "建议：继续记录分钟级下单时间、价差指令标记和止盈止损阈值，这样可以进一步验证是否存在固定 Z 分数阈值、时间止盈或波动率过滤机制。",
            "若后续要扩展报告，可加入更长样本区间、各价差滚动夏普、持仓天数分布和按品种的盈亏归因。",
        ],
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("附录：逐日品种净敞口汇总", level=1)
    add_table(
        document,
        result.product_netting,
        [
            ("日期", "settlement_date"),
            ("品种", "product"),
            ("多头", "long_lots"),
            ("空头", "short_lots"),
            ("净手数", "net_lots"),
            ("盯市盈亏", "mtm"),
            ("保证金", "margin"),
        ],
    )

    document.save(REPORT_PATH)
    print(f"[OK] Word 报告已生成: {REPORT_PATH}")


def main() -> None:
    ensure_output_dirs()
    configure_matplotlib()
    account, trades, positions, closed = load_data_from_db()
    print("[INFO] Fetching market data from database…")
    market_hist = fetch_market_history(trades)
    trade_mkt = build_trade_market_table(trades, market_hist)
    if trade_mkt["settle"].isna().any():
        missing = trade_mkt.loc[trade_mkt["settle"].isna(), "instrument"].dropna().unique().tolist()
        print(f"[WARN] 以下合约未能匹配到行情数据，以成交价替代结算价: {missing}", flush=True)
        # Fall back to trade price so downstream analysis doesn't break on NaN.
        trade_mkt["settle"] = trade_mkt["settle"].fillna(trade_mkt["price"])
    print("[INFO] Running analysis…")
    result = compose_analysis(account, trades, positions, closed, market_hist, trade_mkt)
    write_report(result)
    print(f"[INFO] 图表输出目录: {CHART_DIR}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"\n[FATAL] {type(exc).__name__}: {exc}", flush=True)
        _traceback.print_exc(file=sys.stdout)
        sys.stdout.flush()
        sys.exit(1)
